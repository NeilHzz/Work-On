from __future__ import annotations

from pathlib import Path

from docx import Document


BASE = Path.cwd()
WORK_DIR = next(p for p in BASE.iterdir() if p.is_dir() and p.name.startswith("03_"))
EN_DOCX = WORK_DIR / "0_Manuscript" / "manuscript260608v2.docx"
CN_DOCX = WORK_DIR / "0_Manuscript_CN" / "manuscript260608v2.docx"

EN_INTRO = (
    "An eggshell has to be strong from the outside but breakable from the inside. "
    "During incubation, it protects the developing embryo, maintains a controlled biological compartment, supplies calcium for embryonic growth, and provides the mineralized space in which development proceeds. "
    "Avian eggshell formation is one of the fastest biomineralization processes in nature, assembling a mechanically competent calcitic shell within a narrow uterine time window. (1–3) "
    "This rapid mineralization depends on eggshell matrix proteins, which coordinate calcium delivery, nucleation, crystal growth, and shell architecture rather than acting as passive scaffolds. (1, 2, 4, 5) "
    "Hatching then reverses the mechanical problem: the shell that protected the embryo must be locally fractured from within. "
    "The egg tooth loads a restricted region of the inner shell rather than the whole shell at once, and comparable hatching-assist structures recur across egg-laying amniotes. (6–11) "
    "Biologically meaningful variation in hatching performance may therefore reside not only in the hatching tool, but also in the shell material that the tool acts on. "
    "Yet it remains unclear how species with comparable hatching mechanics generate distinct shell states across ecological and developmental contexts. (12–16) "
    "The mammillary layer provides the key entry point because matrix-guided calcite nucleation first sets the spacing and continuity inherited by later shell units, and because this layer is the first mechanically consequential shell region encountered during inside-out hatching. "
    "The mechanistic question is therefore direct: when the shell-breaking interface is comparable, which molecular regulators at the mammillary layer account for distinct eggshell states across species? (1, 2, 4, 6, 17, 18)"
)

CN_INTRO = (
    "蛋壳必须能够从外部提供足够强度，又必须能从内部被胚胎局部打破。"
    "在孵化过程中，蛋壳保护发育中的胚胎，维持受控的生物安全隔室，为胚胎生长提供钙来源，并构成胚胎发育所依赖的矿化空间。"
    "鸟类蛋壳形成是自然界最快速的生物矿化过程之一，需要在狭窄的子宫时间窗口内组装出具有力学能力的方解石壳。 (1–3) "
    "这一快速矿化过程依赖蛋壳基质蛋白，这些蛋白协调钙递送、成核、晶体生长和壳结构形成，而不是作为被动支架存在。 (1, 2, 4, 5) "
    "孵化随后反转了这一力学问题：曾经保护胚胎的蛋壳必须从内部发生局部破裂。"
    "卵齿作用于内壳的有限区域，而不是一次性加载整个蛋壳；类似的孵化辅助结构也反复出现在产卵羊膜动物中。 (6–11) "
    "因此，孵化性能的生物学差异不仅可能存在于破壳工具本身，也可能存在于该工具所作用的蛋壳材料中。"
    "然而，具有可比孵化机制的物种如何在不同生态和发育背景下形成不同蛋壳状态，仍不清楚。 (12–16) "
    "乳突层提供了关键入口，因为基质引导的方解石成核首先在此建立后续壳单元继承的间距和连续性，同时它也是内向外孵化过程中首先遇到的具有力学后果的壳层。"
    "因此，机制问题是：当破壳界面具有可比性时，乳突层中的哪些分子调控因子解释了跨物种蛋壳状态的差异？ (1, 2, 4, 6, 17, 18)"
)


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def update(path: Path, text: str) -> None:
    doc = Document(str(path))
    replace_paragraph_text(doc.paragraphs[7], text)
    doc.save(str(path))


def main() -> None:
    update(EN_DOCX, EN_INTRO)
    update(CN_DOCX, CN_INTRO)
    print(EN_DOCX)
    print(CN_DOCX)


if __name__ == "__main__":
    main()
