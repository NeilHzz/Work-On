"""
Science Advances 格式 — 英文版
manuscript_results_sa.docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"E:\Data\Desktop\Work On\manuscript_results_sa.docx"

doc = Document()

# ── 页面 ─────────────────────────────────────────────────────────────────
s = doc.sections[0]
s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.54)
s.page_width  = Cm(21.0)
s.page_height = Cm(29.7)

# ── 行号（连续，每页重置）
_lnNum = OxmlElement("w:lnNumType")
_lnNum.set(qn("w:countBy"), "1")
_lnNum.set(qn("w:restart"), "continuous")
_lnNum.set(qn("w:start"), "1")
s._sectPr.append(_lnNum)

# ── 辅助 ─────────────────────────────────────────────────────────────────
FONT = "Times New Roman"   # Science Advances: use universal fonts

def _set_font(rPr, name):
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)
    rPr.insert(0, rFonts)

def fmt(run, size=11, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    _set_font(rPr, FONT if not italic else "Times New Roman")

def spacing(p, before=0, after=120, line=24):
    """line in pt; double-spaced for Science Advances initial submission"""
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement("w:spacing")
    e.set(qn("w:before"),   str(before))
    e.set(qn("w:after"),    str(after))
    e.set(qn("w:line"),     str(line * 20))
    e.set(qn("w:lineRule"), "auto")
    pPr.append(e)

def para(text, bold=False, italic=False, size=11,
         before=0, after=120, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    r = p.add_run(text)
    fmt(r, size=size, bold=bold, italic=italic)
    return p

def mixed(parts, before=0, after=120, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    for text, bold, italic in parts:
        r = p.add_run(text)
        fmt(r, bold=bold, italic=italic)
    return p

def head(text, size=11):
    """Bold subheading — no terminal period, sentence case"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spacing(p, before=240, after=60)
    r = p.add_run(text)
    fmt(r, size=size, bold=True)
    return p

def headm(parts, size=11):
    """Bold subheading with mixed italic support. parts = [(text, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spacing(p, before=240, after=60)
    for text, italic in parts:
        r = p.add_run(text)
        fmt(r, size=size, bold=True, italic=italic)
    return p

def cite(p, numbers):
    """Add inline (parenthesis) citation numbers per Science Advances style.
    Single: (1)  Multiple: (1, 2)  Consecutive 3+: (1–3)  Mixed: (1, 3–5, 7)
    """
    if not numbers:
        return
    sn = sorted(numbers)
    groups = []
    i = 0
    while i < len(sn):
        j = i
        while j + 1 < len(sn) and sn[j + 1] == sn[j] + 1:
            j += 1
        if j - i >= 2:          # 3+ consecutive → en-dash range
            groups.append(f"{sn[i]}\u2013{sn[j]}")
        elif j - i == 1:        # 2 consecutive → comma
            groups.append(f"{sn[i]}, {sn[j]}")
        else:
            groups.append(str(sn[i]))
        i = j + 1
    r = p.add_run(" (" + ", ".join(groups) + ")")
    fmt(r, size=11)
    rPr = r._r.get_or_add_rPr()
    _set_font(rPr, FONT)

# ════════════════════════════════════════════════════════════════════════════
# Science Advances 必需元素：Title / Short title / Authors / Teaser
# ════════════════════════════════════════════════════════════════════════════

# Title (≤135 characters)
para(
    "N-glycan structural class on ovalbumin encodes a composite adaptive response "
    "to developmental strategy and calcium ecology in three avian species",
    bold=True, size=14, before=0, after=160, align=WD_ALIGN_PARAGRAPH.LEFT
)

# Short title (≤50 characters)
para("Avian OVAL glycan class, ecology, and eggshell architecture",
     bold=False, size=11, after=60, align=WD_ALIGN_PARAGRAPH.LEFT)

# Authors & Affiliations (placeholder)
para(
    "[Author names, affiliations, ORCID IDs, and corresponding author to be "
    "completed before submission.]",
    bold=False, size=10, before=80, after=80, align=WD_ALIGN_PARAGRAPH.LEFT
)

# Abstract — skipped per user instruction (to be added before submission)
para("[Abstract: ≤50 words, single paragraph — to be completed before submission.]",
     bold=False, size=10, before=80, after=80, align=WD_ALIGN_PARAGRAPH.LEFT)

# Teaser (≤125 characters, one sentence for non-specialist readers)
para(
    "Teaser: Species-specific N-glycan class on egg-white ovalbumin tunes "
    "eggshell nucleation density and mechanical resistance across the precocial–altricial axis.",
    bold=False, italic=True, size=10, before=80, after=160, align=WD_ALIGN_PARAGRAPH.LEFT
)

# ════════════════════════════════════════════════════════════════════════════
# Introduction
# ════════════════════════════════════════════════════════════════════════════
para("Introduction", bold=True, size=14, before=0, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

# §1 — Background value
p_s1a = para(
    "The avian eggshell is a mineralised bioceramic structure that simultaneously "
    "protects the embryo from mechanical load and microbial challenge, regulates "
    "respiratory gas exchange, and constitutes the primary calcium reserve for "
    "skeletal development of the chick. Organised from outer to inner surface, "
    "the shell comprises four structurally distinct layers: an outer proteinaceous "
    "cuticle, a vertical crystal layer, a palisade layer of columnar calcite "
    "prisms, and a basal mammillary layer composed of periodic organic-matrix "
    "knobs \u2014 the nucleation starting points from which individual calcite units "
    "expand radially outward and the mineralisation origin of the entire shell "
    "architecture; the structural geometry of this mammillary layer determines "
    "the great majority of intact-shell mechanical competence."
)
cite(p_s1a, [1, 4])

p_s1b = para(
    "Avian life-history evolution has diversified reproductive investment along "
    "the precocial\u2013altricial axis: precocial species invest in thick, densely "
    "mineralised shells capable of sustaining the mechanical loads imposed by "
    "ground-nesting environments, whereas altricial species produce thinner "
    "shells consistent with elevated-nest development and reduced structural "
    "demands."
)
cite(p_s1b, [1, 16, 26])

p_s1c = para(
    "The calcium available for rapid eggshell mineralisation differs fundamentally "
    "between terrestrial and aquatic nesting environments: terrestrial species "
    "rely primarily on dietary calcium supplemented by medullary bone mobilisation "
    "and soil-derived carbonate grit, whereas aquatic and semi-aquatic species "
    "can additionally meet a substantial portion of their calcium requirements "
    "through dietary intake of calcium-rich aquatic invertebrates \u2014 providing "
    "a more continuously available mineral supply during the critical overnight "
    "mineralisation phase \u2014 thereby imposing contrasting constraints on maternal "
    "calcium supply capacity and, consequently, on shell thickness and mineral "
    "density."
)
cite(p_s1c, [4, 27])

# §2 — Prior work and its limits
p_intro2 = para(
    "The organic matrix proteins secreted by the uterine epithelium into forming "
    "mammillary cones have been established as central determinants of nucleation "
    "kinetics, crystal growth orientation, and architectural diversity across "
    "species, with ovalbumin (OVAL) and the C-type lectin ovocleidin-116 (OC116) "
    "identified as key Ca\u00b2\u207a-mediated calcification regulators. These proteins carry "
    "N-linked glycans spanning structural classes from compact High-Mannose to "
    "extended Sialylated Complex forms, and prior compositional and structural "
    "analyses have established that glycan class on secreted egg-white proteins "
    "influences protein thermostability, three-dimensional conformation, and the "
    "solvent exposure of surface acidic residues central to Ca\u00b2\u207a coordination. "
    "Furthermore, the altricial\u2013precocial spectrum has been recognised as a major "
    "organising axis of avian reproductive investment, and divergent eggshell "
    "mechanical performances across species spanning this axis have been "
    "characterised by micro-CT morphometry and finite-element simulation. "
    "Structural studies of chicken OVAL glycopeptides have further shown that "
    "N-glycan occupancy at Asn293 alters the solvent-exposed surface area of "
    "adjacent Ca\u00b2\u207a-coordinating residues, yet whether this modulation "
    "influences the rate of Ca\u00b2\u207a saturation during uterine calcification \u2014 "
    "and whether analogous effects operate in other avian species carrying "
    "different glycan classes at orthologous sequons \u2014 has not been examined. "
    "Gene-family evolution analyses across resident avian genomes have "
    "additionally identified lineage-specific shifts in oviduct calcium-transport "
    "repertoires, suggesting that the molecular infrastructure supporting "
    "eggshell calcification continues to diverge under species-specific "
    "selection pressures independent of its phylogenetic conservation. "
    "However, these lines of inquiry have hitherto proceeded independently: "
    "proteomic, glycoproteomic, structural modelling, and biomechanical analyses "
    "have each been conducted within individual species \u2014 predominantly in the "
    "domestic chicken \u2014 without a comparative cross-species framework linking "
    "glycan structural class through Ca\u00b2\u207a-binding site accessibility to mammillary "
    "architecture and shell mechanical performance."
)
cite(p_intro2, [1, 2, 4, 6, 7, 8, 15, 16, 17, 20])

# §3 — Core gap
para(
    "Whether divergence in N-glycan structural class across species spanning "
    "contrasting developmental strategies and calcium acquisition ecologies "
    "translates systematically into differences in mammillary microarchitecture "
    "and eggshell mechanical performance remains uncharacterised. Specifically, "
    "it is unresolved whether the pronounced taxonomic specificity of N-glycan "
    "class on OVAL\u2014compact High-Mannose in Galliformes, neutral complex in "
    "Anseriformes, and Sialylated Complex\u2013Hybrid dominant in "
    "Columbiformes\u2014reflects independent selection pressures from developmental "
    "strategy and ecological calcium availability, their interaction, or "
    "confounding phylogenetic history; and whether this glycan-level divergence "
    "drives or merely correlates with inter-species differences in mammillary "
    "nucleation efficiency and shell mechanical performance. Prior research has "
    "examined eggshell proteome composition, glycan structure, and shell "
    "biomechanics largely within individual species, and existing "
    "glycoproteomic surveys have been performed almost exclusively in the "
    "domestic chicken, leaving inter-species variation in OVAL glycan class "
    "mechanistically unexplained. The mechanistic continuum from glycan-mediated "
    "Ca\u00b2\u207a modulation through nucleation efficiency to tissue-scale load "
    "distribution thus remains unaddressed in any integrated cross-species "
    "framework. Resolving this gap requires a study design that simultaneously "
    "captures inter-species variation in glycan class, quantifies the structural "
    "consequences of that variation for Ca\u00b2\u207a-binding site accessibility at the "
    "protein surface, and traces the effects through to mammillary density and "
    "whole-shell mechanical performance across ecologically contrasting species. "
    "Such a framework has the added benefit of separating the effects of "
    "phylogenetic relatedness from those of ecological strategy, because the "
    "three focal orders span two major avian superorders and represent "
    "independently derived ecological specialisations."
)

# §4 — This study
p_intro4 = mixed([
    ("Here, we integrate micro-CT morphometry, comparative eggshell matrix "
     "proteomics with gene-family evolution analysis, intact glycopeptide mass "
     "spectrometry, Re-Glyco structural ensemble modelling, and finite-element "
     "simulation in three species \u2014 ", False, False),
    ("Gallus gallus", False, True),
    (", ", False, False),
    ("Anas platyrhynchos", False, True),
    (", and ", False, False),
    ("Columba livia", False, True),
    (" \u2014 representing, respectively, the precocial terrestrial ground-nesting, "
     "precocial semi-aquatic, and altricial elevated-nesting ecological guilds. "
     "By traversing from whole-shell morphology through proteome-wide orthogroup "
     "comparison and atomic-level electrostatic modelling to tissue-scale "
     "contact-stress simulation, we demonstrate that species-specific N-glycan "
     "class composition on OVAL encodes a composite adaptive response to "
     "developmental strategy and ecological calcium availability, and that this "
     "glycan-level divergence parallels systematic differences in mammillary "
     "density, geometry, and resistance to contact shear.", False, False),
])
cite(p_intro4, [3, 11, 12, 16])

# ═══════════════════════════════════════════════════════════════════════════
# "Results" section label
# ═══════════════════════════════════════════════════════════════════════════
para("Results", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

para("Unless otherwise stated, all values are expressed as mean \u00b1 s.d. "
     "All statistical tests are two-tailed.",
     bold=False, size=11, before=0, after=120)
# ════════════════════════════════════════════════════════════════════════════
# § Species selection — ecological and developmental niche analysis
# ════════════════════════════════════════════════════════════════════════════
head("Three avian orders occupy non-overlapping ecological and developmental niches")

p_ss1 = para(
    "To establish a systematic ecological rationale for species selection, we applied "
    "principal component analysis (PCA) to three AVONET trait fields — primary "
    "lifestyle, habitat use, and trophic niche — across >10,000 avian species, "
    "supplemented with order-level developmental mode scores on the fully precocial–"
    "fully altricial continuum. PC1 explained 71.7% of total variance and "
    "represented an aquatic-association gradient. K-means clustering (k = 3) "
    "resolved three well-separated clusters (silhouette coefficient = 0.814): "
    "terrestrial precocial, aquatic precocial, and terrestrial altricial "
    "(Supplementary Fig. 1a)."
)
cite(p_ss1, [22, 23])

p_ss2 = mixed([
    ("Mapping these clusters onto a time-calibrated avian phylogeny revealed a "
     "conserved correspondence between ecological niche and developmental mode "
     "(Supplementary Fig. 1b). Ancient ground-nesting lineages — Palaeognathae, "
     "Galloanserae, and basal Gruiformes — were exclusively precocial, consistent "
     "with nest-predation pressure as the ancestral selective force maintaining "
     "early locomotor independence. Fully altricial development was concentrated "
     "within core landbird clades (Afroaves and Australaves), where elevated-nest "
     "security may have relaxed early-locomotion constraints and released "
     "developmental time for neural expansion. Crucially, no aquatic-associated "
     "order evolved extreme altriciality: Aequornithia and Anseriformes were "
     "locked into precocial or semi-precocial modes, consistent with "
     "flood- and predation-exposed nest sites demanding early waterborne mobility. "
     "Orders at the land\u2013water interface \u2014 notably Charadriiformes and "
     "Pelecaniformes \u2014 exhibited elevated ecological discordance between body "
     "form and trophic habitat, marking them as evolutionary transition states "
     "rather than stable archetypes.", False, False),
])
cite(p_ss2, [15, 24, 25])

p_ss3 = mixed([
    ("Within this ecological framework, ", False, False),
    ("Gallus gallus", False, True),
    (" (Galliformes), ", False, False),
    ("Anas platyrhynchos", False, True),
    (" (Anseriformes), and ", False, False),
    ("Columba livia", False, True),
    (" (Columbiformes) occupy three separate corners of the ecological space "
     "(Supplementary Fig. 1a), each showing negligible intra-cluster ecological "
     "discordance. This near-zero discordance confirms that the three model "
     "species are pure archetypes of their respective ecological strategies "
     "rather than edge-case or transitional ecotypes. Their systematic comparison "
     "therefore maximally isolates the molecular determinants of aquatic "
     "adaptation and developmental strategy divergence from confounding niche-"
     "boundary effects, providing a principled basis for the integrative "
     "cross-species analyses that follow.", False, False),
])
cite(p_ss3, [22])

# ════════════════════════════════════════════════════════════════════════════
head("Mammillary layer microstructure and density differ systematically among three avian species")

mixed([
    ("We used micro-CT imaging to characterize the mammillary layer architecture "
     "across the three species and identified distinct morphological types (Fig. 1a\u2013c). "
     "In ", False, False),
    ("G. gallus", False, True),
    (", mammillary tips presented a smooth, nearly planar inner surface "
     "profile, with regularly spaced dome-shaped projections of uniform "
     "height (Fig. 1a). In ", False, False),
    ("A. platyrhynchos", False, True),
    (", the inner surface was markedly irregular and corrugated: mammillary "
     "projections were distributed at variable heights, producing a jagged "
     "cross-sectional profile in which adjacent peaks and valleys were "
     "clearly distinguishable (Fig. 1b). ", False, False),
    ("C. livia", False, True),
    (" displayed a conical mammillary morphology, with each mammilla "
     "presenting as a discrete, sharpened cone whose angular profile was "
     "distinct from the rounded domes of the precocial species (Fig. 1c). "
     "Three-dimensional surface reconstructions confirmed these "
     "cross-sectional observations: the chicken inner surface appeared "
     "fine-textured and isotropic, the duck surface intermediate and "
     "topographically complex, and the pigeon surface dominated by "
     "evenly distributed conical features.", False, False),
])

p_s0b = mixed([
    ("We quantified three mammillary parameters and found significant "
     "inter-species differences in all (Fig. 1d\u2013f). "
     "Mammilla density was highest in ", False, False),
    ("G. gallus", False, True),
    (" (171.36 ± 5.63 per mm²), significantly exceeding both ", False, False),
    ("A. platyrhynchos", False, True),
    (" (155.22 ± 8.63 per mm²) and ", False, False),
    ("C. livia", False, True),
    (" (158.27 ± 11.39 per mm²; Tukey HSD, both p < 0.05), "
     "whereas duck and pigeon densities were not significantly different "
     "from each other (Fig. 1d). "
     "Column unit volume followed an inverse rank: ", False, False),
    ("A. platyrhynchos", False, True),
    (" mammillae were the largest "
     "(5.04 ± 0.28 × 10⁻³ mm³), followed by ", False, False),
    ("G. gallus", False, True),
    (" (4.06 ± 0.13 × 10⁻³ mm³) and ", False, False),
    ("C. livia", False, True),
    (" (2.54 ± 0.19 × 10⁻³ mm³; Fig. 1e). "
     "Unit volume ratio—the fraction of the mammillary column volume "
     "attributable to its organic matrix core relative to total "
     "column volume—was highest in ", False, False),
    ("C. livia", False, True),
    (" (0.5321 ± 0.0389), intermediate in ", False, False),
    ("A. platyrhynchos", False, True),
    (" (0.4413 ± 0.0249), and lowest in ", False, False),
    ("G. gallus", False, True),
    (" (0.3975 ± 0.0127; Fig. 1f). "
     "Our data indicate that chicken mammillae are numerous and "
     "compact with the smallest per-mammilla organic matrix proportion, "
     "duck mammillae are fewer but volumetrically the largest, and pigeon "
     "mammillae are conical with the highest organic matrix ratio\u2014consistent "
     "with a gradient of decreasing mineral deposition efficiency from precocial "
     "ground-nesting through semi-aquatic to altricial elevated-nesting "
     "species. Because mammillary architecture is directly templated "
     "by uterine eggshell matrix proteins, these phenotypic "
     "differences prompted a systematic comparative proteomics investigation.", False, False),
])
cite(p_s0b, [1, 2])

# ════════════════════════════════════════════════════════════════════════════
# S_prot  Eggshell matrix proteome orthogroup analysis
# ════════════════════════════════════════════════════════════════════════════
head("Eggshell matrix proteome orthogroup analysis defines a conserved core and lineage-restricted repertoires")

p_sprot_a = mixed([
    ("We performed shotgun proteomics of eggshell matrix fractions and "
     "used OrthoFinder for orthogroup assignment, identifying 2620, 2921, and "
     "3219 protein orthogroups in ", False, False),
    ("G. gallus", False, True),
    (", ", False, False),
    ("A. platyrhynchos", False, True),
    (", and ", False, False),
    ("C. livia", False, True),
    (" respectively (Fig. 2a). Of these, 1997 orthogroups were conserved "
     "across all three species, constituting the core eggshell matrix "
     "proteome independent of developmental strategy or nesting ecology. "
     "Pairwise-only orthogroups numbered 180 (", False, False),
    ("G. gallus", False, True),
    ("–", False, False),
    ("A. platyrhynchos", False, True),
    ("), 434 (", False, False),
    ("G. gallus", False, True),
    ("–", False, False),
    ("C. livia", False, True),
    ("), and 716 (", False, False),
    ("A. platyrhynchos", False, True),
    ("–", False, False),
    ("C. livia", False, True),
    ("). Species-exclusive orthogroups were most numerous in ", False, False),
    ("C. livia", False, True),
    (" (72), followed by ", False, False),
    ("A. platyrhynchos", False, True),
    (" (28) and ", False, False),
    ("G. gallus", False, True),
    (" (9), indicating that the pigeon eggshell harbours the broadest "
     "lineage-restricted protein complement.", False, False),
])
cite(p_sprot_a, [2])

mixed([
    ("We applied GO enrichment analysis to the three pairwise-shared "
     "orthogroup sets and found functional profiles aligned with "
     "inter-species ecological contrasts (Figs. 2b\u2013d). The ", False, False),
    ("A. platyrhynchos", False, True),
    ("–", False, False),
    ("C. livia", False, True),
    (" set—proteins shared between duck and pigeon but absent in chicken—was "
     "the most significantly enriched for calcium ion binding and metal ion "
     "binding (MF; both p < 10⁻²⁵), alongside Wnt signalling pathway and "
     "signal transduction (BP; Fig. 2c). The convergent enrichment of "
     "calcium-binding capacity in two species that rely exclusively on "
     "dietary or waterborne calcium (duck: aquatic filter-feeding; "
     "pigeon: seed-pecking at elevated sites) but not in soil-calcium-"
     "supplemented chicken reflects divergent molecular strategies for "
     "calcium acquisition from the environment. The ", False, False),
    ("G. gallus", False, True),
    ("–", False, False),
    ("A. platyrhynchos", False, True),
    (" set—shared by the two precocial Galloanserae but absent in "
     "the altricial pigeon—was enriched for adaptive immune response "
     "and spermatogenesis (BP; Fig. 2b), consistent with the conserved "
     "precocial reproductive programme and requirement for immune competence "
     "at hatch. The ", False, False),
    ("G. gallus", False, True),
    ("–", False, False),
    ("C. livia", False, True),
    (" set—proteins shared by chicken and pigeon but absent in duck—"
     "was most strongly enriched for protein transport and intracellular "
     "protein transport (BP) and integral component of membrane (CC; "
     "Fig. 2d), reflecting shared oviduct secretory machinery demands in "
     "two species that nest on fixed, non-aquatic substrates regardless "
     "of developmental strategy.", False, False),
])

mixed([
    ("We next examined GO enrichment of species-exclusive eggshell matrix "
     "proteins and identified lineage-specific functional specialisations (Fig. 2e). "
     "The nine ", False, False),
    ("G. gallus", False, True),
    ("-exclusive orthogroups were prominently enriched for protein "
     "N-linked glycosylation (BP)\u2014identifying glycan-processing capacity "
     "as a uniquely elaborated functional module absent from the proteomes "
     "of both non-precocial species. In contrast, ", False, False),
    ("A. platyrhynchos", False, True),
    ("-exclusive proteins (n\u202f=\u202f28) were dominated by regulation of immune "
     "response (the largest GO bubble in the dataset; p\u202f<\u202f10\u207b\u00b3\u2075) "
     "and B cell activation, alongside response to iron ion\u2014consistent with "
     "the heightened pathogen exposure and elevated mineral metabolic "
     "demands of an aquatic foraging strategy. The largest species-specific "
     "complement belonged to ", False, False),
    ("C. livia", False, True),
    ("-exclusive proteins (n\u202f=\u202f72), which were enriched for nervous system "
     "development, ubiquitin-dependent protein catabolism, and "
     "proteolysis\u2014reflecting the post-hatch developmental complexity of "
     "altricial neonatal maturation. "
     "Taken together, the lineage-specific GO signatures reveal a species-characteristic "
     "molecular fingerprint: glycan-processing elaboration in chicken, "
     "immune and iron-metabolic adaptation in duck, and developmental "
     "proteolysis in pigeon. The emergence of N-linked glycosylation as a ", False, False),
    ("Gallus", False, True),
    ("-exclusive enrichment term motivated a detailed comparative "
     "glycoproteomics investigation of the three species.", False, False),
])

# ════════════════════════════════════════════════════════════════════════════
# S1  Phylogenomics & gene-family evolution
# ════════════════════════════════════════════════════════════════════════════
head("Contrasting gene-family dynamics across precocial and altricial lineages")

p_s1a = mixed([
    ("We reconstructed whole-proteome maximum-likelihood phylogenies (JTT+CAT "
     "substitution model) for three avian species representing two reproductive "
     "guilds and distinct ecological niches: the terrestrial, ground-nesting "
     "precocial species ", False, False),
    ("Gallus gallus", False, True),
    (" (chicken), the semi-aquatic, spatulate-billed precocial species ", False, False),
    ("Anas platyrhynchos", False, True),
    (" (Mallard duck), and the elevated-nesting altricial species ", False, False),
    ("Columba livia", False, True),
    (" (rock pigeon). The analysis recovered the canonical Galloanserae topology in "
     "which the two precocial species form a sister clade (branch lengths: ", False, False),
    ("G. gallus", False, True),
    (" 0.054, ", False, False),
    ("A. platyrhynchos", False, True),
    (" 0.042 substitutions/site) relative to the outgroup ", False, False),
    ("C. livia", False, True),
    (" (0.064 substitutions/site; Fig. 1). "
     "The three species span a gradient of ecological calcium availability: ", False, False),
    ("Gallus", False, True),
    (" forages terrestrially with direct soil-calcium access; ", False, False),
    ("Anas", False, True),
    (" inhabits mineral-rich aquatic and semi-aquatic environments with "
     "continual access to calcium from aquatic food sources; and ", False, False),
    ("Columba", False, True),
    (", nesting at elevation on cliff ledges or urban structures, "
     "relies exclusively on dietary calcium without soil-contact supplementation.", False, False),
])
cite(p_s1a, [3, 4])

mixed([
    ("To quantify gene-family evolution, we applied CAFE5 to OrthoFinder "
     "orthogroup counts mapped onto the time-calibrated three-species phylogeny "
     "and found that the ratio of expanding to contracting gene-family clusters "
     "differed markedly among lineages, with ", False, False),
    ("G. gallus", False, True),
    (" dominated by contractions, ", False, False),
    ("A. platyrhynchos", False, True),
    (" balanced, and ", False, False),
    ("C. livia", False, True),
    (" dominated by expansions (Fig. 4a).", False, False),
])

p_s1b = mixed([
    ("Our CAFE5 analysis revealed asymmetric expansion\u2013"
     "contraction profiles across the three lineages (Fig. 4b). ", False, False),
    ("G. gallus", False, True),
    (" underwent net contraction: 6 clusters (36 proteins) expanded versus "
     "64 clusters (421 proteins) contracted. GO enrichment of contracted "
     "families identified immunoglobulin production (GO:0002377; 48 proteins; "
     "p = 5.6 × 10⁻⁵⁵) and immune response (GO:0006955; 10 proteins; "
     "p = 2.5 × 10⁻¹⁷) as top terms, consistent with domestication-associated "
     "immunoglobulin gene reduction. ", False, False),
    ("A. platyrhynchos", False, True),
    (" showed an intermediate profile (14 expansions, 30 contractions); "
     "contracted families were enriched for motor activity "
     "(GO:0003774; p = 5.0 × 10⁻⁶) and collagen catabolism. ", False, False),
    ("C. livia", False, True),
    (" displayed net expansion: 75 clusters (432 proteins) expanded while "
     "only 8 contracted. Expanded families were enriched for transmembrane "
     "transport (GO:0055085; p = 8.6 × 10⁻⁵), Rho-protein signal regulation, "
     "and synapse assembly.", False, False),
])
cite(p_s1b, [5])

# ════════════════════════════════════════════════════════════════════════════
# S2+S3  Glycoproteome network and protein–glycan abundance coupling
# ════════════════════════════════════════════════════════════════════════════
head("The egg-white glycoproteome is organised into a conserved core and lineage-specific sectors, with selective protein\u2013glycan abundance decoupling at key eggshell proteins")

p_s2a = mixed([
    ("We performed intact glycopeptide mass spectrometry (IGP-MS) combined with "
     "orthology-based network analysis and found a hierarchically structured "
     "egg-white glycoproteome across all three species (Fig. 3). "
     "A conserved three-species core of 33\u201358 glycoproteins per species occupied "
     "the central network rings, predominantly decorated with Complex-Sialylated "
     "(177 glycoproteins; 63 distinct glycan structures) and "
     "Complex-Fucosylated glycans (265 glycoproteins; 65 structures). "
     "High-Mannose glycans accounted for the highest single-protein count "
     "(323 glycoproteins; 10 structures), with stronger representation in ", False, False),
    ("A. platyrhynchos", False, True),
    (" (165 proteins) and ", False, False),
    ("C. livia", False, True),
    (" (127 proteins) than in ", False, False),
    ("G. gallus", False, True),
    (" (31 proteins). Lineage-specific proteins (One-species and Singleton clusters) "
     "were most numerous in ", False, False),
    ("A. platyrhynchos", False, True),
    (" (162 unique glycoproteins) and localised to the outermost network sectors, "
     "frequently carrying Pauci-mannose and Hybrid structures absent from "
     "the conserved core.", False, False),
])
cite(p_s2a, [6, 7])

p_s2b = mixed([
    ("To examine quantitative coupling between protein and glycan-site abundances, "
     "we applied a second-pass strict BlastP ortholog filter to identify "
     "high-confidence cross-species homologues for glycan\u2013protein co-quantification. "
     "We computed Spearman rank correlations between log\u2082-transformed protein "
     "intensities and glycan-site intensities and found significant positive "
     "associations in all three species: \u03c1 = 0.419 (", False, False),
    ("G. gallus", False, True),
    ("; n = 33 sites; p = 0.015), "
     "\u03c1 = 0.419 (", False, False),
    ("A. platyrhynchos", False, True),
    ("; n = 190; p = 1.8 \u00d7 10\u207b\u2079), and "
     "\u03c1 = 0.473 (", False, False),
    ("C. livia", False, True),
    ("; n = 144; p = 2.1 \u00d7 10\u207b\u2079) (Figs. 11\u201313), "
     "consistent with the established positive coupling between "
     "protein expression level and glycan-site occupancy across tissues and species.", False, False),
])
cite(p_s2b, [18])

p_s2c = mixed([
    ("We next performed bivariate analysis of orthologous egg-white proteins in "
     "all three pairwise species comparisons, plotting protein log\u2082FC against "
     "glycan log\u2082FC (Figs. 4\u20136). Under the null expectation that glycosylation scales "
     "proportionally with protein expression, orthologous proteins should cluster "
     "along the identity diagonal; displacement below the diagonal "
     "(glycan log\u2082FC < protein log\u2082FC) indicates that glycan-site abundance is "
     "specifically depleted relative to protein level\u2014reflecting active "
     "post-translational remodelling of glycosylation independent of, or counter "
     "to, translational output\u2014whereas displacement above the diagonal "
     "(glycan log\u2082FC > protein log\u2082FC) indicates that glycan-site abundance "
     "is maintained or enriched despite reduced protein expression, pointing to "
     "selective retention or enhanced modification of glycosylated proteoforms. "
     "In the ", False, False),
    ("Gallus", False, True),
    ("/", False, False),
    ("Columba", False, True),
    (" comparison (Fig. 4), we found pronounced off-diagonal displacement for "
     "three major egg-white glycoproteins. Ovalbumin (OVAL; protein log\u2082FC "
     "\u22120.73, glycan log\u2082FC \u22125.63) "
     "and ovocleidin-116 (OC116; protein log\u2082FC +2.59, glycan log\u2082FC \u22128.35) "
     "both fell markedly below the diagonal, demonstrating that glycan-site "
     "abundance at these loci is specifically suppressed\u2014and in the case of "
     "OC116, counter to the direction of protein expression change, in which "
     "the protein is substantially upregulated yet its glycan-site abundance "
     "is dramatically reduced. In contrast, transferrin (TRFE; protein log\u2082FC "
     "\u22124.78, glycan log\u2082FC \u22123.65) fell above the diagonal, indicating that "
     "glycosylation at this site is relatively preserved despite substantially "
     "reduced protein abundance. The extreme below-diagonal displacement of OC116 "
     "is especially notable: OC116 is an eggshell-specific calcite-nucleating "
     "C-type lectin that directly templates calcium carbonate crystal growth "
     "during eggshell mineralisation, and exhibits the highest intraspecies "
     "amino-acid sequence variability among all major eggshell proteins and the "
     "most pronounced glycan suppression across all three pairwise comparisons, "
     "indicative of strong lineage-specific evolutionary remodelling of its "
     "sequence and glycan decoration. "
     "In the ", False, False),
    ("Gallus", False, True),
    ("/", False, False),
    ("Anas", False, True),
    (" comparison (Fig. 5), a similar pattern of below-diagonal displacement "
     "was observed: OVAL (protein log\u2082FC +2.41, glycan log\u2082FC \u22124.17) and "
     "OC116 (protein log\u2082FC +5.20, glycan log\u2082FC \u221211.68) both fell well "
     "below the diagonal, confirming that glycan suppression at these loci "
     "in ", False, False),
    ("G. gallus", False, True),
    (" relative to ", False, False),
    ("A. platyrhynchos", False, True),
    (" cannot be attributed to reduced protein abundance. "
     "TRFE again remained above the diagonal (protein log\u2082FC \u22122.58, "
     "glycan log\u2082FC \u22121.77), consistent with selective glycan retention "
     "independent of protein abundance changes. "
     "In the ", False, False),
    ("Anas", False, True),
    ("/", False, False),
    ("Columba", False, True),
    (" comparison (Fig. 6), the direction of OC116 displacement reversed: "
     "OC116 fell above the diagonal (protein log\u2082FC \u22122.60, glycan log\u2082FC +3.33), "
     "indicating that OC116 glycan abundance is specifically enriched in "
     " ", False, False),
    ("A. platyrhynchos", False, True),
    (" relative to ", False, False),
    ("C. livia", False, True),
    (" despite lower protein abundance\u2014opposite to its "
     "below-diagonal position in the two comparisons involving ", False, False),
    ("G. gallus", False, True),
    (". OVAL and TRFE clustered near the diagonal in this comparison "
     "(OVAL: protein log\u2082FC \u22123.14, glycan log\u2082FC \u22121.46; "
     "TRFE: protein log\u2082FC \u22122.20, glycan log\u2082FC \u22121.88), indicating "
     "proportional co-variation in abundance across the aquatic\u2013altricial "
     "transition. Together, the three bivariate comparisons demonstrate that "
     "glycan-site abundance at the major eggshell proteins is subject to "
     "independent post-translational regulation in each lineage, and that OC116 "
     "in particular undergoes context-dependent glycan remodelling that reverses "
     "direction across the precocial\u2013altricial ecological axis.", False, False),
])
cite(p_s2c, [19, 21])

# ════════════════════════════════════════════════════════════════════════════
# S4  Glycan compositional profiling
# ════════════════════════════════════════════════════════════════════════════
head("N-glycan structural class composition of ovalbumin differs markedly among the three species")

p_s4a = mixed([
    ("We profiled N-glycan structural class compositions of four egg-white "
     "proteins by IGP-MS and identified species-specific patterns (Figs. 7\u201310). "
     "OVAL from ", False, False),
    ("G. gallus", False, True),
    (" was dominated by High-Mannose glycans "
     "(60.7%), with Neutral Complex/Hybrid (30.9%) and Fucosylated (8.4%) fractions. "
     "OVAL from ", False, False),
    ("A. platyrhynchos", False, True),
    (" carried exclusively Neutral Complex/Hybrid glycans (100%). "
     "OVAL from ", False, False),
    ("C. livia", False, True),
    (" was the most compositionally diverse, with Sialylated Complex/Hybrid "
     "(49.4%), Neutral Complex/Hybrid (38.1%), and High-Mannose (12.3%). "
     "The three-species OVAL glycan gradient thus proceeds from "
     "High-Mannose-dominant (terrestrial ground-nesting precocial), "
     "through exclusively neutral complex (semi-aquatic precocial), "
     "to sialylated-dominant (elevated-nesting altricial).", False, False),
])
cite(p_s4a, [8])

p_s4b = mixed([
    ("We detected OC116, an eggshell calcite-nucleating C-type lectin, in "
     "all three species. In ", False, False),
    ("G. gallus", False, True),
    (", OC116 carried High-Mannose (51.9%) and Neutral Complex/Hybrid (48.1%) "
     "in near-equal proportions; in ", False, False),
    ("A. platyrhynchos", False, True),
    (", Neutral glycans predominated (88.3%) with minor Fucosylated (8.5%) "
     "and Sialylated (2.6%) fractions. "
     "In ", False, False),
    ("C. livia", False, True),
    (", OC116 exhibited a markedly different glycan composition dominated by "
     "Sialylated Complex/Hybrid (40.2%) and Fucosylated Complex/Hybrid (31.0%), "
     "with smaller contributions of High-Mannose (15.1%), "
     "Paucimannose/Truncated (4.5%), and Neutral Complex/Hybrid (9.2%) fractions\u2014"
     "a profile substantially more complex than in both precocial species and "
     "consistent with the strong below-diagonal glycan suppression of OC116 in "
     "the bivariate analyses. "
     "TRFE glycan profiles diverged similarly: ", False, False),
    ("G. gallus", False, True),
    (" TRFE was exclusively Neutral Complex/Hybrid (100%), whereas ", False, False),
    ("A. platyrhynchos", False, True),
    (" TRFE was enriched for Fucosylated (48.0%) and Sialylated (22.4%) classes "
     "alongside Neutral (28.4%); ", False, False),
    ("C. livia", False, True),
    (" TRFE was composed of Neutral (42.7%) and Sialylated (49.0%) in "
     "near-equal fractions. "
     "The Gallus-specific glycoprotein ovocalyxin-17 (OC17, absent in duck and "
     "pigeon) was predominantly Sialylated (81.1%) with secondary "
     "Fucosylated (10.6%) and Neutral (6.1%) fractions.", False, False),
])
cite(p_s4b, [1, 2, 9, 10])

# ════════════════════════════════════════════════════════════════════════════
# S5  ReGlyco
# ════════════════════════════════════════════════════════════════════════════
headm([("N-glycan conformation restricts Ca\u00b2\u207a-binding site accessibility on ovalbumin in ", False), ("Columba", True)])

p_s5a = mixed([
    ("To assess the biophysical consequences of divergent OVAL glycan compositions, "
     "we performed re-glycosylation structural ensemble modelling (Re-Glyco) on "
     "homology-modelled OVAL orthologues (50 conformations per glycan type; "
     "total: 50 models for ", False, False),
    ("G. gallus", False, True),
    (" N293; 150 for ", False, False),
    ("A. platyrhynchos", False, True),
    (" N97; 700 for ", False, False),
    ("C. livia", False, True),
    (" N97) and computed Adaptive Poisson–Boltzmann Solver (APBS) electrostatic potentials on "
     "each glycan-aware model versus a matched deglycosylated (apo) structure "
     "(Figs. ReGlyco-1 to ReGlyco-4).", False, False),
])
cite(p_s5a, [11, 12])

mixed([
    ("We found that ", False, False),
    ("C. livia", False, True),
    (" N-glycans adopted significantly more extended conformations than those "
     "of both precocial species. Radius of gyration (Rg): ", False, False),
    ("C. livia", False, True),
    (" 9.13 ± 0.92 Å vs. ", False, False),
    ("G. gallus", False, True),
    (" 7.50 ± 0.37 Å and ", False, False),
    ("A. platyrhynchos", False, True),
    (" 7.61 ± 0.43 Å (Mann–Whitney U test, two-tailed; all pairwise p < 0.001; "
     "ReGlyco-Fig. 1a). End-to-end distances: ", False, False),
    ("C. livia", False, True),
    (" 22.30 ± 4.12 Å vs. ", False, False),
    ("G. gallus", False, True),
    (" 18.05 ± 4.05 Å and ", False, False),
    ("A. platyrhynchos", False, True),
    (" 17.52 ± 1.75 Å (p < 0.001). "
     "Minimum glycan–protein Cα contact distance was shortest in ", False, False),
    ("C. livia", False, True),
    (" (3.00 ± 0.46 Å vs. ", False, False),
    ("G. gallus", False, True),
    (" 3.58 ± 0.63 Å; p < 0.001).", False, False),
])

mixed([
    ("We next analysed Ca\u00b2\u207a-binding hotspots (surface Asp/Glu residues with APBS "
     "potential < \u22125 kT/e) and found species-dependent glycan-induced reduction "
     "(ReGlyco-Fig. 4b). "
     "In ", False, False),
    ("G. gallus", False, True),
    (", hotspot count was marginally higher in the glycosylated state "
     "(N_hot = 15.0) than in the apo form (14.0; difference not significant; "
     "n = 1 structure). "
     "In ", False, False),
    ("A. platyrhynchos", False, True),
    (", glycosylation reduced mean N_hot to 12.7 versus apo 14.0 "
     "(n = 3 ensemble structures). "
     "In ", False, False),
    ("C. livia", False, True),
    (", glycosylation significantly reduced N_hot to 10.8 ± 0.56 versus "
     "apo 13.0 (one-sample t-test; t₁₃ = −9.3; n = 14; p < 0.05). "
     "Total Asp/Glu solvent-accessible surface area (SASA) was also "
     "significantly reduced in glycosylated ", False, False),
    ("C. livia", False, True),
    (" OVAL (3,433 ± 18 Å² vs. apo 3,473 Å²; paired t-test; p < 0.01), "
     "but not in ", False, False),
    ("G. gallus", False, True),
    (" (3,630 vs. 3,648 Å²; p > 0.05) or ", False, False),
    ("A. platyrhynchos", False, True),
    (" (3,369 ± 13 vs. 3,404 Å²; p = 0.04).", False, False),
])

mixed([
    ("We found that per-hotspot mean SASA after glycosylation was highest in ", False, False),
    ("G. gallus", False, True),
    (" (51.52 ± 0.20 Å²), intermediate in ", False, False),
    ("C. livia", False, True),
    (" (49.13 ± 0.87 Å²), and lowest in ", False, False),
    ("A. platyrhynchos", False, True),
    (" (45.83 ± 0.27 Å²) "
     "(Mann–Whitney U test, two-tailed; all pairwise p < 0.001; ReGlyco-Fig. 3b). "
     "Hotspot fraction of total surface Asp/Glu residues was highest in ", False, False),
    ("G. gallus", False, True),
    (" (0.841 ± 0.003), intermediate in ", False, False),
    ("A. platyrhynchos", False, True),
    (" (0.780 ± 0.017), and lowest in ", False, False),
    ("C. livia", False, True),
    (" (0.744 \u00b1 0.009; all pairwise p\u202f<\u202f0.001; ReGlyco-Fig. 3c). "
     "We further found that glycosylation significantly shifted "
     "the median surface electrostatic potential of ", False, False),
    ("C. livia", False, True),
    (" OVAL to more negative values "
     "(glycosylated \u22121.29 kT/e vs. apo \u22121.03 kT/e; "
     "Wilcoxon signed-rank test; p\u202f<\u202f0.05; ReGlyco-Fig. 4d), "
     "whereas no significant shift was detected in ", False, False),
    ("G. gallus", False, True),
    (" (\u22120.81 vs. \u22120.79 kT/e; p\u202f>\u202f0.05) or ", False, False),
    ("A. platyrhynchos", False, True),
    (" (−0.10 vs. −0.58 kT/e). "
     "Together, these results demonstrate that the High-Mannose glycan of "
     "chicken OVAL does not significantly reduce Ca²⁺ hotspot count or SASA, "
     "whereas the extended Sialylated Complex/Hybrid glycan of pigeon OVAL "
     "significantly reduces both hotspot count and carboxylate SASA and "
     "shifts the surface electrostatic potential toward more negative values. "
     "This glycan-mediated restriction of Ca²⁺-binding site accessibility "
     "in ", False, False),
    ("C. livia", False, True),
    (" correlates with the elevated-nesting ecology of this species, "
     "in which egg formation occurs without soil-calcium supplementation, "
     "and with its altricial reproductive strategy, which does not require "
     "the extensive eggshell mineralisation characteristic of precocial species.", False, False),
])

# ════════════════════════════════════════════════════════════════════════════
# S6  Biomechanical simulation
# ════════════════════════════════════════════════════════════════════════════
head("Finite-element simulation reveals species-specific eggshell resistance to hatching-tooth impact")

p_s6a = mixed([
    ("To assess whether differences in mammillary architecture translate to "
     "measurable differences in eggshell mechanical resistance, we performed "
     "explicit finite-element analysis (FEA; LS-DYNA) on micro-CT-derived eggshell "
     "fragment models. A cone-geometry impactor simulating the caruncle-borne egg "
     "tooth was driven into a disc-shaped eggshell fragment (model diameter "
     "D = 2.0 mm) at nine parametric lateral offset positions (3 × 3 grid; "
     "0.5 mm spacing), yielding n = 9 independent contact-shear-stress "
     "time-courses per species. To enable direct cross-species comparison "
     "at the material level, we extracted the peak contact shear stress \u03c4_max "
     "directly from the FEA element output at each offset position and then "
     "computed the species mean \u00b1 s.d. across the nine positions "
     "(Supplementary Figs. S-Biomech-1 to S-Biomech-3; measured shell thicknesses: ", False, False),
    ("G. gallus", False, True),
    (" 0.29 mm, ", False, False),
    ("A. platyrhynchos", False, True),
    (" 0.35 mm, ", False, False),
    ("C. livia", False, True),
    (" 0.19 mm).", False, False),
])
cite(p_s6a, [16])

p_s6b = mixed([
    ("One-way ANOVA of F_max across the nine offset positions revealed a "
     "significant three-tier hierarchy "
     "(F\u2082,\u2082\u2084 = 127.52, p < 1.0 \u00d7 10\u207b\u00b9\u00b2; Duncan's multiple range test, DMRT, \u03b1 = 0.05): ",
     False, False),
    ("G. gallus", False, True),
    (" 1.117 \u00b1 0.110 N (letter c) > ", False, False),
    ("A. platyrhynchos", False, True),
    (" 0.898 \u00b1 0.090 N (b) > ", False, False),
    ("C. livia", False, True),
    (" 0.485 \u00b1 0.039 N (a), with all pairwise differences significant (Fig. Biomech-1a). "
     "Strikingly, \u03c4_max collapsed to a two-tier pattern "
     "(ANOVA F\u2082,\u2082\u2084 = 14.46, p = 8.9 \u00d7 10\u207b\u2074; DMRT \u03b1 = 0.05): ",
     False, False),
    ("G. gallus", False, True),
    (" \u03c4_max = 551.6 \u00b1 108.8 MPa (b) significantly exceeded both ",
     False, False),
    ("A. platyrhynchos", False, True),
    (" 404.0 \u00b1 39.6 MPa (a) and ", False, False),
    ("C. livia", False, True),
    (" 393.0 \u00b1 35.2 MPa (a), while these two species were "
     "statistically indistinguishable from each other (Fig. Biomech-1b).", False, False),
])
cite(p_s6b, [16, 17])

mixed([
    ("The divergence between F_max and τ_max rankings demonstrates that the "
     "higher raw contact force required to breach ", False, False),
    ("A. platyrhynchos", False, True),
    (" eggshell is attributable primarily to its greater thickness (0.35 mm "
     "vs. 0.19 mm in ", False, False),
    ("C. livia", False, True),
    (") rather than to superior material-level resistance per unit area. "
     "Conversely, the 36\u201340% elevation of \u03c4_max in ", False, False),
    ("G. gallus", False, True),
    (" relative to the two non-precocial species is an intrinsic material "
     "property at the contact interface that persists independently of shell geometry. "
     "The grouping of \u03c4_max into two tiers\u2014", False, False),
    ("G. gallus", False, True),
    (" alone in the high group, ", False, False),
    ("A. platyrhynchos", False, True),
    (" and ", False, False),
    ("C. livia", False, True),
    (" jointly in the low group\u2014exactly recapitulates the mammilla-density "
     "grouping observed by micro-CT (Fig. CT-1). "
     "Together, these simulation results directly link the macroscopic "
     "mineralisation phenotype to microscale material-level mechanical resistance.", False, False),
])

# ════════════════════════════════════════════════════════════════════════════
# Discussion
# ════════════════════════════════════════════════════════════════════════════
para("Discussion", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

p_disc_mam1 = mixed([
    ("Eggshell mammillary layer architecture has been linked to mineralisation "
     "efficiency in the domestic chicken, but systematic multi-species comparison "
     "in an ecological context has not previously been reported. "
     "Our micro-CT imaging identified three distinct mammillary architectures—smooth "
     "and densely packed in ", False, False),
    ("G. gallus", False, True),
    (", irregular and corrugated in ", False, False),
    ("A. platyrhynchos", False, True),
    (", and conical with a high organic-matrix fraction in ", False, False),
    ("C. livia", False, True),
    ("\u2014each representing the integrated phenotypic output of the eggshell "
     "biomineralisation programme. Because mammillary architecture is "
     "physically templated by organic matrix proteins secreted in the "
     "uterus prior to calcification, systematic differences in mammillary "
     "morphotype implicate corresponding differences in matrix protein "
     "composition, abundance, or post-translational modification. "
     "The combination of maximal mammilla density, minimal per-mammilla "
     "volume, and lowest organic-matrix ratio in ", False, False),
    ("G. gallus", False, True),
    (" is consistent with efficient mineralisation around numerous small "
     "nucleation centres, reflecting rapid calcium deposition during the "
     "approximately 20-hour uterine transit characteristic of precocial egg formation. "
     "Conversely, the conical morphology and elevated unit volume ratio "
     "of ", False, False),
    ("C. livia", False, True),
    (" suggest that a greater fraction of column volume remains as "
     "unmineralised organic matrix—an architectural feature compatible with "
     "the reduced mineralisation demand of an altricial species whose "
     "hatchling does not require a mechanically rigid calcium reservoir "
     "at the moment of emergence. "
     "This three-way morphological characterisation extends prior single-species "
     "observations into a comparative framework spanning fundamentally different "
     "ecological and developmental modes. The near-monotonic correspondence of "
     "mammillary density, unit volume, and organic-matrix fraction with the "
     "precocial\u2013altricial axis implies that mammillary layer geometry constitutes "
     "a primary adaptive interface between the molecular programme of shell "
     "formation and the mechanical demands of the hatching environment, rather "
     "than a rigidly conserved structural invariant subject only to phylogenetic "
     "drift.", False, False),
])
cite(p_disc_mam1, [1, 2, 20])

p_disc_mam2 = mixed([
    ("Our multi-level molecular evidence\u2014eggshell proteome "
     "orthogroup analysis, gene-family dynamics, comparative glycoproteomics, "
     "and Re-Glyco structural modelling\u2014converges on a mechanistic model "
     "in which N-glycan composition on OVAL, the dominant egg-white calcium-"
     "binding glycoprotein, modulates Ca\u00b2\u207a loading efficiency and "
     "thereby controls the pace of mammillary nucleation. "
     "Under this model, the compact High-Mannose glycan of chicken OVAL "
     "leaves Ca²⁺-binding Asp/Glu surface residues sterically and "
     "electrostatically accessible, enabling rapid saturating Ca²⁺ "
     "binding and the subsequent OVAL unfolding event that triggers onset "
     "of eggshell calcification. The attenuated Ca²⁺ accessibility of "
     "pigeon OVAL—conferred by the bulkier, more electronegative "
     "Sialylated Complex/Hybrid glycan—would slow this saturation step, "
     "reducing nucleation efficiency and producing the sparse, conical "
     "mammillary phenotype. The intermediate neutral-complex glycan of "
     "duck OVAL occupies an in-between position: sufficient Ca²⁺ "
     "accessibility for precocial mineralisation (supported by continual "
     "dietary calcium from aquatic foraging) but producing the irregular, volumetrically "
     "enlarged mammillae consistent with a less tightly controlled "
     "nucleation process. The enrichment of calcium ion binding capacity "
     "in the ", False, False),
    ("A. platyrhynchos", False, True),
    ("–", False, False),
    ("C. livia", False, True),
    (" shared eggshell proteome, combined with the ", False, False),
    ("Gallus", False, True),
    ("-exclusive N-linked glycosylation enrichment in species-specific proteins, "
     "further supports the interpretation that chicken has evolved a unique "
     "glycan-mediated strategy for calcium management that underpins its "
     "high-density mammillary architecture.", False, False),
])
cite(p_disc_mam2, [1, 2, 4, 27])

p_disc_oval = para(
    "The mechanistic model we propose\u2014in which N-glycan class on OVAL tunes "
    "Ca\u00b2\u207a access to surface Asp/Glu clusters and thereby regulates the onset "
    "of mammillary nucleation\u2014is consistent with broader structural principles "
    "of glycan-mediated surface regulation on secreted proteins. N-linked "
    "glycans can physically occlude adjacent protein surface patches to extents "
    "that depend sensitively on branch number, chain length, and terminal-residue "
    "chemistry: terminal sialyl groups markedly extend the hydrodynamic envelope "
    "of the glycan and increase its electrostatic repulsion of divalent cations, "
    "whereas compact high-mannose cores present a smaller and electroneutral "
    "shielding volume. For OVAL, which must transiently bind Ca\u00b2\u207a, undergo "
    "Ca\u00b2\u207a-triggered conformational change, and then unfold to present a "
    "nucleation surface on the nascent mammillary cone organic template, the "
    "efficiency of the initial Ca\u00b2\u207a-saturation step is rate-limiting: any "
    "steric or electrostatic impediment at the binding surface will delay "
    "saturation and, consequently, nucleation onset. Our APBS electrostatic "
    "mapping confirms this logic: chicken OVAL carries the least shielded "
    "surface at Ca\u00b2\u207a-coordinating-residue clusters, duck OVAL occupies an "
    "intermediate position, and pigeon OVAL is most shielded\u2014a gradient that "
    "precisely mirrors the decreasing mammilla density sequence identified by "
    "micro-CT. The parallel between computed surface electrostatics and "
    "observed tissue-scale architecture across three phylogenetically and "
    "ecologically diverse species lends strong support to the view that "
    "glycan-mediated Ca\u00b2\u207a-binding regulation is a causal, rather than "
    "coincidental, determinant of mammillary packing."
)

p_disc_mech = mixed([
    ("Our finite-element simulation results provide direct functional validation "
     "of this morphological model. We found that peak contact shear stress "
     "\u03c4_max extracted from the FEA element output was significantly higher in ", False, False),
    ("G. gallus", False, True),
    (" (551.6 \u00b1 108.8 MPa) relative to ", False, False),
    ("A. platyrhynchos", False, True),
    (" and ", False, False),
    ("C. livia", False, True),
    (" (~400 MPa for both), demonstrating that chicken eggshell material "
     "provides greater intrinsic resistance to cone-contact shear per unit "
     "area, consistent with the denser mineral lattice arising from the "
     "higher mammilla nucleation centre density. Because \u03c4_max is derived "
     "from the element-level stress field rather than a geometry-normalised "
     "force, this elevated resistance is a true material property of the "
     "calcite composite. Importantly, the DMRT "
     "grouping of τ_max—", False, False),
    ("G. gallus", False, True),
    (" alone in tier 'b', ", False, False),
    ("A. platyrhynchos", False, True),
    (" and ", False, False),
    ("C. livia", False, True),
    (" jointly in tier 'a'—exactly recapitulates the mammilla-density "
     "grouping from micro-CT. That ", False, False),
    ("A. platyrhynchos", False, True),
    (" and ", False, False),
    ("C. livia", False, True),
    (" converge to the same τ_max despite markedly different mammillary "
     "geometries (irregular duck vs. conical pigeon) and different "
     "N-glycan classes on OVAL implies that, below the high-density "
     "threshold achieved by ", False, False),
    ("G. gallus", False, True),
    (", variations in mammillary shape do not independently confer "
     "superior material-level resistance. Together, our results "
     "establish a coherent chain from N-glycan composition → Ca²⁺-binding "
     "efficiency → mammilla nucleation density → eggshell mechanical "
     "resistance, providing a molecular-to-functional explanation for the "
     "species-specific hatching performance inferred from the simulation.", False, False),
])
cite(p_disc_mech, [1, 2, 16])

p_disc_limits = para(
    "Several methodological boundaries of the present study should be "
    "acknowledged. Our structural ensemble modelling was performed with "
    "the single most abundant glycoform at each glycopeptide sequon, whereas "
    "eggshell glycoproteins in vivo carry heterogeneous glycan mixtures at "
    "each site. The degree to which averaging over a glycoform population "
    "alters predicted Ca\u00b2\u207a accessibility relative to the dominant species "
    "alone remains to be established by single-glycan-resolution structural "
    "approaches not yet available at proteome scale. Our finite-element "
    "simulation treated each shell as a spatially uniform material defined "
    "by mean layer thicknesses and mammillary densities measured from "
    "micro-CT sections; incorporating the documented spatial gradients in "
    "palisade column orientation and local surface curvature would permit "
    "more accurate prediction of site-specific fracture probability under "
    "distributed incubational loads. The APBS electrostatic calculations "
    "assumed an ionic milieu approximating mean uterine fluid composition, "
    "but the precise Ca\u00b2\u207a activity, local pH, and competing anion "
    "concentrations at the surface of a rapidly calcifying mammillary cone "
    "are not characterised for any of the three focal species, introducing "
    "quantitative uncertainty into cross-species comparisons of SASA and "
    "electrostatic potential values. Our glycoproteomic and morphometric data "
    "were obtained from freshly laid eggs of captive populations maintained "
    "under standardised nutritional conditions; intraspecific variation in "
    "glycan composition across wild laying seasons, clutch position within "
    "a laying cycle, or imposed dietary calcium restriction may modulate "
    "the glycan landscape in ways not captured here. Finally, while the "
    "concordance across glycan class, Ca\u00b2\u207a hotspot accessibility, mammillary "
    "density, and \u03c4_max is consistent with a causal chain, demonstrating "
    "directionality\u2014distinguishing whether glycan class divergence drove "
    "architectural evolution or co-evolved alongside independent structural "
    "adaptation\u2014requires experimental manipulation of OVAL glycosylation "
    "in defined in vitro mineralisation assays, which constitutes a priority "
    "goal for future investigation. Together, these caveats define a clear "
    "experimental agenda: single-glycan-resolution Ca\u00b2\u207a-binding assays, "
    "spatially resolved mechanical testing of individual mammillary columns, "
    "and reconstitution of OVAL-mediated nucleation with defined glycoforms "
    "will be necessary to elevate the integrative model presented here from "
    "a well-supported comparative inference to a mechanistically confirmed "
    "causal pathway."
)

p_d1 = mixed([
    ("The contrasting gene-family dynamics we observed are consistent with "
     "established patterns of domestication-driven genome evolution. "
     "We found that the pronounced contraction of immunoglobulin-related "
     "gene families in ", False, False),
    ("G. gallus", False, True),
    (" aligns with the well-documented reduction of immune-gene diversity "
     "under artificial selection, in which relaxed pathogen pressure in "
     "managed environments renders expansive immune repertoires selectively "
     "neutral. The enrichment of expanded gene families in ", False, False),
    ("C. livia", False, True),
    (" for transmembrane transport, Rho-protein signal regulation, and "
     "synapse assembly is consistent with the elevated developmental "
     "plasticity required by altricial neonates, which must undergo "
     "rapid and coordinated organ maturation ex ovo. "
     "Together, the opposing genomic trajectories of precocial and altricial "
     "lineages establish a lineage-specific molecular context that predicts "
     "corresponding divergence in protein expression and post-translational "
     "modification programmes, including N-glycosylation.", False, False),
])
cite(p_d1, [14, 15])

mixed([
    ("Evolutionary stratification of glycoproteomes into a conserved core and "
     "lineage-specific peripheral layers has been described in mammalian "
     "secretory systems, but an analogous comparative framework for avian "
     "egg-white has not been established. "
     "Our egg-white glycoprotein network analysis revealed a radial "
     "architecture\u2014with conserved trispecies glycoproteins at the centre "
     "and species-restricted proteins at the periphery\u2014indicating that "
     "evolutionary innovation in glycan structure preferentially occurs at "
     "the lineage-specific layer of the glycoproteome. The disproportionate "
     "enrichment of High-Mannose glycans in ", False, False),
    ("A. platyrhynchos", False, True),
    (" and ", False, False),
    ("C. livia", False, True),
    (" relative to ", False, False),
    ("G. gallus", False, True),
    (" may reflect a greater proportion of glycoproteins that have recently "
     "entered the secretory pathway or that experience reduced glycan-processing "
     "enzyme activity in the magnum. Alternatively, the aquatic and "
     "elevated-nesting niches of these two species may impose different "
     "functional demands on lectin-mediated protein folding or secretion, "
     "for which High-Mannose structures are well suited. "
     "This post-translational glycan autonomy implies that adaptively relevant "
     "differences in mammillary nucleation kinetics can evolve independently "
     "of the transcriptional programme controlling glycoprotein abundance, "
     "substantially expanding the evolutionary degrees of freedom for functional "
     "diversification between species with highly conserved egg-white protein "
     "coding sequences.", False, False),
])

mixed([
    ("Quantitative decoupling between protein expression level and glycan-site "
     "occupancy has been documented across multiple secretory tissues, "
     "yet the magnitude and site-specificity of such decoupling in the "
     "avian oviduct has not previously been characterised. "
     "We observed substantial divergence between protein abundance log\u2082FC "
     "and glycan log\u2082FC for OVAL, OC116, and TRFE, indicating that N-glycan "
     "occupancy at these three loci is regulated post-translationally, "
     "independently of mRNA or protein level. Such decoupling could arise through "
     "site-specific differences in oligosaccharyltransferase accessibility, "
     "competition between co-translational glycosylation and folding kinetics, "
     "or species-specific differences in glycan-processing enzyme expression "
     "in the oviduct. The glycan-suppressed state of OVAL and OC116 in ", False, False),
    ("G. gallus", False, True),
    (" relative to ", False, False),
    ("C. livia", False, True),
    (" is particularly noteworthy given that both proteins are quantitatively "
     "major components of the egg white and participate directly in "
     "eggshell calcification. The partial preservation of TRFE glycan "
     "abundance despite lower protein levels in ", False, False),
    ("G. gallus", False, True),
    (" may indicate that transferrin glycosylation is under stronger "
     "positive selection in this species, possibly to maintain "
     "iron sequestration or antimicrobial functions in the "
     "calcium-rich mineralisation microenvironment. "
     "More generally, the quantitative decoupling we observe between protein "
     "abundance and glycan occupancy across all three focal proteins suggests "
     "that the oviduct glycocalyx operates under a level of post-transcriptional "
     "regulatory control not previously appreciated in comparative studies of "
     "avian eggshell proteomics, and that glycan abundance constitutes an "
     "independent dimension of eggshell proteome variation that cannot be "
     "inferred from quantitative proteomics data alone.", False, False),
])

p_d4 = mixed([
    ("Prior work has established that ovalbumin glycosylation varies among "
     "avian species, but how glycan structural class relates to reproductive "
     "ecology and calcium microenvironment has not been systematically addressed. "
     "We found that the three-species OVAL glycan gradient\u2014High-Mannose dominant "
     "in terrestrial ground-nesting precocial chicken, exclusively neutral complex "
     "in semi-aquatic precocial duck, and sialylated dominant in elevated-nesting "
     "altricial pigeon\u2014does not map exclusively onto the precocial\u2013altricial axis, "
     "indicating a composite influence of developmental strategy and ecological "
     "calcium availability. "
     "The compact High-Mannose glycan of chicken OVAL minimises steric "
     "occlusion of surface acidic residues in a species that has ready "
     "access to soil-derived and dietary calcium for eggshell mineralisation. "
     "For ", False, False),
    ("A. platyrhynchos", False, True),
    (", the aquatic foraging niche\u2014with a spatulate bill specialised for "
     "filter-feeding on calcium-rich aquatic invertebrates\u2014provides continual "
     "access to dietary calcium, potentially relaxing the selective constraint "
     "on glycan-mediated Ca\u00b2\u207a modulation while still sustaining "
     "the precocial biomineralisation programme. "
     "The emergence of elongated Sialylated Complex/Hybrid glycans on "
     "pigeon OVAL may represent an evolutionary response to the dietary-only "
     "calcium constraint imposed by elevated nesting, where molecular-level "
     "regulation of Ca\u00b2\u207a binding could be advantageous during egg formation. "
     "Notably, the exclusive neutral-complex glycan class of duck OVAL\u2014with "
     "no detectable High-Mannose or Sialylated Complex/Hybrid structures\u2014 "
     "distinguishes the semi-aquatic precocial niche from both flanking "
     "strategies and implies that Anseriformes have evolved a distinctly "
     "intermediate glycan programme rather than a simple mixture of the two "
     "extremes. This exclusivity argues against a straightforward phylogenetic "
     "explanation: Galliformes and Anseriformes are closely related sister "
     "lineages within Galloanserae, yet their OVAL glycan classes are as "
     "divergent as either is from Columbiformes, indicating that the "
     "semi-aquatic ecology of Anseriformes constitutes an independent "
     "selective force on glycan class.", False, False),
])
cite(p_d4, [1, 4, 27])

p_d5 = mixed([
    ("Computational modelling of glycan conformational ensembles has established "
     "that glycan chain length and charge composition modulate surface accessibility "
     "of protein binding sites, yet how these effects influence Ca\u00b2\u207a-binding "
     "efficiency in eggshell glycoproteins has not previously been examined. "
     "Our Re-Glyco structural modelling results indicate that the High-Mannose "
     "glycan of chicken OVAL does not substantially impede Ca\u00b2\u207a coordination "
     "at surface Asp/Glu residues\u2014consistent with the requirement for efficient "
     "Ca\u00b2\u207a loading by OVAL preceding its conformational change and the "
     "subsequent initiation of eggshell calcification in the uterus. "
     "For duck OVAL, the neutral complex glycan occupies an intermediate "
     "conformational regime; the spatulate-billed, filter-feeding ecology of ", False, False),
    ("A. platyrhynchos", False, True),
    (" provides sufficient dietary calcium from aquatic invertebrate sources to support precocial "
     "mineralisation even with moderate glycan-mediated site modulation. "
     "The statistically significant reduction in Ca\u00b2\u207a hotspot count, "
     "carboxylate SASA, and surface electrostatic potential associated with "
     "glycosylation in ", False, False),
    ("C. livia", False, True),
    (" OVAL indicates that the extended Sialylated Complex/Hybrid glycan "
     "both sterically and electrostatically attenuates Ca\u00b2\u207a accessibility "
     "at the protein surface. "
     "This glycan-mediated attenuation may reflect selection for tighter "
     "regulation of Ca\u00b2\u207a sequestration during egg formation in a species "
     "whose elevated-nesting ecology restricts passive soil-calcium supplementation "
     "and whose altricial reproductive strategy does not require "
     "extensive eggshell calcification. "
     "Together, our multi-level evidence from proteomics, glycoproteomics, "
     "and structural ensemble modelling indicates that the N-glycan landscape "
     "of avian egg-white proteins encodes a composite adaptive response "
     "to both developmental strategy (precocial versus altricial chick maturity) "
     "and ecological calcium microenvironment (terrestrial, aquatic, and "
     "elevated-nesting habitats). "
     "While here we focused on three avian species spanning contrasting "
     "reproductive strategies and calcium ecologies, the principle that "
     "N-glycan composition on secreted proteins encodes adaptive responses "
     "to mineral acquisition ecology is likely to apply broadly across "
     "oviparous vertebrates and, more generally, to any system in which "
     "glycan-mediated modulation of Ca\u00b2\u207a-binding proteins is subject "
     "to ecological selection.", False, False),
])
cite(p_d5, [1, 4, 12, 27])

p_d6 = para(
    "The present findings carry implications for understanding eggshell "
    "molecular diversity across the precocial\u2013altricial spectrum. "
    "The approximately 10,000 living bird species span widely varying "
    "calcium ecologies\u2014from calcium-rich terrestrial and aquatic invertebrate "
    "diets to calcium-poor elevated or cavity-nesting environments\u2014and the "
    "composite developmental\u2013ecological signal we document across these "
    "three orders makes specific and testable predictions for lineages "
    "not sampled here. Fully precocial ground-nesting "
    "orders such as Struthioniformes and basal terrestrial Galliformes, which "
    "exploit calcium-rich substrates, are predicted to exhibit High-Mannose "
    "glycan-dominated OVAL profiles with maximal Ca\u00b2\u207a hotspot exposure "
    "and the highest mammillary densities, consistent with producing a "
    "large precocial egg under calcium-replete conditions. Obligate altricial "
    "cavity-nesters\u2014including many Passeriformes and Piciformes\u2014which "
    "occupy structurally protected nesting environments and face no premium "
    "on thick shells, are predicted to show the most extensive Sialylated "
    "Complex enrichment and lowest mammillary densities. "
    "The convergent \u03c4_max fracture toughness of duck and pigeon despite "
    "their markedly different mammillary geometries further suggests that "
    "altricial and semi-precocial shells converge on a minimum-sufficiency "
    "mechanical threshold, implying that selection for fracture resistance "
    "decouples from selection for calcification load once a viability floor "
    "is satisfied. This predicts that shell mechanical performance should be "
    "more variable among precocial lineages\u2014where nest-site hazards drive "
    "ongoing optimising selection\u2014than among altricial lineages, where "
    "structural adequacy is stabilised near a minimum. "
    "The nutritional ecology of calcium acquisition constitutes an additional "
    "axis whose interaction with the precocial\u2013altricial gradient remains "
    "unexplored. Wading and coastal orders such as Charadriiformes combine "
    "predominantly precocial developmental programmes with calcium-rich "
    "aquatic and estuarine invertebrate diets, generating a provisioning "
    "regime distinct from both the soil-calcium terrestrial precocial niche "
    "and the calcium-sparse elevated-nesting altricial niche. The OVAL glycan "
    "architecture of such lineages constitutes a direct test of whether "
    "habitat-mediated calcium access or developmental programme is the "
    "dominant determinant of glycan class. Systematic comparative "
    "glycoproteomic surveys spanning families varying independently in "
    "nesting ecology and chick developmental mode\u2014made tractable by the "
    "LFQ-proteomics workflow applied here\u2014will resolve whether the "
    "High-Mannose-to-Sialylated Complex gradient scales across the full "
    "precocial\u2013altricial axis or reflects habitat-specific calcium ecologies "
    "superimposed on a more conserved developmental template, and whether "
    "reproductive mode and ecological calcium availability act as co-equal "
    "selective forces or one dominates in shaping the molecular architecture "
    "of avian eggshell glycoproteins."
)

# ════════════════════════════════════════════════════════════════════════════
# Methods
# ════════════════════════════════════════════════════════════════════════════
para("Methods", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

head("Biological materials")

p_m_bio = mixed([
    ("Six freshly laid eggs from each of three avian species were collected within "
     "24 h of oviposition. ", False, False),
    ("Gallus gallus", False, True),
    (" (domestic chicken) eggs were obtained from the Poultry Resources Conservation "
     "Farm, China Agricultural University (Beijing, China); ", False, False),
    ("Columba livia", False, True),
    (" (rock pigeon) eggs were provided by Prof. Chang Yu, College of Veterinary "
     "Medicine, China Agricultural University; and ", False, False),
    ("Anas platyrhynchos", False, True),
    (" (Mallard duck) eggs were supplied by Jinxing Duck Industry (Beijing, China). "
     "All eggs were stored at 4\u00b0C until processing.", False, False),
])

head("Eggshell matrix protein extraction")

mixed([
    ("Eggshell matrix proteins were extracted from the eggshell mammillary layer "
     "(EML) by an established EDTA demineralisation protocol. Eggs were rinsed with "
     "deionised water and placed in sterile sealed bags. For ", False, False),
    ("G.\u00a0gallus", False, True),
    (" and ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (", the eggshell cuticle layer (ECL) was removed prior to EML "
     "extraction by treatment with 15 mL of 5% EDTA (0.13 mol/L, pH 7.6) supplemented "
     "with 2-mercaptoethanol (10 mmol/L) for 30 min at 20\u00b0C, with gentle manual "
     "kneading to separate the ECL; shells were subsequently rinsed with deionised "
     "water. EML proteins from all three species were then solubilised under the same "
     "EDTA\u20132-mercaptoethanol conditions with the extraction duration extended to 12 h "
     "at 20\u00b0C. The resulting suspension was centrifuged at 1,000 \u00d7 g for 15 min; the "
     "pellet was resuspended and centrifuged a second time, and the pooled supernatant "
     "was stored at \u221280\u00b0C until analysis.", False, False),
])

head("Micro-CT imaging and mammillary morphometry")

para(
    "Two eggshell fragments (each approximately 4\u20135 mm\u00b2) were excised from the "
    "equatorial region of each species and scanned with a Phoenix V|tome|x\u00a0M "
    "microfocus CT system (GE Sensing and Inspection Technologies GmbH, Wunstorf, "
    "Germany) at 85 kV and 160 \u03bcA with no beam filter; scan settings were held "
    "constant across all specimens. Three-dimensional reconstructions were generated "
    "in 3D Slicer by threshold-based segmentation restricted to a cylindrical region "
    "of interest of 1 mm radius. Acquisition noise was suppressed by a 5 \u00d7 5 \u00d7 5 "
    "median filter, followed by largest-island isolation and 9 \u00d7 9 \u00d7 9 hole-filling. "
    "Surface models were exported as STL files and reverse-engineered in Geomagic "
    "Wrap by sequential de-noising (strength 2), triangle simplification to "
    "approximately 300,000 faces, mesh re-gridding at 0.01 mm, iterative defect "
    "correction to zero residual faults, and organic parametric surface fitting at "
    "minimum tolerance. Three morphometric parameters were quantified per specimen: "
    "mammilla density (count per mm\u00b2), column unit volume (mm\u00b3), and unit volume "
    "ratio (organic matrix core volume relative to total column volume; "
    "n\u202f=\u202f2 fragments per species)."
)

head("Shotgun proteomics of eggshell matrix proteins")

para(
    "For each species, two eggs were pooled per biological replicate and three "
    "independent replicates were prepared (n\u202f=\u202f3 per species). Proteins were "
    "extracted by resuspension in lysis buffer (1% SDS, 1% protease inhibitor "
    "cocktail), sonication on ice, and clarification by centrifugation at "
    "12,000 \u00d7 g at 4\u00b0C for 10 min; protein concentration was determined with a "
    "BCA assay kit. Proteins were precipitated with pre-cooled acetone (5 volumes, "
    "\u221220\u00b0C, 2 h), washed twice with acetone, and redissolved in 200 mM TEAB. "
    "Disulfide bonds were reduced with 5 mM dithiothreitol (56\u00b0C, 30 min) and "
    "alkylated with 11 mM iodoacetamide (room temperature, 15 min, dark). Proteins "
    "were digested overnight with sequencing-grade trypsin (enzyme:protein ratio "
    "1:50) and desalted with Strata X SPE columns."
)

mixed([
    ("Desalted peptides were dissolved in 0.1% formic acid and separated on a "
     "home-made 15-cm \u00d7 100-\u03bcm i.d. reversed-phase C18 analytical column connected "
     "to a Vanquish Neo UPLC system (Thermo Fisher Scientific) at 400 nl/min over "
     "a 22.6-min gradient (4\u201399% solvent B; 0.1% formic acid in 80% acetonitrile). "
     "Separated peptides were analysed on an Orbitrap Astral mass spectrometer "
     "(Thermo Fisher Scientific) with a nano-electrospray ionisation source (1,900 V). "
     "Full-MS spectra were acquired in the Orbitrap at 240,000 resolution over "
     "380\u2013980 m/z; MS/MS fragments were acquired in the Astral analyser at 80,000 "
     "resolution using HCD fragmentation (NCE\u202f=\u202f25%), fixed first mass 150 m/z, "
     "AGC target 500%, and maximum injection time 3 ms. DIA data were processed with "
     "DIA-NN v1.8 against species-specific reference proteomes \u2014 ", False, False),
    ("G.\u00a0gallus", False, True),
    (" (43,711 entries), ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (" (91,801 entries), and ", False, False),
    ("C.\u00a0livia", False, True),
    (" (17,309 entries), all downloaded August 2024 \u2014 "
     "concatenated with reverse decoy sequences. "
     "Trypsin/P cleavage was specified with up to one missed cleavage; N-terminal "
     "methionine excision and carbamidomethylation of Cys were set as fixed "
     "modifications. Protein and peptide FDR were each controlled at < 1%.", False, False),
])

head("Intact glycopeptide mass spectrometry")

para(
    "N-glycopeptides were enriched from tryptic digests by hydrophilic interaction "
    "liquid chromatography (HILIC). Peptide digests were redissolved in loading "
    "buffer (80% ACN, 5% TFA), loaded onto a HILIC column, washed three times with "
    "loading buffer, and glycopeptides were eluted twice with 0.1% TFA, 50 mM "
    "ammonium bicarbonate, and 50% ACN. Eluates were desalted with C18 Zip-Tips "
    "and vacuum-dried. Glycopeptide fractions were separated on the same nano-LC "
    "platform using a 34-min gradient (4\u201399% B at 400 nl/min). Full-MS spectra were "
    "acquired at 240,000 resolution over 700\u20132,000 m/z; MS/MS spectra were acquired "
    "at 80,000 resolution with fixed first mass 120 m/z, cycle time 0.6 s, AGC "
    "target 100%, intensity threshold 25,000 ions/s, and maximum injection time 5 ms. "
    "Raw DDA data were processed with MSFragger v3.4 against the same species-specific "
    "reference proteomes as above, with strict trypsin cleavage (up to 2 missed "
    "cleavages), peptide length 7\u201350 residues, fixed carbamidomethylation of Cys, "
    "variable N-terminal acetylation and Met oxidation, and default MSFragger "
    "glycosylation mass offsets. Protein, peptide, and PSM FDR were each controlled "
    "at < 1%. N-glycan structural classes were assigned using the Oxford nomenclature "
    "into six categories: High-Mannose, Paucimannose/Truncated, Neutral Complex/"
    "Hybrid, Fucosylated Complex/Hybrid, Sialylated Complex/Hybrid, and Other. "
    "Per-protein per-species class abundance was computed as the fraction of total "
    "glycan-site signal intensity."
)

head("Comparative proteome analysis and gene-family evolution")

p_m_ortho = mixed([
    ("Protein sequences from all three species were assigned to orthogroups using "
     "OrthoFinder (all-versus-all BlastP; E-value threshold 1 \u00d7 10\u207b\u00b9\u2070; "
     "MCL inflation parameter 2.0), yielding a 3,250-orthogroup background set for "
     "all downstream enrichment analyses. Divergence times were taken from published "
     "timetrees: ", False, False),
    ("G. gallus", False, True),
    ("\u2013", False, False),
    ("A. platyrhynchos", False, True),
    (" 83.37 Mya and ", False, False),
    ("G. gallus", False, True),
    ("\u2013", False, False),
    ("C. livia", False, True),
    (" 90.84 Mya. Gene-family expansion and contraction rates were inferred with "
     "CAFE5 using the time-calibrated phylogeny. GO enrichment of pairwise-shared, "
     "species-exclusive, expanding, and contracting orthogroup sets was performed "
     "using the OrthoVenn3 web platform (https://orthovenn3.bioinfotoolkits.net) "
     "against the 3,250-orthogroup background; terms with adjusted p\u202f<\u202f0.05 were "
     "considered significant.", False, False),
])
cite(p_m_ortho, [3, 5])

head("Cross-species glycoprotein ortholog identification")

mixed([
    ("High-confidence cross-species orthologues of four target eggshell glycoproteins "
     "(OVAL, OC116, TRFE, OC17) were identified by BlastP of each reference ", False, False),
    ("G.\u00a0gallus", False, True),
    (" sequence against the non-reference-species proteomes (E-value "
     "threshold 1 \u00d7 10\u207b\u2075; maximum 500 target sequences; 250 reported alignments). "
     "Candidate hits were retained at average maximum sequence identity \u2265 0.80; "
     "where query and subject non-overlapping HSP counts were unequal, a relaxed "
     "threshold of \u2265 0.50 was applied. Final UniProt ortholog identifiers used for "
     "downstream structural and quantitative analyses are listed in Supplementary "
     "Table\u00a01.", False, False),
])

head("N-glycan structural ensemble modelling")

p_m_reglyco = mixed([
    ("OVAL ortholog protein structures (AlphaFold2-predicted models) were "
     "accessed through the GlycoShape platform (glycoshape.org) via UniProt "
     "accession identifiers. "
     "Experimentally detected N-glycan compositions from IGP-MS were mapped "
     "to the GlycoShape glycan library by monoisotopic mass matching "
     "(tolerance \u00b10.5\u00a0Da), using per-residue masses of 203.0794\u00a0Da (HexNAc), "
     "162.0528\u00a0Da (Hex), 291.0954\u00a0Da (NeuAc), 146.0579\u00a0Da (dHex), and "
     "132.0423\u00a0Da (Pen) with an 18.0106\u00a0Da water correction; matched entries "
     "were retrieved as GlyTouCan accession identifiers. "
     "Full conformational ensembles were then generated with the GlycoShape "
     "Re-Glyco Ensemble tool (glycoshape.org/ensemble), which restores missing "
     "glycans by aligning them to torsion angles from Privateer crystallographic "
     "standards and sampling conformations from the GlycoShape molecular-dynamics "
     "ensemble library. "
     "For each protein, a session was created via the GlycoShape API to identify "
     "available N-glycosylation sequons in the structural model; each matched "
     "glycan was then submitted as an independent modelling job and attached to "
     "the target sequon \u2014 ", False, False),
    ("G.\u00a0gallus", False, True),
    (" N293 and ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (" / ", False, False),
    ("C.\u00a0livia", False, True),
    (" N97 \u2014 with ensemble size 50, random seed 42, and PDB output format. "
     "This procedure yielded 50 glycoprotein models for ", False, False),
    ("G.\u00a0gallus", False, True),
    (" (1 matched glycan type), 150 for ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (" (3 types), and 700 for ", False, False),
    ("C.\u00a0livia", False, True),
    (" (14 types, "
     "including four NeuAc positional isomers resolved from the GlycoShape library). "
     "Ensemble geometric descriptors were calculated for each model from atomic "
     "coordinates using BioPython: the radius of gyration (Rg) of all glycan heavy "
     "atoms, the end-to-end distance of the glycan chain, and the minimum distance "
     "between any glycan heavy atom and protein C\u03b1 atoms (minimum C\u03b1 contact "
     "distance). Per-structure summary statistics (mean\u00a0\u00b1\u00a0s.d.) and pairwise species "
     "comparisons (Mann\u2013Whitney\u00a0U test, two-tailed) were performed for each descriptor.", False, False),
])
cite(p_m_reglyco, [11])

head("Electrostatic potential calculation")

p_m_apbs = para(
    "Electrostatic surface potentials were computed for each Re-Glyco ensemble "
    "model and a matched deglycosylated (apo) reference using APBS v3.4.1. Atomic "
    "partial charges and radii were assigned with PDB2PQR using the CHARMM36 force "
    "field and PROPKA protonation at pH 7.4; glycan heavy-atom partial charges were "
    "assigned from the GLYCAM06 parameter set (Kirschner et al., 2008, J. Comput. "
    "Chem. 29:622). Solvent-accessible surface areas were calculated by the "
    "Shrake\u2013Rupley algorithm; surface residues were defined by relative ASA "
    "\u2265 0.25. Ca\u00b2\u207a-binding electrostatic hotspots were defined as surface Asp or "
    "Glu residues with APBS potential < \u22125 kT/e. Ensemble-level metrics reported "
    "include hotspot count (N_hot), per-hotspot mean SASA, hotspot fraction of total "
    "surface Asp/Glu, and median surface electrostatic potential."
)
cite(p_m_apbs, [12])

head("Finite-element analysis")

p_m_fea = mixed([
    ("Eggshell surface models derived from micro-CT were imported into LS-DYNA "
     "(Ansys) for explicit dynamic finite-element analysis (unit system: "
     "mm/kg/N/s). The eggshell was assigned elasto-plastic material properties "
     "(Young's modulus E\u202f=\u202f3.0 \u00d7 10\u00b9\u2070 Pa; yield strength \u03c3y\u202f=\u202f1.5 \u00d7 10\u2077 Pa; "
     "tangent modulus 0; maximum equivalent plastic strain at failure 0.05). "
     "The impactor, simulating the caruncle-borne egg tooth, was a frustum "
     "(base radius 0.1 mm; top radius 0.5 mm; height 0.5 mm) assigned IRON-ARMCO "
     "explicit material properties. Frictional contact between impactor and eggshell "
     "was set at \u03bc\u202f=\u202f0.2. Shell mesh element sizes were 0.05 mm (", False, False),
    ("G. gallus", False, True),
    ("), 0.05 mm (", False, False),
    ("A. platyrhynchos", False, True),
    ("), and 0.03 mm (", False, False),
    ("C. livia", False, True),
    ("), ensuring \u2265 6 element layers across the shell cross-section; the impactor "
     "was meshed at 0.1 mm. An initial velocity of 50,000 mm/s was applied to the "
     "impactor; symmetric fixed-support boundary conditions were applied to four "
     "lateral faces of each eggshell disc fragment (diameter 2.0 mm). Analyses "
     "ran for 1.0 \u00d7 10\u207b\u2074 s (time-step safety factor 0.7; automatic mass scaling; "
     "minimum time step 1 \u00d7 10\u207b\u2078 s; double-precision arithmetic; 100 equidistant "
     "output intervals). Resultant contact forces (RCFORC) were recorded at "
     "1.0 \u00d7 10\u207b\u2076 s intervals. The impactor was positioned at nine lateral-offset "
     "locations on a 3 \u00d7 3 grid (0.5 mm spacing) per species, from which peak "
     "contact force (F_max) and peak contact shear stress (\u03c4_max) were extracted "
     "per position.", False, False),
])
cite(p_m_fea, [16, 17])

head("Statistical analysis")

mixed([
    ("All values are expressed as mean \u00b1 s.d. All statistical tests are two-tailed "
     "and a p-value < 0.05 was considered statistically significant throughout. "
     "Mammillary morphometric parameters were compared among species by one-way ANOVA "
     "followed by Duncan's multiple range test (DMRT; \u03b1\u202f=\u202f0.05) "
     "(n\u202f=\u202f2 fragments per species). Glycan ensemble geometric descriptors (Rg, "
     "end-to-end distance, minimum glycan\u2013protein contact distance) and per-ensemble "
     "hotspot SASA metrics were compared among species by one-way ANOVA followed by "
     "Duncan's multiple range test (DMRT; \u03b1\u202f=\u202f0.05). "
     "Glycosylation-induced reduction in N_hot within ", False, False),
    ("C.\u00a0livia", False, True),
    (" was assessed by one-sample t-test versus the apo reference value (t\u2081\u2083); "
     "total Asp/Glu SASA differences between glycosylated and apo structures were "
     "evaluated by one-sample t-test against the apo reference value; shifts in "
     "median surface electrostatic potential were assessed by one-sample t-test "
     "against the apo reference value. Protein\u2013glycan abundance "
     "coupling was quantified by Spearman rank correlation of log\u2082-transformed "
     "protein and glycan-site intensities. Finite-element simulation outcomes "
     "(F_max, \u03c4_max) were compared among species by one-way ANOVA with Duncan's "
     "multiple range test (DMRT; \u03b1\u202f=\u202f0.05). All statistical analyses were "
     "conducted in Python using scipy.stats and statsmodels.", False, False),
])

# ════════════════════════════════════════════════════════════════════════════
# References
# ════════════════════════════════════════════════════════════════════════════
para("References", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

_refs = [
    # Science Advances citation format: F. Lastname, F. Lastname, ..., Title. J. Abbrev. Vol, pages (year).
    # Note: Full author lists required by SA; "et al." entries below need expanding before final submission.
    "1. J. Gautron, M. T. Hincke, A. B. Rodr\u00edguez-Navarro, Y. Nys, Avian eggshell biomineralization: "
    "an update on its structure, mineralogy and protein tool kit. BMC Mol. Cell Biol. 22, 11 (2021).",
    "2. M. L. H. Rose, M. T. Hincke, Protein constituents of the eggshell: eggshell-specific matrix "
    "proteins. Cell. Mol. Life Sci. 66, 2707\u20132719 (2009).",
    "3. J. Stiller, [full author list \u2014 to be completed], Complexity of avian evolution revealed by "
    "family-level genomes. Nature (2024).",
    "4. Y. Nys, N. Le Roy, Calcium homeostasis and eggshell biomineralization in female chicken. "
    "Adv. Protein Chem. Struct. Biol. (2018).",
    "5. F. K. Mendes, D. Vanderpool, B. Fulton, M. W. Hahn, CAFE 5 models variation in evolutionary "
    "rates among gene families. Bioinformatics 36, 5516\u20135518 (2021).",
    "6. F. Geng, Y. Huang, Y. Chen, D. Yao, Z. Jiang, M. Huang, Identification of N-glycosites in "
    "chicken egg white proteins using an omics strategy. J. Agric. Food Chem. 65, 5357\u20135364 (2017).",
    "7. D. Dai, [full author list \u2014 to be completed], Proteomic and N-glycosylation analysis of fertile "
    "egg white during storage and incubation in chickens. Poult. Sci. 104, 104526 (2025).",
    "8. D. J. Harvey, M. Rudd, R. A. Bateman, R. H. Bordoli, R. E. Donovan, J. B. Howes, "
    "Composition of N-linked carbohydrates from ovalbumin and co-purified glycoproteins. "
    "J. Am. Soc. Mass Spectrom. 11, 564\u2013571 (2000).",
    "9. K. Yamashita, T. Tachibana, T. Nakayama, M. Kitamura, Y. Ito, A. Kobata, Structural study of "
    "the carbohydrate moiety of hen ovomucoid. J. Biol. Chem. 257, 12809\u201312814 (1982).",
    "10. J. P. Reyes-Grajeda, A. Moreno, A. Romero, Crystal structure of ovocleidin-17, a major protein "
    "of the calcified Gallus gallus eggshell. J. Biol. Chem. 279, 40876\u201340881 (2004).",
    "11. Y.-X. Tsai, [full author list \u2014 to be completed], Rapid simulation of glycoprotein structures "
    "by grafting and steric exclusion of glycan conformer libraries. Cell 187, 1296\u20131311 (2024).",
    "12. E. Jurrus, D. Engel, K. Star, K. Monson, J. Brandi, L. E. Felberg, D. H. Brookes, "
    "L. Wilson, J. Chen, K. Liles, M. Chun, P. Li, D. W. Gohara, T. Dolinsky, R. Konecny, "
    "D. R. Koes, J. E. Nielsen, T. Head-Gordon, W. Geng, R. Krasny, G.-W. Wei, M. J. Holst, "
    "J. A. McCammon, N. A. Baker, Improvements to the APBS biomolecular solvation software suite. "
    "Protein Sci. 27, 112\u2013128 (2018).",
    "13. A. Bar, Calcium transport in strongly calcifying laying birds: the role of calbindin and plasma "
    "membrane calcium ATPase. Comp. Biochem. Physiol. A Mol. Integr. Physiol. 152, 447\u2013469 (2009).",
    "14. C. Kern, [full author list \u2014 to be completed], Functional annotations of three domestic animal "
    "genomes provide vital resources for comparative and agricultural research. "
    "Nat. Commun. 12, 1821 (2021).",
    "15. I. B. R. Scheiber, [full author list \u2014 to be completed], The importance of the "
    "altricial\u2013precocial spectrum for social complexity in mammals and birds: a review. "
    "Front. Zool. 14, 3 (2017).",
    "16. F. Liu, X. Jiang, Z. Chen, L. Wang, Mechanical design principles of avian eggshells for "
    "survivability. Acta Biomater. 178, 233\u2013243 (2024).",
    "17. D. Athanasiadou, W. Jiang, D. Goldbaum, A. Saleem, K. S. Bhatt, H. S. Michelmore, "
    "R. Marchessault, M. T. Hincke, M. D. McKee, Nanostructure, osteopontin, and mechanical "
    "properties of calcitic avian eggshell. Sci. Adv. 4, eaar3219 (2018).",
    "18. L. Zeng, X. Shi, L. Xuan, J. Zheng, Comparative N-glycoproteomic investigation of eggshell "
    "cuticle and mineralized layer proteins. J. Agric. Food Chem. 71, 10448\u201310458 (2023).",
    "19. M. T. Hincke, Y. Nys, J. Gautron, K. Mann, A. B. Rodr\u00edguez-Navarro, M. D. McKee, "
    "Molecular cloning and ultrastructural localization of the core protein of an eggshell matrix "
    "proteoglycan, ovocleidin-116. J. Biol. Chem. 274, 32915\u201332923 (1999).",
    "20. A. B. Rodr\u00edguez-Navarro, [full author list \u2014 to be completed], Amorphous calcium carbonate "
    "controls avian eggshell mineralization: a new paradigm for understanding rapid eggshell "
    "calcification. J. Struct. Biol. 190, 291\u2013303 (2015).",
    "21. K. Mann, M. T. Hincke, Y. Nys, Isolation of ovocleidin-116 from chicken eggshells, correction "
    "of its amino acid sequence and identification of disulfide bonds and glycosylated Asn. "
    "Matrix Biol. 21, 383\u2013387 (2002).",
    "22. J. A. Tobias, [full author list \u2014 to be completed], AVONET: morphological, ecological and "
    "geographical data for all birds. Ecol. Lett. 25, 581\u2013597 (2022).",
    "23. J. M. Starck, R. E. Ricklefs, Avian Growth and Development: Evolution within the "
    "Altricial\u2013Precocial Spectrum (Oxford Univ. Press, New York, 1998).",
    "24. R. O. Prum, [full author list \u2014 to be completed], A comprehensive phylogeny of birds (Aves) "
    "using targeted next-generation DNA sequencing. Nature 526, 569\u2013573 (2015).",
    "25. E. D. Jarvis, [full author list \u2014 to be completed], Whole-genome analyses resolve early "
    "branches in the tree of life of modern birds. Science 346, 1320\u20131331 (2014).",
    "26. X. Chen, [full author list \u2014 to be completed], Comparative study of eggshell antibacterial "
    "effectivity in precocial and altricial birds using Escherichia coli. "
    "PLoS ONE 14, e0220054 (2019).",
    "27. G. L. Krapu, Nutrition of female dabbling ducks during reproduction, in Waterfowl and Wetlands: "
    "An Integrated Review, G. A. Swanson, G. L. Krapu, T. A. Bookhout, Eds. "
    "(North Central Section, The Wildlife Society, Madison, 1979).",
]

for ref_text in _refs:
    p_ref = doc.add_paragraph(style="Normal")
    p_ref.paragraph_format.space_before = Pt(0)
    p_ref.paragraph_format.space_after = Pt(4)
    p_ref.paragraph_format.left_indent = Pt(18)
    p_ref.paragraph_format.first_line_indent = Pt(-18)
    r_ref = p_ref.add_run(ref_text)
    r_ref.font.size = Pt(9)
    rPr = r_ref._r.get_or_add_rPr()
    _set_font(rPr, FONT)

# ─────────────────────────────────────────────────────────────────────────
# Acknowledgments
# ─────────────────────────────────────────────────────────────────────────
para("Acknowledgments", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

para(
    "Author contributions: [to be completed]. "
    "Competing interests: The authors declare that they have no competing interests. "
    "Data availability: All data needed to evaluate the conclusions in the paper are "
    "present in the paper and/or the Supplementary Materials. "
    "Raw mass spectrometry data and proteomics search results have been deposited in a "
    "public repository [accession number to be provided upon acceptance]. "
    "Funding: [Funding sources to be listed before submission.]",
    bold=False, size=11, before=0, after=120
)

# ─────────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"[OK]  {OUT}")
