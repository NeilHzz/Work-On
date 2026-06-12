from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


ROOT = Path.cwd()
WRITING = next(p for p in ROOT.iterdir() if p.is_dir() and p.name.startswith("03_"))
SUPP_DOC_SRC = WRITING / "0_Supplementary_materials" / "supplementary_materials260605.docx"
SUPP_DOC_OUT = WRITING / "0_Supplementary_materials" / "supplementary_materials260608.docx"
TABLE_SRC = ROOT / "Supplementary" / "Tables"
TABLE_OUT = WRITING / "Supplementary Table"
MAIN_DOC = WRITING / "0_Manuscript" / "manuscript260608v2.docx"
MAIN_DOC_CN = WRITING / "0_Manuscript_CN" / "manuscript260608v2.docx"


FIGURE_CAPTIONS = {
    "Fig. S1.": (
        "Fig. S1. Sensitivity analysis of the macroecological species-selection framework.",
        "Distributions of variance explained (R²) and cluster silhouette coefficients from 500 randomized perturbation iterations applied to the AVONET-based principal-component space used for species selection. The focal species were Gallus gallus, Anas platyrhynchos, and Columba livia. Categorical ecological variables were numerically encoded, and each iteration shifted all encoding weights independently within ±30% of the original values. The narrow distributions around the baseline values indicate that focal-species separation was stable under alternative encoding schemes.",
    ),
    "Fig. S2.": (
        "Fig. S2. Avian phylogenetic context and comparative-axis coding for the focal species.",
        "Order-level avian phylogenetic relationships are shown with heatmap tracks for aquatic association, developmental mode, and lifestyle-habitat discordance. Colored order labels indicate the broader comparative frame used for species selection. The focal lineages occupy functional positions that only partly overlap with phylogeny, supporting a comparison that separates ecological-developmental state from ancestry.",
    ),
    "Fig. S3.": (
        "Fig. S3. Shared and lineage-restricted eggshell matrix orthogroups across the three species.",
        "OrthoFinder-based analysis resolved the eggshell matrix proteomes into a large three-species shared core, smaller pairwise-shared subsets, and lineage-restricted subsets. The bar plots summarize the number of species-specific orthogroups and the number of orthogroups assigned to each comparison set. The large shared core supports downstream analyses on a common matrix-protein background rather than on wholesale protein replacement.",
    ),
    "Fig. S4.": (
        "Fig. S4. Maximum-likelihood phylogenetic tree reconstructed from single-copy orthologs.",
        "The three-species tree was inferred by IQ-TREE from a concatenated alignment of single-copy orthologous protein sequences. Branch lengths indicate substitutions per site, and internal node labels show ultrafast bootstrap support from 1000 replicates. The topology places Gallus and Anas as sister lineages and Columba as the outgroup in the focal comparison.",
    ),
    "Fig. S5.": (
        "Fig. S5. GO enrichment across pairwise-shared and species-specific eggshell matrix orthogroups.",
        "The upper heatmap summarizes GO terms enriched in pairwise-shared orthogroup sets (GnA, Gallus-Anas; GnC, Gallus-Columba; AnC, Anas-Columba). The lower heatmap summarizes enriched GO terms in species-specific orthogroups from Gallus, Anas, and Columba. Colors denote GO category, including biological process, cellular component, and molecular function. The Gallus-specific set included protein N-linked glycosylation among enriched biological-process terms.",
    ),
    "Fig. S6.": (
        "Fig. S6. Glycan-class profiles of recurrent eggshell matrix proteins.",
        "Stacked bar plots show the relative abundance of detected glycan classes for OVAL, OC116, TRFE, and OC17 across chicken, duck, and pigeon. Glycan classes include high-mannose, paucimannose-truncated, neutral complex/hybrid, fucosylated complex/hybrid, and sialylated complex/hybrid structures. These protein-specific profiles identify OVAL as the clearest cross-species glycan-state contrast in the shared eggshell matrix background.",
    ),
    "Fig. S7.": (
        "Fig. S7. OVAL surface-potential distributions and residue-level APBS potential maps.",
        "Panel A shows species-level distributions of surface APBS potential for OVAL structural ensembles. Brackets indicate statistical comparisons among species. Panel B maps residue-level APBS potentials along the OVAL sequence for glycosylated and apo models. Rows are grouped by species and model state, and colors encode electrostatic potential values. The maps provide the residue-resolved electrostatic context used to interpret Ca²⁺-accessible surface differences.",
    ),
    "Fig. S8.": (
        "Fig. S8. CAFE5 gene-family expansion and contraction across the three focal species.",
        "The species tree is annotated with lineage-specific gene-family expansion (red) and contraction (blue) events inferred by CAFE5 using the divergence-time tree. Numbers at internal nodes indicate estimated ancestral gene-family sizes, and numbers on branches indicate net gene-family change. Pie charts summarize the proportion of expanded and contracted families for each lineage. Only gene families with per-family Viterbi p < 0.05 are shown.",
    ),
    "Fig. S9.": (
        "Fig. S9. Functional enrichment associated with gene-family turnover.",
        "Alluvial plots connect species, gene-family turnover direction, and enriched Gene Ontology terms inferred from expanded and contracted gene families. Flow colors distinguish expansion and contraction signals. Terminal blocks summarize enriched biological-process, cellular-component, and molecular-function terms for each lineage-specific turnover class.",
    ),
    "Fig. S10.": (
        "Fig. S10. Re-Glyco ensemble analysis of OVAL Ca²⁺ hotspot exposure and glycan shielding.",
        "Panel A shows total exposed Ca²⁺ hotspot counts across species-specific OVAL structural ensembles using the exposed-SASA threshold shown above the plot. Panel B shows glycan-shielded hotspot counts, defined by the glycan-induced SASA change threshold shown above the plot. Letters denote post hoc group differences after ANOVA. Panel C shows hotspot-count trajectories across 50 conformation models for each species-level ensemble. Together, these panels show that chicken retained the highest exposed Ca²⁺ hotspot state, whereas pigeon carried stronger glycan shielding and duck occupied an intermediate range.",
    ),
    "Fig. S11.": (
        "Fig. S11. Finite-element Y-force and contact-stress time courses across offset loading positions.",
        "Panels A and B show chicken simulations, panels C and D show duck simulations, and panels E and F show pigeon simulations. For each species, the 3 × 3 panels show contact Y-force time courses across nine lateral target-plate offsets, with peak force and contact-stress markers annotated on each curve. The paired summary panel overlays the nine Y-force trajectories for the same species. Solid and dashed curves show the paired force and contact-stress readouts used to derive peak F_max and τ_max values reported in the main text and Fig. 5.",
    ),
}


TABLES = [
    {
        "src": "SuppTable1_Protein_MS.xlsx",
        "out": "Table_S1_Protein_MS.xlsx",
        "number": "Table S1",
        "title": "Protein mass spectrometry data for eggshell matrix proteomes.",
        "main_text": "Supports the eggshell matrix proteome comparison and the shared-protein background used for cross-species analysis.",
        "contents": "Species-level protein and peptide quantification sheets for Anas, Columba, and Gallus.",
    },
    {
        "src": "SuppTable2_Glycan_MS.xlsx",
        "out": "Table_S2_Glycan_MS.xlsx",
        "number": "Table S2",
        "title": "Intact glycopeptide mass spectrometry data for eggshell matrix proteins.",
        "main_text": "Supports the glycan-class comparison of OVAL, OC116, TRFE, and OC17.",
        "contents": "Species-level intact glycopeptide and glycosite quantification sheets.",
    },
    {
        "src": "SuppTable3_Ortholog_GO_CAFE5.xlsx",
        "out": "Table_S3_Ortholog_GO_CAFE5.xlsx",
        "number": "Table S3",
        "title": "Ortholog identification, GO enrichment, and CAFE5 gene-family analysis.",
        "main_text": "Contains the target UniProt ortholog identifiers used for downstream structural and quantitative analyses.",
        "contents": "BlastP screening parameters, target ortholog identifiers, orthogroup input tables, GO enrichment results, and CAFE5 output.",
    },
    {
        "src": "SuppTable4_JointAnalysis.xlsx",
        "out": "Table_S4_Protein_Glycan_JointAnalysis.xlsx",
        "number": "Table S4",
        "title": "Integrated protein abundance and glycan abundance analysis.",
        "main_text": "Supports protein-glycan abundance integration across the three species.",
        "contents": "Species-level integrated protein intensity, glycan intensity, glycosite position, and log2 abundance summaries.",
    },
    {
        "src": "SuppTable5_ReGlyco_Ensemble_Results.xlsx",
        "out": "Table_S5_ReGlyco_Ensemble_Results.xlsx",
        "number": "Table S5",
        "title": "Re-Glyco ensemble structural and APBS analysis results.",
        "main_text": "Supports OVAL structural ensemble, surface accessibility, hotspot, and electrostatic analyses.",
        "contents": "Ensemble summary, APBS glycan-aware output, ensemble SASA output, and apo APBS output.",
    },
    {
        "src": "SuppTable6_ReGlyco_Ensemble_Stats.xlsx",
        "out": "Table_S6_ReGlyco_Ensemble_Stats.xlsx",
        "number": "Table S6",
        "title": "Re-Glyco glycan conformational statistics.",
        "main_text": "Supports glycan geometry and conformational-state interpretation for OVAL structural ensembles.",
        "contents": "Species-level glycan IDs, source folders, torsion-angle statistics, and conformational summary values.",
    },
    {
        "src": "SuppTable7_FEA.xlsx",
        "out": "Table_S7_FEA.xlsx",
        "number": "Table S7",
        "title": "Finite-element reaction force and contact-stress data.",
        "main_text": "Supports Fig. 5 and the hatching-relevant local shell-response comparison.",
        "contents": "Species-level Y-force summaries, offset-specific time-course data, peak-force comparisons, tau_max comparisons, and raw Duncan-test input.",
    },
]


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.text = text


def rewrite_supplementary_doc() -> None:
    doc = Document(SUPP_DOC_SRC)

    for para in doc.paragraphs:
        if para.text.strip() == "Cross-species OVAL glycan states connect mammillary-layer organisation to hatching-favourable eggshell mechanics":
            replace_paragraph(para, "OVAL glycan states link eggshell matrix chemistry to avian shell-breaking mechanics")
        elif "centred" in para.text:
            replace_paragraph(para, para.text.replace("centred", "centered"))

    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        key = next((k for k in FIGURE_CAPTIONS if para.text.strip().startswith(k)), None)
        if not key:
            continue
        title, legend = FIGURE_CAPTIONS[key]
        replace_paragraph(para, title)
        if i + 1 < len(paragraphs):
            replace_paragraph(paragraphs[i + 1], legend)

    doc.save(SUPP_DOC_OUT)


def write_table_description(ws, info: dict[str, str]) -> None:
    ws.title = "Table_Description"
    ws.append(["Field", "Description"])
    rows = [
        ("Table number", info["number"]),
        ("Table title", info["title"]),
        ("Main-text link", info["main_text"]),
        ("Contents", info["contents"]),
        ("Source file", info["src"]),
        ("Rearranged file", info["out"]),
    ]
    for row in rows:
        ws.append(row)

    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 110
    ws.freeze_panes = "A2"


def reorder_tables() -> None:
    if TABLE_OUT.exists():
        shutil.rmtree(TABLE_OUT)
    TABLE_OUT.mkdir(parents=True)

    for info in TABLES:
        src = TABLE_SRC / info["src"]
        out = TABLE_OUT / info["out"]
        wb = load_workbook(src)
        if "Table_Description" in wb.sheetnames:
            del wb["Table_Description"]
        desc = wb.create_sheet("Table_Description", 0)
        write_table_description(desc, info)
        wb.save(out)


def fix_main_table_reference(path: Path, old: str, new: str) -> bool:
    if not path.exists():
        return False
    doc = Document(path)
    changed = False
    for para in doc.paragraphs:
        if old in para.text:
            replace_paragraph(para, para.text.replace(old, new))
            changed = True
    if changed:
        doc.save(path)
    return changed


def main() -> None:
    rewrite_supplementary_doc()
    reorder_tables()
    fixed_en = fix_main_table_reference(MAIN_DOC, "Supplementary Table\u00a01", "Supplementary Table\u00a03")
    fixed_en = fix_main_table_reference(MAIN_DOC, "Supplementary Table 1", "Supplementary Table 3") or fixed_en
    fixed_cn = fix_main_table_reference(MAIN_DOC_CN, "Supplementary Table\u00a01", "Supplementary Table\u00a03")
    fixed_cn = fix_main_table_reference(MAIN_DOC_CN, "Supplementary Table 1", "Supplementary Table 3") or fixed_cn

    print(f"Supplementary document: {SUPP_DOC_OUT}")
    print(f"Supplementary tables: {TABLE_OUT}")
    print(f"Main manuscript table reference fixed: {fixed_en}")
    print(f"CN manuscript table reference fixed: {fixed_cn}")


if __name__ == "__main__":
    main()
