from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path(__file__).resolve().parent / "OVAL_CaCO3_experiment_plan.docx"


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)


def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def add_paragraph(document: Document, text: str, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.4
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_number(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_heading(document: Document, text: str, level=1):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return paragraph


def add_reference_list(document: Document, references):
    for idx, ref in enumerate(references, start=1):
        add_number(document, f"{idx}. {ref}")


def build_document() -> None:
    doc = Document()
    set_page_layout(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("OVAL糖基化/去糖基化 CaCO3 体外矿化实验计划")
    set_run_font(title_run, size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("适用于当前稿件的并行补实验与投稿准备")
    set_run_font(subtitle_run, size=11)

    add_heading(doc, "一、实验目标", level=1)
    add_paragraph(
        doc,
        "本实验用于直接测试天然糖基化 OVAL 与 PNGase F 去糖后的 OVAL 是否会在 CaCO3 体外矿化过程中表现出不同的成核动力学、亚稳相稳定能力和终末晶体形态。该实验的价值不在于再次证明 OVAL 存在，而在于把稿件中的 glycan-state divergence 从比较性证据推进到更接近功能验证的层面。",
    )
    add_paragraph(
        doc,
        "建议将本实验作为当前高水平投稿尝试期间并行推进的补实验。主线使用直接混合动力学矿化实验，辅助使用蒸气扩散法展示晶体 habit 变化；CaCl2 溶液中的蛋白构象测量可以作为 supplementary control，而不应替代矿化功能实验。",
    )

    add_heading(doc, "二、核心实验组", level=1)
    groups = [
        "天然糖基化 chicken OVAL。",
        "PNGase F 去糖后的 OVAL。",
        "Mock 处理 OVAL：与去糖组完全相同的缓冲液、孵育时间和温度，但不加入 PNGase F。",
        "无蛋白空白组。",
        "可选：BSA 对照组，用于表明观察到的效应不是任意蛋白的普遍效应。",
    ]
    for item in groups:
        add_bullet(doc, item)

    add_heading(doc, "三、关键文献锚点", level=1)
    add_paragraph(doc, "下面这些文献分别提供了 OVAL 直接调控 CaCO3 矿化、蛋壳基质蛋白体外矿化方法学和蛋壳矿化 ACC 机制这三类支撑。")
    references = [
        "Wolf, S. E. et al. 2011. The evolution of biomimetic and biologically inspired chemistry. Journal of the American Chemical Society. DOI: 10.1021/ja202622g. 该文是 OVAL 稳定 amorphous calcium carbonate 的直接机制锚点。",
        "Wang, X. et al. 2009. Ovalbumin-stabilized metastable phases of calcium carbonate. Journal of Physical Chemistry B. DOI: 10.1021/jp810281f. 该文支持利用相态和形态读出评估 OVAL 对 vaterite 和 calcite 转化的影响。",
        "Wang, X. et al. 2010. Effects of ovalbumin on CaCO3 precipitation kinetics and morphology. Journal of Physical Chemistry B. DOI: 10.1021/jp1008237. 该文最接近本实验的直接混合动力学设计。",
        "Pipich, V. et al. 2008. Time-resolved analysis of OVAL-mediated calcium carbonate nucleation. Journal of the American Chemical Society. DOI: 10.1021/ja801798h. 该文支撑对早期成核动力学进行连续读数。",
        "Hincke, M. T. 1995. Ovalbumin as a component of chicken eggshell matrix. Connective Tissue Research. DOI: 10.3109/03008209509010814. 该文把 chicken OVAL 与 eggshell matrix 直接关联。",
        "Gautron, J. et al. 1996. Hen eggshell soluble matrix alters in vitro CaCO3 precipitation rate and crystal morphology. British Poultry Science. DOI: 10.1080/00071669608417914. 该文证明蛋壳基质蛋白的体外矿化验证在逻辑上是成立的。",
        "Lakshminarayanan, R. et al. 2003. Eggshell matrix proteins and in vitro mineralization. Journal of Biological Chemistry. 该文提供蛋壳基质体外矿化的经典方法学框架。",
        "Lakshminarayanan, R. et al. 2006. Transient ACC precursor in eggshell biomineralization. Biomacromolecules. DOI: 10.1021/bm0605412. 该文支持将 ACC 稳定视为蛋白调控矿化的重要层面。",
        "Chien, Y.-C. et al. 2008. Matrix proteins alter calcite crystal habit. Journal of Structural Biology. DOI: 10.1016/j.jsb.2008.04.008. 该文可作为晶体 habit 改变的形态学方法学先例。",
        "Rodríguez-Navarro, A. B. et al. 2015. Avian eggshell mineralization and ACC. Journal of Structural Biology. DOI: 10.1016/j.jsb.2015.04.014. 该文适合在讨论部分连接 avian eggshell biomineralization 与 ACC 机制。",
    ]
    add_reference_list(doc, references)

    add_heading(doc, "四、详细操作流程", level=1)
    add_heading(doc, "4.1 蛋白准备", level=2)
    protein_prep = [
        "优先采购 chicken egg white 来源的天然 ovalbumin，不使用重组蛋白。起始批次建议先做 glycoprotein stain 或 lectin blot，确认蛋白确实带 N-糖。",
        "用 20 mM HEPES、150 mM NaCl、pH 7.4 配制 2 mg/mL OVAL 储备液。4°C 轻柔溶解 1 至 2 小时，随后以 12000 g 离心 10 分钟去除聚集体。",
        "使用 BCA 对同一批蛋白定量后，再分成 native、PNGase F、mock 三份，避免批次差异。",
    ]
    for item in protein_prep:
        add_number(doc, item)

    add_heading(doc, "4.2 PNGase F 去糖与 mock 对照", level=2)
    deglyco = [
        "推荐先采用非变性去糖条件：50 mM sodium phosphate，pH 7.5；OVAL 终浓度 1 mg/mL；PNGase F 用量每 mg 蛋白 500 至 1000 U；37°C 孵育 24 小时。若切除不足，可延长到 36 至 48 小时。",
        "Mock 组完全复制上述缓冲液组成、时间和温度，但不加入 PNGase F。",
        "不建议一开始使用 SDS 或强高温变性法去糖，因为那样很难区分观察到的是糖链效应还是蛋白变性效应。",
    ]
    for item in deglyco:
        add_number(doc, item)

    add_heading(doc, "4.3 去糖验证", level=2)
    validation = [
        "至少完成 SDS-PAGE 加 glycoprotein stain，确认去糖组糖染信号明显下降。",
        "同步观察 SDS-PAGE 条带轻微下移。由于 OVAL 只有一个 N-glycan，条带位移通常有限，但仍具有参考价值。",
        "如果平台允许，可增加 intact mass 或 LC-MS 作为更强验证，但不是当前最小可执行方案的必需项。",
    ]
    for item in validation:
        add_number(doc, item)

    add_heading(doc, "4.4 缓冲液置换", level=2)
    exchange_items = [
        "去糖结束后必须进行 buffer exchange，目标是彻底去除 phosphate。原因是 phosphate 会直接与 Ca2+ 反应，严重干扰 CaCO3 矿化。",
        "推荐使用 10 kDa 或 30 kDa 超滤管进行 3 至 4 次置换，目标缓冲液可选 10 mM HEPES，pH 7.8，或超纯水后在反应前调入工作液。",
        "最终将蛋白浓度调整为 0.2、0.5、1.0 mg/mL 三档储备液，用于后续初筛。",
    ]
    for item in exchange_items:
        add_number(doc, item)

    add_heading(doc, "4.5 主实验一：直接混合动力学矿化", level=2)
    kinetics_items = [
        "分别配制新鲜的 40 mM CaCl2 和 40 mM Na2CO3 储液，室温平衡 20 至 30 分钟，现配现用。",
        "将 native、deglyco、mock 蛋白分别配制到 40 mM CaCl2 中，使蛋白浓度为 200 μg/mL。与等体积 40 mM Na2CO3 混合后，最终体系为 20 mM CaCl2、20 mM Na2CO3、100 μg/mL 蛋白。",
        "建议先做 25、50、100、200 μg/mL 的浓度初筛。如果时间有限，先完成 100 μg/mL 与 blank 的比较。",
        "在 96 孔板中每孔加入 100 μL 含蛋白的 40 mM CaCl2，快速加入 100 μL 40 mM Na2CO3 启动反应，立即在酶标仪读取 OD570 或 OD600。",
        "读数间隔设为每 30 秒一次，总时长 30 至 60 分钟，温度保持 25°C。每组至少 6 个技术重复，整套实验独立重复 3 次。",
        "最终提取 lag time、最大斜率、30 分钟 AUC 和终点浊度四个动力学指标。",
    ]
    for item in kinetics_items:
        add_number(doc, item)

    add_heading(doc, "4.6 主实验二：平行管终点样品收集", level=2)
    endpoint_items = [
        "平行采用离心管体系，每个反应体积 1 mL：500 μL 含蛋白的 40 mM CaCl2 加 500 μL 40 mM Na2CO3。",
        "建议设置 5 分钟、30 分钟、120 分钟、24 小时四个时间点。前两个时间点用于观察早期颗粒和亚稳相，24 小时用于观察成熟晶体形态和终末晶型。",
        "每个时间点以 12000 g 离心 5 分钟，弃上清，用超纯水洗 2 次、无水乙醇洗 1 次，室温风干后用于 SEM、XRD 或 FTIR。",
    ]
    for item in endpoint_items:
        add_number(doc, item)

    add_heading(doc, "4.7 终点表征", level=2)
    characterization = [
        "SEM：将颗粒重悬于少量无水乙醇，滴加到硅片或导电样品台，自然干燥后喷金。每组至少采集 5 个随机视野，并用 ImageJ 统计粒径、长宽比和圆度；每组至少统计 200 个颗粒。",
        "XRD：优先用于区分 calcite 与 vaterite 的相态差异。如果样品量不足，可退而求其次采用 Raman 或 FTIR。",
        "FTIR 或 Raman：建议同时准备 calcite 标样和 vaterite 富集标样，以便解释特征峰比例变化。",
    ]
    for item in characterization:
        add_number(doc, item)

    add_heading(doc, "4.8 辅助实验：蒸气扩散法", level=2)
    vapor_diffusion = [
        "在密闭培养皿或 24 孔板中放置 NH4HCO3 作为 CO2 和碳酸源。",
        "在盖玻片或小反应皿中加入 100 至 200 μL 含 10 mM CaCl2 和蛋白的滴液，25°C 静置 24 至 48 小时。",
        "取出后直接进行光镜或 SEM 观察。该方法适合展示晶体 habit 的变化，但不建议作为正文唯一主结果。",
    ]
    for item in vapor_diffusion:
        add_number(doc, item)

    add_heading(doc, "4.9 补充实验：CaCl2 中蛋白构象测量", level=2)
    structural_control = [
        "建议作为 supplementary figure，目的是回答去糖后是否只是让 OVAL 发生明显变性。",
        "可选技术包括 CD、内源荧光、DLS 和 DSC。建议使用 OVAL 0.2 mg/mL，10 mM HEPES，pH 7.8，CaCl2 梯度为 0、1、5、10 mM。",
        "如果构象变化不大，但矿化行为明显改变，则有助于支持糖链状态本身参与矿化调控。",
    ]
    for item in structural_control:
        add_number(doc, item)

    add_heading(doc, "五、统计与结果判读", level=1)
    interpretation = [
        "动力学参数建议对 lag time、max slope、AUC 和终点浊度做 one-way ANOVA 加 Tukey 事后检验。",
        "SEM 定量建议每组 3 个独立样品，每个样品至少统计 200 个颗粒。",
        "相态分析建议结合 XRD 峰面积或 Raman/FTIR 特征峰比值做组间比较。",
        "预期最有价值的结果模式是：native OVAL 更能延缓或改写沉淀动力学，并更容易保留 ACC 或 vaterite 特征；deglyco OVAL 更容易转向规则 calcite 或更成熟的终末晶体。",
        "不应在实验开始前预设单一方向，因为 OVAL 在不同浓度和离子强度下可能既表现为促进，也可能表现为延迟。真正重要的是不同组之间是否形成系统性分离。",
    ]
    for item in interpretation:
        add_bullet(doc, item)

    add_heading(doc, "六、最小可执行版本", level=1)
    minimal = [
        "实验组限定为 native OVAL、PNGase F OVAL、mock、blank 四组。",
        "主浓度先做 100 μg/mL。",
        "完成直接混合动力学曲线。",
        "完成 24 小时终点 SEM。",
        "在 XRD 与 FTIR 中至少完成一项。",
        "至少完成一次 glycoprotein stain 验证去糖。",
    ]
    for item in minimal:
        add_bullet(doc, item)
    add_paragraph(doc, "如果当前目标是在不拖慢投稿的前提下尽快获得一个可用于返修或二审补件的功能验证结果，这一版本已经足够。")

    add_heading(doc, "七、常见风险点", level=1)
    pitfalls = [
        "未去除 phosphate 即进入矿化体系，会直接干扰 CaCO3 反应。",
        "采用强变性去糖却没有设置恰当对照，会让结论退化为构象效应而不是糖链效应。",
        "只做终点 SEM 而不做动力学，会削弱机制说服力。",
        "只做动力学而没有相态和形态读出，会削弱矿化层面的直接证据。",
        "只测试单一蛋白浓度，容易因为条件偶然性错过最有分辨率的窗口。",
    ]
    for item in pitfalls:
        add_bullet(doc, item)

    add_heading(doc, "八、建议的图版结构", level=1)
    figure_plan = [
        "Figure panel 1：去糖验证，含 SDS-PAGE 和 glycoprotein stain。",
        "Figure panel 2：动力学曲线及 lag time、max slope、AUC 定量。",
        "Figure panel 3：XRD 或 Raman/FTIR 的相态比较。",
        "Figure panel 4：SEM 晶体形态图与颗粒统计。",
        "Supplementary panel：CaCl2 中 OVAL 构象测量。",
    ]
    for item in figure_plan:
        add_bullet(doc, item)

    doc.save(OUT)


if __name__ == "__main__":
    build_document()