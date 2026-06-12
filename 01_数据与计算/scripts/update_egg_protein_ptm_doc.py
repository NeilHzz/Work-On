from docx import Document

DOC_PATH = r"e:\Data\Desktop\Work On\eggshell_matrix_ptm_direct_related_refs.docx"

NEW_HEADING = "三、扩展到鸡蛋蛋白的直接修饰研究"
NEW_NOTE = "注：本节将范围从蛋壳基质蛋白扩展到鸡蛋白、卵黄、卵黄膜及代表性鸡蛋蛋白的糖基化、磷酸化等翻译后修饰研究，文内仅保留 DOI。"

REFERENCES = [
    "Cavallero GJ, Landoni M, Couto AS. 2020. In depth N-glycoproteomics shows glyco-features of chicken egg white. Food Bioscience. DOI: 10.1016/j.fbio.2020.100590",
    "Landoni M, Rodriguez M, Couto AS. 2024. Unveiling the presence of O-glycosylation in different glycoproteins present in chicken egg white. Food Bioscience. DOI: 10.1016/j.fbio.2024.103938",
    "Xiao J, Wang J, Cheng L, Gao S, Li S, Qiu N, Li H, Peng L, Geng F. 2020. A puzzle piece of protein N-glycosylation in chicken egg: N-glycoproteome of chicken egg vitelline membrane. International Journal of Biological Macromolecules. DOI: 10.1016/j.ijbiomac.2020.08.193",
    "Sun H, Qiu N, Keast R, Wang H, Li B, Huang Q, Li S. 2019. Comparative Quantitative Phosphoproteomic Analysis of the Chicken Egg during Incubation Based on Tandem Mass Tag Labeling. Journal of Agricultural and Food Chemistry. DOI: 10.1021/acs.jafc.9b04638",
    "Meng Y, Diao C, Qiu N, Mine Y, Keast R, Meng S, Zhu C. 2021. Comparative N-glycoproteomic analysis of Tibetan and lowland chicken fertilized eggs: Implications on proteins biofunction and species evolution. Journal of Food Biochemistry. DOI: 10.1111/jfbc.14006",
    "Wallace RA, Morgan JP. 1986. Chromatographic resolution of chicken phosvitin. Multiple macromolecular species in a classic vitellogenin-derived phosphoprotein. Biochemical Journal. DOI: 10.1042/BJ2400871",
    "Miller MS, Mas MT, White HB. 1984. Highly phosphorylated region of chicken riboflavin-binding protein: chemical characterization and phosphorus-31 NMR studies. Biochemistry. DOI: 10.1021/bi00298a027",
    "Suzuki T, Kitajima K, Emori Y, Inoue Y, Inoue S. 1997. Site-specific de-N-glycosylation of diglycosylated ovalbumin in hen oviduct by endogenous peptide:N-glycanase as a quality control system for newly synthesized proteins. Proceedings of the National Academy of Sciences of the United States of America. DOI: 10.1073/pnas.94.12.6244",
    "Nau F, Pasco M, Desert C, Molle D, Croguennec T, Guerin-Dubiard C. 2005. Identification and Characterization of Ovalbumin Gene Y in Hen Egg White. Journal of Agricultural and Food Chemistry. DOI: 10.1021/jf048369l",
]


def ensure_not_already_present(document: Document) -> bool:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == NEW_HEADING:
            return False
    return True


def main() -> None:
    document = Document(DOC_PATH)

    if not ensure_not_already_present(document):
        print("Heading already present; no changes made.")
        return

    document.add_heading(NEW_HEADING, level=1)
    document.add_paragraph(NEW_NOTE)
    for reference in REFERENCES:
        document.add_paragraph(reference)

    document.save(DOC_PATH)
    print(f"Updated {DOC_PATH} with {len(REFERENCES)} references.")


if __name__ == "__main__":
    main()
