from __future__ import annotations

from pathlib import Path

from docx import Document


BASE = Path.cwd()
WORK_DIR = next(p for p in BASE.iterdir() if p.is_dir() and p.name.startswith("03_"))
EN_DOCX = WORK_DIR / "0_Manuscript" / "manuscript260608v2.docx"
CN_DOCX = WORK_DIR / "0_Manuscript_CN" / "manuscript260608v2.docx"

OLD_EN = (
    "Our study extends current understanding of avian eggshell biomineralization by linking glycan-state variation "
    "on conserved matrix proteins to mammillary-layer organization and local shell-breaking mechanics."
)
NEW_EN = (
    "The avian eggshell is built by matrix-guided mineralization, but the molecular surface states that connect "
    "nucleation-site formation to hatching-relevant mechanics remain poorly resolved."
)

OLD_CN = "本研究通过将保守基质蛋白上的糖链状态变异与乳突层组织和局部破壳力学联系起来，拓展了我们对鸟类蛋壳生物矿化的理解。"
NEW_CN = "鸟类蛋壳由基质引导的矿化过程构建而成，但连接成核位点形成与孵化相关力学的分子表面状态仍缺乏清晰解析。"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def update(path: Path, old: str, new: str) -> None:
    doc = Document(str(path))
    paragraph = doc.paragraphs[4]
    text = paragraph.text
    if old not in text:
        raise RuntimeError(f"Opening sentence not found in {path}")
    replace_paragraph_text(paragraph, text.replace(old, new, 1))
    doc.save(str(path))


def main() -> None:
    update(EN_DOCX, OLD_EN, NEW_EN)
    update(CN_DOCX, OLD_CN, NEW_CN)
    print(EN_DOCX)
    print(CN_DOCX)


if __name__ == "__main__":
    main()
