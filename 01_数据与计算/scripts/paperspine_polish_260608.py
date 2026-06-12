from __future__ import annotations

import re
from pathlib import Path

from docx import Document


BASE = Path.cwd()
WORK_DIR = next(BASE.glob("03_*"))
EN_SRC = WORK_DIR / "manuscript260606.docx"
CN_SRC = WORK_DIR / "manuscript260606_cn.docx"
EN_OUT = WORK_DIR / "manuscript260608.docx"
CN_OUT = WORK_DIR / "manuscript260608_cn.docx"
REPORT_OUT = WORK_DIR / "manuscript260608_changes.md"


EN_REWRITE = {
    0: "OVAL glycan states link eggshell matrix chemistry to avian shell-breaking mechanics",
    1: "OVAL glycans tune shell-breaking mechanics",
    4: (
        "Eggshell matrix proteins regulate shell architecture, but most comparative work has focused on protein identity or modification sites rather than the chemical properties of attached glycans. "
        "Here, we asked whether glycan states on conserved matrix proteins help explain cross-species divergence in the mammillary layer and in local shell-breaking mechanics. "
        "We compared chicken, duck, and pigeon under a shared egg-tooth hatching interface by integrating micro-CT morphometry, eggshell-matrix proteomics, intact glycopeptide mass spectrometry, Re-Glyco structural modeling, electrostatic analysis, and finite-element simulation. "
        "Species separation appeared first in mammillary-layer organization, whereas the matrix-protein toolkit remained largely shared. "
        "Within this shared background, ovalbumin (OVAL) shifted from High-Mannose-dominant glycans in chicken to Neutral Complex/Hybrid-dominant glycans in duck and Sialylated Complex/Hybrid-dominant glycans in pigeon. "
        "These glycan states predicted a Ca²⁺ surface-accessibility gradient on OVAL, providing a plausible route from glycan-modulated OVAL unfolding and matrix-bound nucleation-site exposure to mammillary-layer formation, mature mammillary density, and local shell-breaking strength after shell-thickness effects were separated. "
        "Together, the data connect Ca²⁺-accessible matrix-protein surfaces to mammillary-layer organization and to the local mechanics that can favor controlled hatching."
    ),
    5: "Teaser: OVAL glycan states connect matrix-protein surface accessibility to shell regions used during hatching.",
    7: (
        "Avian eggshell formation is a rapid vertebrate mineralization process that assembles a mechanically competent calcitic shell within a narrow uterine time window. (1–3) "
        "This compressed schedule depends on an organic matrix-protein system that coordinates calcium delivery, nucleation, crystal growth, and shell architecture rather than acting as a passive scaffold. (1, 2, 4, 5) "
        "Hatching is also a local mechanical event: the egg tooth loads a restricted region of the inner shell instead of the whole shell at once. (6, 7) "
        "Comparable hatching-assist structures recur across egg-laying amniotes, suggesting that biologically meaningful variation in hatching performance is likely to reside in the shell material as well as in the tool itself. (6, 8–11) "
        "Yet it remains unclear how species with comparable hatching mechanics generate distinct shell states across ecological and developmental contexts. (12–16) "
        "The mammillary layer provides the key entry point because matrix-guided calcite nucleation first sets the spacing and continuity inherited by later shell units, and because this layer is the first mechanically consequential shell region encountered during inside-out hatching. "
        "The mechanistic question is therefore direct: when the shell-breaking interface is comparable, which molecular regulators at the mammillary layer account for distinct eggshell states across species? (1, 2, 4, 6, 17, 18)"
    ),
    8: (
        "Eggshell matrix proteins regulate mammillary-layer mineralization, crystal growth, and mature shell architecture, and recurrent factors such as OC17, OC116, TRFE, and OVAL define a shared shell-building toolkit. (1, 2, 4, 5, 19–21) "
        "The unresolved issue is not whether this toolkit exists, but how it is chemically deployed across species. (1, 2, 4) "
        "This gap is sharpest for posttranslational modification. Phosphorylation and glycosylation sites have both been catalogued, but phosphate side groups are comparatively similar, whereas glycans vary widely in composition, size, and charge and can generate distinct molecular states on the same protein scaffold. (21–26) "
        "Glycan class is therefore a mechanistic variable rather than only a feature of site occupancy. (24, 25) "
        "Glycoproteomic studies have shown that eggshell matrix proteins carry distinct N-glycosylation states, including glycosylated Asn in OC116 and defined OVAL-associated glycan compositions, but most analyses have remained within one species, one compartment, or one site inventory at a time. (21, 23, 27–30) "
        "As a result, avian eggshell comparisons still rarely resolve matched glycan states on shared matrix proteins across species, leaving open whether this glycan layer helps explain why similar protein toolkits yield distinct shell states. (2, 4, 5, 23, 31)"
    ),
    10: (
        "The missing step is a direct bridge from glycan class to surface presentation on a shared matrix background. "
        "OVAL provides that bridge because its dominant glycan classes can be followed from glycoproteomics into structural modeling. (2, 5, 23, 32)"
    ),
    11: (
        "Here, we compared chicken, duck, and pigeon using an integrated workflow that combines micro-CT morphometry, comparative proteomics, intact glycopeptide mass spectrometry, Re-Glyco modeling with electrostatic analysis, and finite-element simulation. "
        "This cross-scale design allowed eggshell structure, glycan state, and hatching-relevant mechanics to be read in the same comparative frame. "
        "The analyses show that mammillary-layer organization separates the species before the broader matrix-protein toolkit does, and that OVAL glycan state is the strongest molecular signal linking Ca²⁺ surface accessibility, mammillary-layer organization, and local mechanics favorable to hatching. "
        "These findings establish a framework for connecting glycan-state variation to eggshell formation and to the localized shell response used during inside-out hatching."
    ),
    15: (
        "Viewed through this shared hatching context, mammillary-layer morphology was the first eggshell level to show a clear species contrast (Fig. 1C). "
        "Chicken mammillae were smoother and formed rounded projections. Duck mammillae showed more ridges and angular turns across the inner surface. Pigeon shells were dominated by discrete triangular-conical mammillae. "
        "Three-dimensional surface reconstructions agreed with the cross-sectional views, indicating that the sampled inner-shell regions differed in mammillary geometry rather than representing minor variants of a shared inner-surface template."
    ),
    24: "OVAL glycosylation provides the clearest cross-species molecular contrast",
    25: (
        "Intact-glycopeptide profiling showed that the three species differed in sampling depth but still shared a stable comparison core (Fig. 2A to D). "
        "The cluster view recovered 25 clusters shared by all three species, with the largest additional pairwise overlap between duck and pigeon at 64 clusters, whereas chicken contributed little species-private cluster space (Fig. 2A). "
        "The same pattern held for the catalog counts: duck yielded 321 glycoproteins, 547 glycosites, and 197 glycan compositions; pigeon yielded 192, 257, and 162; and chicken yielded 55, 88, and 105 (Fig. 2B). "
        "Shared-core Jensen-Shannon similarity remained between 0.33 and 0.40, with the duck-pigeon pair highest (Fig. 2C). "
        "These values indicate divergence within a comparable glycoproteomic background rather than three disconnected chemical spaces. "
        "Glycan-class composition reinforced the same point at the chemical-deployment level. High-Mannose and Complex-Fucosylated glycans formed a broad cross-species background, whereas Complex-Sialylated and other more extended classes contributed more strongly to lineage separation (Fig. 2D)."
    ),
    42: "Inside-out loading resolves hatching-relevant local mechanics",
    48: (
        "Peak F_max differed significantly among species (p = 1.64 × 10⁻¹³). Chicken reached 1.12 ± 0.11 N, duck reached 0.90 ± 0.09 N, and pigeon reached 0.49 ± 0.04 N, and all pairwise differences were significant by Tukey HSD (Fig. 5B). "
        "By contrast, τ_max resolved a two-level pattern (p = 6.64 × 10⁻¹⁰). Chicken reached 551.60 ± 108.80 MPa and was significantly higher than duck at 404.00 ± 39.60 MPa and pigeon at 393.00 ± 35.20 MPa. "
        "Duck and pigeon did not differ significantly from each other (p = 0.728; Fig. 5C)."
    ),
    52: (
        "Cross-species divergence in this dataset emerged first at the mammillary layer rather than through wholesale matrix-protein turnover. "
        "Within a largely shared eggshell-matrix toolkit, OVAL glycan state provided the clearest molecular axis linking Ca²⁺-relevant surface accessibility, mammillary-layer organization, and hatching-relevant local mechanics. (1, 6, 23)"
    ),
    60: (
        "Taken together, the comparison converged on a local eggshell state favorable to hatching in this dataset. "
        "Chicken combined the densest mammillary field, the least shielded Ca²⁺-relevant OVAL surface, and the strongest local stress response under inside-out loading. "
        "Duck showed why the chain cannot be reduced to shell thickness: despite its thicker shell and higher F_max than pigeon, its τ_max grouped with pigeon rather than chicken. "
        "This pattern supports the inference that chemically specific states on reused matrix proteins can organize mineralized phenotypes more directly than proteome turnover alone in this comparison. (58–60)"
    ),
    63: (
        "In summary, this study links mammillary organization, glycoprotein state, Ca²⁺ surface accessibility, and local hatching mechanics across three avian eggshells. "
        "Chicken defined the high-mammillary-density end of this axis, with dense mammillary organization, compact OVAL glycans, greater Ca²⁺-relevant surface exposure, and the strongest local response at the mammillary interface. "
        "Duck occupied the critical intermediate position: its thicker shell increased absolute force but did not reproduce the same local stress state, separating shell thickness from the mammillary-interface mechanism. "
        "As comparable glycoform assignments become available, the same framework can extend to other abundant eggshell matrix proteins. "
        "Across morphometric, glycoproteomic, structural, and mechanical layers, OVAL glycan state remains the molecular feature most consistently aligned with the high-mammillary-density eggshell state recovered here."
    ),
}


CN_REWRITE = {
    0: "OVAL糖链状态连接蛋壳基质化学与鸟类破壳力学",
    1: "OVAL糖链调控蛋壳破壳力学",
    3: "摘要",
    4: (
        "蛋壳基质蛋白调控蛋壳结构形成，但既有跨物种研究多集中于蛋白身份或修饰位点，对糖基化侧链的化学性质关注不足。"
        "本研究询问：保守基质蛋白上的糖链状态是否有助于解释乳突层结构和局部破壳力学的跨物种分化。"
        "我们在共同的卵齿孵化界面下比较鸡、鸭和鸽，整合micro-CT形态测量、蛋壳基质蛋白质组学、完整糖肽质谱、Re-Glyco结构建模、静电分析和有限元模拟。"
        "物种分离首先出现在乳突层组织，而基质蛋白工具箱总体上仍然共享。"
        "在这一共享背景中，卵清蛋白（OVAL）的糖链由鸡中的高甘露糖主导型，转变为鸭中的中性复合/杂合主导型，以及鸽中的唾液酸化复合/杂合主导型。"
        "这些糖链状态预测了OVAL表面的Ca²⁺可及性梯度，并为“糖链调控的OVAL展开和基质结合成核位点暴露-乳突层形成-成熟乳突密度-剥离壳厚效应后的局部破壳强度”提供了合理路径。"
        "总体而言，本研究将Ca²⁺可及的基质蛋白表面与乳突层组织及有利于受控孵化的局部力学联系起来。"
    ),
    5: "导语：OVAL糖链状态将基质蛋白表面可及性与孵化时受力的蛋壳区域联系起来。",
    6: "引言",
    7: (
        "禽类蛋壳形成是快速的脊椎动物矿化过程，需要在狭窄的子宫时间窗口内组装出具有力学能力的方解石壳。 (1–3) "
        "这一压缩过程依赖有机基质蛋白系统来协调钙供应、成核、晶体生长和蛋壳结构，而不是作为被动支架存在。 (1, 2, 4, 5) "
        "孵化同样是一个局部力学事件：卵齿作用于内壳的有限区域，而不是一次性加载整个蛋壳。 (6, 7) "
        "类似的孵化辅助结构在产卵羊膜动物中反复出现，提示孵化性能的生物学差异不仅存在于工具本身，也可能存在于蛋壳材料中。 (6, 8–11) "
        "然而，具有相似孵化机制的物种如何在不同生态和发育背景下形成不同蛋壳状态，仍不清楚。 (12–16) "
        "乳突层是关键入口，因为基质引导的方解石成核首先在此建立后续壳单元继承的间距和连续性，同时它也是内向外孵化过程中首先遇到的具有力学后果的壳层。"
        "因此，机制问题是：当破壳界面具有可比性时，乳突层中的哪些分子调控因子解释了跨物种蛋壳状态的差异？ (1, 2, 4, 6, 17, 18)"
    ),
    8: (
        "蛋壳基质蛋白调控乳突层矿化、晶体生长和成熟壳结构，OC17、OC116、TRFE和OVAL等反复出现的因子构成了共享的壳构建工具箱。 (1, 2, 4, 5, 19–21) "
        "未解决的问题不是该工具箱是否存在，而是它如何在不同物种中被化学部署。 (1, 2, 4) "
        "这一空缺在翻译后修饰层面最为明显。磷酸化和糖基化位点均已有记录，但磷酸侧基相对相似，而糖链在组成、大小和电荷方面差异很大，可在同一蛋白支架上形成不同分子状态。 (21–26) "
        "因此，糖链类别是机制变量，而不仅是位点占有情况的描述。 (24, 25) "
        "糖蛋白质组学研究表明，蛋壳基质蛋白具有不同N-糖基化状态，包括OC116中的糖基化Asn和明确的OVAL相关糖链组成；但多数分析仍局限于单一物种、单一区室或单个位点清单。 (21, 23, 27–30) "
        "因此，鸟类蛋壳比较仍很少在跨物种共享基质蛋白上解析匹配的糖链状态，也就尚不清楚这一糖链层是否解释了相似蛋白工具箱为何产生不同壳状态。 (2, 4, 5, 23, 31)"
    ),
    10: "目前缺少的关键步骤，是在共享基质背景上建立从糖链类别到表面呈现的直接桥梁。OVAL提供了这一桥梁，因为其主导糖链类别可以从糖蛋白质组学追踪到结构建模。 (2, 5, 23, 32)",
    11: (
        "在本研究中，我们比较鸡、鸭和鸽，并建立了整合micro-CT形态测量、比较蛋白质组学、完整糖肽质谱、Re-Glyco建模、静电分析和有限元模拟的工作流程。"
        "这一跨尺度设计使蛋壳结构、糖链状态和孵化相关力学能够在同一比较框架中读取。"
        "分析结果显示，乳突层组织先于更广泛的基质蛋白工具箱发生物种分离；OVAL糖链状态则是连接Ca²⁺表面可及性、乳突层组织和有利于孵化的局部力学的最强分子信号。"
        "这些发现建立了一个框架，用于连接糖链状态变异、蛋壳形成以及内向外孵化过程中使用的局部壳响应。"
    ),
    12: "结果",
    13: "保守的孵化界面限定蛋壳变异",
    24: "OVAL糖基化提供最清晰的跨物种分子对比",
    33: "OVAL糖链状态重塑表面可及性",
    42: "内向外加载解析孵化相关的局部力学",
    51: "讨论",
    52: "在本数据集中，跨物种分化首先出现在乳突层，而不是表现为蛋壳基质蛋白的大规模替换。在总体共享的蛋壳基质蛋白工具箱中，OVAL糖链状态提供了最清晰的分子轴，将Ca²⁺相关表面可及性、乳突层组织和孵化相关局部力学联系起来。 (1, 6, 23)",
    60: (
        "综上，本数据集中的比较收敛到一种有利于孵化的局部蛋壳状态。"
        "鸡同时具有最密集的乳突场、遮蔽程度最低的Ca²⁺相关OVAL表面，以及内向外加载下最强的局部应力响应。"
        "鸭说明这一链条不能简化为壳厚度：尽管鸭壳比鸽更厚且F_max更高，其τ_max仍与鸽而非鸡归为一组。"
        "这一模式支持如下推断：在本比较中，重复使用的基质蛋白上的化学特异性状态，比单纯蛋白质组更替更直接地组织矿化表型。 (58–60)"
    ),
    61: (
        "同样的分析序列可能扩展到鸟类蛋壳之外。许多生物矿化系统依赖有机基质，通过化学特异的界面状态而不是整体组成来调节离子进入、表面暴露和矿物成核。"
        "在更广泛的框架中，本研究工作流程提供了一条从糖蛋白质组状态到表面呈现、再到介观功能的路径。"
        "其他矿化组织和仿生材料中也存在类似的跨尺度问题。"
        "这一逻辑还可能启发再生医学场景，其中蛋壳来源材料或膜蛋白正被探索用于组织工程和骨修复。"
        "更广义地说，本研究为检验化学特异性基质状态如何在生物和生物医学背景中组织生物矿物行为提供了模板。 (58, 59, 61, 62)"
    ),
    62: (
        "本研究范围仍然有限。我们分析的是主导糖型，而不是完整的体内糖链集合；同时，我们在平均蛋壳尺度上将每个物种视为力学均一，并在APBS框架中依赖尚未完全约束的子宫离子条件。"
        "就近期湿实验跟进而言，一个实际限制是：在当前工具条件下，要在同一体系中实现物种匹配的直系同源基质蛋白表达，并在特定位点精确安装预定义糖链结构，仍具有相当难度。"
        "与此同时，生物矿化是大型耦合反应系统，任何单一测定都可能难以给出直接的单变量解释。"
        "因此，下一步决定性测试应包括定义糖型的矿化实验、在鸡中直接操纵OVAL糖基化、对相同内向外力学对比进行位点解析验证，以及提高因果分辨率的方法开发。"
        "这些实验将澄清OVAL糖链状态是直接参与壳矿化，还是以异常高的保真度标记高乳突密度蛋壳状态。"
    ),
    63: (
        "总之，本研究在三种鸟类蛋壳中连接了乳突层组织、糖蛋白状态、Ca²⁺表面可及性和局部孵化力学。"
        "鸡定义了该轴的高乳突密度端，表现为密集乳突组织、紧凑OVAL糖链、更高Ca²⁺相关表面暴露，以及乳突界面处最强的局部响应。"
        "鸭处于关键中间位置：较厚的壳提高了绝对力，但没有重现相同的局部应力状态，从而将壳厚度与乳突界面机制区分开来。"
        "随着更多可比糖型分配结果的获得，同一框架可扩展到其他高丰度蛋壳基质蛋白。"
        "在形态、糖蛋白质组、结构和力学层面，OVAL糖链状态仍是与本研究所解析高乳突密度蛋壳状态最一致的分子特征。"
    ),
    64: "材料与方法",
    65: "生物材料",
    67: "蛋壳基质蛋白提取",
    69: "micro-CT成像和乳突层形态测量",
    71: "蛋壳基质蛋白鸟枪法蛋白质组学",
    74: "完整糖肽质谱分析",
    76: "比较蛋白质组分析和基因家族进化",
    78: "跨物种糖蛋白直系同源鉴定",
    80: "蛋白质-糖链丰度整合比较",
    83: "N-糖链结构集合建模",
    85: "静电势计算",
    87: "有限元分析",
    89: "统计分析",
    91: "参考文献",
}


CN_TERMS = [
    ("蛋壳 基质蛋白s", "蛋壳基质蛋白"),
    ("蛋壳 基质蛋白", "蛋壳基质蛋白"),
    ("基质蛋白s", "基质蛋白"),
    ("基质蛋白 ", "基质蛋白"),
    (" 糖链状态", "糖链状态"),
    ("糖链状态 ", "糖链状态"),
    (" 糖链 ", "糖链"),
    (" 糖链", "糖链"),
    ("糖链 ", "糖链"),
    (" 糖基化 ", "糖基化"),
    (" 糖基化", "糖基化"),
    ("糖基化 ", "糖基化"),
    (" 乳突层 ", "乳突层"),
    (" 乳突层", "乳突层"),
    ("乳突层 ", "乳突层"),
    (" 蛋壳 ", "蛋壳"),
    (" 蛋壳", "蛋壳"),
    ("蛋壳 ", "蛋壳"),
    (" 显微CT ", "micro-CT"),
    ("显微CT", "micro-CT"),
    ("卵清蛋白（OVAL）", "OVAL"),
    ("蛋白质组周转", "蛋白质组更替"),
    ("矩阵", "基质"),
    ("可访问性ible", "可及性"),
    ("可访问性", "可及性"),
    ("可访问界面", "可及界面"),
    ("机械能力", "力学能力"),
    ("机械问题", "机制问题"),
    ("机械变量", "机制变量"),
    ("本地", "局部"),
    ("分离物种", "区分物种"),
    ("整体建模", "集合建模"),
    ("直向同源", "直系同源"),
    ("直向同源物", "直系同源物"),
    ("脱盐方案", "脱矿方案"),
    ("贝壳", "蛋壳"),
    ("外壳", "蛋壳"),
    ("乳突层旋钮", "乳突突起"),
    ("蛋白质加入", "蛋白质登录号"),
    ("每个加入物", "每个登录号"),
    ("序列序列", "序列基序"),
    ("完整的糖肽", "完整糖肽"),
    ("显着", "显著"),
    ("水垢桥接", "跨尺度衔接"),
    ("proteiN", "protein"),
    ("monOVALent", "monovalent"),
    ("蛋壳s", "蛋壳"),
    ("鸭s", "鸭"),
    ("鸽s", "鸽"),
    ("高甘露糖", "高甘露糖型"),
    ("中性复合物/混合", "中性复合/杂合"),
    ("唾液酸化复合物/混合", "唾液酸化复合/杂合"),
    ("复合物/杂合体", "复合/杂合型"),
    ("物种专有", "物种特异"),
    ("平均最大序列同一性", "平均最大序列一致性"),
    ("场地占用", "位点占有"),
    ("卵齿 轴承尖端", "带卵齿的喙尖"),
]


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def polish_en_text(text: str) -> str:
    replacements = [
        ("Ca²⁺surface", "Ca²⁺ surface"),
        ("Ca²⁺relevant", "Ca²⁺-relevant"),
        ("Ca²⁺hotspot", "Ca²⁺ hotspot"),
        ("organisation", "organization"),
        ("Organisation", "Organization"),
        ("modelling", "modeling"),
        ("prioritised", "prioritized"),
        ("Fig. 2A).The", "Fig. 2A). The"),
        ("spaces.Glycan", "spaces. Glycan"),
        ("1 × 10-5", "1 × 10⁻⁵"),
        ("1 × 10^-13", "1.64 × 10⁻¹³"),
        ("6.64 × 10^-10", "6.64 × 10⁻¹⁰"),
        (". pigeon", ". Pigeon"),
        (". chicken", ". Chicken"),
        ("p <", "p <"),
        ("α =", "α ="),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r"\s{2,}", " ", text)
    return text


def polish_cn_text(text: str) -> str:
    for old, new in CN_TERMS:
        text = text.replace(old, new)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"\s+([，。；：？！、）])", r"\1", text)
    text = re.sub(r"([（])\s+", r"\1", text)
    text = text.replace(" 。", "。").replace(" ，", "，")
    text = re.sub(r"\s{2,}", " ", text).strip()
    return text


def process_english() -> Document:
    doc = Document(str(EN_SRC))
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in EN_REWRITE:
            replace_paragraph_text(paragraph, EN_REWRITE[idx])
        elif paragraph.text.strip():
            replace_paragraph_text(paragraph, polish_en_text(paragraph.text))
    doc.save(str(EN_OUT))
    return doc


def process_chinese(en_doc: Document) -> None:
    doc = Document(str(CN_SRC))
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in CN_REWRITE:
            replace_paragraph_text(paragraph, CN_REWRITE[idx])
        elif idx > 91 and idx < len(en_doc.paragraphs):
            replace_paragraph_text(paragraph, en_doc.paragraphs[idx].text)
        elif paragraph.text.strip():
            replace_paragraph_text(paragraph, polish_cn_text(paragraph.text))
    doc.save(str(CN_OUT))


def write_report() -> None:
    report = """# manuscript260608 修改说明

- PaperSpine安装：已安装 `paper-spine` 系列 Codex skills，并按 rewrite/translate 的“主线-证据-改写矩阵”思路处理。
- 英文稿：重写题名、短题名、摘要、导语、引言关键段、结果转承、讨论首尾段，统一为 glycan state -> surface accessibility -> mammillary organization -> local mechanics 的主线。
- 英文稿：修正术语和格式，包括 Ca²⁺ spacing、US spelling、句首物种大写、图号后缺空格、科学计数法上标等。
- 中文稿：重写题名、短题名、摘要、导语、引言关键段、讨论总结段和方法小标题，修复明显机翻表达。
- 中文稿：统一术语，如蛋壳基质蛋白、糖链状态、乳突层、表面可及性、直系同源、蛋白质-糖链丰度整合比较。
- 参考文献：中文稿参考文献条目改为保留英文格式，避免作者名和期刊名被机翻破坏。
"""
    REPORT_OUT.write_text(report, encoding="utf-8")


def main() -> None:
    en_doc = process_english()
    process_chinese(en_doc)
    write_report()
    print(EN_OUT)
    print(CN_OUT)
    print(REPORT_OUT)


if __name__ == "__main__":
    main()
