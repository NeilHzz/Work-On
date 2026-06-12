from __future__ import annotations

import re
from pathlib import Path

from docx import Document


BASE = Path.cwd()
WORK_DIR = next(p for p in BASE.iterdir() if p.is_dir() and p.name.startswith("03_"))
EN_SRC = WORK_DIR / "manuscript260608_citation_fixed.docx"
CN_SRC = WORK_DIR / "0_Manuscript_CN" / "manuscript260608_cn.docx"
EN_DIR = WORK_DIR / "0_Manuscript"
CN_DIR = WORK_DIR / "0_Manuscript_CN"
EN_OUT = EN_DIR / "manuscript260608v2.docx"
CN_OUT = CN_DIR / "manuscript260608v2.docx"
REPORT = WORK_DIR / "manuscript260608v2_changes.md"


EN_REPLACEMENTS = {
    4: (
        "Our study extends current understanding of avian eggshell biomineralization by linking glycan-state variation on conserved matrix proteins to mammillary-layer organization and local shell-breaking mechanics. "
        "We asked whether glycan states on conserved matrix proteins help explain cross-species divergence in the mammillary layer under a shared egg-tooth hatching interface. "
        "We compared chicken, duck, and pigeon by integrating micro-CT morphometry, eggshell-matrix proteomics, intact glycopeptide mass spectrometry, Re-Glyco structural modeling, electrostatic analysis, and finite-element simulation. "
        "Species separation appeared first in mammillary-layer organization, whereas the matrix-protein toolkit remained largely shared. "
        "Within this shared background, ovalbumin (OVAL) shifted from High-Mannose-dominant glycans in chicken to Neutral Complex/Hybrid-dominant glycans in duck and Sialylated Complex/Hybrid-dominant glycans in pigeon. "
        "Chicken OVAL retained the greatest Ca²⁺-accessible surface, supporting a model in which OVAL in a Ca²⁺-rich uterine-fluid environment can reach the Ca²⁺ load required for unfolding earlier, expose matrix-bound nucleation sites more efficiently, and generate a denser mammillary layer. "
        "This denser mammillary state was associated with the highest local shell-breaking response after shell-thickness effects were separated. "
        "Together, the data connect Ca²⁺-accessible matrix-protein surfaces to mammillary-layer organization and to the local mechanics that can favor controlled hatching."
    ),
    37: (
        "Matched glycosylated-versus-apo comparisons then showed that glycan addition changed the number of Ca²⁺-relevant hotspot residues and the exposed carboxylate surface most clearly in pigeon (Fig. 4K and L; Fig. S10). "
        "With panel B in mind, glycosylation preserves or hides the same Ca²⁺-relevant patches rather than creating new ones: the colored patches remain reachable, whereas the black patches are the same sites after shielding. "
        "Duck shifted in the same direction without a resolved structure-level significance call, and chicken could be assessed only descriptively because one glycosylated structure was available. "
        "Fig. 4K to N then collapsed the same comparison to the whole-interface level. Across those panels, chicken preserved the most accessible Ca²⁺-relevant surface, pigeon shifted the largest share into a glycan-affected state, and duck trended toward the lower-accessibility side but did not separate from pigeon or apo references uniformly across metrics. "
        "This ordering supports the structural premise that chicken OVAL would reach the Ca²⁺ load required for conformational opening more readily in a Ca²⁺-rich uterine-fluid environment, thereby exposing matrix-bound nucleation sites earlier or more efficiently. "
        "A higher density of available nucleation sites provides a direct route to denser mammillary-layer formation and to the stronger local response measured under inside-out loading. "
        "Taken together, Fig. 4A to N link glycan-dependent separation, glycan geometry, interface masking, Ca²⁺-relevant accessibility, and the mechanistic premise for testing whether the resulting mammillary organization affects local shell-breaking mechanics. (2, 5, 40)"
    ),
    46: (
        "Fig. 5. Local finite-element loading connects egg-tooth contact geometry to hatching-relevant shell mechanics. "
        "(A) Dorsal beak views and matched micro-CT shell-fragment finite-element models for chicken, pigeon, and duck from top to bottom. Each model shows the conical egg-tooth impactor and a representative local stress field on the reconstructed shell geometry. "
        "(B) Mean contact-force time courses across nine impact positions, with shaded ±1σ envelopes, and the corresponding peak contact force (F_max) distribution. "
        "(C) Mean contact shear-stress time courses across the same nine positions, with shaded ±1σ envelopes, and the corresponding peak shear stress (τ_max) distribution. "
        "In the peak-value summaries, star markers denote peaks that differ significantly by Tukey HSD, whereas dot markers denote peaks without significant pairwise difference. "
        "Box-plot points denote individual impact positions (n = 9 per species), bars show mean ± s.d., p values above the plots are one-way ANOVA omnibus p values, and different letters indicate Tukey HSD groupings at p < 0.05. "
        "Simulations used identical loading and material settings so that the comparisons isolate geometry-dependent local response."
    ),
    49: (
        "The difference between F_max and τ_max clarified the duck result. Its higher raw contact force was driven mainly by greater shell thickness (0.35 mm versus 0.19 mm in pigeon), not by a stronger local stress response at the mammillary interface. "
        "In other words, the thicker duck shell could carry more total contact force, but it did not show greater local resistance after the response was normalized to the contact-stress scale. "
        "By contrast, chicken exhibited a 36-40% increase in τ_max relative to the two other species, indicating a stronger hatching-relevant local stress response independent of shell thickness. "
        "This high-versus-low grouping, with chicken alone in the high group and duck together with pigeon in the low group, matched the grouping recovered for mammillary density by Tukey HSD (Fig. 1D). "
        "The mechanics therefore retained the contrast already recovered from mammillary organization and OVAL accessibility."
    ),
    54: (
        "Mammillary-layer mineralization mode remains the central structural level in the interpretation. (1, 17) "
        "In this model, compact chicken OVAL glycans preserve the greatest Ca²⁺-accessible acidic surface. Under the Ca²⁺-rich uterine-fluid conditions of shell mineralization, this surface state would allow OVAL to reach the Ca²⁺ load required for conformational opening more rapidly than more heavily shielded glycoforms. "
        "Earlier OVAL opening would expose matrix-bound nucleation sites sooner or more efficiently, increasing the density of early calcite nucleation events and producing a more compact mammillary field. "
        "Once early calcite crystal units are established, later eggshell regions inherit the spacing logic set in that first mineralization window. A dense mammillary field can therefore alter matrix retention, mineral continuity, local stress redistribution, and mature morphology. (1, 46, 47) "
        "That emphasis is consistent with earlier eggshell studies that place the mammillary layer at the intersection of crystal nucleation and matrix control. The present comparison extends that line by linking the layer to a cross-species glycan-state readout rather than to shell-quality descriptors alone. (1, 17, 48, 49) "
        "Recent poultry omics studies increasingly tie age, shell-gland transcription, extracellular-vesicle cargo, and other whole-shell quality traits to eggshell phenotype. Those descriptors, however, usually remain broader than the proximate material layer isolated here. (18, 50–52) "
        "For that reason, mammillary organization is not merely another shell trait, but the earliest material context in which matrix chemistry can plausibly bias later mechanical outcome. (1, 4) "
        "That structural position is why mammillary organization is the first cross-species difference that can be read as potentially consequential rather than merely descriptive. (1, 17)"
    ),
    57: (
        "Re-Glyco and APBS analyses provide the structural bridge in the argument. Across species, this produced a glycan-state gradient: compact chicken glycans preserved the most accessible acidic surface, neutral complex/hybrid duck glycans imposed an intermediate constraint, and extended sialylated pigeon glycans generated the strongest steric and electrostatic shielding. "
        "Earlier in vitro and structural work had already suggested that OVAL conformation and electrostatics matter during mineralization, but matched glycoform-resolved surface ensembles had not been compared across bird species. (2, 40) "
        "The relevant implication is not simply that glycans cover the protein surface, but that they alter how quickly a shared matrix protein may satisfy the Ca²⁺-loading requirement for unfolding in a mineralizing uterine environment. "
        "Chicken therefore represents the high-accessibility endpoint, in which OVAL would be expected to expose nucleation-relevant surfaces earlier, whereas pigeon represents the shielded endpoint, in which OVAL opening and nucleation-site presentation would be delayed or reduced. "
        "Although this result does not establish direct causality, it supports a restrained inference: different glycan states on the same matrix protein can alter the chemical surface presented to the mineralizing environment and may thereby contribute to the structural divergence observed here."
    ),
    58: (
        "The mechanical comparison targets inside-out loading during hatching rather than conventional outside compression or whole-shell fracture. That choice is critical: eggshell thickness inflates absolute failure load, whereas τ_max is less thickness-confounded and more directly reports stress transfer through the mammillary interface. (6, 45, 54) "
        "Sun et al. likewise showed from longitudinal and latitudinal measurements that eggshell thickness varies across the whole egg, with the circumferential zone around the blunt end being locally thinnest, reinforcing the need to separate global shell-thickness buffering from local hatching-interface mechanics. (55) "
        "The analysis therefore asks whether the inner mammillary interface retains the same contrast already inferred from matrix chemistry and morphology. (6, 45, 54, 56) "
        "Duck makes this separation explicit: its thicker shell increased F_max but did not recreate the high-τ_max state observed in chicken. "
        "This distinction separates thickness buffering and developmental background from the material pathway emphasized here. Eggshell thickness, body size, broad reproductive ecology, and lineage history all contribute background structure. (43, 57) "
        "Yet thickness-only explanations do not account for τ_max, and diffuse lineage-divergence explanations do not explain why the same ordering recurs across glycan class, electrostatic accessibility, mammillary organization, and hatching-relevant mechanics. (6, 44) "
        "What recurs is the same alignment: high OVAL Ca²⁺ accessibility, dense mammillary organization, and high τ_max in chicken, with duck and pigeon converging toward lower local-response states through different structural routes. "
        "In this framework, OVAL glycan state is the most directly readable molecular layer through which that state becomes mechanically interpretable."
    ),
    60: (
        "Taken together, the comparison converged on a local eggshell state favorable to hatching in this dataset. Chicken combined the densest mammillary field, the least shielded Ca²⁺-relevant OVAL surface, and the strongest local stress response under inside-out loading. "
        "This chain is coherent with a Ca²⁺-loading model: higher OVAL surface accessibility would promote earlier unfolding in the Ca²⁺-rich uterine-fluid environment, more efficient nucleation-site exposure, denser mammillary-layer formation, and greater local shell-breaking resistance. "
        "Duck showed why the chain cannot be reduced to shell thickness: despite its thicker shell and higher F_max than pigeon, its τ_max grouped with pigeon rather than chicken. "
        "From an evolutionary perspective, the sampled precocial-to-altricial contrast is consistent with the possibility that changes in life history and niche altered the required balance between shell protection and hatchling escape. "
        "In altricial birds such as pigeon, reduced hatchling force capacity could favor a shell state that is easier to break locally; altered glycosylation of key matrix proteins may contribute to that state by making OVAL harder to unfold, reducing nucleation-site density, and producing a more open mammillary layer. "
        "This interpretation remains a hypothesis to be tested across broader avian sampling, but it explains why chemically specific states on reused matrix proteins may organize mineralized phenotypes more directly than proteome turnover alone in this comparison. (58–60)"
    ),
    63: (
        "In summary, this study links mammillary organization, glycoprotein state, Ca²⁺ surface accessibility, and local hatching mechanics across three avian eggshells. Chicken defined the high-mammillary-density end of this axis, with compact OVAL glycans, greater Ca²⁺-relevant surface exposure, a predicted lower threshold for Ca²⁺-dependent OVAL opening, denser nucleation-site formation, and the strongest local response at the mammillary interface. "
        "Duck occupied the critical intermediate position: its thicker shell increased absolute force but did not reproduce the same local stress state, separating shell thickness from the mammillary-interface mechanism. "
        "Pigeon occupied the more shielded, lower-local-response endpoint, consistent with a model in which altered OVAL glycosylation delays unfolding and supports a less dense mammillary architecture that may reduce the structural difficulty of hatching for weaker altricial young. "
        "As comparable glycoform assignments become available, the same framework can extend to other abundant eggshell matrix proteins. Across morphometric, glycoproteomic, structural, and mechanical layers, OVAL glycan state remains the molecular feature most consistently aligned with the high-mammillary-density eggshell state recovered here."
    ),
    66: (
        "Fertilized eggs were collected from three avian lines: seven eggs from Chahua pink-shell laying hens, seven eggs from Shaoxing spotted green-shell ducks, and 19 eggs from White King pigeons. "
        "Gallus gallus eggs were obtained from the Poultry Resources Conservation Farm, China Agricultural University (Beijing, China); Columba livia eggs were obtained from the College of Veterinary Medicine, China Agricultural University (Beijing, China); and Anas platyrhynchos eggs were supplied by Jinxing Duck Industry (Beijing, China). "
        "All eggs were stored at 16°C for 7 d under breeder-egg holding conditions before analysis."
    ),
}


CN_REPLACEMENTS = {
    4: (
        "本研究通过将保守基质蛋白上的糖链状态变异与乳突层组织和局部破壳力学联系起来，拓展了我们对鸟类蛋壳生物矿化的理解。"
        "我们询问：在共同的卵齿孵化界面下，保守基质蛋白上的糖链状态是否有助于解释乳突层和局部破壳力学的跨物种分化。"
        "我们比较鸡、鸭和鸽，整合micro-CT形态测量、蛋壳基质蛋白质组学、完整糖肽质谱、Re-Glyco结构建模、静电分析和有限元模拟。"
        "物种分离首先出现在乳突层组织，而基质蛋白工具箱总体上仍然共享。"
        "在这一共享背景中，OVAL糖链由鸡中的高甘露糖主导型，转变为鸭中的中性复合/杂合主导型，以及鸽中的唾液酸化复合/杂合主导型。"
        "鸡OVAL保留了最高的Ca²⁺可及表面，支持这样一个模型：在富含Ca²⁺的子宫液环境中，鸡OVAL可更早满足解折叠所需的钙载荷，更高效地暴露基质结合成核位点，并形成更致密的乳突层。"
        "这种更致密的乳突层状态在剥离壳厚效应后，对应最高的局部破壳响应。"
        "总体而言，本研究将Ca²⁺可及的基质蛋白表面与乳突层组织及有利于受控孵化的局部力学联系起来。"
    ),
    37: (
        "糖基化与apo参考的匹配比较进一步显示，糖链添加对鸽中Ca²⁺相关热点残基数量和暴露羧酸盐表面的改变最为明显（图4K和L；图S10）。"
        "结合图4B理解，糖基化并不是创造新的Ca²⁺相关斑块，而是在保留或遮蔽同一组斑块：彩色区域表示仍可及的斑块，黑色区域表示被遮蔽后的同一位点。"
        "鸭也朝同一方向移动，但没有解析出结构水平显著性；鸡由于只有一种糖基化结构可用，只能进行描述性评估。"
        "图4K至N将同一比较推进到整体界面水平。在这些指标中，鸡保留了最高的Ca²⁺相关表面可及性，鸽将最大比例表面转入糖链影响状态，鸭则趋向较低可及性一侧，但并未在所有指标上与鸽或apo参考稳定分离。"
        "这一排序支持一个结构前提：在富含Ca²⁺的子宫液环境中，鸡OVAL更容易达到构象打开所需的钙载荷，从而更早或更高效地暴露基质结合成核位点。"
        "更高密度的可用成核位点为更致密的乳突层形成提供了直接路径，也为内向外加载下更强的局部响应提供了结构基础。"
        "综合来看，图4A至N将糖链相关分化、糖链几何、界面遮蔽、Ca²⁺相关可及性以及乳突层组织影响局部破壳力学的机制前提联系起来。 (2, 5, 40)"
    ),
    46: (
        "图5. 局部有限元加载将卵齿接触几何与孵化相关蛋壳力学联系起来。"
        "(A) 从上到下为鸡、鸽和鸭的背侧喙视图及匹配的micro-CT壳片段有限元模型。每个模型显示锥形卵齿冲击器以及重建壳几何上的代表性局部应力场。"
        "(B) 九个冲击位置的平均接触力时间曲线，阴影表示±1σ，并给出相应峰值接触力（F_max）分布。"
        "(C) 相同九个位置的平均接触剪应力时间曲线，阴影表示±1σ，并给出相应峰值剪应力（τ_max）分布。"
        "在峰值汇总图中，星形标记表示经Tukey HSD检验存在显著差异的峰值，点标记表示无显著成对差异的峰值。"
        "箱线图点表示单个冲击位置（每个物种n = 9），柱表示平均值±标准差，图上方p值为单因素方差分析总体p值，不同字母表示p < 0.05下的Tukey HSD分组。"
        "模拟采用相同加载和材料设置，使比较能够隔离几何依赖的局部响应。"
    ),
    49: (
        "F_max和τ_max之间的差异解释了鸭的结果。鸭较高的原始接触力主要由更大的蛋壳厚度驱动（0.35 mm，而鸽为0.19 mm），并不代表其乳突界面具有更强的局部应力响应。"
        "换言之，较厚的鸭壳可以承载更高的总接触力，但在接触应力尺度上并未表现出更强的局部抗力。"
        "相比之下，鸡的τ_max较其他两个物种提高36-40%，说明其具有更强的孵化相关局部应力响应，且该响应不依赖壳厚度。"
        "这种高低分组，即鸡单独位于高响应组、鸭和鸽位于低响应组，与Tukey HSD得到的乳突密度分组一致（图1D）。"
        "因此，力学结果保留了乳突层组织和OVAL可及性中已经出现的对比。"
    ),
    54: (
        "乳突层矿化方式仍是本研究解释中的核心结构层。 (1, 17) "
        "在这一模型中，紧凑的鸡OVAL糖链保留最高的Ca²⁺可及酸性表面。在蛋壳矿化所处的富Ca²⁺子宫液环境中，这种表面状态可使OVAL比遮蔽更强的糖型更快达到构象打开所需的钙载荷。"
        "更早的OVAL打开会更早或更高效地暴露基质结合成核位点，提高早期方解石成核事件密度，并形成更致密的乳突层场。"
        "一旦早期方解石晶体单元建立，后续蛋壳区域就会继承第一矿化窗口中确定的间距逻辑。致密乳突场因此可能改变基质保留、矿物连续性、局部应力再分配和成熟形态。 (1, 46, 47) "
        "这一解释与早期将乳突层置于晶体成核和基质控制交叉点的蛋壳研究一致。本研究进一步将该层与跨物种糖链状态读数联系起来，而不仅限于蛋壳质量描述。 (1, 17, 48, 49) "
        "近期家禽组学研究越来越多地将年龄、壳腺转录、细胞外囊泡货物以及其他全壳质量性状与蛋壳表型联系起来，但这些描述通常比本研究分离出的近端材料层更宽泛。 (18, 50–52) "
        "因此，乳突层组织不仅是另一种壳特征，而是基质化学最可能影响后续力学结果的最早材料背景。 (1, 4) "
        "这也解释了为什么乳突层组织是第一个可被解读为具有潜在功能后果，而不仅仅是描述性差异的跨物种差异。 (1, 17)"
    ),
    57: (
        "Re-Glyco和APBS分析为论证提供了结构桥梁。在不同物种中，这形成了糖链状态梯度：紧凑的鸡糖链保留最高的酸性表面可及性，中性复合/杂合鸭糖链施加中等约束，而延展的唾液酸化鸽糖链产生最强的空间和静电遮蔽。"
        "早期体外和结构研究已提示OVAL构象和静电性质在矿化过程中具有重要性，但匹配糖型解析的表面集合尚未在鸟类物种间比较。 (2, 40) "
        "这里的关键含义并不只是糖链覆盖蛋白表面，而是糖链会改变同一基质蛋白在矿化子宫环境中满足Ca²⁺载荷并发生解折叠的速度。"
        "因此，鸡代表高可及性端点，预期OVAL可更早暴露成核相关表面；鸽则代表遮蔽端点，预期OVAL打开和成核位点呈现会被延迟或降低。"
        "尽管该结果不能单独证明直接因果关系，但它支持一个克制推论：同一基质蛋白上的不同糖链状态可以改变呈现给矿化环境的化学表面，并可能由此参与本研究观察到的结构分化。"
    ),
    58: (
        "力学比较针对孵化期间的内向外加载，而不是传统外压或全壳断裂。该选择很关键：蛋壳厚度会提高绝对失效载荷，而τ_max受厚度混杂较小，更直接反映通过乳突界面的应力传递。 (6, 45, 54) "
        "Sun等也通过纵向和纬向测量显示，蛋壳厚度在整枚蛋上存在变化，钝端周围区域局部最薄，进一步说明需要区分全局壳厚缓冲和局部孵化界面力学。 (55) "
        "因此，分析的核心问题是内侧乳突界面是否保留了从基质化学和形态中已经推断出的同一对比。 (6, 45, 54, 56) "
        "鸭清楚地显示了这种分离：其较厚的壳提高了F_max，但没有重现鸡中观察到的高τ_max状态。"
        "这一差异将壳厚缓冲和发育背景与本文强调的材料路径区分开来。蛋壳厚度、体型、广义繁殖生态和谱系历史都会构成背景结构。 (43, 57) "
        "然而，仅用厚度无法解释τ_max，仅用弥散的谱系分化也无法解释为什么相同排序反复出现在糖链类别、静电可及性、乳突层组织和孵化相关力学中。 (6, 44) "
        "反复出现的是同一组对齐关系：鸡具有高OVAL Ca²⁺可及性、致密乳突层组织和高τ_max，而鸭和鸽则通过不同结构路径收敛到较低局部响应状态。"
        "在这一框架中，OVAL糖链状态是使该蛋壳状态具有力学解释性的最直接分子层。"
    ),
    60: (
        "综上，本数据集中的比较收敛到一种有利于孵化的局部蛋壳状态。鸡同时具有最密集的乳突场、遮蔽程度最低的Ca²⁺相关OVAL表面，以及内向外加载下最强的局部应力响应。"
        "这一链条与Ca²⁺载荷模型一致：更高的OVAL表面可及性可在富Ca²⁺子宫液环境中促进更早解折叠、更高效的成核位点暴露、更致密的乳突层形成，以及更高的局部破壳阻力。"
        "鸭说明这一链条不能简化为壳厚度：尽管鸭壳比鸽更厚且F_max更高，其τ_max仍与鸽而非鸡归为一组。"
        "从进化角度看，本研究采样的早成-晚成对比与一种可能性相一致：生活史和生态位变化改变了蛋壳保护与雏鸟出壳能力之间的平衡。"
        "在鸽等晚成鸟中，较弱的雏鸟出壳力量可能有利于局部更易破壳的蛋壳状态；关键基质蛋白糖基化状态的改变可能通过使OVAL更难解折叠、降低成核位点密度并形成更疏松的乳突层来参与这一状态。"
        "这一解释仍需更广泛鸟类采样检验，但它解释了为什么在本比较中，重复使用的基质蛋白上的化学特异性状态，可能比单纯蛋白质组更替更直接地组织矿化表型。 (58–60)"
    ),
    63: (
        "总之，本研究在三种鸟类蛋壳中连接了乳突层组织、糖蛋白状态、Ca²⁺表面可及性和局部孵化力学。"
        "鸡定义了该轴的高乳突密度端，表现为紧凑OVAL糖链、更高Ca²⁺相关表面暴露、预测更低的Ca²⁺依赖OVAL打开阈值、更密集的成核位点形成，以及乳突界面处最强的局部响应。"
        "鸭处于关键中间位置：较厚的壳提高了绝对力，但没有重现相同的局部应力状态，从而将壳厚度与乳突界面机制区分开来。"
        "鸽则位于更强遮蔽、较低局部响应端点，这与一个模型相符：OVAL糖基化改变延迟其解折叠，并支持较低密度乳突结构，从而可能降低较弱晚成雏鸟的结构性破壳难度。"
        "随着更多可比糖型分配结果的获得，同一框架可扩展到其他高丰度蛋壳基质蛋白。"
        "在形态、糖蛋白质组、结构和力学层面，OVAL糖链状态仍是与本研究所解析高乳突密度蛋壳状态最一致的分子特征。"
    ),
    66: (
        "从三个禽系中收集受精蛋：茶花粉壳蛋鸡7个蛋、绍兴斑点绿壳鸭7个蛋、白王鸽19个蛋。"
        "Gallus gallus鸡蛋获自中国农业大学家禽资源保护场（中国北京）；Columba livia种蛋获自中国农业大学兽医学院（中国北京）；Anas platyrhynchos鸭蛋由金星鸭实业（中国北京）提供。"
        "分析前，所有蛋均在16°C种蛋保存条件下保存7天。"
    ),
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def get_reference_texts(doc: Document) -> dict[int, str]:
    refs = {}
    in_refs = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "References":
            in_refs = True
            continue
        if in_refs and text:
            match = re.match(r"^(\d+)\.\s+.+$", text)
            if match:
                refs[int(match.group(1))] = text
    return refs


def apply_replacements(doc: Document, replacements: dict[int, str]) -> None:
    for idx, text in replacements.items():
        if idx < len(doc.paragraphs):
            replace_paragraph_text(doc.paragraphs[idx], text)


def sync_cn_references(cn_doc: Document, en_refs: dict[int, str]) -> None:
    in_refs = False
    for paragraph in cn_doc.paragraphs:
        text = paragraph.text.strip()
        if text in {"参考文献", "References"}:
            in_refs = True
            continue
        if not in_refs or not text:
            continue
        match = re.match(r"^(\d+)[\.\。]\s+.+$", text)
        if match:
            number = int(match.group(1))
            if number in en_refs:
                replace_paragraph_text(paragraph, en_refs[number])


def main() -> None:
    EN_DIR.mkdir(parents=True, exist_ok=True)
    CN_DIR.mkdir(parents=True, exist_ok=True)

    en_doc = Document(str(EN_SRC))
    apply_replacements(en_doc, EN_REPLACEMENTS)
    en_doc.save(str(EN_OUT))

    cn_doc = Document(str(CN_SRC))
    apply_replacements(cn_doc, CN_REPLACEMENTS)
    sync_cn_references(cn_doc, get_reference_texts(en_doc))
    cn_doc.save(str(CN_OUT))

    REPORT.write_text(
        "\n".join(
            [
                "# manuscript260608v2 修改说明",
                "",
                "- 输出英文版到 `0_Manuscript/manuscript260608v2.docx`。",
                "- 输出中文版到 `0_Manuscript_CN/manuscript260608v2.docx`。",
                "- 摘要首句改为中性扩展句，避免强转折语气。",
                "- Fig. 5B/C 图注明确：星形标记代表显著差异峰值，点标记代表无显著差异峰值。",
                "- 改写 duck F_max 与 τ_max 解释，明确“较高总接触力不等于更强局部单位应力响应”。",
                "- 在摘要、结果和讨论补入 Ca²⁺可及性、OVAL解折叠、成核位点密度、乳突层致密度和破壳力学之间的逻辑链。",
                "- 在讨论中以克制语气补入早成-晚成生态/生活史解释。",
                "- Materials and Methods 删除产蛋中期，规范品种名，并删除鸽子来源中的个人姓名。",
                "- 中文稿参考文献同步为英文 DOI 版。",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    print(EN_OUT)
    print(CN_OUT)
    print(REPORT)


if __name__ == "__main__":
    main()
