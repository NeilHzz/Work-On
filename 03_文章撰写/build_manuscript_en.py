"""
Science Advances manuscript generator.
Builds manuscript260602v2.docx from the scripted manuscript source.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re
from shared_references import REFS

OUT = str(Path(__file__).with_name("0_Manuscript") / "manuscript_submitted.docx")
FIG_BASE = Path(__file__).resolve().parent.parent / "02_可视化" / "00_正文与补充材料图片" / "main_composed"

REF_TEXTS = {}
for ref_text in REFS:
    match = re.match(r"^(\d+)\.\s+(.*)$", ref_text)
    if not match:
        raise ValueError(f"Invalid reference entry: {ref_text}")
    REF_TEXTS[int(match.group(1))] = match.group(2)

CITATION_ORDER = []
CITATION_MAP = {}

doc = Document()

# 鈹€鈹€ 椤甸潰 鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€
s = doc.sections[0]
s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)
s.page_width  = Inches(8.5)
s.page_height = Inches(11)

# 鈹€鈹€ 琛屽彿锛堣繛缁紝姣忛〉閲嶇疆锛?
_lnNum = OxmlElement("w:lnNumType")
_lnNum.set(qn("w:countBy"), "1")
_lnNum.set(qn("w:restart"), "continuous")
_lnNum.set(qn("w:start"), "1")
s._sectPr.append(_lnNum)

# 鈹€鈹€ 杈呭姪 鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€鈹€
FONT = "Times New Roman"   # Science Advances: use universal fonts

# Latin taxon names should be italicized consistently in running text.
TAXON_PATTERN = re.compile(
    r"(Gallus\s+gallus|Anas\s+platyrhynchos|Columba\s+livia|"
    r"G\.\s*[\u00A0\s]*gallus|A\.\s*[\u00A0\s]*platyrhynchos|C\.\s*[\u00A0\s]*livia|"
    r"\bGallus\b|\bAnas\b|\bColumba\b)"
)
COMMON_SPECIES_NAMES = {"chicken", "duck", "pigeon"}

def _set_font(rPr, name):
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)
    rPr.insert(0, rFonts)

def fmt(run, size=11, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    _set_font(rPr, FONT if not italic else "Times New Roman")

def add_text_with_taxon_italics(p, text, size=11, bold=False, italic=False):
    if not p.runs:
        text = text.lstrip()
    if text.strip().lower() in COMMON_SPECIES_NAMES:
        italic = False
    if italic:
        r = p.add_run(text)
        fmt(r, size=size, bold=bold, italic=True)
        return

    parts = TAXON_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        run_italic = bool(TAXON_PATTERN.fullmatch(part))
        r = p.add_run(part)
        fmt(r, size=size, bold=bold, italic=run_italic)

def spacing(p, before=0, after=120, line=24):
    """line in pt; double-spaced for Science Advances initial submission"""
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement("w:spacing")
    e.set(qn("w:before"),   str(before))
    e.set(qn("w:after"),    str(after))
    e.set(qn("w:line"),     str(line * 20))
    e.set(qn("w:lineRule"), "auto")
    pPr.append(e)

def para(text, bold=False, italic=False, size=11,
         before=0, after=120, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    add_text_with_taxon_italics(p, text, size=size, bold=bold, italic=italic)
    return p

def runpara(parts, size=11, before=0, after=120, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    for text, superscript in parts:
        r = p.add_run(text.lstrip() if not p.runs else text)
        fmt(r, size=size)
        r.font.superscript = superscript
    return p

def mixed(parts, before=0, after=120, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    for text, bold, italic in parts:
        if not p.runs:
            text = text.lstrip()
        add_text_with_taxon_italics(p, text, bold=bold, italic=italic)
    return p

def head(text, size=11):
    """Bold subheading without a terminal period, sentence case."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spacing(p, before=240, after=60)
    r = p.add_run(text)
    fmt(r, size=size, bold=True)
    return p

def headm(parts, size=11):
    """Bold subheading with mixed italic support. parts = [(text, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spacing(p, before=240, after=60)
    for text, italic in parts:
        r = p.add_run(text)
        fmt(r, size=size, bold=True, italic=italic)
    return p

def _citation_text(numbers):
    if not numbers:
        return ""
    mapped_numbers = []
    for number in numbers:
        if number not in REF_TEXTS:
            raise KeyError(f"Reference {number} not found in REFS")
        if number not in CITATION_MAP:
            CITATION_MAP[number] = len(CITATION_ORDER) + 1
            CITATION_ORDER.append(number)
        mapped_numbers.append(CITATION_MAP[number])
    sn = sorted(set(mapped_numbers))
    groups = []
    i = 0
    while i < len(sn):
        j = i
        while j + 1 < len(sn) and sn[j + 1] == sn[j] + 1:
            j += 1
        if j - i >= 2:
            groups.append(f"{sn[i]}\u2013{sn[j]}")
        elif j - i == 1:
            groups.append(f"{sn[i]}, {sn[j]}")
        else:
            groups.append(str(sn[i]))
        i = j + 1
    return " (" + ", ".join(groups) + ")"

def _add_citation_run(p, numbers):
    citation_text = _citation_text(numbers)
    if not citation_text:
        return
    trailing = ""
    if p.runs:
        text = p.runs[-1].text
        stripped = text.rstrip()
        if stripped.endswith((".", ";", ":", "?", "!")):
            trailing = stripped[-1]
            p.runs[-1].text = stripped[:-1] + text[len(stripped):]
    r = p.add_run(citation_text)
    fmt(r, size=11, italic=True)
    rPr = r._r.get_or_add_rPr()
    _set_font(rPr, FONT)
    if trailing:
        r = p.add_run(trailing)
        fmt(r, size=11)
        rPr = r._r.get_or_add_rPr()
        _set_font(rPr, FONT)

def spara(sentences, before=0, after=120, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    for text, numbers in sentences:
        if not p.runs:
            text = text.lstrip()
        add_text_with_taxon_italics(p, text, size=11)
        _add_citation_run(p, numbers)
    return p

def smixed(sentences, before=0, after=120, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    for parts, numbers in sentences:
        for text, bold, italic in parts:
            if not p.runs:
                text = text.lstrip()
            add_text_with_taxon_italics(p, text, bold=bold, italic=italic)
        _add_citation_run(p, numbers)
    return p

def cite(p, numbers):
    _add_citation_run(p, numbers)

def add_centered_figure(image_name, width_cm, before=120, after=60):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=before, after=after, line=12)
    r = p.add_run()
    r.add_picture(str(FIG_BASE / image_name), width=Cm(width_cm))
    return p

def add_main_figure_legend(label, title, caption_parts, before=0, after=160):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    spacing(p, before=before, after=after)
    r = p.add_run(label + " ")
    fmt(r, bold=True)
    r = p.add_run(title + " ")
    fmt(r, bold=True)
    for text, bold, italic in caption_parts:
        add_text_with_taxon_italics(p, text, bold=bold, italic=italic)
    return p

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
# Science Advances front matter.
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲

# Title (鈮?35 characters)
para(
    "OVAL glycan states link eggshell matrix chemistry to shell architecture and avian shell-breaking mechanics",
    bold=True, size=14, before=0, after=160, align=WD_ALIGN_PARAGRAPH.LEFT
)

# Short title (鈮?0 characters)
para("OVAL glycans link shell architecture and mechanics",
     bold=False, size=11, after=60, align=WD_ALIGN_PARAGRAPH.LEFT)

runpara([
    ("Lin Xuan", False), ("1", True), (", Yaqi Li", False), ("1", True),
    (", Jiajie Yang", False), ("1", True), (", Yu Liu", False), ("1", True),
    (", Chengyu Zhang", False), ("1", True), (", Qiulian Wang", False), ("1", True),
    (", Lingsen Zeng", False), ("2,3", True), (", Tongyao Li", False), ("1", True),
    (", Wenbin Zhou", False), ("1", True), (", Guiyun Xu", False), ("1", True),
    (", and Jiangxia Zheng", False), ("1*", True),
], size=11, before=80, after=40, align=WD_ALIGN_PARAGRAPH.LEFT)
runpara([
    ("1", True),
    (" National Engineering Laboratory for Animal Breeding and MOA Key Laboratory of Animal Genetics and Breeding, College of Animal Science and Technology, China Agricultural University, No. 2 Yuanmingyuan West Road, Haidian District, Beijing 100193, China.", False),
], size=11, before=0, after=40, align=WD_ALIGN_PARAGRAPH.LEFT)
runpara([
    ("2", True),
    (" Animal Breeding and Genomics, Wageningen University & Research, 6708 PB, Wageningen, The Netherlands.", False),
], size=11, before=0, after=40, align=WD_ALIGN_PARAGRAPH.LEFT)
runpara([
    ("3", True),
    (" State Key Laboratory of Genome and Multi-omics Technologies, Shenzhen Branch, Guangdong Laboratory of Lingnan Modern Agriculture, Key Laboratory of Livestock and Poultry Multi-omics of MARA, Agricultural Genomics Institute at Shenzhen, Chinese Academy of Agricultural Sciences, Shenzhen, 518124, China.", False),
], size=11, before=0, after=40, align=WD_ALIGN_PARAGRAPH.LEFT)
runpara([
    ("*", True),
    ("Corresponding author. Email: jxzheng@cau.edu.cn", False),
], size=11, before=0, after=120, align=WD_ALIGN_PARAGRAPH.LEFT)

para("Abstract", bold=True, size=11, before=80, after=40,
     align=WD_ALIGN_PARAGRAPH.LEFT)

para(
    "Eggshell matrix proteins are key regulators of eggshell structural formation, and existing studies have mapped posttranslational sites, leaving glycan side-chain properties less resolved. Using matched multi-layer analyses, we asked how glycan states on conserved matrix proteins map onto cross-species shell divergence. "
    "We compared chicken, duck, and pigeon under a conserved egg-tooth interface by integrating morphometric, glycoproteomic, structural, electrostatic, and finite-element analyses. "
    "Cross-species separation emerged first in mammillary-layer organisation within a largely shared matrix-protein toolkit. "
    "OVAL shifted from High Mannose-dominant glycans in chicken through Neutral Complex/Hybrid glycans in duck to Sialylated Complex/Hybrid glycans in pigeon. "
    "These glycan states predicted a Ca²⁺surface-accessibility gradient on OVAL that provides a plausible route from glycan-modulated OVAL unfolding and matrix-bound nucleation-site exposure to mammillary-layer formation, mature mammillary density, and, after separating shell-thickness effects, local shell-breaking structural strength. "
    "Together, the data refine current understanding of avian eggshell formation by connecting Ca²⁺accessible matrix-protein surfaces to mammillary-layer organisation and to local shell-breaking mechanics.",
    bold=False, size=10, before=0, after=80, align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

para("Teaser", bold=True, size=11, before=80, after=40,
     align=WD_ALIGN_PARAGRAPH.LEFT)
para(
    "OVAL glycan states reveal how eggshell matrix chemistry can shape local physical eggshell architecture used during shell breaking.",
    bold=False, italic=False, size=10, before=0, after=160,
    align=WD_ALIGN_PARAGRAPH.LEFT
)

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
# Introduction
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
para("Introduction", bold=True, size=14, before=0, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

# 搂1 鈥?Background value
p_s1a = spara([
    (" An avian eggshell forms under a demanding paradox: it must mineralize quickly, protect the embryo, and still yield to a localized force from within.", [1, 16]),
    (" This balance is achieved during one of the fastest known vertebrate CaCO₃ biomineralization processes, which assembles a mechanically competent calcitic shell within a narrow uterine time window.", [4, 13]),
    (" The process is not a simple precipitation of calcium carbonate. It depends on an organic matrix-protein system that coordinates calcium delivery, nucleation, crystal growth, and shell architecture.", [1, 2, 4, 29]),
    (" The functional test of that architecture arrives at shell breaking, when force is applied locally rather than across the whole shell.", [16]),
    (" Although avian beaks differ widely in shape, the egg-tooth remains the conserved tool that concentrates this inside-out load onto a small region of the inner shell.", [16, 86]),
    (" Comparable egg-tooth or egg-tooth-like structures also recur beyond birds, underscoring the repeated use of localized shell-opening devices across egg-laying amniotes.", [82, 83, 84, 85]),
    (" The central question is therefore how matrix chemistry formed during rapid mineralization becomes translated into the local physical eggshell architecture used during inside-out shell breaking.", [1, 28, 57]),
])

# 搂2 鈥?Prior work and its limits
p_intro2 = spara([
    (" The mammillary layer is where this question becomes experimentally tractable.", [1, 28]),
    (" It is the first calcified layer established on the shell membrane, where matrix-guided nucleation sets the spacing and continuity inherited by later shell units.", [1, 4, 28, 57]),
    (" It is also the inner structural layer that faces the egg-tooth during shell breaking, placing early mineralization and late mechanical function on the same material axis.", [16, 86]),
    (" Species with similar egg-tooth-based shell-breaking geometry can nevertheless produce distinct shell states across ecology and development.", [15, 26, 39]),
    (" Macroecological work on eggshell structure and texture further shows that shell phenotypes vary across avian life-history settings.", [40, 41]),
    (" The mammillary layer therefore offers a natural readout of how a rapidly mineralized shell stores species-specific mechanical information.", [1, 4, 28]),
])

p_intro_sig = spara([
    (" Eggshell matrix proteins are positioned to regulate this conversion from chemistry to physical eggshell architecture.", [1, 2, 4]),
    (" They guide mammillary-layer mineralization, crystal growth, and mature shell architecture.", [1, 2, 4, 29]),
    (" Proteomic studies have also identified a broadly shared shell-building toolkit.", [10, 19, 21]),
    (" What remains much less clear is how the chemical states of shared matrix proteins instruct the first physical eggshell architecture.", [18, 21]),
    (" This gap is sharpest for posttranslational modification.", [80]),
    (" Phosphorylation and glycosylation sites are increasingly catalogued in eggshell matrices.", [17, 18, 21]),
    (" Broader mass-spectrometry advances make such inventories tractable, but most studies still describe landscapes rather than mechanisms.", [49, 50]),
    (" Layer-resolved work has shown that the same protein can occupy different modification states across shell compartments, yet the route from those states to mammillary-layer formation remains largely unresolved.", [18]),
    (" Avian eggshell comparisons therefore rarely connect matched glycan states on shared matrix proteins to physical eggshell architecture across species.", [18, 29, 66]),
])

p_intro_gap = spara([
    (" Glycans provide a plausible chemical route from shared protein identity to distinct material behaviour.", [44, 78, 81]),
    (" Studies of glycan-shielded protein surfaces show that glycans can reshape exposure and dynamics, rather than simply decorate proteins.", [42, 43]),
    (" Related glycan-shield studies further show that these effects can alter protein motion, access, and recognition.", [61, 63, 72]),
    (" Unlike phosphate groups, glycans vary widely in composition, size, branching, and charge, and can reshape accessible protein surfaces.", [49, 50]),
    (" Glycan class is therefore a potential mechanistic variable, not only a feature of site occupancy.", [49]),
    (" OVAL provides a tractable test case because it is abundant, mineralization-relevant, and linked to Ca²⁺-responsive conformational behaviour during early shell formation.", [4, 11, 18, 29]),
    (" The missing step is a direct bridge from OVAL glycan class to Ca²⁺-accessible surface presentation on a shared matrix background.", [4, 18, 29]),
    (" If such a bridge exists, it could connect glycan-modulated OVAL unfolding, nucleation-site exposure, mammillary-layer formation, and local shell-breaking mechanics.", [4, 16, 28]),
])

# 搂4 鈥?This study
p_intro4 = smixed([
    ([('Here, we compared ', False, False),
      ('chicken', False, True),
      (', ', False, False),
      ('duck', False, True),
      (', and ', False, False),
      ('pigeon', False, True),
                        ('.', False, False)], []),
    ([(' We anchored the comparison to a conserved egg-tooth shell-breaking interface and followed a single axis across morphology, matrix proteomics, intact glycopeptides, glycoform-resolved structural modelling, electrostatics, and finite-element mechanics.', False, False)], []),
    ([(' This design asks whether dominant OVAL glycan states can link Ca²⁺ surface accessibility to mammillary-layer organisation and local shell-breaking response.', False, False)], []),
    ([(' In doing so, the study tests whether a rapidly mineralized shell can encode later shell-breaking mechanics through chemically specific states on a shared matrix protein.', False, False)], []),
    ([(' The resulting chain links OVAL glycan state to local physical eggshell architecture and to the local mechanical response used during shell breaking.', False, False)], []),
])

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺?
# "Results" section label
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺?
para("Results", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
# 搂 Species selection 鈥?ecological and developmental niche analysis
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
head("The conserved shell-breaking interface constrains shell variation")

p_ss1 = smixed([
    ([(' We placed extant birds into a comparative space using 10,993 AVONET species records and selected three deliberately separated model species from that space (Fig. 1A).', False, False)], [16, 22, 41]),
    ([(' This broader mapping prioritized two egg-relevant axes least likely to be secondary: nesting environment and offspring developmental state. Those axes span terrestrial-to-aquatic habitat use and a continuum from more precocial to more altricial young.', False, False)], [15, 22, 23]),
        ([(' Within this comparison space, ', False, False),
            ('chicken', False, True),
            (', ', False, False),
            ('duck', False, True),
            (', and ', False, False),
            ('pigeon', False, True),
                (' were therefore chosen near contrasting regions of those ecological-developmental gradients, reducing blurring by intermediate combinations.', False, False)], [3, 22, 23, 41]),
    ([(' Selection-space perturbation controls are provided in Fig. S1. This functional grouping only partly overlaps with phylogeny. Chicken and duck remain closely related precocial taxa but separate along the habitat axis, whereas pigeon anchors the altricial end of the comparison (Fig. S2).', False, False)], [3, 22, 23]),
    ([(' The design preserves shared ancestry while separating life histories within the same frame.', False, False)], []),
    ([(' The focal species differed in beak-tip geometry, but the egg-tooth remained a similarly localized dorsal breaker in all three species and therefore pointed to the same inside-out shell-breaking event (Fig. 1B).', False, False)], [16, 37, 82, 86]),
    ([(' With that interface held constant, the next question is which eggshell layer first separates the species.', False, False)], []),
])

mixed([
    ("Viewed through that shared shell-breaking context, the first eggshell level to show a clear contrast in the present comparison was mammillary-layer morphology (Fig. 1C). In ", False, False),
    ("chicken", False, True),
    (", mammillae were smoother overall and formed rounded projections. In ", False, False),
    ("duck", False, True),
    (", mammillae showed more ridges and angular turns across the inner surface. ", False, False),
    ("pigeon", False, True),
    (" was dominated by discrete triangular-conical mammillae. Three-dimensional surface reconstructions agreed with the cross-sectional views, indicating that the sampled inner-shell regions differed in mammillary geometry rather than representing minor variants of a shared inner-surface template.", False, False),
])

p_s0b = mixed([
    ("Quantification resolved the sampled regions in two related but not identical ways (Fig. 1D). Mammillary knob density was highest in ", False, False),
    ("chicken", False, True),
    (" (171.36 ± 5.63 per mm2), exceeding both ", False, False),
    ("duck", False, True),
    (" (155.22 ± 8.63 per mm2) and ", False, False),
    ("pigeon", False, True),
    (" (158.27 ± 11.39 per mm2), while duck and pigeon remained similar to each other. By contrast, crystal-unit proportion was highest in ", False, False),
    ("pigeon", False, True),
    (" (0.53 ± 0.04), intermediate in ", False, False),
    ("duck", False, True),
    (" (0.44 ± 0.02), and lowest in ", False, False),
    ("chicken", False, True),
    (" (0.40 ± 0.01). Within this scanned-fragment comparison, chicken showed the highest local mammillary density, whereas pigeon devoted the largest share of shell volume to crystal units grown from individual mammillary knobs. Duck remained intermediate in crystal-unit proportion while resembling pigeon in density. The two metrics did not collapse into one monotonic axis, but together they indicated that a mammillary-level contrast was already detectable before later shell traits were considered. Because this layer is the earliest structural level linked to eggshell mechanics and matrix control, we then asked a narrower question: did the observed contrast reflect wholesale toolkit replacement or differential use of a largely shared system?", False, False),
])
cite(p_s0b, [1, 4, 28])

doc.add_page_break()
add_centered_figure("Fig1_composed.png", width_cm=10.1, before=0, after=20)
add_main_figure_legend(
    "Fig. 1.",
    "Comparative species space, shell-breaking interface, and mammillary morphology in three model birds.",
    [
        ("(A) Three-dimensional AVONET comparison space built from 10,993 species records. Axes summarize aquatic association, lifestyle-habitat discordance, and developmental mode. Colors denote avian orders, and gray boxes indicate the sampled regions for ", False, False),
        ("Gallus", False, False),
        (", ", False, False),
        ("Anas", False, False),
        (", and ", False, False),
        ("Columba", False, False),
        (". (B) Lateral head views (top) and dorsal beak views (bottom) showing the egg-tooth-bearing tip in the three species. (C) Representative micro-CT sections and three-dimensional inner-surface reconstructions of the mammillary layer. Scale bars, 100 μm. (D) Box plots of mammillary density and unit-volume ratio across species. Points denote nine non-overlapping subfragments from one scanned fragment per species. P values were calculated by one-way ANOVA, and different letters indicate Tukey HSD groupings.", False, False),
    ],
    before=20,
    after=80,
)
doc.add_page_break()

p_sprot_bg = spara([
    ("Orthogroup analysis resolved the three eggshell matrix proteomes into a large shared core with smaller pairwise-shared and lineage-restricted complements (Fig. S3). At the overall level, the proteome still followed broad ancestry (Fig. S4), arguing against wholesale replacement of the eggshell-matrix toolkit.", []),
])

p_sprot_go = spara([
        (" The shared core therefore became the relevant comparison frame. The next question was whether the species separated through wholesale proteome turnover or through different glycan states on shared matrix proteins.", []),
])

p_sprot_focus = spara([
    (" The chicken-exclusive set was simultaneously enriched for protein N-linked glycosylation (BP; Fig. S5), shifting the comparison from protein presence to chemical deployment. The retained shared core thus became the relevant molecular background. Glycosylation on shared proteins emerged as a proximate candidate layer for explaining divergence in mammillary organisation and downstream shell behaviour. The linked supplementary outputs follow the downstream logic from recurrent matrix-protein glycosylation profiles (Fig. S6) and OVAL structural context (Fig. S7) to CAFE5 gene-family turnover (Fig. S8) and turnover-associated enrichment (Fig. S9).", [18]),
    (" Most recurrent eggshell matrix proteins emphasized in earlier studies were recovered in the broader proteomic and glycoproteomic background here, indicating substantial agreement with prior eggshell-matrix work.", [1, 2, 4, 29]),
    (" The present dataset also broadened that comparative background for OC17, OC116, and related matrix components.", [10, 19, 21]),
])

head("OVAL glycosylation gives the clearest cross-species molecular contrast")

p_s2a = spara([
    ("Intact-glycopeptide profiling showed that the three species differed in sampling depth but still shared a stable comparison core (Fig. 2A to D; Tables S1 and S2). The cluster view recovered 25 clusters shared by all three species, with the largest additional pairwise overlap between duck and pigeon at 64 clusters, whereas chicken contributed little species-private cluster space (Fig. 2A).", []),
    (" The same pattern held for the catalog counts: duck yielded 321 glycoproteins, 547 glycosites, and 197 glycan compositions; pigeon yielded 192, 257, and 162; and chicken yielded 55, 88, and 105 (Fig. 2B). Protein-level MS outputs are provided in Table S1, and glycopeptide and glycosite MS outputs are provided in Table S2. Shared-core Jensen-Shannon similarity remained between 0.33 and 0.40, with the duck-pigeon pair highest (Fig. 2C). These values indicate divergence within a still comparable glycoproteomic background rather than three disconnected chemical spaces.", []),
    (" Glycan-class composition reinforced the same point at the chemical-deployment level. High-Mannose and Complex-Fucosylated glycans formed a broad cross-species background, whereas Complex-Sialylated and other more extended classes contributed more strongly to lineage separation (Fig. 2D).", []),
])

add_centered_figure("Fig2_composed.png", width_cm=14.6)
add_main_figure_legend(
    "Fig. 2.",
    "Shared-core glycoproteome architecture and glycan-class deployment across species.",
    [
        ("(A) Species-partitioned cluster counts in the glycoproteomic dataset. (B) Numbers of detected glycoproteins, glycosites, and glycan compositions in each species. (C) Shared-core Jensen-Shannon similarity among species. (D) Species-level distribution of glycan classes (High Mannose, Pauci-mannose, Hybrid, Complex-Plain, Complex-Fucosylated, Complex-Sialylated, and Other). (E) Ortholog-glycan chord map linking eggshell glycoproteins to dominant glycan classes in Gallus, Anas, and Columba, with highlighted matrix proteins retained for downstream analyses.", False, False),
    ],
)

p_s2b = mixed([
    ("A stricter BlastP-based filter retained an orthologous glycoprotein subset suitable for structural comparison and summarized that shared candidate space in Fig. 2E and Table S3. Using ", False, False),
    ("chicken", False, True),
    (" as the reference, non-reference candidates were retained only when the mean E-value was below 1 × 10⁻⁵ and sequence identity met the final comparability thresholds. The resulting target-ortholog table is provided in Table S3. This restricted the downstream comparison to high-confidence orthologs. Under that stricter mapping, OC17 was glycosylated only in chicken, whereas OC116, TRFE, and OVAL all retained glycosylation signals across the three species and served as shared anchors. Among them, OVAL showed a clear cross-species glycan shift and was prioritised for structural analysis.", False, False),
])
p_s2c = spara([
    ("Integrated protein and glycan abundance profiles further identified OVAL as the shared protein most closely aligned with the cross-species eggshell differences. In chicken, the protein-abundance and glycan-abundance space showed weak overall coupling, and OVAL occupied a comparatively modest glycan-burden position among the highlighted matrix proteins (Fig. 3A). In duck, the same analysis showed stronger positive protein-glycan coupling, with OVAL retaining high protein abundance while shifting toward higher glycan output (Fig. 3B). In pigeon, protein-glycan coupling was again positive, and OVAL carried the strongest glycan burden among the three species while remaining within the shared matrix-protein background (Fig. 3C).", []),
    (" Pairwise enrichment plots then showed why OVAL remained the most interpretable discriminator. In the chicken-versus-pigeon plane, OVAL shifted toward the glycan-rich side of the comparison, indicating that its glycan change exceeded the corresponding protein-abundance change (Fig. 3D). In the chicken-versus-duck plane, OVAL again deviated from the equal-change diagonal, separating chicken from duck through glycan enrichment rather than simple protein abundance (Fig. 3E). In the duck-versus-pigeon plane, OVAL remained displaced from protein-glycan equivalence, preserving a glycan-state ordering beyond the chicken-centered comparisons (Fig. 3F).", []),
    (" Among the highlighted eggshell-matrix proteins, OVAL therefore remained abundant in all three species but differed sharply in glycan burden: relatively modest in chicken, stronger in duck, and strongest in pigeon. OC116 and TRFE remained informative shared proteins, but neither separated bulk protein abundance from glycan output as consistently as OVAL.", []),
    (" Intact-glycopeptide assignments placed OVAL along a coherent cross-species progression. Chicken carried compact High-Mannose glycans, duck was enriched for Neutral Complex/Hybrid glycans, and pigeon carried more extended Sialylated Complex/Hybrid glycans. Protein-glycan joint-analysis tables are provided in Table S4, and the recurrent matrix-protein glycosylation profiles are summarized in Fig. S6. Taken together, the six Fig. 3 panels place OVAL among the shared proteins whose glycosylation aligns most closely with the phenotype ordering recovered here.", []),
])

p_s2d = spara([
    ("Because those OVAL glycan classes differ strongly in size and charge distribution, the more informative comparative variable was OVAL surface accessibility rather than OVAL abundance alone. OVAL ensemble geometry and APBS-derived surface-potential context are summarized in Fig. S7 and Tables S5 and S6.", []),
    (" Broader gene-family expansion, contraction, and turnover-enrichment analyses are provided in Figs. S8 and S9, separating lineage-scale background from the OVAL-focused structural comparison.", []),
    (" The relevant feature was how much of the acidic OVAL interface remained chemically reachable once decorated by different glycans. Ortholog control, abundance decoupling, and glycan-class progression together left OVAL as the shared candidate that was most directly comparable, chemically specific, and structurally actionable.", []),
])

add_centered_figure("Fig3_composed.png", width_cm=15.5)
add_main_figure_legend(
    "Fig. 3.",
    "Ortholog-restricted abundance-glycan analysis identifies OVAL as the leading shared discriminator.",
    [
        ("(A to C) Proteotype coevolution plots of log2-transformed protein abundance versus glycan abundance within Gallus, Anas, and Columba. Insets report Spearman's rho and two-sided p values. Highlighted labels indicate retained matrix proteins (OVAL, OC116, TRFE, and OC17). (D to F) Pairwise two-dimensional glycan-protein enrichment plots for Gallus versus Columba, Gallus versus Anas, and Anas versus Columba. The dashed diagonal indicates equal protein and glycan change.", False, False),
    ],
)

head("OVAL glycan state reshapes surface accessibility")

p_s3a = spara([
    ("OVAL was selected for structural analysis because it remained shared, chemically distinct, and directly comparable across species. Dominant glycosylated OVAL ensembles and matched apo references were rebuilt to test whether the three species differed mainly through glycan-dependent surface behavior rather than through backbone sequence alone.", [4, 11]),
    (" Representative rebuilt glycan conformations and species-specific surface maps showed that the dominant glycans occupied different spatial envelopes on the same folded protein scaffold (Fig. 4A and B). In panel A, points 1 and 2 mark the glycan end-to-end vector, point 3 marks the glycan centroid, point 4 marks the protein Cα centroid, and the translucent sphere shows the radius-of-gyration envelope centered on point 3. Panel B then translates that geometry into surface exposure by coloring Ca²⁺relevant regions and blackening the same regions after shielding.", []),
])

p_s3b = spara([
    (" Pigeon first separated by occupying the largest overall glycan envelope, as shown by the higher radius of gyration in Fig. 4C and the longer end-to-end distances in Fig. 4E. In panel A terms, that means a longer 1-2 span and a larger sphere around point 3, with point 4 serving as the protein anchor that the glycan is measured against. That expansion occurred together with closer local approaches to the backbone in Fig. 4D and a broader glycan-protein distance distribution in Fig. 4F, consistent with extended glycans that sample a larger envelope while still folding back toward the OVAL surface.", []),
    (" Chicken defined the opposite endpoint, with compact glycans and the weakest geometric intrusion into the acidic interface. Duck remained closer to chicken in radius of gyration and overall glycan-protein spacing, but it separated from both species in end-to-end span and minimum backbone approach. Fig. 4C to F therefore translate glycan-class progression into a shielding geometry: a smaller sphere and shorter 1-2 span in panel A indicate a more compact glycan state, whereas a larger sphere and longer 1-2 span indicate a more extended, surface-engaging state.", []),
])

p_s3c = spara([
    (" Fig. 4G to J resolved the same acidic interface at progressively stricter levels. Viewed against panel B, the colored patches are the Ca²⁺relevant regions that remain accessible, and black marks the same regions after shielding. Fig. 4G therefore measures gross interface shielding, Fig. 4H asks what fraction of candidate acidic residues remained hotspots, Fig. 4I measures the surface area retained by hotspot residues, and Fig. 4J counts the subset of Ca²⁺hotspots that remained both electrostatically favorable and physically reachable.", []),
    (" Interface shielding increased stepwise from chicken to duck to pigeon, and the same ordering was retained across hotspot surface area, hotspot fraction, and net accessible Ca²⁺hotspots. Together, these panels show progressive masking of the shared acidic OVAL face during early mineralization.", []),
])

p_s3d = spara([
    (" Matched glycosylated-versus-apo comparisons then showed that glycan addition changed the number of Ca²⁺-relevant hotspot residues and the exposed carboxylate surface most clearly in pigeon (Fig. 4K and L; Fig. S10). With panel B in mind, glycosylation preserves or hides the same Ca²⁺-relevant patches rather than creating new ones: the colored patches remain reachable, whereas the black patches are the same sites after shielding. Duck shifted in the same direction without a resolved structure-level significance call, and chicken could be assessed only descriptively because one glycosylated structure was available. This pattern is more consistent with a glycan-imposed shift in the acidic surface presented at mineralization onset than with a generic sequence effect alone.", []),
    (" Fig. 4K to N then collapsed the same comparison to the whole-interface level. Across those panels, chicken preserved the most accessible Ca²⁺-relevant surface, pigeon shifted the largest share into a glycan-affected state, and duck trended toward the lower-accessibility side but did not separate from pigeon or apo references uniformly across metrics.", []),
    (" Chicken therefore retained the highest inferred Ca²⁺-capturing capacity and the state most compatible with earlier Ca²⁺-responsive opening of OVAL at mineralization onset. Duck and pigeon moved toward the lower-accessibility side from different structural backgrounds. The same ordering matched the phenotype sequence: chicken combined the densest mammillary field with the strongest local response to inside-out loading, whereas duck and pigeon converged toward the lower-response side. Taken together, Fig. 4A to N link glycan-dependent separation, glycan geometry, interface masking, and Ca²⁺-relevant accessibility on a shared matrix protein.", [4, 11, 29]),
    (" This ordering places OVAL glycan-dependent Ca²⁺accessibility upstream of glycan-modulated OVAL unfolding and nucleation-site exposure, providing the structural premise for testing whether the resulting mammillary organisation affects local shell-breaking mechanics.", []),
])

doc.add_page_break()
add_centered_figure("Fig4_composed.png", width_cm=15.2, before=0, after=20)
add_main_figure_legend(
    "Fig. 4.",
    "OVAL glycan classes reshape surface geometry and Ca²⁺relevant interface accessibility.",
    [
        ("(A) Representative rebuilt OVAL-glycan conformations on the protein surface. The numbered markers indicate the start and end of the glycan end-to-end vector (1 and 2), the glycan centroid (3), and the protein Cα centroid (4). The translucent sphere denotes the glycan radius-of-gyration envelope centered on the glycan centroid. (B) Species-specific surface maps showing glycan positions in color and Ca²⁺relevant regions in color, with shielded Ca²⁺relevant regions shown in black. (C to F) Glycan radius of gyration, minimum glycan-backbone distance, glycan end-to-end distance, and glycan-protein distance across species. (G to J) Interface shielding, hotspot fraction, hotspot residue SASA, and net accessible Ca²⁺hotspots. (K to N) Glycosylated-versus-apo comparisons of Ca²⁺hotspot residue counts, carboxylate surface accessibility, Ca²⁺hotspot accessibility, and Ca²⁺hotspot residue SASA. Species contrasts for ensemble-derived metrics used two-sided Mann-Whitney U tests. Glycosylated-versus-apo structure-level contrasts used one-sample Wilcoxon signed-rank tests when structure-level variation was present.", False, False),
    ],
    before=20,
    after=80,
)
doc.add_page_break()

head("Inside-out loading resolves local shell-breaking mechanics")

p_s4a = mixed([
    ("We next asked whether the OVAL accessibility-to-mammillary sequence recovered above was reproduced in local shell-breaking mechanics. Finite-element testing translated the shared egg-tooth interface into an explicit inside-out loading design. Fig. 5A shows the species-specific dorsal beak views and the matched micro-CT-derived finite-element shell fragments used for loading. Because the meshes preserved species-specific shell geometry, the analysis remained anchored to the same mammillary context identified morphologically. Impact loading was sampled across multiple offset positions on the eggshell fragments, yielding independent contact-force and contact-shear-stress time courses for each species (Fig. 5B and C). We recorded both raw peak contact force (F_max) and peak contact shear stress (τ_max) so that thickness-driven force effects could be separated from the local stress response at the mammillary contact interface. Peak τ_max was used as the direct readout of local shell-breaking response. Species means ± s.d. were calculated across the sampled positions (Fig. S11; Table S7; eggshell thicknesses: ", False, False),
    ("chicken", False, True),
    (" 0.29 mm, ", False, False),
    ("duck", False, True),
    (" 0.35 mm, and ", False, False),
    ("pigeon", False, True),
    (" 0.19 mm).", False, False),
])

doc.add_page_break()
add_centered_figure("Fig5_composed.png", width_cm=15.5, before=0, after=20)
add_main_figure_legend(
    "Fig. 5.",
    "Local finite-element loading connects egg-tooth contact geometry to shell-breaking mechanics.",
    [
        ("(A) Dorsal beak views and matched micro-CT shell-fragment finite-element models for ", False, False),
        ("chicken", False, True),
        (", ", False, False),
        ("duck", False, True),
        (", and ", False, False),
        ("pigeon", False, True),
        (" from top to bottom. Each model shows the conical egg-tooth impactor and a representative local stress field on the reconstructed shell geometry. (B) Mean contact-force time courses across nine impact positions, with shaded ±1σ envelopes, and the corresponding peak contact force (F_max) distribution. Star markers indicate the mean peak value on each time-course trace. (C) Mean contact shear-stress time courses across the same nine positions, with shaded ±1σ envelopes, and the corresponding peak shear stress (τ_max) distribution. Star markers indicate the mean peak value on each trace. Box-plot points denote individual impact positions (n = 9 per species), bars show mean ± s.d., p values above the plots are one-way ANOVA omnibus p values, and different letters indicate Tukey HSD groupings at p < 0.05. Underlying finite-element values are reported in Table S7. Simulations used identical loading and material settings so that the comparisons isolate geometry-dependent local response.", False, False),
    ],
    before=20,
    after=80,
)
doc.add_page_break()

mixed([
    ("Peak F_max differed significantly among species (p = 1.64 × 10⁻¹³). ", False, False),
    ("Chicken", False, True),
    (" reached 1.12 ± 0.11 N, ", False, False),
    ("duck", False, True),
    (" reached 0.90 ± 0.09 N, and ", False, False),
    ("pigeon", False, True),
    (" reached 0.49 ± 0.04 N, and all pairwise differences were significant by Tukey HSD (Fig. 5B). By contrast, τ_max resolved a two-level pattern (p = 6.64 × 10⁻¹⁰). ", False, False),
    ("Chicken", False, True),
    (" reached 0.613 ± 0.061 MPa and was significantly higher than ", False, False),
    ("duck", False, True),
    (" at 0.413 ± 0.041 MPa and ", False, False),
    ("pigeon", False, True),
    (" at 0.406 ± 0.033 MPa. The latter two species did not differ significantly from each other (Tukey HSD, p = 0.957; Fig. 5C).", False, False),
])

mixed([
    ("The difference between F_max and τ_max clarified the duck result. Its higher raw contact force was driven mainly by greater shell thickness (0.35 mm versus 0.19 mm in pigeon), not by a stronger local stress response at the mammillary interface. In other words, the thicker duck shell could carry more total contact force, but it did not show greater local resistance after the response was normalized to the contact-stress scale. By contrast, ", False, False),
    ("Chicken", False, True),
    (" exhibited a 36-40% increase in τ_max relative to the two other species, indicating a stronger local shell-breaking stress response independent of shell thickness. This high-versus-low grouping, with ", False, False),
    ("chicken", False, True),
    (" alone in the high group and ", False, False),
    ("duck", False, True),
    (" together with ", False, False),
    ("pigeon", False, True),
    (" in the low group, matched the grouping recovered for mammillary density by Tukey HSD (Fig. 1D). The mechanics therefore retained the contrast already recovered from mammillary organisation and OVAL accessibility.", False, False),
])

mixed([
    ("Whole-shell fracture force alone could make duck appear mechanically superior to chicken because of its greater shell thickness, despite the absence of the same high-density mammillary state. By focusing instead on the local stress response at the micro-CT-derived mammillary interface, τ_max removes that ambiguity. It shows that the high-density chicken state remains distinct, whereas duck and pigeon converge at a lower local-response level. This functional readout preserved the same asymmetry already visible in earlier sections and linked glycosylation-associated differences to local shell-breaking mechanics across the three model species.", False, False),
])

cite(p_s4a, [16, 37, 69])

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
# Discussion
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
para("Discussion", bold=True, size=14, before=320, after=160,
    align=WD_ALIGN_PARAGRAPH.LEFT)

p_disc_mam1 = smixed([
        ([('Cross-species divergence in this dataset was resolved most clearly at the mammillary layer, rather than through wholesale matrix-protein turnover. ', False, False),
            ('Within this shared eggshell-matrix toolkit, OVAL glycan state provided the most direct molecular axis linking Ca²⁺-relevant surface accessibility, mammillary-layer organisation, and shell-breaking mechanics. The Discussion therefore treats each result as one step in a matrix-chemistry-to-physical-eggshell-architecture-to-shell-breaking chain.', False, False)], [1, 16, 18]),
])

p_disc_regulator = spara([
    ("The three-species design asks whether the same matrix-to-mechanics axis persists across ecological and developmental gradients.", [15, 39]),
    (" Eggshell traits vary along continuous axes rather than a binary label. Nesting environment spans terrestrial to aquatic conditions, and offspring condition ranges from more precocial to more altricial young. Neither dimension is captured by a simple yes-or-no partition.", [3, 23]),
    (" Duck is especially informative in this design. It retains a broadly precocial developmental condition and the greatest shell thickness in this comparison. Yet it shifts toward an intermediate OVAL glycan state and accessibility profile, while its τ_max outcome converges with pigeon rather than chicken. Duck therefore separates shell-thickness buffering from the Ca²⁺-accessibility-mammillary-mechanics axis emphasized here. The design thus tests the hypothesis against a thickness-rich counterexample, not only across species labels.", [3, 15, 22, 23]),
])

p_disc_axis = spara([
    ("Mammillary-layer mineralization remains the central structural level in this interpretation because it is where matrix chemistry can first shape physical eggshell architecture.", [1, 28]),
    (" In this model, glycan-dependent exposure of acidic OVAL surfaces could influence Ca²⁺ access and OVAL unfolding during the earliest matrix-mineral interaction. These changes could alter nucleation-site presentation, mammillary-layer formation, mature mammillary density, and local shell-breaking mechanics.", []),
    (" Once early calcite units are established, subsequent eggshell regions inherit the spacing logic set during this first mineralization window. A dense mammillary field can therefore influence matrix retention, mineral continuity, local stress redistribution, and mature morphology.", [1, 30, 36]),
    (" This emphasis is consistent with earlier eggshell studies that place the mammillary layer at the intersection of crystal nucleation and matrix control. The present comparison extends that view by linking this layer to a cross-species glycan-state readout rather than to shell-quality descriptors alone.", [1, 28, 31, 32]),
    (" Recent poultry omics studies increasingly connect age, shell-gland transcription, extracellular-vesicle cargo, and whole-shell quality traits to eggshell phenotype. Those descriptors, however, remain broader than the proximate material layer isolated here.", [33, 52, 57, 70]),
    (" Mammillary organisation is therefore not simply another shell trait. It is the earliest physical eggshell context in which matrix chemistry can plausibly bias later mechanical outcome.", [1, 2]),
    (" This structural position makes mammillary organisation the first physical-architecture readout of the proposed glycan-dependent matrix mechanism.", [1, 28]),
])

p_disc_mam2 = spara([
    ("Among the molecular layers examined here, OVAL N-glycan architecture most closely tracked the cross-species structural contrast.", []),
    (" Orthogroup turnover, gene-family change, and glycoprotein-network divergence remain relevant, but they mainly define the comparative background. OVAL glycan state is more directly interpretable because it is shared across species, chemically resolved, and located on an abundant matrix protein already implicated in mineralization.", [4, 18, 27]),
    (" Earlier studies had already placed OVAL among abundant eggshell glycoproteins and mineralization candidates and had shown that eggshell-matrix proteins can occupy distinct N-glycosylation states. The advance here is an ortholog-resolved comparison of which glycan states align with phenotype and how those assignments carry into physical eggshell architecture and mechanics.", [4, 18]),
    (" Earlier chicken studies also established a glycosite foundation for OVAL and identified glycosylated Asn in OC116. The present dataset resolved dominant glycan classes on corresponding OVAL ortholog sequons used for structural modelling (chicken N293; duck and pigeon N97), extending site detection into cross-species glycan-class interpretation.", [8, 18, 21]),
    (" OVAL is therefore useful not because it is unique, but because it provides the clearest molecular readout through which the proposed chain can be tested across species.", []),
])

p_disc_other = spara([
    ("The non-OVAL signals remain important because they define the shared matrix background against which the OVAL axis becomes visible.", []),
    (" OC116 and TRFE were informative shared proteins, whereas OC17 appeared glycosylated only in chicken and may reflect a more lineage-restricted mineralization program.", []),
    (" This pattern is consistent with earlier work assigning functional importance to OC17, OC116, and ovotransferrin-related matrix components.", [10, 19, 29]),
    (" It also preserves the value of glycoproteomic studies that established experimentally accessible site-level inventories.", [18, 21]),
    (" The present comparison adds a cross-species layer to that foundation: shared orthologs, dominant glycan classes, and their consequences for surface presentation.", []),
    (" Thus, the toolkit remains multicomponent, while OVAL remains the clearest entry point into the chemistry-physical-eggshell-architecture-mechanics chain tested here.", []),
])

p_disc_oval = spara([
    ("Re-Glyco and APBS analyses provide the structural bridge for this argument.", []),
    (" Across species, they resolved a glycan-state gradient. Compact chicken glycans preserved the most accessible acidic surface, neutral complex/hybrid duck glycans imposed an intermediate constraint, and extended sialylated pigeon glycans generated the strongest steric and electrostatic shielding.", []),
    (" Earlier in vitro and structural work suggested that OVAL conformation and electrostatics matter during mineralization. Matched glycoform-resolved surface ensembles, however, had not been compared across bird species.", [4, 11]),
    (" Glycan-state variation is therefore resolved here as a physically interpretable surface difference. Although this result does not establish direct causality, it supports the hypothesized conversion of glycan class into a mineral-facing Ca²⁺-accessibility state.", []),
])

p_disc_mech = spara([
    ("The mechanical comparison was designed to test the endpoint of the proposed chain under inside-out shell-breaking load, rather than under conventional outside compression or whole-shell fracture.", []),
    (" This distinction matters because shell thickness inflates absolute failure load, whereas τ_max is less thickness-confounded and more directly reports stress transfer through the mammillary interface.", [16, 34, 69]),
    (" Sun et al. similarly showed that eggshell thickness varies across the whole egg. The circumferential zone around the blunt end was locally thinnest, reinforcing the need to separate global shell-thickness buffering from local shell-breaking-interface mechanics.", [90]),
    (" The finite-element analysis therefore asks whether the inner mammillary interface preserves the same contrast inferred from matrix chemistry and morphology.", [16, 34, 35, 69]),
    (" Duck is the critical control for this interpretation. Its thicker shell increased F_max but did not recreate the high-τ_max state observed in chicken, separating thickness buffering from the material pathway emphasized here.", []),
    (" Eggshell thickness, body size, reproductive ecology, and lineage history still define the background design space.", [3, 14]),
    (" However, these factors do not explain why the same ordering recurs across glycan class, Ca²⁺ surface accessibility, mammillary organisation, and τ_max under inside-out loading.", [16, 37]),
    (" Duck and pigeon bound this mechanical axis from opposite directions: duck has greater shell thickness and intermediate OVAL accessibility, whereas pigeon has a thinner shell and more extended OVAL glycans, yet both converge at low τ_max.", []),
    (" This contrast makes chicken a useful reference for linking glycan-dependent matrix behaviour to a local eggshell state with concentrated mechanical response at the shell-breaking interface.", [4, 18]),
    (" The finite-element result therefore tests the final link of the hypothesis: whether glycan-associated physical eggshell architecture tracks the local response used during shell breaking.", []),
])

p_disc_evo = para(
    "A second interpretive issue is that the mammillary layer can be partly resorbed during late incubation and shell breaking. This possibility does not remove the relevance of the present comparison because the quantified descriptors were mammillary density and crystal-unit organisation. These features remain embedded in the shell even when part of the innermost material has been absorbed. The same consideration guided the mechanical readout. We emphasized the second characteristic peak rather than the first because the earliest force excursion is more strongly dominated by initial morphology-dependent contact. The later peak more faithfully reflects stress transmission through the shell wall as a whole. This limitation therefore bounds the interpretation while preserving the structural link required by the proposed mechanism."
)

p_disc_function = para(
    "Taken together, the comparison converged on a local eggshell state associated with stronger shell-breaking mechanics in this dataset. Chicken combined the densest mammillary field, the least shielded Ca²⁺-relevant OVAL surface, and the strongest local stress response under inside-out loading. Duck showed why the chain cannot be reduced to shell thickness. Despite its thicker shell and higher F_max than pigeon, its τ_max grouped with pigeon rather than chicken. This pattern provides the functional synthesis of the hypothesis: chemically specific states on reused matrix proteins can align with mineralized phenotypes more directly than proteome turnover alone."
)
cite(p_disc_function, [67, 73, 74])

p_disc_biomineral = para(
    "The same analytical sequence may extend beyond avian eggshells, but its main value here is to show how a matrix chemical state can be followed across scales. Many biomineralization systems use organic matrices to regulate ion access, surface exposure, and mineral nucleation through chemically specific interfaces. The workflow used here therefore offers a compact template: glycoproteomic state, surface presentation, and mesoscale function. Similar logic may be useful in other mineralized tissues, biomimetic materials, and eggshell-derived regenerative materials. This broader implication, however, remains anchored to the eggshell mechanism resolved here."
)
cite(p_disc_biomineral, [64, 67, 68, 73])

p_disc_future = para(
    "The present scope remains bounded. We analyzed dominant glycoforms rather than the full in vivo glycan ensemble. We also treated each species as mechanically uniform at the scale of the mean eggshell and relied on incompletely constrained uterine ionic conditions in the APBS framework. Immediate wet-lab follow-up faces a practical limitation: current toolkits still make it difficult to combine species-matched expression of orthologous matrix proteins with precise installation of predefined glycan structures at specific glycosylation sites. In parallel, biomineralization operates as a large coupled reaction system, which limits straightforward single-variable interpretation from any one assay. The next decisive tests are therefore defined-glycoform mineralization assays, direct manipulation of OVAL glycosylation in chicken, and site-resolved validation of the same inside-out mechanical contrast, together with methods that improve causal resolution. These experiments would test the causal version of the hypothesis by distinguishing an active OVAL glycan mechanism from a high-mammillary-density marker."
)

p_disc_close = para(
    "In summary, this study links mammillary organisation, glycoprotein state, Ca²⁺ surface accessibility, physical eggshell architecture, and local shell-breaking mechanics across three avian eggshells. Chicken defined the high-mammillary-density end of that axis, with dense mammillary organisation, compact OVAL glycans, greater Ca²⁺-relevant surface exposure, and the strongest local response at the mammillary interface. Duck occupied the critical intermediate position. Its thicker shell increased absolute force but did not reproduce the same local stress state, separating shell thickness from the mammillary-interface mechanism. As comparable glycoform assignments become available, the same framework can extend to other abundant eggshell matrix proteins. Across morphometric, glycoproteomic, structural, and mechanical layers, OVAL glycan state remains the most consistently aligned molecular feature of the matrix-chemistry-to-physical-eggshell-architecture-to-shell-breaking chain recovered here."
)

p_disc_limits = p_disc_close

# Methods
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
para("Materials and Methods", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

head("Biological materials")

p_m_bio = mixed([
    ("Fertilized eggs were collected from three avian lines: seven eggs from Chahua Chicken, seven eggs from Shaoxing Duck, and 19 eggs from White King Pigeon. ", False, False),
    ("Gallus gallus", False, True),
    (" eggs were obtained from the Poultry Resources Conservation "
     "Farm, China Agricultural University (Beijing, China); ", False, False),
    ("Columba livia", False, True),
    (" eggs were provided by the College of Veterinary Medicine, China Agricultural University; and ", False, False),
    ("Anas platyrhynchos", False, True),
    (" eggs were obtained from a commercial layer-duck farm in Shandong Province, China. All eggs were stored at 16\u00b0C for 7 d under breeder-egg holding conditions before analysis. No live animals were used in this study.", False, False),
])

head("Eggshell matrix protein extraction")

mixed([
    ("Eggshell matrix proteins were extracted from the eggshell mammillary layer "
     "(EML) by an established EDTA demineralization protocol. Eggs were rinsed with "
    "deionized water and placed in sterile sealed bags. One egg from each species was reserved for micro-CT analysis; the remaining shells were used for matrix-protein extraction. The remaining six chicken eggs and six duck eggs were divided into three two-egg extraction units per species, whereas the remaining 18 pigeon eggs were divided into three six-egg extraction units. These pooled extraction units were used as the matched biological samples for both shotgun proteomics and intact glycopeptide analysis. For ", False, False),
    ("G.\u00a0gallus", False, True),
    (" and ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (", the eggshell cuticle layer (ECL) was removed prior to EML "
     "extraction by treatment with 15 mL of 5% EDTA (0.13 mol/L, pH 7.6) supplemented "
     "with 2-mercaptoethanol (10 mmol/L) for 30 min at 20\u00b0C, with gentle manual "
    "kneading to separate the ECL; eggshells were subsequently rinsed with deionized "
     "water. No discrete cuticle layer was recognized in pigeon eggshells, so ", False, False),
    ("C.\u00a0livia", False, True),
    (" shells were rinsed with distilled water and entered the same extraction workflow directly. EML proteins from all three species were then solubilized under the same "
     "EDTA\u20132-mercaptoethanol conditions with the extraction duration extended to 12 h "
     "at 20\u00b0C. The resulting suspension was centrifuged at 1,000 \u00d7 g for 15 min; the "
     "pellet was resuspended and centrifuged a second time, and the pooled supernatant "
        "was stored at \u221280\u00b0C until analysis. Species were processed in parallel under the same extraction chemistry so that downstream differences were less likely to reflect handling drift.", False, False),
])

head("Micro-CT imaging and mammillary morphometry")

p_m_ct = para(
    "One eggshell fragment (4 mm \u00d7 4 mm) from each species was excised from the equatorial region near the blunt-half midpoint and scanned using a Phoenix V|tome|x\u00a0M "
    "microfocus CT system (GE Sensing and Inspection Technologies GmbH, Wunstorf, "
    "Germany). The X-ray source was operated at 85 kV tube voltage and 160 \u03bcA current without a beam filter; scan settings were held "
    "constant across all specimens, and a total of 1,800 projection images were collected for each eggshell. Reconstructed image volumes were exported as 16-bit unsigned isotropic datasets with a sampling distance of approximately 0.003836 mm along the x, y, and z axes (about 3.84 \u03bcm voxel size). After scanning, each fragment was subdivided evenly into nine subfragments for downstream regional quantification. Volumetric data were reconstructed in 3D Slicer, and the eggshell region was segmented with the Threshold module in the Segmentation workflow using operator-guided threshold correction after automatic selection. The same operator reviewed all three species side by side; grayscale thresholds were fine-tuned to retain all true shell voxels while preserving visible pore openings. Acquisition noise was suppressed by a 5 \u00d7 5 \u00d7 5 "
    "median filter, followed by largest-island isolation and 9 \u00d7 9 \u00d7 9 hole-filling. "
    "Within each subfragment, three morphometric parameters were then calculated from the labelmap. "
    "The segmented shell model was first duplicated as a single-copy volume and solid-filled with the Fill Holes operation; "
    "subtraction of the original shell model from the filled solid yielded the mammillary interspace layer, and closed voids appearing in that interspace plane were defined as mammillary knobs. "
    "Mammillary density was calculated as mammilla count divided by analysis-unit area. "
    "Total eggshell volume in the same analysis unit was obtained directly from the labelmap, and mean column-unit volume was defined as total shell volume divided by mammilla count; "
    "column-unit volume fraction was then calculated as mean column-unit volume divided by the total eggshell volume of the corresponding analysis unit. "
    "Because the columnar units initiated by mammillae are arranged as repetitive and approximately even planar units in normal avian eggshell microstructure, these parameters were treated as local average representatives of whole-shell organisation. The nine observations per species were non-overlapping regional subsamples from one scanned fragment and were used to quantify within-fragment spatial variation around that local mean rather than to represent nine independent eggs. The same segmentation and post-processing workflow was applied to all scans so that species contrasts reflected morphology rather than reconstruction settings."
)
cite(p_m_ct, [1, 30])

head("Shotgun proteomics of eggshell matrix proteins")

para(
    "For shotgun proteomics, eggshell-matrix extracts from the post-CT sample set were analyzed as three pooled biological samples per species, with each chicken and duck sample comprising two eggs and each pigeon sample comprising six eggs. Proteins were "
    "extracted by resuspension in lysis buffer (1% SDS, 1% protease inhibitor "
    "cocktail), sonication on ice, and clarification by centrifugation at "
    "12,000 \u00d7 g at 4\u00b0C for 10 min; protein concentration was determined with a "
    "BCA assay kit. Proteins were precipitated with pre-cooled acetone (5 volumes, "
    "\u221220\u00b0C, 2 h), washed twice with acetone, and redissolved in 200 mM TEAB. "
    "Disulfide bonds were reduced with 5 mM dithiothreitol (56\u00b0C, 30 min) and "
    "alkylated with 11 mM iodoacetamide (room temperature, 15 min, dark). They "
    "were digested overnight with sequencing-grade trypsin (enzyme:protein ratio "
    "1:50) and desalted with Strata X SPE columns. No offline HPLC fractionation was performed before LC-MS/MS analysis. For proteome acquisition, 20 \u03bcg peptide was injected per run."
)

mixed([
    ("Desalted peptides were dissolved in mobile phase A (0.1% formic acid in water) and separated on a "
     "home-made 15-cm \u00d7 100-\u03bcm i.d. reversed-phase C18 analytical column connected "
     "to a Vanquish Neo UPLC system (Thermo Fisher Scientific). Mobile phase B consisted of 0.1% formic acid in 80% acetonitrile, and the flow rate was maintained at 400 nl/min. The gradient program was 4% B for 0\u20130.5 min, 4\u20138% B for 0.5\u20130.6 min, 8\u201322.5% B for 0.6\u201313.6 min, 22.5\u201335% B for 13.6\u201320.5 min, 35\u201355% B for 20.5\u201320.9 min, 55\u201399% B for 20.9\u201321.4 min, and 99% B for 21.4\u201322.6 min. "
    "After UPLC separation, peptides were introduced into a nano-electrospray ionization source operated at 1,900 V and analyzed on an Orbitrap Astral mass spectrometer (Thermo Fisher Scientific). "
     "Full-MS spectra were acquired in the Orbitrap at 240,000 resolution over "
    "380\u2013980 m/z; MS/MS fragments were acquired in the Astral analyzer at 80,000 "
     "resolution using HCD fragmentation in DIA mode (NCE\u202f=\u202f25%), fixed first mass 150 m/z, "
     "AGC target 500%, and maximum injection time 3 ms. DIA data were processed with "
     "DIA-NN v1.8 against species-specific reference proteomes \u2014 ", False, False),
    ("G.\u00a0gallus", False, True),
    (" (43,711 entries), ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (" (91,801 entries), and ", False, False),
    ("C.\u00a0livia", False, True),
    (" (17,309 entries), all downloaded August 2024 \u2014 "
     "concatenated with reverse decoy sequences. "
     "Using species-matched reference databases reduced the risk that peptide assignments would be biased toward the best-annotated proteome. "
     "Trypsin/P cleavage was specified with up to one missed cleavage; N-terminal "
     "methionine excision and carbamidomethylation of Cys were set as fixed "
     "modifications. Protein and peptide FDR were each controlled at < 1%.", False, False),
])

head("Intact glycopeptide mass spectrometry")

para(
    "For intact glycopeptide analysis, the same three pooled biological samples per species were carried forward, and 200 \u03bcg peptide digest from each pooled sample was used as input for enrichment. N-glycopeptides were enriched from tryptic digests by hydrophilic interaction "
    "liquid chromatography (HILIC). Peptide digests were redissolved in loading "
    "buffer (80% ACN, 5% TFA), loaded onto a HILIC column, washed three times with "
    "loading buffer, and glycopeptides were eluted twice with 0.1% TFA, 50 mM "
    "ammonium bicarbonate, and 50% ACN. Eluates were desalted with C18 Zip-Tips "
    "and vacuum-dried. Glycopeptide fractions were separated on the same nano-LC "
    "platform using a 34-min gradient (4\u201399% B at 400 nl/min). Full-MS spectra were "
    "acquired at 240,000 resolution over 700\u20132,000 m/z; MS/MS spectra were acquired "
    "at 80,000 resolution with fixed first mass 120 m/z, cycle time 0.6 s, AGC "
    "target 100%, intensity threshold 25,000 ions/s, and maximum injection time 5 ms. "
    "Raw DDA data were processed with MSFragger v3.4 against the same species-specific "
    "reference proteomes as above, with strict trypsin cleavage (up to 2 missed "
    "cleavages), peptide length 7\u201350 residues, fixed carbamidomethylation of Cys, "
    "variable N-terminal acetylation and Met oxidation, and default MSFragger "
    "glycosylation mass offsets. Protein, peptide, and PSM FDR were each controlled "
    "at < 1%. N-glycan structural classes were assigned using the Oxford nomenclature "
    "into six categories: High-Mannose, Paucimannose/Truncated, Neutral Complex/"
    "Hybrid, Fucosylated Complex/Hybrid, Sialylated Complex/Hybrid, and Other. "
    "Per-protein per-species class abundance was computed as the fraction of total "
    "glycan-site signal intensity. Structural classes were assigned after site-level identification so that shared glycan usage could be compared on a common protein background across species."
)

head("Comparative proteome analysis and gene-family evolution")

p_m_ortho = mixed([
    ("Protein sequences from all three species were assigned to orthogroups using "
     "OrthoFinder (all-versus-all BlastP; E-value threshold 1 \u00d7 10\u207b\u00b9\u2070; "
     "MCL inflation parameter 2.0), yielding a 3,250-orthogroup background set for "
     "all downstream enrichment analyses. Divergence times were taken from published "
     "timetrees: ", False, False),
    ("G. gallus", False, True),
    ("\u2013", False, False),
    ("A. platyrhynchos", False, True),
    (" 83.37 Mya and ", False, False),
    ("G. gallus", False, True),
    ("\u2013", False, False),
    ("C. livia", False, True),
    (" 90.84 Mya. Gene-family expansion and contraction rates were inferred with "
     "CAFE5 using the time-calibrated phylogeny. GO enrichment of pairwise-shared, "
     "species-exclusive, expanding, and contracting orthogroup sets was performed "
     "using the OrthoVenn3 web platform (https://orthovenn3.bioinfotoolkits.net) "
     "against the 3,250-orthogroup background; terms with adjusted p\u202f<\u202f0.05 were "
     "considered significant.", False, False),
])
cite(p_m_ortho, [3, 5, 14])

head("Cross-species glycoprotein ortholog identification")

mixed([
    ("High-confidence cross-species orthologues of four target eggshell glycoproteins "
     "(OVAL, OC116, TRFE, OC17) were identified by BlastP of each reference ", False, False),
    ("G.\u00a0gallus", False, True),
    (" sequence against the non-reference-species proteomes (E-value "
     "threshold 1 \u00d7 10\u207b\u2075; maximum 500 target sequences; 250 reported alignments). "
     "Candidate hits were retained at average maximum sequence identity \u2265 0.40; "
     "where query and subject non-overlapping HSP counts were unequal, a relaxed "
     "maximum-identity threshold of \u2265 0.40 was applied. Final UniProt ortholog identifiers used for "
     "downstream structural and quantitative analyses are listed in Table S3.", False, False),
])

head("Integrated protein-glycan abundance comparison")

para(
    "For Fig. 3A to C, protein-level and glycan-site quantification tables were "
    "loaded separately for each species from the Protein_quant and Site_quant sheets. "
    "When the Number Comparable field was present, protein entries and glycan-site entries with values < 2 were excluded so that only features detected in at least two of the three pooled biological samples were retained for downstream comparable-protein analysis. Mean protein intensity "
    "was calculated across all species-matched intensity columns for each accession, "
    "and mean glycan intensity was calculated for each quantified glycosylation site "
    "across the corresponding site-intensity columns. Only positive-intensity entries "
    "were retained. Glycan-site rows were then inner-joined to protein rows by protein "
    "accession so that each point represented one quantified glycosylated sequon with "
    "a matched protein-abundance measurement. Protein and glycan intensities were "
    "log2-transformed, and within-species protein-glycan coupling was summarized by "
    "Spearman rank correlation with two-sided p values. OVAL, OC116, TRFE, and OC17 "
    "were highlighted by the strict ortholog assignments summarized in Fig. 2E and Table S3, and labels "
    "were annotated with the corresponding glycosylated Asn position.")

para(
    "For Fig. 3D to F, pairwise glycan-protein enrichment plots were built from "
    "ortholog-mapped protein and glycan abundance differences between species. Protein "
    "abundance for each accession was defined as the mean of nonzero replicate "
    "intensities after excluding proteins with Number Comparable < 2 when available. "
    "Glycan abundance was defined at the protein level as the sum of mean nonzero site "
    "intensities across all quantified glycosylation sites assigned to that accession after the same comparable-feature filtering. "
    "Gallus-versus-Anas and Gallus-versus-Columba comparison spaces were built from "
    "blastp outfmt 6 mappings, retaining the best hit per query when the mean E value "
    "was <= 1 × 10⁻⁵ and the average sequence identity was >= 0.40; when query and "
    "subject had different numbers of non-overlapping HSPs, the maximum identity "
    "threshold >= 0.40 was applied instead. The Anas-versus-Columba plane was bridged "
    "through shared Gallus orthologs that passed the same filter in both datasets. For "
    "each retained ortholog pair, the x coordinate was calculated as log2(I_ref) - "
    "log2(I_comp) and the y coordinate as log2(G_ref) - log2(G_comp), where I and G "
    "denote protein and glycan abundance, respectively. The y = x diagonal therefore "
    "marked matched protein-glycan change, whereas displacement toward the glycan-rich "
    "side identified proteins whose glycan shift exceeded the corresponding change in "
    "bulk protein abundance.")

head("N-glycan structural ensemble modelling")

p_m_reglyco = mixed([
    ("OVAL ortholog protein structures (AlphaFold2-predicted models) were "
     "accessed through the GlycoShape platform (glycoshape.org) via UniProt "
     "accession identifiers. "
     "Experimentally detected N-glycan compositions from IGP-MS were mapped "
     "to the GlycoShape glycan library by monoisotopic mass matching "
     "(tolerance \u00b10.5\u00a0Da), using per-residue masses of 203.0794\u00a0Da (HexNAc), "
     "162.0528\u00a0Da (Hex), 291.0954\u00a0Da (NeuAc), 146.0579\u00a0Da (dHex), and "
    "132.0423\u00a0Da (Pen) with an 18.0106\u00a0Da water correction; matched entries "
    "were retrieved as GlyTouCan accession identifiers. All experimentally detected glycoforms that could be matched to the GlycoShape library were retained for downstream modelling rather than pre-filtering to a dominant subset. "
     "Full conformational ensembles were then generated with the GlycoShape "
     "Re-Glyco Ensemble tool (glycoshape.org/ensemble), which restores missing "
     "glycans by aligning them to torsion angles from Privateer crystallographic "
     "standards and sampling conformations from the GlycoShape molecular-dynamics "
    "ensemble library. "
    "For each OVAL ortholog, glycan modelling was constrained to the experimentally detected glycopeptide site identified in the present dataset; each matched "
    "glycan was then submitted as an independent modelling job via the GlycoShape API and attached to the measured target sequon \u2014 ", False, False),
    ("G.\u00a0gallus", False, True),
    (" N293 and ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (" / ", False, False),
    ("C.\u00a0livia", False, True),
    (" N97 \u2014 with ensemble size 50, random seed 42, and PDB output format. "
     "This procedure yielded 50 glycoprotein models for ", False, False),
    ("G.\u00a0gallus", False, True),
    (" (1 matched glycan type), 150 for ", False, False),
    ("A.\u00a0platyrhynchos", False, True),
    (" (3 types), and 700 for ", False, False),
    ("C.\u00a0livia", False, True),
    (" (14 types, "
     "including four NeuAc positional isomers resolved from the GlycoShape library). Using the same random seed and ensemble size for every submitted glycan maintained matched sampling depth across species. "
     "Ensemble geometric descriptors were calculated for each model from atomic "
     "coordinates using BioPython: the radius of gyration (Rg) of all glycan heavy "
     "atoms, the end-to-end distance of the glycan chain, and the minimum distance "
     "between any glycan heavy atom and protein C\u03b1 atoms (minimum C\u03b1 contact "
     "distance). Per-structure summary statistics (mean\u00a0\u00b1\u00a0s.d.) and pairwise species "
     "comparisons (Mann\u2013Whitney\u00a0U test, two-tailed) were performed for each descriptor and are reported in Tables S5 and S6.", False, False),
])
cite(p_m_reglyco, [11, 56, 65])

head("Electrostatic potential calculation")

p_m_apbs = para(
    "Electrostatic surface potentials were computed for each Re-Glyco ensemble "
    "model and a matched apo reference with glycans removed using APBS v3.4.1. Atomic "
    "partial charges and radii were assigned with PDB2PQR using the CHARMM36 force "
    "field and PROPKA protonation at pH 7.4; glycan heavy-atom partial charges were "
    "assigned from published GLYCAM06 values. APBS input grids were generated automatically from the PQR bounding box with 10 \u00c5 padding and a target grid spacing of 0.5 \u00c5; focusing grids were set to 70% of the coarse-grid lengths. The nonlinear Poisson\u2013Boltzmann equation was solved with single-ion boundary conditions, 0.15 mol/L monovalent salt (ion radii 2.0 and 1.8 \u00c5 for cation and anion, respectively), protein dielectric 2.0, solvent dielectric 78.54, solvent-accessible surface definition smol, charge discretization spl0, solvent probe radius 1.4 \u00c5, spline window 0.3 \u00c5, surface density 10.0, and temperature 298.15 K. Solvent-accessible surface areas were calculated by the "
    "Shrake\u2013Rupley algorithm; surface residues were defined by relative ASA "
    "\u2265 0.25. Ca\u00b2\u207a-binding electrostatic hotspots were defined as surface Asp or "
    "Glu residues with APBS potential < \u22125 kT/e. Ensemble-level metrics reported "
    "include hotspot count (N_hot), per-hotspot mean SASA, hotspot fraction of total "
    "surface Asp/Glu, and median surface electrostatic potential. The same electrostatic threshold and surface-definition rules were applied to all models to preserve cross-species comparability."
)
cite(p_m_apbs, [12, 42, 43])

head("Finite-element analysis")

p_m_fea = smixed([
    ([
        ("The region of interest used for downstream finite-element analysis was defined as a cylindrical volume of 1 mm radius during micro-CT reconstruction. Surface models derived from micro-CT were first exported as STL files and reverse-engineered in Geomagic Wrap for finite-element pre-processing by sequential de-noising (strength 2), triangle simplification to approximately 300,000 faces, mesh re-gridding at 0.01 mm, iterative defect correction to zero residual faults, and organic parametric surface fitting at minimum tolerance. The resulting eggshell surface models were then imported into Ansys Workbench 2023 R1 and solved with the explicit LS-DYNA module (unit system: mm/kg/N/s). ", False, False),
    ], []),
    ([
        ("To isolate structure-driven mechanical differences, identical eggshell material properties were assigned across species; parameter values were adopted from a previous avian eggshell elasticity study rather than re-estimated separately for each species.", False, False),
    ], [89]),
    ([
        (" In the solver keyword deck, the eggshell was modeled with *MAT_PLASTIC_KINEMATIC and *SECTION_SOLID, using density 2770 kg/m³, Young's modulus 3.0 \u00d7 10\u00b9\u2070 Pa, Poisson's ratio 0.33, yield strength 1.5 \u00d7 10\u2077 Pa, tangent modulus 0, and maximum equivalent plastic strain at failure 0.05. This explicit impact setup followed the general logic of crash-deformation simulations, but was rescaled to the local eggshell loading geometry studied here. The impactor, used to simulate the egg tooth, was a frustum (base radius 0.1 mm; top radius 0.5 mm; height 0.5 mm) assigned the library IRON-ARMCO explicit material and meshed as a separate solid part. Contact between the impactor and eggshell was defined with *CONTACT_AUTOMATIC_SURFACE_TO_SURFACE using a friction coefficient of 0.2. Eggshell mesh element sizes were 0.05 mm (", False, False),
        ("G. gallus", False, True),
        ("), 0.05 mm (", False, False),
        ("A. platyrhynchos", False, True),
        ("), and 0.03 mm (", False, False),
        ("C. livia", False, True),
        ("), ensuring at least six element layers across the eggshell cross-section; the impactor was meshed at 0.1 mm. The impactor was driven by an imposed initial velocity of 50,000 mm/s along the loading axis, whereas one boundary set on the fragment perimeter was fully fixed in all translational and rotational degrees of freedom. Analyses ran for 1.0 \u00d7 10\u207b\u2074 s with a time-step safety factor of 0.7, erosion enabled, a minimum time step of 1 \u00d7 10\u207b\u2078 s, and automatic mass scaling; solver outputs included GLSTAT, SPCFORC, RCFORC, NCFORC, BNDOUT, NODOUT, MATSUM, ELOUT, JNTFORC, and DEFORC at 1.0 \u00d7 10\u207b\u2077 s intervals, with D3PLOT and INTFOR written every 1.0 \u00d7 10\u207b\u2076 s. For the positional loading analysis, the impactor was sampled at nine lateral offsets arranged on a 3 \u00d7 3 grid with 0.5 mm spacing. Across these nine cases, only the impactor coordinates were translated; material definitions, contact settings, boundary conditions, fragment size, and all other solver controls were held constant. Peak contact force (F_max) and peak contact shear stress (\u03c4_max) were extracted for each position and are reported in Table S7. Sampling nine offsets allowed local positional heterogeneity to be measured without changing fragment size or loading geometry between species.", False, False),
    ], [16, 37]),
])

head("Statistical analysis")

mixed([
    ("All values are expressed as mean \u00b1 s.d. All statistical tests were two-tailed, "
     "and p < 0.05 was considered statistically significant throughout. "
    "Normality was evaluated by the Shapiro\u2013Wilk test and homogeneity of variance by Levene's test before parametric between-species analyses. Mammillary morphometric parameters were compared among species by one-way ANOVA "
    "followed by Tukey's honestly significant difference test (Tukey HSD; \u03b1\u202f=\u202f0.05). These nine morphometric observations per species were non-overlapping subfragments from one scanned fragment and should therefore be interpreted as within-fragment regional replicates rather than as nine independent biological samples. The same assumption checks supported one-way ANOVA with Tukey HSD for finite-element outcomes (F_max and \u03c4_max), with the underlying values reported in Table S7. In contrast, glycan ensemble geometric descriptors (Rg, end-to-end distance, minimum glycan\u2013protein contact distance) and hotspot-derived ensemble metrics did not satisfy normality and/or homoscedasticity across species, so pairwise species contrasts for these variables were evaluated with two-sided Mann\u2013Whitney U tests. "
     "Glycosylation-induced reduction in N_hot within ", False, False),
    ("C.\u00a0livia", False, True),
    (" was assessed by one-sample Wilcoxon signed-rank test versus the apo reference value; total Asp/Glu SASA at the whole-interface level was summarized descriptively because the structure-level values were invariant across ", False, False),
    ("C.\u00a0livia", False, True),
    (" glycoforms; shifts in median surface electrostatic potential were assessed by one-sample Wilcoxon signed-rank test against the apo reference value. No multiple-testing correction was applied, and no outliers were removed. All statistical analyses were "
     "conducted in Python using scipy.stats and statsmodels.", False, False),
])

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
# References
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
para("References", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)

for source_number in CITATION_ORDER:
    p_ref = doc.add_paragraph(style="Normal")
    p_ref.paragraph_format.space_before = Pt(0)
    p_ref.paragraph_format.space_after = Pt(4)
    p_ref.paragraph_format.left_indent = Pt(18)
    p_ref.paragraph_format.first_line_indent = Pt(-18)
    ref_text = f"{CITATION_MAP[source_number]}. {REF_TEXTS[source_number]}"
    r_ref = p_ref.add_run(ref_text)
    r_ref.font.size = Pt(9)
    rPr = r_ref._r.get_or_add_rPr()
    _set_font(rPr, FONT)

para("Acknowledgments", bold=True, size=14, before=320, after=160,
     align=WD_ALIGN_PARAGRAPH.LEFT)
para(
    "We thank J. Chang from the College of Veterinary Medicine, China Agricultural "
    "University, for providing pigeon egg materials. We thank X. Ye for drawing the bird "
    "icons and realistic bird images. We thank B. Tan, Z. Huang, and X. Li for valuable comments "
    "on the manuscript."
)

mixed([
    ("Funding:", True, False),
    (" This work was supported by the National Key Research and Development Program of China "
     "(2022YFD1300100), the China Agriculture Research Systems (CARS-40), and the National "
     "Key Research and Development Program of China (2021YFD1200803).", False, False),
])

mixed([
    ("Author contributions:", True, False),
    (" conceptualization: L.X., L.Z., G.X., and J.Z. methodology: L.X., L.Z., X.S., and J.Z. validation: L.X., Yaqi L., J.Y., Yu L., C.Z., T.L., W.Z., and J.Z. formal analysis: L.X. investigation: L.X., Yaqi L., Yu L., C.Z., and L.Z. data curation: L.X., Yaqi L., J.Y., Yu L., and C.Z. software: L.X. and J.Y. resources: G.X. and J.Z. writing—original draft: L.X. writing—review and editing: L.X., Yaqi L., J.Y., Q.W., T.L., G.X., and J.Z. visualization: J.Y., Q.W., and J.Z. supervision: G.X. and J.Z. funding acquisition: G.X. and J.Z.", False, False),
])

mixed([
    ("Competing interests:", True, False),
    (" The authors declare they have no competing interests.", False, False),
])

mixed([
    ("Data, code, and materials availability:", True, False),
    (" All data and code needed to evaluate and reproduce the results in the paper are present "
     "in the paper and/or the Supplementary Materials. Materials generated in this study are "
     "available from J.Z. on reasonable request (jxzheng@cau.edu.cn).", False, False),
])

doc.save(OUT)
print(f"[OK]  {OUT}")



