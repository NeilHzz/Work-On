"""
Generate a Science Advances cover letter for the current manuscript.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(__file__).with_name("0_Cover letter") / "cover_letter260615.docx"
FONT = "Times New Roman"


def set_font(run, size=11, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), FONT)
    r_pr.insert(0, r_fonts)


def spacing(paragraph, before=0, after=120, line=16):
    p_pr = paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:spacing")
    element.set(qn("w:before"), str(before))
    element.set(qn("w:after"), str(after))
    element.set(qn("w:line"), str(line * 20))
    element.set(qn("w:lineRule"), "auto")
    p_pr.append(element)


def para(doc, text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=120):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    spacing(paragraph, after=after)
    run = paragraph.add_run(text)
    set_font(run, bold=bold)
    return paragraph


doc = Document()
section = doc.sections[0]
section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(2.54)

para(doc, "June 15, 2026")
para(doc, "Science Advances Editorial Office")
para(doc, "Dear Editors,")

para(
    doc,
    'We are pleased to submit our manuscript entitled "OVAL glycan states link eggshell matrix chemistry to shell architecture and avian shell-breaking mechanics" for consideration as a Research Article in Science Advances.',
)

para(
    doc,
    "Avian eggshells offer a concise biological test of a broader materials question: how a conserved organic matrix can produce mineralized structures that are both protective and locally breakable. In eggshell formation, this problem is concentrated at the mammillary layer, where matrix-guided nucleation first organizes the physical eggshell architecture later encountered during shell breaking.",
)

para(
    doc,
    "The manuscript builds a cross-scale argument for this problem. We first show that species divergence emerges at the mammillary layer before broad turnover of the eggshell-matrix protein toolkit. We then identify OVAL glycan state as the clearest chemically interpretable contrast within that shared matrix background.",
)

para(
    doc,
    "Using intact glycopeptide mass spectrometry, Re-Glyco structural modelling, and electrostatic analysis, we connect OVAL glycan-state variation to calcium-accessible surface presentation. This supports a model in which glycan-dependent surface accessibility can influence OVAL opening, nucleation-site exposure, and mammillary-layer organization.",
)

para(
    doc,
    "The inside-out finite-element analysis completes the argument by moving the endpoint from whole-shell strength to local stress transfer at the mammillary interface. The resulting framework links shared matrix chemistry, glycan-state variation, physical eggshell architecture, and localized shell-breaking mechanics.",
)

para(
    doc,
    "We believe the study is well suited to Science Advances because it addresses a broad biomineralization question with a mechanistic, cross-scale design. Rather than presenting a descriptive comparison of avian eggshells, the manuscript provides a testable framework for how posttranslational matrix states organize material phenotypes across molecular, mesoscale, and mechanical levels.",
)

para(
    doc,
    "All authors have approved the manuscript and its submission. The work is original and is not under consideration elsewhere. Conflicts of interest, funding information, and data and code availability statements are provided in the manuscript.",
)

para(doc, "Sincerely,")
para(doc, "Jiangxia Zheng")
para(doc, "National Engineering Laboratory for Animal Breeding and MOA Key Laboratory of Animal Genetics and Breeding, College of Animal Science and Technology, China Agricultural University")
para(doc, "jxzheng@cau.edu.cn")

doc.save(OUT)
print(f"[OK] {OUT}")
