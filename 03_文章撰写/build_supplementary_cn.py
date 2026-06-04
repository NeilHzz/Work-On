"""
Science Advances 格式 — 中文版补充材料
输出: supplementary_materials_cn.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FIG_BASE = Path(r"D:\system_folder\Desktop\Work On\Supplementary\Figures")
MAIN_FIG_BASE = Path(__file__).resolve().parent.parent / "Figure260421"
PANEL_FIG_BASE = Path(__file__).resolve().parent.parent / "02_可视化" / "Figure" / "PNG"
OUT = str(Path(__file__).with_name("supplementary_materials260605_cn.docx"))

doc = Document()


# 页面设置
s = doc.sections[0]
s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.54)
s.page_width = Cm(21.0)
s.page_height = Cm(29.7)

SCI_F = "Times New Roman"
BODY = "SimSun"


def _set_font(rPr, latin_name=SCI_F, east_asia_name=BODY):
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), latin_name)
    rFonts.set(qn("w:hAnsi"), latin_name)
    rFonts.set(qn("w:cs"), latin_name)
    rFonts.set(qn("w:eastAsia"), east_asia_name)
    rPr.insert(0, rFonts)


def fmt(run, size=11, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    _set_font(run._r.get_or_add_rPr())


def spacing(p, before=0, after=120, line=24):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement("w:spacing")
    e.set(qn("w:before"), str(before))
    e.set(qn("w:after"), str(after))
    e.set(qn("w:line"), str(line * 20))
    e.set(qn("w:lineRule"), "auto")
    pPr.append(e)


def keep_with_next(p):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement("w:keepNext")
    pPr.append(e)


def para(text, bold=False, italic=False, size=11, before=0, after=120,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, before=before, after=after)
    r = p.add_run(text)
    fmt(r, size=size, bold=bold, italic=italic)
    return p


def fig_title(label, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    spacing(p, before=200, after=120)
    keep_with_next(p)
    r_label = p.add_run(label + " ")
    fmt(r_label, bold=True)
    r_title = p.add_run(title)
    fmt(r_title, bold=True)
    return p


def fig_caption(text, before=0, after=200):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    spacing(p, before=before, after=after)
    r = p.add_run(text)
    fmt(r)
    return p


def add_image(img_path, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=60, after=60, line=12)
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    return p


def add_images_row(img_paths, width_cm=7.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=60, after=60, line=12)
    for i, img_path in enumerate(img_paths):
        p.add_run().add_picture(str(img_path), width=Cm(width_cm))
        if i < len(img_paths) - 1:
            p.add_run("  ")
    return p


# 封面
para("补充材料", bold=True, size=14,
     before=0, after=240, align=WD_ALIGN_PARAGRAPH.CENTER)

para(
    "跨物种 OVAL 糖链状态将乳突层组织与有利于破壳的蛋壳力学联系起来",
    italic=False, size=11, before=0, after=360,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

para("[作者、单位和通讯作者信息与正文一致]",
     size=10, before=0, after=240, align=WD_ALIGN_PARAGRAPH.CENTER)

para("本 PDF 文件包括：", size=11, before=0, after=60,
     align=WD_ALIGN_PARAGRAPH.LEFT)

for line in [
    "补充文本 1 至 2",
    "图 S1 至 S11",
    "",
    "本文其他补充材料包括：",
    "表 S1 至 S7",
]:
    para(line, size=11, before=0, after=40, align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()


# 补充文本
para("补充文本", bold=True, size=12,
     before=0, after=240, align=WD_ALIGN_PARAGRAPH.LEFT)

para("补充文本 1. 物种选择与敏感性分析。",
     bold=True, size=11, before=0, after=60, align=WD_ALIGN_PARAGRAPH.LEFT)

fig_caption(
    "基于 AVONET 生态性状得分（体重、喙长、喙宽、喙深、跗跖长、翼长、Kipps 距离、手翼指数、尾长，以及经数值编码的主要生活方式、栖息地和营养生态位）的主成分分析，将 Gallus gallus、Anas platyrhynchos 和 Columba livia 分离到鸟类生态空间中三个彼此独立的区域（图S1），分别对应陆栖地面营巢的早成型、半水生早成型和高位营巢的晚成型生活史策略。这组三物种被选中，是为了在共同出壳框架下同时覆盖发育和生态两个比较维度，而不是追求单一的系统发育或外部形态差异。为检验这种分离是否依赖于分类变量的数值编码，我们进行了 500 次随机扰动迭代，在每次迭代中对全部编码权重在原始值的 ±30% 范围内独立变动。方差解释度和聚类 silhouette 分数始终紧密集中在未扰动的基线值附近（图S1），说明物种分组对合理范围内的编码变化保持稳健。"
)

para("补充文本 2. 蛋壳基质蛋白组正交组分析。",
     bold=True, size=11, before=240, after=60, align=WD_ALIGN_PARAGRAPH.LEFT)

fig_caption(
    "基于蛋白组学鉴定到的蛋壳基质蛋白，经 OrthoFinder 正交分析流程整理后，在 G. gallus、A. platyrhynchos 和 C. livia 中分别得到 2,620、2,921 和 3,219 个正交组（图S3）。其中，1,997 个正交组为三物种共享，构成保守的蛋壳基质蛋白核心。两两共享但并非三物种共同共享的正交组数量分别为 180（G. gallus–A. platyrhynchos）、434（G. gallus–C. livia）和 716（A. platyrhynchos–C. livia）；谱系限制性集合分别包括鸡、鸭和鸽的 9、28 和 72 个正交组。该结构说明，后续比较建立在共享基质背景之上，而不是建立在工具箱整体替换之上。"
)

fig_caption(
    "成对共享集合的 Gene Ontology（GO）富集显示，功能分层更多沿生态轴线展开，而不是简单追随系统发育邻近性（图S5）。A. platyrhynchos–C. livia 共享集合最显著富集于钙离子结合和金属离子结合，以及 Wnt 信号通路和信号转导；G. gallus–A. platyrhynchos 共享集合则富集于适应性免疫反应和精子发生。谱系限制性 GO 信号进一步强化了这一对比，最突出的是 G. gallus 特异集合中保留了蛋白 N-连接糖基化。"
)

fig_caption(
    "A. platyrhynchos 特异集合（n = 28）主要富集于免疫反应调控、B 细胞激活和铁反应；C. livia 特异集合（n = 72）则富集于神经系统发育、泛素依赖性蛋白质分解代谢和蛋白水解。与之相比，G. gallus 特异集合虽然规模较小（9 个正交组），但保留了蛋白 N-连接糖基化这一关键信号。"
)

fig_caption(
    "基因家族扩张与收缩分析（CAFE5）进一步显示谱系分化具有明显不对称性：G. gallus 总体表现为净家族收缩，A. platyrhynchos 居中，而 C. livia 表现为净扩张（图S8和S9）。鸡中收缩的家族富集于免疫相关功能；鸽中扩张的家族则富集于跨膜转运、Rho 信号和突触相关过程。这些蛋白组层面的模式说明三种蛋壳形成系统之间存在广泛分化，同时共享核心工具箱仍被保留。"
)

doc.add_page_break()

para("补充图", bold=True, size=12,
     before=0, after=240, align=WD_ALIGN_PARAGRAPH.LEFT)

# 图S1
fig_title("图S1.", "验证宏观生态物种选择框架的敏感性分析。")
add_image(FIG_BASE / "SuppFig1_Species_Selection" / "Sensitivity_Analysis_Results.png", width_cm=15.5)
fig_caption(
    "在基于 AVONET 的主成分空间中进行 500 次随机扰动迭代后，方差解释度（R²）和聚类 silhouette 系数的分布。目标物种为 Gallus gallus、Anas platyrhynchos 和 Columba livia。分类生态变量采用数值编码；每次迭代均在原始值 ±30% 范围内对全部编码权重进行独立随机扰动。两项指标集中于基线值附近，支持物种分组在不同编码方案下保持稳定。"
)


# 图S2
doc.add_page_break()
fig_title("图S2.", "目标物种在更广鸟类比较框架中的系统位置与比较轴热图。")
add_image(PANEL_FIG_BASE / "Fig1B.png", width_cm=15.5)
fig_caption(
    "代表性鸟类类群的系统发育关系及水生关联（X）、发育方式（Z）和生态不一致性（Y）三条比较轴的热图。彩色目级标签表示物种选择所使用的更广比较框架。目标谱系覆盖的功能轴线只与系统发育关系部分重合。"
)


# 图S3
doc.add_page_break()
fig_title("图S3.", "三物种共享与谱系限制性蛋壳基质正交组的 Venn 图。")
add_image(FIG_BASE / "SuppFig2_Venn_Orthogroups" / "Fig_venn_orthogroups.png", width_cm=12.0)
fig_caption(
    "三种蛋壳基质蛋白组的 OrthoFinder 正交组分析。数字表示三物种共享核心、两两共享集合和谱系限制性集合中的正交组数量。大的共享核心支持在共同蛋白背景上进行跨物种比较。"
)


# 图S4
doc.add_page_break()
fig_title("图S4.", "基于单拷贝直系同源物重建的三个目标物种的最大似然系统发育树。")
add_image(FIG_BASE / "SuppFig3_Phylo_Tree" / "Fig_phylo_tree.png", width_cm=14.0)
fig_caption(
    "基于单拷贝直系同源蛋白序列串联比对并由 IQ-TREE 推断的系统发育树。分支长度表示每个位点替换数。内部节点显示 1000 次重复得到的 ultrafast bootstrap 支持度。"
)


# 图S5
doc.add_page_break()
fig_title("图S5.", "物种特异与成对共享蛋壳基质蛋白集合的 GO 富集。")
add_image(FIG_BASE / "SuppFig4_GO_Enrichment" / "图2.jpg", width_cm=16.0)
fig_caption(
    "上半部分显示三个两两共享正交组区域（GnA，Gallus–Anas；GnC，Gallus–Columba；AnC，Anas–Columba）的 GO 富集条目；下半部分显示三个物种特异正交组集合（Gallus、Anas、Columba）的 GO 富集条目。颜色区分 GO 类别：生物过程（BP）、细胞组分（CC）和分子功能（MF）。G. gallus 特异集合包括蛋白 N-连接糖基化这一富集生物过程。"
)


# 图S6
doc.add_page_break()
fig_title("图S6.", "反复出现的蛋壳基质蛋白的蛋白特异性糖基化谱。")
add_image(PANEL_FIG_BASE / "Fig4D_G.png", width_cm=15.5)
fig_caption(
    "OVAL、OC116、TRFE 和 OC17 等反复出现的蛋壳基质蛋白在鸡、鸭和鸽中的糖链类别组成。堆叠柱表示每个蛋白-物种组合中检测到的不同糖链类别的相对贡献。"
)


# 图S7
doc.add_page_break()
fig_title("图S7.", "反复出现的蛋壳基质蛋白糖型谱及 OVAL 结构系综表面电势背景。")
add_image(PANEL_FIG_BASE / "FigS7.png", width_cm=15.8)
fig_caption(
    "(A) 糖基化与 apo OVAL 结构系综在三个物种中的表面电势分布。(B) 糖基化 OVAL 模型及配对 apo 参考结构的逐结构表面电势图。表面电势由 APBS 计算得到。"
)


# 图S8
doc.add_page_break()
fig_title("图S8.", "三个物种的 CAFE5 基因家族扩张与收缩。")
add_image(FIG_BASE / "SuppFig5_CAFE5_Gene_Family_Turnover" / "Fig_cafe5_expansion_contraction.png", width_cm=14.0)
fig_caption(
    "系统发育树上标注了利用物种分化时间树并由 CAFE5 推断的谱系特异性基因家族扩张（红色）与收缩（蓝色）事件。节点上的数字表示估计的祖先基因家族大小，分支上的数字表示净变化。仅展示每个家族 p < 0.05（Viterbi p 值）的基因家族。"
)


# 图S9
doc.add_page_break()
fig_title("图S9.", "基因家族周转对应的谱系偏向功能富集。")
add_image(PANEL_FIG_BASE / "Fig2H.png", width_cm=16.0)
fig_caption(
    "基于扩张和收缩基因家族得到的谱系、周转方向与 GO 富集条目之间的 alluvial 汇总图。流线颜色区分扩张与收缩信号，右侧端点概括各谱系相关的生物过程、细胞组分和分子功能条目。"
)
# 图S10
doc.add_page_break()
fig_title("图S10.", "OVAL 糖链几何及 apo/糖基化对照的 Re-Glyco 系综分析。")
add_image(FIG_BASE / "SuppFig7_Glycosylation_Hotspot" / "Fig_hotspot_ensemble_1.png", width_cm=15.5)
fig_caption(
    "(A) 三种物种特异性 OVAL–糖链复合物在构象系综重复中的糖链回转半径（Rg）分布，按物种着色（G. gallus 为橙色，A. platyrhynchos 为蓝色，C. livia 为绿色）。(B) 相同三种复合物的糖链端到端距离分布。(C) 各物种糖基化与 apo OVAL 结构逐构象的 Ca²⁺ 热点计数（Nhot）比较。Apo 结构为移除 N-糖链后的参考状态。Panel C 在存在结构层面变异时，采用相对于 apo 参考的一样本 Wilcoxon 符号秩检验评估。"
)


# 图S11
doc.add_page_break()
fig_title("图S11.", "各物种九个偏移位置的有限元反力时间历程。")
add_images_row([
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "chicken_rcforc_3x3.png",
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "chicken_rcforc_yforce.png",
], width_cm=7.5)
add_images_row([
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "duck_rcforc_3x3.png",
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "duck_rcforc_yforce.png",
], width_cm=7.5)
add_images_row([
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "pigeon_rcforc_3x3.png",
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "pigeon_rcforc_yforce.png",
], width_cm=7.5)
fig_caption(
    "(A–C) G. gallus（A）、A. platyrhynchos（B）和 C. livia（C）在九个参数化冲击位置（3 × 3 横向偏移网格）上的接触力（F）时间历程曲线。每条曲线代表一次模拟，显示从接触开始到峰值接触力的过程。内嵌图给出了各物种九个位置的峰值接触力（F_max）分布。(D–F) 对应的 Y 方向反力（FY）时间历程。由九次重复计算得到的物种峰值接触力（F_max）和峰值接触剪切应力（τ_max）的均值 ± s.d. 已在正文和图5中报告。模拟采用 LS-DYNA（Ansys）显式动力有限元分析完成。蛋壳厚度设置为基于 micro-CT 测得的物种特异性数值。"
)


doc.save(OUT)
print(f"Saved -> {OUT}")
