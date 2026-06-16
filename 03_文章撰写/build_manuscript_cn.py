"""
Strict Chinese translation generator.
Produces a sentence-aligned Chinese manuscript from the latest English manuscript.
"""

from pathlib import Path
import json
import re
import time
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from deep_translator import GoogleTranslator

EN_DOC = Path(__file__).with_name("0_Manuscript") / "manuscript_submitted.docx"
OUT_DOC = Path(__file__).with_name("0_Manuscript_CN") / "manuscript_submitted_cn.docx"
CACHE_FILE = Path(__file__).with_name(".translation_cache_cn.json")
INCLUDE_IMAGES = False
INCLUDE_REFERENCES = False

translator = GoogleTranslator(source="en", target="zh-CN")
try:
    _cache: dict[str, str] = json.loads(CACHE_FILE.read_text(encoding="utf-8"))
except (FileNotFoundError, json.JSONDecodeError):
    _cache = {}

# Locked terminology keeps sentence-level translation stable across reruns.
TERM_LOCKS: list[tuple[str, str]] = [
    ("OVAL glycan states link eggshell matrix chemistry to shell architecture and avian shell-breaking mechanics", "OVAL糖链状态将蛋壳基质化学与壳结构和鸟类破壳力学联系起来"),
    ("OVAL glycans link shell architecture and mechanics", "OVAL糖链连接壳结构与力学"),
    ("Teaser", "一句话亮点"),
    ("Introduction", "引言"),
    ("Materials and Methods", "材料与方法"),
    ("Acknowledgments", "致谢"),
    ("Funding:", "资助："),
    ("Author contributions:", "作者贡献："),
    ("Competing interests:", "利益冲突："),
    ("Data, code, and materials availability:", "数据、代码和材料可用性："),
    ("Cross-species OVAL glycan states reveal a matrix mechanism for avian shell-breaking mechanics", "跨物种OVAL糖链状态揭示鸟类破壳力学的基质机制"),
    ("Eggshell matrix proteins are key regulators of eggshell structural formation, and existing studies have generated rich posttranslational-modification site maps, leaving the side-chain properties of glycosylation modifications as an important layer for further investigation.", "蛋壳基质蛋白是蛋壳结构形成的关键调控因子，已有研究已经建立了丰富的翻译后修饰位点图谱，而糖基化修饰侧链性质仍是有待进一步解析的重要层面。"),
    ("Ca²⁺accessibility-mammillary-mechanics axis", "Ca²⁺可及性-乳突层-力学轴"),
    ("Ca²⁺accessibility", "Ca²⁺可及性"),
    ("Ca²⁺access", "Ca²⁺可及性"),
    ("Ca²⁺accessible matrix-protein surfaces", "Ca²⁺可及的基质蛋白表面"),
    ("rapid vertebrate mineralization process", "快速的脊椎动物矿化过程"),
    ("rapid mineralization process", "快速矿化过程"),
    ("organic matrix-protein system", "有机基质蛋白系统"),
    ("calcium delivery", "钙供应"),
    ("matrix-guided calcite nucleation", "基质引导的方解石成核"),
    ("spacing and continuity", "间距和连续性"),
    ("later shell units", "后续蛋壳结构单元"),
    ("glycan-modulated OVAL unfolding", "糖链调控的OVAL解折叠"),
    ("OVAL unfolding", "OVAL解折叠"),
    ("matrix-bound nucleation-site exposure", "基质结合成核位点暴露"),
    ("nucleation-site exposure", "成核位点暴露"),
    ("nucleation-site presentation", "成核位点呈现"),
    ("mammillary-layer formation", "乳突层形成"),
    ("mature mammillary density", "成熟乳突层密度"),
    ("local shell-breaking structural strength", "局部破壳结构强度"),
    ("Sun et al.", "Sun等"),
    ("matrix behaviour", "基质行为"),
    ("matrix behavior", "基质行为"),
    ("ovalbumin (OVAL)", "卵清蛋白（OVAL）"),
    ("ovalbumin", "卵清蛋白"),
    ("OVAL", "卵清蛋白（OVAL）"),
    ("OVAL glycan states", "OVAL糖链状态"),
    ("OVAL glycan state", "OVAL糖链状态"),
    ("OC116", "OC116"),
    ("TRFE", "TRFE"),
    ("OC17", "OC17"),
    ("Re-Glyco", "Re-Glyco"),
    ("GlycoShape", "GlycoShape"),
    ("GlyTouCan", "GlyTouCan"),
    ("AlphaFold2", "AlphaFold2"),
    ("APBS", "APBS"),
    ("PDB2PQR", "PDB2PQR"),
    ("MSFragger", "MSFragger"),
    ("DIA-NN", "DIA-NN"),
    ("OrthoFinder", "OrthoFinder"),
    ("OrthoVenn3", "OrthoVenn3"),
    ("CAFE5", "CAFE5"),
    ("Gallus gallus", "Gallus gallus"),
    ("Anas platyrhynchos", "Anas platyrhynchos"),
    ("Columba livia", "Columba livia"),
    ("G. gallus", "G. gallus"),
    ("A. platyrhynchos", "A. platyrhynchos"),
    ("C. livia", "C. livia"),
    ("chicken", "鸡"),
    ("duck", "鸭"),
    ("pigeon", "鸽"),
    ("mammillary-layer", "乳突层"),
    ("mammillary layer", "乳突层"),
    ("mammillary", "乳突层"),
    ("matrix-protein", "基质蛋白"),
    ("matrix protein", "基质蛋白"),
    ("glycoprotein", "糖蛋白"),
    ("glycopeptides", "糖肽"),
    ("glycopeptide", "糖肽"),
    ("glycan-state", "糖链状态"),
    ("glycan states", "糖链状态"),
    ("glycan state", "糖链状态"),
    ("N-glycan", "N-糖链"),
    ("glycans", "糖链"),
    ("glycan", "糖链"),
    ("glycosylation", "糖基化"),
    ("inside-out loading", "由内向外加载"),
    ("inside-out", "由内向外"),
    ("egg-tooth", "卵齿"),
    ("Local finite-element loading connects egg-tooth contact geometry to species-specific shell resistance", "局部有限元加载将卵齿接触几何与物种特异性蛋壳阻力联系起来"),
    ("finite-element", "有限元"),
    ("micro-CT", "显微CT"),
    ("shell resistance", "蛋壳阻力"),
    ("interface resistance", "界面阻力"),
    ("local resistance", "局部阻力"),
    ("egg shell", "蛋壳"),
    ("eggshell", "蛋壳"),
    ("radius of gyration", "回转半径"),
    ("shell-breaking resistance", "破壳阻力"),
    ("shell-breaking-associated", "破壳相关"),
    ("shell-breaking", "破壳"),
    ("local shell-breaking mechanics", "局部破壳力学"),
    ("shell-breaking mechanics", "破壳力学"),
    ("shell-breaking mechanics", "破壳力学"),
    ("shell mechanics", "蛋壳力学"),
    ("local stress response", "局部应力响应"),
    ("mammillary-interface mechanism", "乳突层界面机制"),
    ("hotspots", "热点"),
    ("hotspot", "热点"),
    ("Ca²⁺relevant", "Ca²⁺相关"),
    ("Ca²⁺responsive", "Ca²⁺响应性"),
    ("Ca²⁺", "Ca²⁺"),
    ("Ca2+", "Ca²⁺"),
]

POST_TRANSLATION_REPLACEMENTS: tuple[tuple[str, str], ...] = (
    ("预告片", "一句话亮点"),
    ("简介", "引言"),
    ("材料和方法", "材料与方法"),
    ("可访问性", "可及性"),
    ("可达性", "可及性"),
    ("蛋白s", "蛋白"),
    ("蛋壳s", "蛋壳"),
    ("糖链s", "糖链"),
    ("基质蛋白s", "基质蛋白"),
    ("糖肽s", "糖肽"),
    ("proteiN", "蛋白"),
    ("阿纳斯", "Anas"),
    ("天鸽", "Columba"),
    ("原鸡", "Gallus"),
    ("外壳", "蛋壳"),
    ("本地机制", "局部力学"),
    ("本地力学", "局部力学"),
    ("可访问性ible", "可及性"),
    ("卵齿轴承尖端", "携带卵齿的喙尖"),
    ("乳突层旋钮密度", "乳突密度"),
    ("乳突层物种密度", "乳突密度"),
    ("单位体积比", "晶体单元体积比"),
    ("糖蛋白质组", "糖蛋白组"),
    ("Intact-糖肽", "完整糖肽"),
    ("完整的糖肽", "完整糖肽"),
    ("完整的糖肽质谱仪", "完整糖肽质谱"),
    ("蛋白质型协同进化图", "蛋白质层面的协同变化图"),
    ("蛋白质型协同进化", "蛋白质层面的协同变化"),
    ("加入物", "登录条目"),
    ("种质", "登录条目"),
    ("整体模型", "系综模型"),
    ("整体建模", "系综建模"),
    ("重建显性糖基化", "重建主要糖基化"),
    ("竞争利益", "利益冲突"),
    ("无竞争利益", "无利益冲突"),
    ("鸽鸡蛋", "鸽蛋"),
    ("常教授", "Chang"),
    ("逼真的鸟类图像", "写实鸟类图片"),
    ("商业蛋鸡鸭农场", "山东某蛋鸭养殖场"),
    ("茶花鸡", "Chahua Chicken"),
    ("绍兴鸭", "Shaoxing Duck"),
    ("白王鸽", "White King Pigeon"),
    ("1.64 × 10^3", "1.64 × 10⁻¹³"),
    ("6.64 × 10⁻1⁰", "6.64 × 10⁻¹⁰"),
    ("平均E值<= 1 ×时保留每个查询的最佳命中10-5，平均序列同一性>=0.40；", "平均 E 值 <= 1 × 10⁻⁵ 且平均序列同一性 >= 0.40 时，保留每个查询的最佳命中；"),
    ("平均E值<= 1 ×时保留每个查询的最佳命中10-5，平均序列同一性>= 0.40；", "平均 E 值 <= 1 × 10⁻⁵ 且平均序列同一性 >= 0.40 时，保留每个查询的最佳命中；"),
)

FIXED_CN_PARAGRAPHS: dict[int, str] = {
    0: "OVAL糖链状态将蛋壳基质化学与壳结构和鸟类破壳力学联系起来",
    1: "OVAL糖链连接壳结构与力学",
    2: "Lin Xuan1, Yaqi Li1, Jiajie Yang1, Yu Liu1, Chengyu Zhang1, Qiulian Wang1, Lingsen Zeng2,3, Tongyao Li1, Wenbin Zhou1, Guiyun Xu1, and Jiangxia Zheng1*",
    3: "1 National Engineering Laboratory for Animal Breeding and MOA Key Laboratory of Animal Genetics and Breeding, College of Animal Science and Technology, China Agricultural University, No. 2 Yuanmingyuan West Road, Haidian District, Beijing 100193, China.",
    4: "*通信作者。邮箱：jxzheng@cau.edu.cn",
    5: "摘要",
    7: "一句话亮点",
    8: "OVAL糖链状态揭示蛋壳基质化学如何塑造破壳时使用的局部蛋壳物理结构。",
    9: "引言",
    10: "鸟类蛋壳是在一个苛刻的悖论中形成的：它必须快速矿化、保护胚胎，同时又必须能够响应来自内部的局部化作用力。这种平衡发生在已知最快的脊椎动物CaCO₃生物矿化过程之一中，该过程在很短的子宫时间窗口内组装出具有力学功能的方解石壳。这个过程并不是简单的碳酸钙沉淀。它依赖有机基质蛋白系统来协调钙供应、成核、晶体生长和壳结构。这一结构的功能检验出现在破壳过程中，此时作用力局部施加，而不是分布在整枚蛋壳上。尽管鸟类喙形状差异很大，卵齿仍然是保守的工具，将这种由内向外的载荷集中到内壳的一个小区域。因此，核心问题是：快速矿化过程中形成的基质化学，如何转化为由内向外破壳时使用的局部蛋壳物理结构。",
    11: "乳突层使这个问题变得可以实验检验。它是在壳膜上建立的第一个钙化层，在这里，基质引导的成核设定了后续蛋壳结构单元所继承的间距和连续性。它也是破壳过程中面对卵齿的内侧结构层，从而把早期矿化和后期力学功能放在同一条材料轴上。具有相似卵齿破壳几何结构的物种，仍然可以在不同生态和发育背景下形成不同的蛋壳状态。因此，乳突层为理解快速矿化的蛋壳如何储存物种特异性力学信息提供了自然读出。",
    12: "蛋壳基质蛋白处于调控这种从化学到蛋壳物理结构转换的关键位置。它们引导乳突层矿化、晶体生长和成熟壳结构，而蛋白质组学研究已经识别出一个广泛共享的壳构建工具包。仍然远不清楚的是，共享基质蛋白的化学状态如何指令蛋壳最初的物理结构。这个缺口在翻译后修饰层面最为突出。磷酸化和糖基化位点正在被不断编目，但大多数研究仍然描述的是景观，而不是机制。分层研究已经显示，同一蛋白可以在不同蛋壳层中处于不同修饰状态，但从这些状态到乳突层形成的路径仍然很大程度上未被解析。因此，鸟类蛋壳比较仍很少把共享基质蛋白上的匹配糖链状态与跨物种蛋壳物理结构联系起来。",
    13: "糖链为从共享蛋白身份到不同材料行为提供了一条合理的化学路径。不同于磷酸基团，糖链在组成、大小、分支和电荷上差异很大，并且能够重塑蛋白质表面的可及性。因此，糖链类别可能是一种机制变量，而不只是位点占据状态的描述。OVAL提供了一个可操作的测试对象，因为它丰度高、与矿化相关，并且与早期蛋壳形成中的Ca²⁺响应性构象行为有关。当前缺失的一步，是在共享基质背景下，从OVAL糖链类别到Ca²⁺可及性表面呈现之间建立直接桥梁。如果这座桥梁存在，它就可能连接糖链调控的OVAL解折叠、成核位点暴露、乳突层形成和局部破壳力学。",
    14: "在本文中，我们比较了鸡、鸭和鸽。我们将比较锚定在保守的卵齿破壳界面上，并沿一条单一轴整合形态学、基质蛋白质组学、完整糖肽、糖型解析结构建模、静电分析和有限元力学。这个设计询问主要OVAL糖链状态是否能够将Ca²⁺表面可及性、乳突层组织和局部破壳响应连接起来。由此，本研究检验了快速矿化的蛋壳是否可以通过共享基质蛋白上的化学特异状态，编码后续的破壳力学。所得链条将OVAL糖链状态与局部蛋壳物理结构以及破壳时使用的局部力学响应联系起来。",
    15: "结果",
    16: "保守的破壳界面限制壳结构变异",
    26: "OVAL糖基化给出最清晰的跨物种分子对比",
    33: "OVAL糖链状态重塑表面可及性",
    41: "由内向外加载解析局部破壳力学",
    49: "讨论",
    50: "本数据集中的跨物种分化最清晰地体现在乳突层，而不是基质蛋白的整体周转。在这一共享的蛋壳基质工具包中，OVAL糖链状态提供了最直接的分子轴，连接Ca²⁺相关表面可及性、乳突层组织和破壳力学。因此，讨论部分将每个结果视为从基质化学到蛋壳物理结构再到破壳力学链条中的一个步骤。",
    51: "三物种设计检验的是同一条基质到力学轴是否能够跨越生态和发育梯度。蛋壳性状沿连续轴变化，而不是由一个二元标签决定。筑巢环境从陆生延伸到水生条件，后代发育状态则从更早成到更晚成连续变化。任何一个维度都不能用简单的是或否划分来概括。鸭在这一设计中特别有信息量。它保留了广义早成发育状态，并具有本比较中最大的壳厚度。然而，它转向中间型OVAL糖链状态和可及性轮廓，而其τ_max结果与鸽而不是鸡收敛。因此，鸭将壳厚度缓冲与本文强调的Ca²⁺可及性-乳突层-力学轴区分开来。由此，该设计不是只跨物种标签检验假说，而是用一个厚壳反例检验假说。",
    52: "乳突层矿化仍然是这一解释中的核心结构层级，因为这是基质化学最早能够塑造蛋壳物理结构的位置。在这一模型中，糖链依赖的酸性OVAL表面暴露可能在最早期基质-矿物相互作用中影响Ca²⁺可及性和OVAL解折叠。这些变化可能进一步改变成核位点呈现、乳突层形成、成熟乳突密度和局部破壳力学。一旦早期方解石单元建立，后续蛋壳区域会继承这一最初矿化窗口设定的间距逻辑。因此，致密的乳突层场可以影响基质保留、矿物连续性、局部应力再分配和成熟形态。这一重点与既往蛋壳研究一致，即乳突层位于晶体成核和基质控制的交汇处。当前比较进一步扩展了这一观点，将乳突层与跨物种糖链状态读出联系起来，而不只是与蛋壳质量描述符联系起来。近年的家禽组学研究越来越多地将年龄、壳腺转录、细胞外囊泡货物和整壳质量性状与蛋壳表型联系起来。然而，这些描述符仍然比本文分离出的近端材料层级更宽泛。因此，乳突层组织并不只是另一项蛋壳性状。它是基质化学可能影响后续力学结果的最早蛋壳物理环境。这一结构位置使乳突层组织成为本文提出的糖链依赖基质机制的第一个物理结构读出。",
    53: "在本文考察的分子层中，OVAL N-糖链结构最紧密地跟踪了跨物种结构对比。直系同源群周转、基因家族变化和糖蛋白网络分化仍然相关，但它们主要定义比较背景。OVAL糖链状态更容易被直接解释，因为它在物种间共享、化学上可解析，并且位于一个已经与矿化相关的高丰度基质蛋白上。既往研究已经将OVAL置于高丰度蛋壳糖蛋白和矿化候选因子之中，也显示蛋壳基质蛋白可以处于不同的N-糖基化状态。本文的进展在于以直系同源物为基础比较哪些糖链状态与表型一致，以及这些分配如何进入蛋壳物理结构和力学解释。早期鸡研究还为OVAL建立了糖基化位点基础，并在OC116中鉴定了糖基化Asn。当前数据集解析了对应OVAL直系同源序列位点上用于结构建模的主要糖链类别（鸡N293；鸭和鸽N97），从而将位点检测拓展到跨物种糖链类别解释。因此，OVAL并非因为唯一而有用，而是因为它提供了最清晰的分子读出，用于跨物种检验本文提出的链条。",
    54: "非OVAL信号仍然重要，因为它们定义了共享基质背景，使OVAL轴更加清晰。OC116和TRFE是有信息量的共享蛋白，而OC17似乎只在鸡中糖基化，并且可能反映更受谱系限制的矿化程序。这一模式与既往赋予OC17、OC116和卵转铁蛋白相关基质组分功能意义的研究一致，也保留了早期糖蛋白组学研究的价值，因为这些研究建立了可实验获取的位点层面清单。当前比较在这一基础上进一步加入跨物种层面：共享直系同源物、主要糖链类别，以及它们对表面呈现的影响。因此，基质工具包仍然是多组分的，而OVAL仍然提供了本文所检验的化学-蛋壳物理结构-力学链条中最清晰的入口。",
    55: "Re-Glyco和APBS分析为这一论证提供了结构桥梁。跨物种比较解析出一个糖链状态梯度。紧凑的鸡糖链保留最可及的酸性表面，中性复合型/杂合型鸭糖链施加中间约束，而延展的唾液酸化鸽糖链产生最强的空间和静电屏蔽。既往体外和结构研究提示，OVAL构象和静电性质在矿化中重要。然而，匹配的糖型解析表面系综此前尚未在鸟类物种间比较。因此，糖链状态变化在这里被解析为物理上可解释的表面差异。尽管这一结果尚不能证明直接因果关系，但它支持本文假说中从糖链类别转化为面向矿化界面的Ca²⁺可及性状态这一步。",
    56: "力学比较旨在由内向外的破壳加载下检验本文提出链条的终点，而不是检验传统的外部压缩或整壳断裂。这一区分很重要，因为蛋壳厚度会提高绝对失效载荷，而τ_max受厚度混杂较少，更直接报告通过乳突层界面的应力传递。Sun等同样显示，整枚蛋的蛋壳厚度存在差异。钝端周围环带局部最薄，这进一步强调需要区分全局壳厚缓冲和局部破壳界面力学。因此，有限元分析询问的是，内侧乳突层界面是否保留了从基质化学和形态学推断出的同一对比。鸭是这一解释的关键对照。其更厚的壳提高了F_max，但没有重现鸡中观察到的高τ_max状态，从而将厚度缓冲与本文强调的材料通路区分开来。蛋壳厚度、体型、繁殖生态和谱系历史仍然定义背景设计空间。然而，这些因素不能解释为什么相同排序会在糖链类别、Ca²⁺表面可及性、乳突层组织和由内向外加载下的τ_max中反复出现。鸭和鸽从相反方向限定了这一力学轴：鸭具有更大的壳厚度和中间型OVAL可及性，而鸽具有更薄的壳和更延展的OVAL糖链，但二者均收敛于较低τ_max。这一对比使鸡成为连接糖链依赖基质行为与一种局部蛋壳状态的有用参照，该状态在破壳界面集中力学响应。因此，有限元结果检验了本文假说的最后一环：糖链相关蛋壳物理结构是否跟踪破壳时使用的局部响应。",
    57: "第二个解释问题是乳突层在孵化后期和破壳过程中可能发生部分吸收。这一可能性并不消除当前比较的相关性，因为本文量化的描述符是乳突密度和晶体单元组织。即使部分最内侧材料被吸收，这些特征仍然嵌入壳体中。同样的考虑也指导了力学读数。我们强调第二个特征峰而不是第一个，因为最早的力偏移更强地受初始形态相关接触支配。后续峰更准确地反映了穿过整个壳壁的应力传递。因此，这一限制限定了解释范围，同时保留了本文机制所需的结构联系。",
    58: "总体来看，该比较在本数据集中收敛到一种与更强破壳力学相关的局部蛋壳状态。鸡结合了最致密的乳突层场、屏蔽最少的Ca²⁺相关OVAL表面以及由内向外加载下最强的局部应力响应。鸭说明这一链条不能被简化为壳厚度。尽管其壳比鸽更厚且F_max更高，但τ_max与鸽而不是鸡分组。这一模式构成了本文假说的功能综合：重复使用的基质蛋白上的化学特异状态，可以比单纯的蛋白质组周转更直接地对应矿化表型。",
    59: "相同的分析序列可能扩展到鸟类蛋壳之外，但其在本文中的主要价值是展示一个基质化学状态如何被跨尺度追踪。许多生物矿化系统依赖有机基质，通过化学特异界面调节离子进入、表面暴露和矿物成核。因此，本文流程提供了一个简洁模板：糖蛋白组状态、表面呈现和介观功能。类似逻辑也可能用于其他矿化组织、仿生材料和蛋壳来源再生材料。然而，这种更广泛意义仍然以本文解析的蛋壳机制为基础。",
    60: "当前研究范围仍然有限。我们分析的是主要糖型，而不是完整的体内糖链系综。我们还在平均蛋壳尺度上将每个物种视为力学均一，并在APBS框架中使用尚未完全约束的子宫离子条件。直接的湿实验后续验证面临一个实践限制：当前工具体系仍然难以同时实现物种匹配的直系同源基质蛋白表达，并在特定糖基化位点精准安装预定义糖链结构。同时，生物矿化作为一个大型耦合反应系统，会限制从任何单一实验中进行直接的单变量解释。因此，下一步决定性测试应包括定义糖型的矿化实验、鸡中OVAL糖基化的直接操控，以及同一由内向外力学对比的位点解析验证，并结合能够提高因果解析度的方法发展。这些实验将检验本文假说的因果版本，即区分主动的OVAL糖链机制和高乳突密度状态标记。",
    61: "总之，本研究将乳突层组织、糖蛋白状态、Ca²⁺表面可及性、蛋壳物理结构和局部破壳力学跨三个鸟类蛋壳联系起来。鸡定义了该轴的高乳突密度端，具有密集的乳突层组织、紧凑的OVAL糖链、更高的Ca²⁺相关表面暴露以及乳突层界面处最强的局部响应。鸭处于关键中间位置。其较厚的壳提高了绝对力，却没有重现相同的局部应力状态，从而将壳厚度与乳突层界面机制区分开来。随着可比糖型分配数据的增加，同一框架可以扩展到其他高丰度蛋壳基质蛋白。在形态测量、糖蛋白组学、结构和力学层面，OVAL糖链状态仍然是本文恢复的基质化学到蛋壳物理结构再到破壳力学链条中最稳定对应的分子特征。",
    62: "材料与方法",
    63: "生物材料",
    65: "蛋壳基质蛋白提取",
    67: "显微CT成像和乳突层形态测量",
    69: "蛋壳基质蛋白的鸟枪法蛋白质组学",
    72: "完整糖肽质谱",
    74: "比较蛋白质组分析和基因家族进化",
    76: "跨物种糖蛋白直系同源物鉴定",
    78: "整合的蛋白-糖链丰度比较",
    81: "N-糖链结构系综建模",
    83: "静电势计算",
    85: "有限元分析",
    87: "统计分析",
    89: "致谢",
    90: "感谢中国农业大学动物医学院 J. Chang 提供鸽蛋材料。感谢 X. Ye 绘制鸟类图标和写实鸟类图片。感谢 B. Tan、Z. Huang 和 X. Li 对稿件提出宝贵意见。",
    91: "资助：本研究得到国家重点研发计划（2022YFD1300100）、中国农业研究系统（CARS-40）和国家重点研发计划（2021YFD1200803）的支持。",
    92: "作者贡献：conceptualization: L.X., L.Z., G.X., and J.Z. methodology: L.X., L.Z., X.S., and J.Z. validation: L.X., Yaqi L., J.Y., Yu L., C.Z., T.L., W.Z., and J.Z. formal analysis: L.X. investigation: L.X., Yaqi L., Yu L., C.Z., and L.Z. data curation: L.X., Yaqi L., J.Y., Yu L., and C.Z. software: L.X. and J.Y. resources: G.X. and J.Z. writing—original draft: L.X. writing—review and editing: L.X., Yaqi L., J.Y., Q.W., T.L., G.X., and J.Z. visualization: J.Y., Q.W., and J.Z. supervision: G.X. and J.Z. funding acquisition: G.X. and J.Z.",
    93: "利益冲突：作者声明无利益冲突。",
    94: "数据、代码和材料可用性：评估和复现本文结果所需的所有数据和代码均包含在正文和/或补充材料中。本研究生成的材料可向 J.Z. 合理索取（jxzheng@cau.edu.cn）。",
}

FIXED_CN_PARAGRAPHS = {
    (index if index < 4 else index + 2): text
    for index, text in FIXED_CN_PARAGRAPHS.items()
}
FIXED_CN_PARAGRAPHS.update({
    4: "2 Animal Breeding and Genomics, Wageningen University & Research, 6708 PB, Wageningen, The Netherlands.",
    5: "3 State Key Laboratory of Genome and Multi-omics Technologies, Shenzhen Branch, Guangdong Laboratory of Lingnan Modern Agriculture, Key Laboratory of Livestock and Poultry Multi-omics of MARA, Agricultural Genomics Institute at Shenzhen, Chinese Academy of Agricultural Sciences, Shenzhen, 518124, China.",
})

ABBREVIATIONS: tuple[str, ...] = (
    "Fig.",
    "Figs.",
    "Table.",
    "Eq.",
    "Dr.",
    "Prof.",
    "G.",
    "A.",
    "C.",
    "s.d.",
    "e.g.",
    "i.e.",
    "vs.",
 )


def _has_english_letters(text: str) -> bool:
    return bool(re.search(r"[A-Za-z]", text))


def _should_translate(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if not _has_english_letters(t):
        return False
    if re.fullmatch(r"[\d\s.,;:()\[\]{}\-+/×=]+", t):
        return False
    return True


def _translate_text(text: str, retries: int = 3) -> str:
    if text in _cache:
        return _cache[text]

    last_err = None
    for i in range(retries):
        try:
            out = translator.translate(text)
            if not out:
                out = text
            _cache[text] = out
            _save_cache()
            return out
        except Exception as e:  # noqa: BLE001
            last_err = e
            time.sleep(0.8 * (i + 1))

    # Fallback: keep original if remote translation fails.
    print(f"[WARN] translation failed, keep original: {text[:80]} :: {last_err}")
    _cache[text] = text
    _save_cache()
    return text


def _save_cache() -> None:
    CACHE_FILE.write_text(
        json.dumps(_cache, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _term_token(index: int) -> str:
    return f"ZXQTERM{index:03d}X"


def _lock_terms_for_translation(text: str) -> tuple[str, dict[str, str]]:
    protected = text
    token_to_cn: dict[str, str] = {}
    ordered_terms = sorted(TERM_LOCKS, key=lambda item: len(item[0]), reverse=True)
    for index, (src_term, cn_term) in enumerate(ordered_terms):
        token = _term_token(index)
        token_to_cn[token] = cn_term
        protected = re.sub(re.escape(src_term), token, protected, flags=re.IGNORECASE)
    return protected, token_to_cn


def _restore_locked_terms(text: str, token_to_cn: dict[str, str]) -> str:
    out = text
    for token, cn_term in token_to_cn.items():
        out = re.sub(re.escape(token), cn_term, out, flags=re.IGNORECASE)
    return out


def _normalize_locked_terms(text: str) -> str:
    out = text
    for src_term, cn_term in TERM_LOCKS:
        out = re.sub(re.escape(src_term), cn_term, out, flags=re.IGNORECASE)
    return out


def _group_translation_texts(texts: list[str], max_count: int, max_chars: int = 4500) -> list[list[str]]:
    groups: list[list[str]] = []
    current: list[str] = []
    current_chars = 0
    for text in texts:
        next_chars = current_chars + len(text) + 20
        if current and (len(current) >= max_count or next_chars > max_chars):
            groups.append(current)
            current = []
            current_chars = 0
        current.append(text)
        current_chars += len(text) + 20
    if current:
        groups.append(current)
    return groups


def _translate_joined_sentences(texts: list[str]) -> list[str]:
    delimiter = "ZXQSENTSEP000X"
    joined = f"\n{delimiter}\n".join(texts)
    translated = translator.translate(joined)
    if not translated:
        raise RuntimeError("empty translation")
    parts = [part.strip() for part in translated.split(delimiter)]
    if len(parts) != len(texts):
        raise RuntimeError("joined sentence split mismatch")
    return parts


def _translate_batch(texts: list[str], batch_size: int = 8) -> list[str]:
    out: list[str] = []
    for i in range(0, len(texts), batch_size):
        chunk = texts[i : i + batch_size]
        cached = [_cache.get(text) for text in chunk]
        if all(item is not None for item in cached):
            out.extend(item or source for item, source in zip(cached, chunk))
            print(f"[cn] translated/cached sentences: {min(i + batch_size, len(texts))}/{len(texts)}", flush=True)
            continue

        missing_positions = [index for index, item in enumerate(cached) if item is None]
        missing_texts = [chunk[index] for index in missing_positions]
        for group in _group_translation_texts(missing_texts, max_count=batch_size):
            try:
                translated = _translate_joined_sentences(group)
                for source, target in zip(group, translated):
                    _cache[source] = target or source
            except Exception:
                for source in group:
                    _translate_text(source)

        _save_cache()
        out.extend(_cache.get(text, text) for text in chunk)
        print(f"[cn] translated/cached sentences: {min(i + batch_size, len(texts))}/{len(texts)}", flush=True)
    return out


def _protect_abbreviations(text: str) -> tuple[str, dict[str, str]]:
    protected = text
    token_to_abbr: dict[str, str] = {}
    for index, abbr in enumerate(ABBREVIATIONS):
        token = f"ZXQABBR{index:03d}X"
        token_to_abbr[token] = abbr
        protected = protected.replace(abbr, token)
    return protected, token_to_abbr


def _restore_abbreviations(text: str, token_to_abbr: dict[str, str]) -> str:
    restored = text
    for token, abbr in token_to_abbr.items():
        restored = restored.replace(token, abbr)
    return restored


def _split_sentences(text: str) -> list[str]:
    s = text.strip()
    if not s:
        return []

    protected, token_to_abbr = _protect_abbreviations(s)
    raw_sentences = re.split(r"(?<=[.!?])\s+", protected)
    sentences = []
    for sentence in raw_sentences:
        restored = _restore_abbreviations(sentence.strip(), token_to_abbr)
        if restored:
            sentences.append(restored)
    return sentences or [s]


def _translate_long_text(text: str) -> str:
    locked_text, token_to_cn = _lock_terms_for_translation(text)
    sentences = _split_sentences(locked_text)
    translated_chunks = _translate_batch(sentences, batch_size=20)
    aligned_sentences = []
    for source, translated in zip(sentences, translated_chunks):
        out = translated.strip() if translated and translated.strip() else source
        aligned_sentences.append(out)
    joined = " ".join(aligned_sentences)
    normalized = _normalize_locked_terms(joined)
    return _restore_locked_terms(normalized, token_to_cn)


def _restore_translated_sentences(
    source_sentences: list[str],
    translated_sentences: list[str],
    token_to_cn: dict[str, str],
) -> str:
    aligned_sentences = []
    for source, translated in zip(source_sentences, translated_sentences):
        out = translated.strip() if translated and translated.strip() else source
        aligned_sentences.append(out)
    joined = " ".join(aligned_sentences)
    normalized = _normalize_locked_terms(joined)
    return _restore_locked_terms(normalized, token_to_cn)


def _replace_paragraph_text_keep_format(p, new_text: str) -> None:
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def _delete_paragraph(p) -> None:
    element = p._element
    element.getparent().remove(element)
    p._p = p._element = None


def _remove_reference_section(doc: Document) -> int:
    if INCLUDE_REFERENCES:
        return 0

    paragraphs = list(doc.paragraphs)
    start = None
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "References":
            start = index
            break

    if start is None:
        return 0

    back_matter_headings = {
        "Acknowledgments",
        "Funding",
        "Author contributions",
        "Competing interests",
        "Data, code, and materials availability",
    }
    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if paragraphs[index].text.strip() in back_matter_headings:
            end = index
            break

    removed = 0
    for paragraph in reversed(paragraphs[start:end]):
        _delete_paragraph(paragraph)
        removed += 1
    return removed


def _remove_image_paragraphs(doc: Document) -> int:
    if INCLUDE_IMAGES:
        return 0

    removed = 0
    for paragraph in list(doc.paragraphs):
        if not (paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:pict")):
            continue
        if paragraph.text.strip():
            for drawing in list(paragraph._p.xpath(".//w:drawing")):
                drawing.getparent().remove(drawing)
            for pict in list(paragraph._p.xpath(".//w:pict")):
                pict.getparent().remove(pict)
        else:
            _delete_paragraph(paragraph)
        removed += 1
    return removed


def _set_cn_fonts(doc: Document) -> None:
    # Ensure East Asian fonts render consistently in translated output.
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Times New Roman"
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "SimSun")


def _postprocess_translated_doc(doc: Document) -> None:
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if not text:
            continue
        for old, new in POST_TRANSLATION_REPLACEMENTS:
            text = text.replace(old, new)
        text = re.sub(
            r"平均\s*E\s*值\s*<=\s*1\s*×\s*时保留每个查询的最佳命中10-5，平均序列同一性\s*>=\s*0\.40；",
            "平均 E 值 <= 1 × 10⁻⁵ 且平均序列同一性 >= 0.40 时，保留每个查询的最佳命中；",
            text,
        )
        text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
        _replace_paragraph_text_keep_format(paragraph, text)

    for index, text in FIXED_CN_PARAGRAPHS.items():
        if index < len(doc.paragraphs):
            _replace_paragraph_text_keep_format(doc.paragraphs[index], text)

    calibrated_paragraphs = {
        46: (
            "不同物种之间的峰值 F_max 存在显著差异（p = 1.64 × 10⁻¹³）。"
            "鸡达到 1.12 ± 0.11 N，鸭达到 0.90 ± 0.09 N，鸽达到 0.49 ± 0.04 N，"
            "且所有成对差异均通过 Tukey HSD 检验达到显著（图 5B）。"
            "相比之下，τ_max 呈现两级模式（p = 6.64 × 10⁻¹⁰）。"
            "鸡达到 0.613 ± 0.061 MPa，显著高于鸭的 0.413 ± 0.041 MPa 和鸽的 0.406 ± 0.033 MPa。"
            "后两个物种之间无显著差异（Tukey HSD，p = 0.957；图 5C）。"
        ),
        47: (
            "F_max 和 τ_max 之间的差异阐明了鸭的结果。其较高的原始接触力主要由更大的壳厚度驱动"
            "（鸭 0.35 mm，鸽 0.19 mm），而不是由乳突层界面处更强的局部应力响应驱动。"
            "换言之，较厚的鸭蛋壳可以承载更高的总接触力，但当响应归一化到接触应力尺度后，"
            "并未表现出更高的局部抗性。相比之下，鸡的 τ_max 相对于另外两个物种提高了 36-40%，"
            "提示其在破壳相关局部应力响应上具有更强能力，且这一差异不依赖于壳厚度。"
            "这种高低分组，即鸡单独位于高组、鸭和鸽位于低组，与乳突密度得到的分组相一致（图 1D）。"
            "因此，力学结果保留了乳突层组织和 OVAL 可及性已经揭示的对比。"
        ),
        64: (
            "从三个禽类品种中采集受精蛋：7 枚 Chahua Chicken 蛋、7 枚 Shaoxing Duck 蛋和 19 枚 White King Pigeon 蛋。"
            "Gallus gallus 鸡蛋获自中国农业大学家禽资源保护场（中国北京）；Columba livia 种蛋由中国农业大学动物医学院提供；"
            "Anas platyrhynchos 鸭蛋获自山东某蛋鸭养殖场。所有蛋在分析前均在 16°C 种蛋保存条件下保存 7 天。"
            "本研究未使用活体动物。"
        ),
        80: (
            "对于图 3D 至 F，根据直系同源映射后的蛋白丰度差异和糖链丰度差异构建成对糖链-蛋白富集图。"
            "每个登录条目的蛋白丰度定义为非零重复强度的平均值；当存在 Number Comparable 字段时，"
            "排除 Number Comparable < 2 的蛋白。糖链丰度在蛋白水平上定义为分配到该登录条目的所有定量糖基化位点的平均非零位点强度之和，"
            "并使用相同的可比特征过滤。Gallus-vs-Anas 和 Gallus-vs-Columba 比较空间由 blastp outfmt 6 映射构建；"
            "当平均 E 值 <= 1 × 10⁻⁵ 且平均序列同一性 >= 0.40 时，保留每个查询的最佳命中；"
            "当查询和目标具有不同数量的非重叠 HSP 时，则应用最大同一性 >= 0.40 的阈值。"
            "Anas-vs-Columba 平面通过在两个数据集中均通过相同过滤的共享 Gallus 直系同源物进行桥接。"
            "对于每个保留的直系同源对，x 坐标计算为 log2(I_ref) - log2(I_comp)，"
            "y 坐标计算为 log2(G_ref) - log2(G_comp)，其中 I 和 G 分别表示蛋白丰度和糖链丰度。"
            "因此，y = x 对角线表示匹配的蛋白-糖链变化，而向糖链富集一侧的偏移则识别糖链变化超过整体蛋白丰度变化的蛋白。"
        ),
    }
    calibrated_paragraphs = {
        (index if index < 4 else index + 2): text
        for index, text in calibrated_paragraphs.items()
    }
    for index, text in calibrated_paragraphs.items():
        if index < len(doc.paragraphs):
            _replace_paragraph_text_keep_format(doc.paragraphs[index], text)


def _translate_paragraphs(paragraphs) -> int:
    changed = 0
    targets = []
    for p in paragraphs:
        src = "".join(r.text for r in p.runs) if p.runs else p.text
        if _should_translate(src):
            targets.append((p, src))

    if not targets:
        return 0

    prepared = []
    all_sentences = []
    for p, src in targets:
        locked_text, token_to_cn = _lock_terms_for_translation(src)
        sentences = _split_sentences(locked_text)
        start = len(all_sentences)
        all_sentences.extend(sentences)
        prepared.append((p, sentences, token_to_cn, start, len(sentences)))

    translated_sentences = _translate_batch(all_sentences, batch_size=30)

    for p, sentences, token_to_cn, start, count in prepared:
        chunk = translated_sentences[start : start + count]
        dst = _restore_translated_sentences(sentences, chunk, token_to_cn)
        _replace_paragraph_text_keep_format(p, dst)
        changed += 1
    return changed


def main() -> None:
    if not EN_DOC.exists():
        raise FileNotFoundError(f"English manuscript not found: {EN_DOC}")

    doc = Document(str(EN_DOC))
    removed_images = _remove_image_paragraphs(doc)
    removed_references = _remove_reference_section(doc)
    changed = 0

    changed += _translate_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                changed += _translate_paragraphs(cell.paragraphs)

    for sec in doc.sections:
        changed += _translate_paragraphs(sec.header.paragraphs)
        changed += _translate_paragraphs(sec.footer.paragraphs)

    _postprocess_translated_doc(doc)
    _set_cn_fonts(doc)
    doc.save(str(OUT_DOC))
    print(f"[OK] translated paragraphs: {changed}")
    print(f"[OK] removed image paragraphs: {removed_images}")
    print(f"[OK] removed reference paragraphs: {removed_references}")
    print(f"[OK] {OUT_DOC}")


if __name__ == "__main__":
    main()
