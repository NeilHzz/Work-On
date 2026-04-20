"""
Polishes both EN and CN SA manuscripts according to Science Advances guide:
- Title ≤ 135 characters
- Short title ≤ 50 characters
- Abstract ≤ 150 words, single paragraph, no citations
- Teaser ≤ 125 characters
- All other SA Research Article format requirements
"""

import docx
from docx import Document
from docx.oxml.ns import qn
import copy
import re

# ─────────────────────────────────────────────
# POLISHED CONTENT
# ─────────────────────────────────────────────

NEW_TITLE_EN = (
    "N-glycan class on ovalbumin encodes a composite adaptive response to "
    "developmental strategy and calcium ecology in three avian species"
)

NEW_SHORT_TITLE_EN = "Avian OVAL glycan, ecology, and eggshell structure"

NEW_TEASER_EN = (
    "N-glycan class on ovalbumin tunes eggshell nucleation density and "
    "mechanical resistance across the precocial\u2013altricial axis."
)

NEW_ABSTRACT_EN = (
    "The avian eggshell mammillary layer determines mechanical competence, "
    "but how egg-white glycoprotein N-glycan structural class varies with "
    "reproductive ecology is unknown. We integrated micro-CT morphometry, "
    "intact glycopeptide proteomics, Re-Glyco structural ensemble modelling, "
    "and finite-element simulation across three avian orders spanning the "
    "precocial\u2013altricial axis. Ovalbumin carries species-specific "
    "N-glycan classes\u2014High-Mannose in Gallus gallus, Neutral Complex in "
    "Anas platyrhynchos, and Sialylated Complex/Hybrid in Columba livia"
    "\u2014that progressively restrict surface Ca\u00b2\u207a-binding site "
    "accessibility. This glycan gradient parallels decreasing mammillary "
    "nucleation density and eggshell shear resistance across the three "
    "species. Proteome orthogroup analyses and gene-family dynamics "
    "independently confirm lineage-specific molecular profiles consistent "
    "with this model. N-glycan class on ovalbumin therefore encodes a "
    "composite adaptive response to developmental strategy and ecological "
    "calcium availability, providing a mechanistic continuum from molecular "
    "modification to tissue-scale eggshell biomechanics."
)

# Chinese versions
NEW_TITLE_CN = (
    "卵白蛋白N-糖链类别编码三种禽类发育策略与生态钙可获性适应响应的复合信号"
)

NEW_SHORT_TITLE_CN = "禽类卵白蛋白糖链类别与蛋壳生态适应"

NEW_TEASER_CN = (
    "卵白蛋白N-糖链类别通过调控蛋壳乳突成核密度与力学抗性，"
    "编码早成\u2013晚成发育轴上的种间适应策略。"
)

NEW_ABSTRACT_CN = (
    "鸟类蛋壳乳突层决定其机械强度，但卵清蛋白N-糖链结构类别如何随繁殖生态学变化"
    "至今缺乏系统研究。本研究在跨越早成\u2013晚成发育轴的三目鸟类中，整合了"
    "显微CT形态测量、完整糖肽质谱、Re-Glyco结构系综建模及有限元仿真分析。"
    "结果显示，卵白蛋白携带物种特异性N-糖链类别：家鸡（Gallus gallus）"
    "以高甘露糖型为主，绿头鸭（Anas platyrhynchos）为纯中性复合型，"
    "岩鸽（Columba livia）以唾液酸化复合/杂合型为主，"
    "三者依次递进地限制蛋白质表面Ca²⁺结合位点的溶剂可及性。"
    "这一糖链梯度与三物种乳突成核密度及蛋壳剪切抗力的递减趋势相吻合。"
    "蛋壳蛋白质组正交组分析与基因家族动态研究独立确认了与该模型相符的"
    "谱系特异性分子特征。上述发现揭示，卵白蛋白N-糖链类别编码了对发育策略与"
    "生态钙可获性的复合适应响应，为从分子修饰到组织尺度蛋壳生物力学建立了"
    "完整的机制链条。"
)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def replace_paragraph_text(para, new_text):
    """Replace ALL runs in a paragraph with a single run carrying the new text.
    Preserves the formatting of the first run."""
    if not para.runs:
        para.text = new_text
        return
    # Copy font from first run
    first_run = para.runs[0]
    first_run.text = new_text
    # Remove remaining runs
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)


def insert_paragraph_after(ref_para, text, style=None):
    """Insert a new paragraph immediately after ref_para."""
    new_para = copy.deepcopy(ref_para._element)
    ref_para._element.addnext(new_para)
    # Find the inserted paragraph object
    doc = ref_para._element.getroottree().getroot()
    # Overwrite text
    from docx.oxml import OxmlElement
    # Clear all <w:r> elements
    for r in new_para.findall(qn('w:r')):
        new_para.remove(r)
    # Add a fresh run
    r_elem = OxmlElement('w:r')
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    if len(text) > 0 and (text[0] == ' ' or text[-1] == ' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    new_para.append(r_elem)
    return new_para


def word_count(text):
    return len(text.split())


# ─────────────────────────────────────────────
# PROCESS ENGLISH MANUSCRIPT
# ─────────────────────────────────────────────

print("=" * 60)
print("PROCESSING ENGLISH MANUSCRIPT")
print("=" * 60)

doc_en = Document("manuscript_results_sa.docx")
paras_en = doc_en.paragraphs

changes_en = []

for i, para in enumerate(paras_en):
    txt = para.text.strip()

    # ── Title
    if txt.startswith("N-glycan structural class on ovalbumin encodes"):
        old = txt
        replace_paragraph_text(para, NEW_TITLE_EN)
        changes_en.append(f"[{i}] Title: {len(old)}→{len(NEW_TITLE_EN)} chars")

    # ── Short title
    elif txt.startswith("Avian OVAL glycan class, ecology, and eggshell architecture"):
        replace_paragraph_text(para, NEW_SHORT_TITLE_EN)
        changes_en.append(f"[{i}] Short title: {len(txt)}→{len(NEW_SHORT_TITLE_EN)} chars")

    # ── Abstract placeholder → real abstract
    elif txt.startswith("[Abstract:"):
        replace_paragraph_text(para, NEW_ABSTRACT_EN)
        wc = word_count(NEW_ABSTRACT_EN)
        changes_en.append(f"[{i}] Abstract: placeholder → {wc} words")

    # ── Teaser
    elif "Species-specific N-glycan class on egg-white ovalbumin tunes eggshell nucleation density" in txt:
        # The teaser para starts with "Teaser: "
        new_teaser_full = "Teaser: " + NEW_TEASER_EN
        replace_paragraph_text(para, new_teaser_full)
        changes_en.append(f"[{i}] Teaser: {len(txt)}→{len(new_teaser_full)} chars")

print("\nEN Changes applied:")
for c in changes_en:
    print(" ", c)

# Validate
print(f"\nValidation:")
print(f"  Title ({len(NEW_TITLE_EN)} chars, limit 135): {'✓' if len(NEW_TITLE_EN) <= 135 else '✗ OVER LIMIT'}")
print(f"  Short title ({len(NEW_SHORT_TITLE_EN)} chars, limit 50): {'✓' if len(NEW_SHORT_TITLE_EN) <= 50 else '✗ OVER LIMIT'}")
print(f"  Teaser ({len(NEW_TEASER_EN)} chars, limit 125): {'✓' if len(NEW_TEASER_EN) <= 125 else '✗ OVER LIMIT'}")
print(f"  Abstract ({word_count(NEW_ABSTRACT_EN)} words, limit 150): {'✓' if word_count(NEW_ABSTRACT_EN) <= 150 else '✗ OVER LIMIT'}")

doc_en.save("manuscript_results_sa_polished.docx")
print("\nSaved: manuscript_results_sa_polished.docx")


# ─────────────────────────────────────────────
# PROCESS CHINESE MANUSCRIPT
# ─────────────────────────────────────────────

print("\n" + "=" * 60)
print("PROCESSING CHINESE MANUSCRIPT")
print("=" * 60)

doc_cn = Document("manuscript_results_sa_cn.docx")
paras_cn = doc_cn.paragraphs

changes_cn = []
abstract_inserted = False

for i, para in enumerate(paras_cn):
    txt = para.text.strip()

    # ── CN Title
    if i == 0 and txt:
        replace_paragraph_text(para, NEW_TITLE_CN)
        changes_cn.append(f"[{i}] Title updated")

    # ── CN Short title
    elif "Short title:" in txt or txt.startswith("Short title"):
        new_short_cn_full = "Short title: " + NEW_SHORT_TITLE_CN
        replace_paragraph_text(para, new_short_cn_full)
        changes_cn.append(f"[{i}] Short title updated")

    # ── Insert CN Abstract after authors placeholder
    elif ("[作者列表" in txt or "作者" in txt or "待补充" in txt) and not abstract_inserted:
        # Insert abstract as new paragraph right after this one
        insert_paragraph_after(para, NEW_ABSTRACT_CN)
        abstract_inserted = True
        changes_cn.append(f"[{i}] CN Abstract inserted after authors paragraph")

    # ── CN Teaser (remove excessively long teaser)
    elif ("乳突成核密度" in txt and "力学抗性" in txt) or \
         ("铔嬩腑鏍稿寲瀵嗗害" in txt) or \
         (txt.startswith("Teaser:") and len(txt) > 50):
        new_teaser_cn_full = "Teaser: " + NEW_TEASER_CN
        replace_paragraph_text(para, new_teaser_cn_full)
        changes_cn.append(f"[{i}] CN Teaser updated")

print("\nCN Changes applied:")
for c in changes_cn:
    print(" ", c)

doc_cn.save("manuscript_results_sa_cn_polished.docx")
print("\nSaved: manuscript_results_sa_cn_polished.docx")

print("\n" + "=" * 60)
print("DONE. Both polished manuscripts saved.")
print("=" * 60)
