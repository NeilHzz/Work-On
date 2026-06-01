"""
Science Advances supplementary materials generator (English).
Output: supplementary_materials_en.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FIG_BASE = Path(r"D:\system_folder\Desktop\Work On\Supplementary\Figures")
MAIN_FIG_BASE = Path(__file__).resolve().parent.parent / "Figure260421"
PANEL_FIG_BASE = Path(__file__).resolve().parent.parent / "02_可视化" / "Figure" / "PNG"
OUT = str(Path(__file__).with_name("supplementary_materials0602.docx"))

doc = Document()

section = doc.sections[0]
section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(2.54)
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)

FONT = "Times New Roman"


def _set_font(run_props, name=FONT):
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), name)
    run_props.insert(0, r_fonts)


def fmt(run, size=11, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    _set_font(run._r.get_or_add_rPr())


def sp(paragraph, before=0, after=80, line=12):
    props = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line * 20))
    spacing.set(qn("w:lineRule"), "auto")
    props.append(spacing)


def para(text, bold=False, italic=False, size=11, before=0, after=80,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    sp(paragraph, before=before, after=after)
    run = paragraph.add_run(text)
    fmt(run, size=size, bold=bold, italic=italic)
    return paragraph


def mpara(parts, before=0, after=80, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    sp(paragraph, before=before, after=after)
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        fmt(run, bold=bold, italic=italic)
    return paragraph


def section_head(text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(paragraph, before=360, after=120)
    run = paragraph.add_run(text)
    fmt(run, size=12, bold=True)
    return paragraph


def st_head(text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(paragraph, before=240, after=80)
    run = paragraph.add_run(text)
    fmt(run, size=11, bold=True)
    return paragraph


def add_image(image_path, width_cm=15.5):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(paragraph, before=60, after=60, line=12)
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return paragraph


def add_images_row(image_paths, width_cm=7.5):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(paragraph, before=60, after=60, line=12)
    for index, image_path in enumerate(image_paths):
        paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
        if index < len(image_paths) - 1:
            paragraph.add_run("  ")
    return paragraph


def fig_title(label, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(paragraph, before=80, after=40)
    run_label = paragraph.add_run(label + " ")
    fmt(run_label, bold=True)
    run_title = paragraph.add_run(title)
    fmt(run_title, bold=True)
    return paragraph


def fig_caption(parts, before=0, after=240):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(paragraph, before=before, after=after)
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        fmt(run, bold=bold, italic=italic)
    return paragraph


para("Supplementary Materials for", size=11, before=0, after=40,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para(
    "OVAL glycan state tracks mammillary organisation and local hatching resistance across avian eggshells",
    bold=True, size=12, before=0, after=160, align=WD_ALIGN_PARAGRAPH.CENTER,
)
para("",
    size=10, before=0, after=300, align=WD_ALIGN_PARAGRAPH.CENTER)

para("This PDF file includes:", bold=True, size=11, before=0, after=60,
     align=WD_ALIGN_PARAGRAPH.LEFT)
for line in [
    "Supplementary Text 1 to 2",
    "Figs. S1 to S11",
    "Table S1 to S7 (uploaded separately as Excel files)",
]:
    bullet = doc.add_paragraph(style="List Bullet")
    bullet.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(bullet, before=0, after=20)
    run = bullet.add_run(line)
    fmt(run, size=11)

doc.add_page_break()

section_head("Supplementary Text")

st_head("Supplementary Text 1. Species selection and sensitivity analysis.")
mpara([
    ("Principal component analysis of AVONET ecological trait scores (Body Mass, Beak Length, Beak Width, Beak Depth, Tarsus Length, Wing Length, Kipps Distance, Hand-Wing Index, Tail Length; plus Primary Lifestyle, Habitat, and Trophic Niche encoded numerically) placed ", False, False),
    ("Gallus gallus", False, True),
    (", ", False, False),
    ("Anas platyrhynchos", False, True),
    (", and ", False, False),
    ("Columba livia", False, True),
    (" in three distinct, non-overlapping regions of avian ecological space (Fig. S1), corresponding to terrestrial ground-nesting precocial, semi-aquatic precocial, and elevated-nesting altricial life-history strategies, respectively. The three-species set was selected to sample both developmental and ecological contrasts within a common hatching framework rather than to maximize any single phylogenetic or morphological separation.", False, False),
])
mpara([
    ("To test whether this separation depended on the numerical encoding of categorical variables, we performed 500 randomized perturbation iterations in which all encoding weights were independently shifted within +/-30% of their original values. Across all iterations, the variance explained by the first two principal components and the cluster silhouette coefficient remained tightly centred on the unperturbed baseline (Fig. S1). The species assignments therefore remained stable under reasonable alternative encodings.", False, False),
], before=80)

st_head("Supplementary Text 2. Eggshell matrix proteome orthogroup analysis.")
mpara([
    ("Proteomics-identified eggshell matrix proteins were organized with an OrthoFinder-based orthology workflow into 2,620, 2,921, and 3,219 orthogroups in ", False, False),
    ("G. gallus", False, True),
    (", ", False, False),
    ("A. platyrhynchos", False, True),
    (", and ", False, False),
    ("C. livia", False, True),
    (", respectively (Fig. S3). The workflow resolved these proteins into a conserved three-species core of 1,997 orthogroups, pairwise-shared subsets of 180 (Gallus-Anas), 434 (Gallus-Columba), and 716 (Anas-Columba), and lineage-restricted sets of 9, 28, and 72 orthogroups for chicken, duck, and pigeon, respectively. The orthogroup structure therefore supports comparison on a shared matrix background rather than wholesale replacement of the eggshell toolkit.", False, False),
])
mpara([
    ("GO enrichment of pairwise-shared sets highlighted ecological rather than purely phylogenetic stratification (Fig. S5). The ", False, False),
    ("A. platyrhynchos", False, True),
    ("-", False, False),
    ("C. livia", False, True),
    ("-shared set was strongly enriched for calcium-ion binding and metal-ion binding, whereas the ", False, False),
    ("G. gallus", False, True),
    ("-", False, False),
    ("A. platyrhynchos", False, True),
    ("-shared set retained precocial-associated functions such as adaptive immune response and spermatogenesis. Lineage-restricted GO signals sharpened the same contrast, most notably by retaining protein N-linked glycosylation in the chicken-specific set.", False, False),
], before=80)
mpara([
    ("Gene-family expansion and contraction inferred by CAFE5 further supported asymmetric lineage divergence (Figs. S8 and S9): ", False, False),
    ("G. gallus", False, True),
    (" showed net family contraction, ", False, False),
    ("A. platyrhynchos", False, True),
    (" was intermediate, and ", False, False),
    ("C. livia", False, True),
    (" showed net expansion. These proteome-level patterns indicate broad lineage divergence while retaining a conserved shared toolkit.", False, False),
], before=80)

doc.add_page_break()
section_head("Figures")

add_image(FIG_BASE / "SuppFig1_Species_Selection" / "Sensitivity_Analysis_Results.png", width_cm=15.5)
fig_title("Fig. S1.", "Sensitivity analysis of the macroecological species-selection framework.")
fig_caption([
    ("Distribution of variance explained (R^2) and cluster silhouette coefficients from 500 randomized perturbation iterations applied to the AVONET-based principal-component space used to select ", False, False),
    ("Gallus gallus", False, True),
    (", ", False, False),
    ("Anas platyrhynchos", False, True),
    (", and ", False, False),
    ("Columba livia", False, True),
    (" as focal species. Categorical ecological variables were numerically encoded, and each iteration introduced independent random shifts to all encoding weights within +/-30% of the original values. The tight concentration of both metrics around the baseline indicates that species-group assignments are robust to the categorical encoding scheme.", False, False),
])

doc.add_page_break()
add_image(PANEL_FIG_BASE / "Fig1B.png", width_cm=15.5)
fig_title("Fig. S2.", "Order-level avian phylogenetic context and comparative-axis heatmaps for the focal species.")
fig_caption([
    ("Phylogenetic relationship of representative avian taxa together with heatmap tracks for aquatic association (X), developmental mode (Z), and lifestyle-habitat discordance (Y). Colored order labels locate the broader comparative frame used for species selection. The positions of the focal lineages show that the chicken, duck, and pigeon comparison spans functional axes that only partly overlap with phylogeny.", False, False),
])

doc.add_page_break()
add_image(FIG_BASE / "SuppFig2_Venn_Orthogroups" / "Fig_venn_orthogroups.png", width_cm=12.0)
fig_title("Fig. S3.", "Three-species Venn diagram of shared and lineage-restricted eggshell matrix orthogroups.")
fig_caption([
    ("OrthoFinder-based orthogroup analysis resolves the three eggshell matrix proteomes into a large three-species shared core, three pairwise-shared subsets, and three lineage-restricted subsets. Numbers indicate orthogroup counts for each subset. The large shared core indicates that cross-species comparison is built on a common protein repertoire rather than on wholesale protein replacement.", False, False),
])

doc.add_page_break()
add_image(FIG_BASE / "SuppFig3_Phylo_Tree" / "Fig_phylo_tree.png", width_cm=14.0)
fig_title("Fig. S4.", "Maximum-likelihood phylogenetic tree of the three focal species reconstructed from single-copy orthologs.")
fig_caption([
    ("Phylogenetic tree inferred by IQ-TREE from a concatenated alignment of single-copy orthologous protein sequences. Branch lengths reflect substitutions per site. Ultrafast bootstrap support values (1000 replicates) are shown at internal nodes. The topology places Galliformes and Anseriformes as sister clades within Galloanseres and Columbiformes as the more distant outgroup, consistent with published avian phylogenies.", False, False),
])

doc.add_page_break()
add_image(FIG_BASE / "SuppFig4_GO_Enrichment" / "图2.jpg", width_cm=16.0)
fig_title("Fig. S5.", "GO enrichment across species-specific and pairwise eggshell matrix protein sets.")
fig_caption([
    ("Top, GO terms enriched in the three pairwise-shared ortholog sets (GnA, Gallus-Anas; GnC, Gallus-Columba; AnC, Anas-Columba). Bottom, GO terms enriched in the three species-specific ortholog sets (Gallus, Anas, Columba). Colors denote GO category: biological process (BP), cellular component (CC), and molecular function (MF). The combined view highlights both ecological signal in the pairwise-shared sets and lineage-restricted signal in the species-specific sets; notably, the ", False, False),
    ("G. gallus", False, True),
    ("-specific set retained protein N-linked glycosylation among its enriched biological-process terms.", False, False),
])

doc.add_page_break()
add_image(PANEL_FIG_BASE / "Fig4D_G.png", width_cm=15.5)
fig_title("Fig. S6.", "Protein-specific glycosylation profiles of recurrent eggshell matrix proteins.")
fig_caption([
    ("Stacked glycan-class profiles for recurrent eggshell matrix proteins across chicken, duck, and pigeon, including OVAL, OC116, TRFE, and OC17. Bars summarize the relative contribution of detected glycan classes for each protein-species combination, providing the protein-level glycosylation background that motivated the focused OVAL structural analysis in the main figures.", False, False),
])

doc.add_page_break()
add_image(PANEL_FIG_BASE / "FigS7.png", width_cm=15.8)
fig_title("Fig. S7.", "Protein-specific glycosylation profiles and surface electrostatic context for OVAL structural ensembles.")
fig_caption([
    ("Panel A shows the species-level surface potential distribution comparing glycosylated and apo OVAL structural ensembles, summarizing the APBS potential landscape that complements the hotspot and accessibility analyses in the main Re-Glyco figure. Panel B shows the per-structure surface electrostatic map across glycosylated OVAL models and matched apo references, providing the ensemble-level context behind the summarized APBS-potential comparison.", False, False),
])

doc.add_page_break()
add_image(FIG_BASE / "SuppFig5_CAFE5_Gene_Family_Turnover" / "Fig_cafe5_expansion_contraction.png", width_cm=14.0)
fig_title("Fig. S8.", "CAFE5 gene-family expansion and contraction across the three species.")
fig_caption([
    ("Phylogenetic tree annotated with lineage-specific gene-family expansion (red) and contraction (blue) events inferred by CAFE5 using the species divergence time tree. Numbers at nodes indicate the estimated ancestral gene-family size; numbers on branches indicate the net change. Only gene families with a per-family ", False, False),
    ("p", False, True),
    (" < 0.05 (Viterbi ", False, False),
    ("p", False, True),
    (" value) are shown. Lineages differed in the turnover of immune- and defense-related gene families, whereas core eggshell matrix families remained broadly conserved.", False, False),
])

doc.add_page_break()
add_image(PANEL_FIG_BASE / "Fig2H.png", width_cm=16.0)
fig_title("Fig. S9.", "Functional enrichment links gene-family turnover to lineage-biased biological processes.")
fig_caption([
    ("Alluvial summary connecting species, gene-family turnover direction, and enriched Gene Ontology terms inferred from expanded and contracted families. Flow colors distinguish expansion and contraction signals, and terminal blocks summarize the enriched biological-process, cellular-component, and molecular-function terms associated with each lineage. The plot complements the CAFE5 turnover tree by showing which functional categories account for the lineage-biased expansion and contraction patterns.", False, False),
])

doc.add_page_break()
add_image(FIG_BASE / "SuppFig7_Glycosylation_Hotspot" / "Fig_hotspot_ensemble_1.png", width_cm=15.5)
fig_title("Fig. S10.", "Re-Glyco ensemble analysis of OVAL glycan geometry and apo-versus-glycosylated states.")
fig_caption([
    ("(A) Distribution of glycan radius of gyration (R", False, False),
    ("g", False, False),
    (") across conformational ensemble replicates for the three species-specific OVAL-glycan complexes, colored by species (", False, False),
    ("G. gallus", False, True),
    (" orange, ", False, False),
    ("A. platyrhynchos", False, True),
    (" blue, ", False, False),
    ("C. livia", False, True),
    (" green). (B) Glycan end-to-end distance distributions across conformations for the same three complexes. (C) Per-conformation Ca²⁺ hotspot count (N", False, False),
    ("hot", False, False),
    (") comparing the glycosylated and apo OVAL structures for each species. ", False, False),
    ("C. livia", False, True),
    (" showed the largest conformational space and the strongest glycan shielding; ", False, False),
    ("G. gallus", False, True),
    (" showed the smallest conformational space and the weakest shielding; ", False, False),
    ("A. platyrhynchos", False, True),
    (" was intermediate. The apo comparison provides an internal control: once N-glycans were removed, cross-species separation in hotspot count largely collapsed. Panel C contrasts were evaluated against the apo reference by one-sample Wilcoxon signed-rank test when structure-level variation was present.", False, False),
])

doc.add_page_break()
add_images_row([
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "chicken_rcforc_3x3.png",
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "chicken_rcforc_yforce.png",
], width_cm=7.5)
add_images_row([
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "duck_rcforc_3x3.png",
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "duck_rcforc_yforce.png",
], width_cm=7.5)
add_images_row([
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "pigeon_rcforc_3x3.png",
    FIG_BASE / "SuppFig8_FEA_Force_Analysis" / "pigeon_rcforc_yforce.png",
], width_cm=7.5)
fig_title("Fig. S11.", "Per-species finite-element reaction-force time courses across all nine offset positions.")
fig_caption([
    ("(A-C) Contact force (", False, False),
    ("F", False, True),
    (") time-course curves for all nine parametric impact positions (3 × 3 lateral offset grid) for ", False, False),
    ("G. gallus", False, True),
    (" (A), ", False, False),
    ("A. platyrhynchos", False, True),
    (" (B), and ", False, False),
    ("C. livia", False, True),
    (" (C). Each curve represents one simulation; curves are shown from the onset of contact to peak force. Insets show the peak contact force (", False, False),
    ("F", False, True),
    ("_max", False, False),
    (") distribution across the nine positions for each species. (D-F) Corresponding Y-direction reaction-force (", False, False),
    ("F", False, True),
    ("Y", False, False),
    (") time courses. Species means ± s.d. of ", False, False),
    ("F", False, True),
    ("_max", False, False),
    (" and peak shear stress (τ", False, False),
    ("_max", False, False),
    (") computed from these nine replicates per species are reported in the main text and Fig. 5. Simulations were run in LS-DYNA (Ansys) using explicit dynamic finite-element analysis, with eggshell thickness set to the species-specific value measured from micro-CT.", False, False),
])

doc.save(OUT)
print(f"Saved -> {OUT}")
