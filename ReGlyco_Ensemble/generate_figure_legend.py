"""
generate_figure_legend.py
生成四张图的 Nature 风格图注 Word 文档
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = r"E:\Data\Desktop\Work On\ReGlyco_Ensemble\Figure_Legends_Nature_Style.docx"
IMG_DIR  = r"E:\Data\Desktop\Work On\ReGlyco_Ensemble"

FIGURES = {
    "Fig_glycan_ensemble_stats.png": "Figure 1",
    "Fig_hotspot_ensemble_1.png":    "Figure 2",
    "Fig_hotspot_ensemble_2.png":    "Figure 3",
    "Fig_ensemble_calcium.png":      "Figure 4",
}

# ── 文档样式辅助 ────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13 if level==1 else 11, bold=True)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_panel_label(doc, label, body):
    """加粗面板标签 + 正常体说明"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent  = Inches(0.3)
    r1 = p.add_run(label + " ")
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(body)
    set_font(r2, size=11)
    return p

def insert_image(doc, img_name, width_inch=6.0):
    path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inch))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 80)
    set_font(run, size=9, color=(180, 180, 180))

# ═══════════════════════════════════════════════════════════════════════════════
# 正文内容
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document()

# 全局页面边距
from docx.oxml import OxmlElement
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 文档标题 ────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_r = title_p.add_run(
    "N-Glycan Conformational Dynamics and Ca²⁺ Binding Accessibility "
    "across Avian Species: Figure Legends and Results Interpretation"
)
set_font(title_r, size=14, bold=True)
doc.add_paragraph()

intro_p = doc.add_paragraph()
intro_run = intro_p.add_run(
    "The following legends and interpretations describe four multi-panel figures "
    "derived from Re-Glyco ensemble modelling of avian eggshell matrix proteins "
    "(Gallus gallus, Anas platyrhynchos, and Columba livia). Each figure addresses "
    "a distinct dimension of the relationship between N-glycosylation, glycan "
    "conformational space, and electrostatic accessibility of Ca²⁺ binding hotspots. "
    "Statistical comparisons were performed using the two-sided Mann–Whitney U test "
    "(inter-species) or one-sample t-test (glycosylated vs. deglycosylated), with "
    "significance thresholds of p < 0.05 (*), p < 0.01 (**), and p < 0.001 (***)."
)
set_font(intro_run, size=11, italic=True)
intro_p.paragraph_format.space_after = Pt(12)

# ════════════════════════════════════════════════════════════════════════════════
# FIGURE 1
# ════════════════════════════════════════════════════════════════════════════════
add_separator(doc)
add_heading(doc, "Figure 1. Conformational diversity of N-linked glycans across 50 ensemble models.")
insert_image(doc, "Fig_glycan_ensemble_stats.png", width_inch=6.2)

add_body(doc,
    "Violin plots depicting the distribution of four glycan geometric descriptors "
    "computed from a 50-member structural ensemble for each species. Each violin "
    "represents the kernel-density estimate of the full conformational sampling; "
    "the central horizontal bar denotes the median. Individual conformation values "
    "are overlaid as jittered scatter points. Pairwise inter-species significance "
    "is indicated above each panel (Mann–Whitney U test; n = 50 conformations for "
    "Gallus, n = 150 for Anas, n = 700 for Columba)."
)

add_panel_label(doc, "(a) Glycan radius of gyration (Rg, Å).",
    "Rg measures the mass-weighted root-mean-square deviation of glycan atoms from "
    "their centroid, providing a scalar measure of glycan compactness. "
    "Gallus displayed a Rg of 7.50 ± 0.37 Å, while Anas was comparable at "
    "7.61 ± 0.43 Å. Columba exhibited a significantly larger Rg of 9.13 ± 0.92 Å "
    "(p < 0.001 vs. both Gallus and Anas), indicating that the Columba N-glycan "
    "adopts substantially more extended conformations. This may reflect differences "
    "in glycan chain length or branching architecture, as well as divergent protein "
    "surface topographies at the N97 glycosylation site."
)

add_panel_label(doc, "(b) End-to-end distance (Å).",
    "The distance between the first and last heavy atoms of the glycan backbone "
    "quantifies the degree of chain extension. Gallus (18.05 ± 4.05 Å) and "
    "Anas (17.52 ± 1.75 Å) showed similar end-to-end distances, whereas Columba "
    "exhibited a markedly larger value (22.30 ± 4.12 Å; p < 0.001). The broader "
    "distribution in Columba—spanning a range comparable to that of Gallus despite "
    "a far larger ensemble—points to high intrinsic glycan flexibility at the N97 "
    "site in this species. Notably, Anas showed a narrower distribution than "
    "Gallus, suggesting a more restrained conformational landscape for the same "
    "glycan attachment site."
)

add_panel_label(doc, "(c) Glycan centroid–protein distance (Å).",
    "This metric captures the distance from the glycan centre-of-mass to the "
    "nearest protein heavy atom, reflecting how far the glycan projects from the "
    "protein surface. All three species clustered around 30 Å "
    "(Gallus 30.20 ± 1.27 Å; Anas 30.18 ± 1.46 Å; Columba 30.89 ± 2.65 Å), "
    "yet Columba showed a significantly broader distribution (p < 0.001 vs. both). "
    "The higher variance in Columba implies that a substantial fraction of its "
    "glycan conformers extend further from the protein, potentially reducing "
    "steric coverage of adjacent surface hotspots in those conformations."
)

add_panel_label(doc, "(d) Minimum distance to Cα (Å).",
    "The minimum distance between any glycan heavy atom and the closest protein "
    "Cα atom reports on the tightness of glycan–protein contact. All species "
    "maintained close proximity (Gallus 3.58 ± 0.63 Å; Anas 3.25 ± 0.49 Å; "
    "Columba 3.00 ± 0.46 Å), consistent with non-covalent packing of the glycan "
    "against the protein backbone. Columba showed the smallest median distance "
    "(p < 0.001 vs. Gallus), suggesting a tighter glycan–protein interface that "
    "may impose steric restraints on local surface accessibility."
)

add_body(doc,
    "Collectively, Figure 1 demonstrates that glycans across all three species "
    "sample a broadly similar spatial envelope (~30 Å from the protein), yet "
    "differ significantly in compactness and extension. Columba N-glycans are "
    "the most conformationally diverse, while Anas glycans are the most "
    "spatially constrained. These differences are expected to modulate the "
    "efficiency and uniformity of Ca²⁺ binding site shielding."
)

# ════════════════════════════════════════════════════════════════════════════════
# FIGURE 2
# ════════════════════════════════════════════════════════════════════════════════
add_separator(doc)
add_heading(doc, "Figure 2. Distribution and conformational trajectory of Ca²⁺ hotspot counts.")
insert_image(doc, "Fig_hotspot_ensemble_1.png", width_inch=6.2)

add_body(doc,
    "This figure characterises the abundance of Ca²⁺ binding hotspot residues "
    "and the influence of glycan conformation on their shielding across 50 "
    "structural ensemble models. Ca²⁺ hotspots are defined as surface-exposed "
    "acidic residues (Asp and Glu) with solvent-accessible surface area (SASA) "
    "> 1 Å², identified from glycan-aware APBS electrostatic calculations. "
    "A hotspot residue is classified as glycan-shielded when the reduction in "
    "SASA upon glycan inclusion (ΔSASA) exceeds 5 Å²."
)

add_panel_label(doc, "(a) Total Ca²⁺ hotspot count (violin + scatter, n per conformation).",
    "Gallus possessed the highest median total hotspot count (~90 per conformation), "
    "significantly exceeding both Anas (~73; p < 0.001) and Columba (~79; p < 0.001). "
    "Gallus hotspot counts showed near-zero variance across conformations "
    "(std = 0.29), reflecting the single-structure ensemble, whereas Anas "
    "(std = 3.87) and Columba (std = 2.84) displayed conformationally dependent "
    "variability. These results indicate that Gallus presents more acidic surface "
    "residues available for Ca²⁺ coordination, suggesting a higher intrinsic "
    "Ca²⁺ binding capacity in the absence of glycan effects."
)

add_panel_label(doc, "(b) Glycan-shielded hotspot count (ΔSASA > 5 Å²).",
    "Despite its higher total hotspot count, Gallus exhibited the largest mean "
    "number of glycan-shielded hotspots (~24 per conformation), indicating that "
    "its N-glycan at N293 effectively occludes a substantial fraction of acidic "
    "residues. Anas (~16.5 shielded) and Columba (~22.3 shielded) showed "
    "intermediate and comparable shielding, respectively. Mann–Whitney U tests "
    "revealed significant pairwise differences in all combinations (p < 0.001), "
    "establishing that glycan shielding capacity is species-specific and linked "
    "to the structural context of the glycosylation site."
)

add_panel_label(doc, "(c) Hotspot count trajectory across 50 conformations.",
    "Temporal traces of total hotspot count plotted against conformation model "
    "index (1–50). Individual structure trajectories are drawn in species-specific "
    "colours, with line opacity proportional to ensemble size; bold lines represent "
    "species means. Gallus (single structure, 50 models) showed negligible "
    "variation, confirming that hotspot identity is structurally invariant for "
    "this protein. Anas and Columba showed fluctuating trajectories, with "
    "Columba displaying the widest inter-model spread. The absence of a temporal "
    "trend in any species confirms that the 50-member ensemble samples "
    "independent conformations rather than a directed pathway, validating the "
    "use of ensemble statistics."
)

# ════════════════════════════════════════════════════════════════════════════════
# FIGURE 3
# ════════════════════════════════════════════════════════════════════════════════
add_separator(doc)
add_heading(doc, "Figure 3. Multidimensional analysis of Ca²⁺ hotspot accessibility and glycan shielding.")
insert_image(doc, "Fig_hotspot_ensemble_2.png", width_inch=6.0)

add_body(doc,
    "Figure 3 dissects the mechanistic basis of glycan-mediated Ca²⁺ modulation "
    "through five complementary metrics, each presented as violin plots with "
    "overlaid individual data points and pairwise significance brackets. "
    "Net accessible hotspot counts and residue SASA are summarised as stacked "
    "bar charts in panels (e) and (f), respectively."
)

add_panel_label(doc, "(a) Interface shielding by glycan (Å²).",
    "Quantifies the total solvent-accessible surface area of hotspot residues "
    "occluded by the glycan, representing the geometric 'footprint' of glycan "
    "coverage. Columba exhibited the highest shielding (mean 5.31 ± 1.32 Å²; "
    "p < 0.001 vs. Gallus 1.40 ± 0.58 Å²), followed by Anas (3.00 ± 0.72 Å²). "
    "The large variance in Columba reflects ensemble-dependent glycan "
    "positioning, consistent with its conformational diversity documented in "
    "Figure 1. These results demonstrate that although Gallus has more total "
    "hotspots, the per-residue shielding area is substantially smaller than "
    "in the waterfowl species."
)

add_panel_label(doc, "(b) Hotspot residue mean SASA (Å²).",
    "Reports the mean solvent-exposed surface area of hotspot residues in the "
    "glycan-present conformation, after accounting for steric occlusion. "
    "Gallus hotspot residues retained the highest mean SASA (51.52 ± 0.20 Å²), "
    "significantly exceeding Anas (45.83 ± 0.27 Å²; p < 0.001) and "
    "Columba (49.13 ± 0.87 Å²; p < 0.001). This indicates that while Gallus "
    "has the broadest shielding footprint in absolute terms, the individual "
    "hotspot residues in waterfowl are more effectively buried per unit area, "
    "pointing to qualitatively distinct shielding geometries."
)

add_panel_label(doc, "(c) Hotspot fraction (hotspots / total acidic candidates).",
    "The proportion of surface Asp/Glu residues that qualify as active hotspots "
    "provides a normalised measure of Ca²⁺ accessibility. Gallus showed the "
    "highest hotspot fraction (0.841 ± 0.003), followed by Anas (0.780 ± 0.017) "
    "and Columba (0.744 ± 0.009; p < 0.001 across all pairs). The lower fraction "
    "in Columba indicates that a greater proportion of its acidic surface residues "
    "are occluded or in unfavourable electrostatic environments, suggesting more "
    "extensive modulation of Ca²⁺ binding capacity through glycan shielding."
)

add_panel_label(doc, "(d) Net accessible Ca²⁺ hotspots (total − shielded).",
    "Net accessible counts represent hotspots available for Ca²⁺ coordination "
    "after subtracting glycan-shielded residues. Gallus retained the highest "
    "net count (66.0 ± 0.45 per conformation), with Anas (56.8 ± 2.04) and "
    "Columba (56.6 ± 2.04) reaching comparable, significantly lower values "
    "(p < 0.001 vs. Gallus). The convergence of Anas and Columba net counts, "
    "despite their different shielding mechanisms, suggests that the extent of "
    "Ca²⁺ site accessibility may be functionally conserved in non-chicken "
    "species."
)

add_panel_label(doc, "(e) Stacked bar chart: Net accessible vs. glycan-shielded hotspot counts (mean ± 95% CI).",
    "Summary of hotspot partitioning across species. In Gallus, 66.0 residues "
    "remain accessible and 24.0 are shielded—representing a 26.7% occlusion "
    "rate. Anas shows 56.8 accessible and 16.5 shielded (22.5% occlusion). "
    "Columba shows 56.6 accessible and 22.3 shielded (28.3% occlusion). "
    "Gallus net counts are significantly higher than both Anas and Columba "
    "(p < 0.001), reflecting differences in total hotspot abundance rather "
    "than differential glycan efficiency."
)

add_panel_label(doc, "(f) Stacked bar chart: Net accessible vs. glycan-shielded residue SASA (Å²).",
    "This panel translates hotspot counts into surface area units, providing a "
    "physically meaningful measure of glycan coverage. Gallus hotspot residues "
    "present 35.4 Å² of net accessible SASA, with only 1.4 Å² shielded by "
    "the glycan. Anas and Columba show net SASA values of 37.3 Å² and 36.2 Å², "
    "respectively, with glycan-shielded contributions of 3.0 Å² and 5.3 Å². "
    "The disproportionately small shielded SASA in Gallus relative to its "
    "large shielded hotspot count (panel e) implies that the N293 glycan "
    "occludes more hotspot residues by count but with comparatively smaller "
    "per-residue SASA reduction, revealing a quantitatively but not qualitatively "
    "dominant shielding mechanism."
)

# ════════════════════════════════════════════════════════════════════════════════
# FIGURE 4
# ════════════════════════════════════════════════════════════════════════════════
add_separator(doc)
add_heading(doc, "Figure 4. Integrated electrostatic landscape and Ca²⁺ accessibility in glycosylated versus deglycosylated proteins.")
insert_image(doc, "Fig_ensemble_calcium.png", width_inch=6.2)

add_body(doc,
    "Figure 4 integrates structural, electrostatic, and SASA-based analyses into "
    "a unified view of how N-glycosylation modulates the Ca²⁺ binding environment "
    "across three avian species. The figure spans a full-width strip chart (a) "
    "and three bottom panels (b–d) directly comparing glycosylated and "
    "deglycosylated (apo) structural states. Glycosylated structures are "
    "represented in saturated species colours; apo structures are shown in the "
    "same hue at reduced opacity or with hatching. Statistical comparisons "
    "between glycosylated distributions and the corresponding single-point apo "
    "value used one-sample t-tests (*: p < 0.05; **: p < 0.01; ***: p < 0.001)."
)

add_panel_label(doc, "(a) Full-width APBS surface potential strip chart (kT/e).",
    "Each row displays the per-residue electrostatic potential (APBS, kT/e) "
    "mapped onto the protein surface for a single structural model, with "
    "residue sequence position on the x-axis and potential value on the colour "
    "axis (blue: negative/anionic; red: positive/cationic; white: neutral). "
    "Rows are ordered by species and glycosylation state: Gallus (G1 apo → G1 "
    "glyco), Anas (A1 apo → A1–A3 glyco), and Columba (C1 apo → C1–C14 glyco). "
    "Glycosylation site residues (N293 in Gallus, N97 in Anas and Columba) are "
    "marked with vertical dashed lines. Across all species, glycosylated models "
    "display broadly similar surface potential patterns to their apo counterparts, "
    "with localised differences near the glycosylation site. The dense negative "
    "band in Gallus around residues 280–310 is consistent with the Ca²⁺ "
    "binding domain of eggshell matrix proteins and appears attenuated in "
    "certain glycosylated conformations, indicating glycan-mediated modulation "
    "of local electrostatics."
)

add_panel_label(doc, "(b) Ca²⁺ hotspot residue count: glycosylated vs. deglycosylated (bar + scatter).",
    "Per-structure counts of acidic surface residues with APBS potential "
    "< −5 kT/e (N_hotspot), presented as mean bars with individual points. "
    "In Gallus, the glycosylated model (N_hotspot = 15.0) showed a marginally "
    "higher count than the apo form (14.0); the difference was not statistically "
    "significant given n = 1 in each condition. Anas glycosylated structures "
    "(mean 12.7 across three models) trended below the apo value of 14.0, "
    "suggesting partial electrostatic neutralisation by glycan screening. "
    "Columba glycosylated models (mean 10.8 across 14 models) were significantly "
    "lower than the apo value of 13.0 (p < 0.05), providing direct evidence that "
    "N-glycosylation in Columba reduces the count of strongly negative surface "
    "sites available for Ca²⁺ coordination."
)

add_panel_label(doc, "(c) Carboxylate surface SASA (Å²): glycosylated vs. deglycosylated.",
    "Total solvent-accessible surface area of all surface-exposed Asp and Glu "
    "residues, calculated from glycan-aware APBS models. The y-axis is truncated "
    "at 1,000 Å² to emphasise within-species differences. Gallus glycosylated "
    "SASA (3,630 Å²) was marginally lower than apo (3,648 Å²; n = 1, no test). "
    "Anas glycosylated structures (3,369 ± 13 Å²) showed a significant reduction "
    "compared with the apo value (3,404 Å²; p < 0.05), consistent with partial "
    "steric occlusion of carboxylate groups by the glycan. Columba exhibited the "
    "largest SASA reduction: glycosylated structures averaged 3,433 ± 18 Å², "
    "significantly below the apo value of 3,473 Å² (p < 0.01). These data "
    "indicate that N-glycans exert a measurable, species-specific reduction in "
    "carboxylate solvent exposure, with the effect being most pronounced in Columba."
)

add_panel_label(doc, "(d) Surface electrostatic potential distribution: glycosylated vs. deglycosylated (violin).",
    "Violin plots of the full distribution of per-residue APBS median potentials "
    "(kT/e) for all surface residues. Solid, opaque violins represent glycosylated "
    "ensembles; semi-transparent outlined violins represent apo structures. "
    "A reference line at −5 kT/e marks the hotspot threshold. "
    "Gallus showed similar median APBS values for glycosylated (−0.81 kT/e) and "
    "apo (−0.79 kT/e) conditions, indicating minimal ensemble-level electrostatic "
    "perturbation by the glycan. Anas displayed a notable shift from apo "
    "(−0.58 kT/e) to glycosylated (−0.10 kT/e), a positive (less anionic) shift "
    "that, while not statistically significant given the limited n = 3 glyco "
    "ensemble, suggests that glycan addition partially offsets the negative "
    "surface charge density. Columba showed the opposite trend: the glycosylated "
    "ensemble (−1.29 kT/e) was more negative than apo (−1.03 kT/e; p < 0.05), "
    "consistent with the glycan indirectly enhancing electrostatic Ca²⁺ "
    "pre-organisation through conformational reorganisation of acidic residues."
)

add_body(doc,
    "Taken together, Figures 1–4 establish a coherent picture in which N-glycan "
    "conformational diversity and surface topology jointly govern the accessibility "
    "of Ca²⁺ binding hotspots at the protein surface. Gallus relies on a high "
    "absolute density of acidic residues with moderate glycan shielding; Anas "
    "presents compact, constrained glycan conformations that reduce electrostatic "
    "negativity; and Columba harbours the most conformationally plastic glycans, "
    "which achieve the greatest per-residue geometric shielding and, paradoxically, "
    "the most negative residual surface potential. These divergent strategies likely "
    "reflect adaptive fine-tuning of Ca²⁺ sequestration during avian eggshell "
    "biomineralisation, with implications for shell microstructure, mechanical "
    "strength, and calcification kinetics."
)

# ── 保存 ────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Word 文档已保存: {OUT_PATH}")
