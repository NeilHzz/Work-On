import docx
from docx.shared import Inches

doc = docx.Document()

doc.add_heading('Supplementary Note 1: Sensitivity and Robustness Analysis of Heuristic Feature Encoding', level=1)

doc.add_heading('Objective', level=2)
doc.add_paragraph('To evaluate the robustness of our macroevolutionary clustering against potential subjective biases introduced by the heuristic numeric encoding of categorical ecological traits (e.g., mapping "Marine" habitat to 1.00 and "Forest" to 0.05), we performed a comprehensive Monte Carlo sensitivity analysis. This ensures that the primary axes of ecological divergence and the resulting functional clusters are biologically meaningful and not artifacts of absolute step-size assumptions (interval assumption).')

doc.add_heading('Perturbation Model Design', level=2)
doc.add_paragraph('We modeled input uncertainty by introducing stochastic noise to the baseline heuristic scores. For each ecological trait value (Primary Lifestyle, Habitat, and Trophic Niche), we applied a Gaussian noise perturbation:')
doc.add_paragraph('V_new = V_orig * (1 + ε * N(0,1))')
doc.add_paragraph('We set the maximum noise amplitude (ε) to ±10%, simulating a scenario where the expert-assigned weights could deviate by up to 10% from their baseline values. We performed 500 independent Monte Carlo iterations. For each iteration, the completely perturbed dataset was subjected to Principal Component Analysis (PCA) and unsupervised K-means clustering (k=3). We recorded the explained variance of the first principal component (PC1) and the overall Silhouette coefficient of the clusters.')

doc.add_heading('Results and Stability Assessment', level=2)
p1 = doc.add_paragraph()
p1.add_run('1. Stability of the Evolutionary Axis (PC1): ').bold = True
p1.add_run('Across all perturbation iterations, PC1 (representing the terrestrial-to-aquatic ecological gradient) consistently explained the dominant proportion of the multidimensional variance. This confirms that the transition from terrestrial to aquatic environments represents a robust, overriding macroevolutionary trend that is insensitive to minor scoring discrepancies.')

p2 = doc.add_paragraph()
p2.add_run('2. Robustness of Functional Clustering: ').bold = True
p2.add_run('The average Silhouette coefficient remained highly stable (consistently > 0.75) under continuous noise injection. The overall topological boundaries distinguishing the three primary evolutionary groups (terrestrial altricial, terrestrial precocial, and aquatic precocial) did not degrade, demonstrating that the intra-group cohesion and inter-group separation are significantly stronger than the analytical noise.')

doc.add_heading('Conclusion', level=2)
doc.add_paragraph('Our sensitivity analysis quantitatively validates that the deep phylogenetic divergences observed in our multidimensional ecological space are statistically robust. The defined heuristic numeric encoding effectively captures the biological reality without artificially predetermining the clustering outcomes.')

doc.add_heading('Supplementary Figures', level=2)
try:
    doc.add_picture('Intro/Sensitivity_Analysis_Results.png', width=Inches(6.0))
except Exception as e:
    doc.add_paragraph('[Image: Sensitivity_Analysis_Results.png goes here]')

fig_cap = doc.add_paragraph()
fig_cap_run = fig_cap.add_run('Supplementary Figure 1 | Monte Carlo sensitivity analysis of the ecological trait space. ')
fig_cap_run.bold = True
fig_cap.add_run('a, Distribution of the explained variance ratio for Principal Component 1 (PC1) across 500 iterations under 10% stochastic noise. b, Distribution of average Silhouette coefficients for the K-means clustering (k=3) under the same perturbation iterations. Both distributions are highly concentrated, validating the structural stability of the underlying macroevolutionary model.')

doc.save('Intro/Supplementary_Note_Sensitivity_NatCommun.docx')
print("Nature Communications formatted Word document created successfully.")
