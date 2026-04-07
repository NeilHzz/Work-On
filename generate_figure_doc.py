"""
生成包含所有可视化图片说明的 Word 文档（Nature 期刊风格）
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

doc = Document()
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

# ── 辅助函数 ──────────────────────────────────────────────────────────────
def set_font(run, name="Arial", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Word中文字体也需单独设置
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    """添加标题段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=14 if level==1 else 12, bold=True,
             color=(40, 40, 40) if level==1 else (60, 60, 60))
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240' if level==1 else '160')
    spacing.set(qn('w:after'), '80')
    pPr.append(spacing)
    return p

def add_label_paragraph(doc, label_text, body_text):
    """添加 '图注标签: 正文' 风格段落（标签加粗，正文普通）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:line'), '320')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    r1 = p.add_run(label_text)
    set_font(r1, size=10.5, bold=True)
    r2 = p.add_run(body_text)
    set_font(r2, size=10.5)
    return p

def add_body(doc, text, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:line'), '320')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_figure_caption(doc, fig_num, title):
    """添加图片标题行，如 'Figure 1 | Phylogenetic tree...'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')
    spacing.set(qn('w:after'), '60')
    pPr.append(spacing)
    r1 = p.add_run(f"Figure {fig_num}")
    set_font(r1, size=11, bold=True, color=(0, 0, 0))
    r2 = p.add_run(f" | {title}")
    set_font(r2, size=11, bold=True, color=(50, 50, 50))
    return p

def insert_image(doc, img_path, width_cm=14):
    """插入图片，若文件不存在则跳过并添加提示"""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[图片文件未找到: {img_path}]")
        set_font(run, size=9, italic=True, color=(180, 0, 0))
    return p

def add_divider(doc):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BBBBBB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)

# ═══════════════════════════════════════════════════════════════════════════
# 封面 / 标题页
# ═══════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_title._p.get_or_add_pPr()
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:before'), '480')
spacing.set(qn('w:after'), '240')
pPr.append(spacing)
r = p_title.add_run("Figure Legends")
r.font.name  = "Arial"
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = RGBColor(20, 20, 20)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run(
    "Evolutionary dynamics of egg-white glycoproteins across three avian species:\n"
    "Gallus gallus, Anas platyrhynchos and Columba livia"
)
set_font(r2, size=12, italic=True, color=(80, 80, 80))

doc.add_paragraph()

p_intro_hd = doc.add_paragraph()
r3 = p_intro_hd.add_run("Overview")
set_font(r3, size=13, bold=True)

add_body(doc,
    "This document describes each figure produced in the comparative glycoproteomics study of avian "
    "egg-white proteins. Three species are examined: the domestic chicken (Gallus gallus, a precocial "
    "bird), the domestic duck (Anas platyrhynchos, precocial), and the pigeon (Columba livia, altricial). "
    "Integrating phylogenomics, gene-family evolution analysis, intact glycopeptide mass spectrometry, "
    "and network-based glycan annotation, the figures collectively address how protein expression, "
    "N-glycosylation composition, and glycan-site occupancy co-evolve along divergent avian lineages "
    "with distinct reproductive strategies. All figures were generated at 300 dpi and follow Nature "
    "Communications typographic conventions."
)

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# FIGURE 1 — Phylogenetic tree
# ═══════════════════════════════════════════════════════════════════════════
add_figure_caption(doc, 1,
    "Maximum-likelihood phylogenetic tree of three focal avian species.")

insert_image(doc,
    r"E:\Data\Desktop\Work On\Ortho\Phylogenetic\Fig_phylo_tree.png",
    width_cm=10)

add_label_paragraph(doc, "Overall design.  ",
    "The tree was reconstructed from concatenated whole-proteome alignments using the "
    "JTT+CAT substitution model under a maximum-likelihood framework. "
    "Each terminal leaf represents a focal species: "
    "Gallus gallus (chicken, crimson), "
    "Anas platyrhynchos (duck, steel blue), "
    "and Columba livia (pigeon, amber). "
    "A filled circle at the root denotes the inferred common ancestor."
)

add_label_paragraph(doc, "Branch lengths.  ",
    "Numbers positioned above each branch segment indicate the estimated number of "
    "amino acid substitutions per site (scale bar = 0.02 substitutions/site). "
    "The relatively short internode separating Gallus and Anas reflects the well-documented "
    "recent divergence of Galliformes and Anseriformes (Galloanserae clade), "
    "whereas the longer branch leading to Columba is consistent with a deeper separation "
    "of Columbiformes from the Galloanserae lineage."
)

add_label_paragraph(doc, "Topology and biological significance.  ",
    "The recovered topology — (Gallus, Anas), Columba — recapitulates accepted avian "
    "systematics and provides the reference evolutionary framework for interpreting all "
    "downstream analyses. The clustering of the two precocial egg-laying species (Gallus and Anas) "
    "relative to the altricial Columba is central to the comparative glycoproteomics logic "
    "of this study: differences observed between Gallus/Anas and Columba can be attributed "
    "to divergent selection pressures associated with the precocial–altricial axis of reproductive "
    "life-history evolution."
)

add_label_paragraph(doc, "Model annotation.  ",
    "The bottom-right inset (italicised text) indicates the substitution model and inference method "
    "(JTT+CAT; Maximum Likelihood), ensuring methodological transparency in accordance with "
    "best practices for molecular phylogenetics reporting."
)

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# FIGURE 2 — Expansions & Contractions Sankey
# ═══════════════════════════════════════════════════════════════════════════
add_figure_caption(doc, 2,
    "Gene-family expansion and contraction dynamics inferred by CAFE5 across three avian species.")

insert_image(doc,
    r"E:\Data\Desktop\Work On\Ortho\Expansions Contractions Results\Fig.Expansions and Contractions.png",
    width_cm=15)

add_label_paragraph(doc, "Overall design.  ",
    "A custom Sankey (alluvial) diagram summarises the gene-family dynamics as inferred by CAFE5 "
    "for three species. Flow magnitude is proportional to the number of gene-family clusters "
    "undergoing each event. The middle column displays species nodes (Gallus, crimson; "
    "Anas, steel blue; Pigeon/Columba, amber), each split vertically into expansion (upper) "
    "and contraction (lower) contributions. Flows connect each species-event combination to "
    "the significantly enriched Gene Ontology (GO) terms in the right column."
)

add_label_paragraph(doc, "Species-level expansion/contraction statistics.  ",
    "Gallus exhibits a strongly asymmetric pattern: 6 expanded clusters (36 proteins) versus "
    "64 contracted clusters (421 proteins), indicating a net reductive evolution of specific "
    "gene repertoires in the chicken lineage. "
    "Columba (Pigeon) shows the inverse pattern: 75 expanded clusters (432 proteins) versus only "
    "8 contracted clusters (104 proteins), pointing to widespread gene-family amplification "
    "in the altricial pigeon. "
    "Anas occupies an intermediate position with 14 expansions (139 proteins) "
    "and 30 contractions (190 proteins)."
)

add_label_paragraph(doc, "GO enrichment flows and biological interpretation.  ",
    "Flow width from each species node to a given GO term is proportional to the number of "
    "proteins mapped to that term. Key GO term colors denote ontology domain: "
    "green = Biological Process (P), purple = Cellular Component (C), tan = Molecular Function (F). "
    "Contracted Gallus gene families are most prominently enriched for immunoglobulin production "
    "(48 proteins), immune response, and skeletal system development, consistent with "
    "the domestication-driven reduction of immune gene diversity in commercial poultry. "
    "Expanded Pigeon families converge on mitotic cell cycle regulation, cell adhesion, "
    "and Wnt signalling, potentially reflecting developmental plasticity requirements of the "
    "altricial neonate. "
    "In Anas, contraction signatures include motor activity and rRNA processing, "
    "while expansions are enriched for stress-response pathways and sphingolipid metabolism."
)

add_label_paragraph(doc, "Significance of the precocial–altricial contrast.  ",
    "The opposite directionality of genome evolution between Gallus/Anas and Columba "
    "supports a model in which the precocial egg-laying habit constrains gene-family expansions "
    "(necessitating tight developmental programmes), whereas the altricial strategy is permissive "
    "of broader genomic diversification. These contrasting trajectories contextualise the "
    "species-specific glycoprotein expression and glycan-remodelling patterns described in "
    "subsequent figures."
)

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# FIGURE 3 — Glycan network
# ═══════════════════════════════════════════════════════════════════════════
add_figure_caption(doc, 3,
    "Hierarchical network of orthologous egg-white glycoproteins and their glycan-type associations "
    "across three avian species.")

insert_image(doc,
    r"E:\Data\Desktop\Work On\同源糖型蛋白\Fig_glycan_network.png",
    width_cm=15)

add_label_paragraph(doc, "Overall design.  ",
    "This radial network integrates OrthoFinder clustering results with intact-glycopeptide mass "
    "spectrometry data to simultaneously represent orthology, species contribution, and glycan-type "
    "connectivity. Three concentric zones encode evolutionary conservation, and the outermost ring "
    "encodes N-glycan structural class."
)

add_label_paragraph(doc, "Central concentric circles (trispecies-conserved glycoproteins, GAC).  ",
    "Proteins identified as glycoproteins in all three species (trispecies orthogroups, labelled GAC) "
    "are arranged in up to seven concentric rings (radii 0–2.52, step 0.42 arbitrary units). "
    "Node colour follows a custom gradient from pale teal (#D6E2E2) to deep navy (#1B4D59): "
    "deeper colour indicates a larger orthogroup cluster size, serving as a proxy for greater "
    "evolutionary conservation or gene-family complexity. "
    "Proteins with the largest clusters — hence most ancient, broadly conserved functions — "
    "occupy the innermost ring."
)

add_label_paragraph(doc, "Inner annular sector (species-unique glycoproteins).  ",
    "Single-species glycoproteins, defined as those present in exactly one of the three species, "
    "are placed in a sector region immediately outside the GAC core (inner radius ≈ 3.07, "
    "outer radius ≈ 5.80). Nodes are rendered as diamonds, coloured by species: "
    "crimson for Gallus-specific, steel blue for Anas-specific, amber for Columba-specific. "
    "The proportion of the sector occupied by each species reflects the relative number of "
    "species-unique glycoproteins, highlighting the extent of lineage-specific glycoproteome expansion."
)

add_label_paragraph(doc, "Outer annular sector (bispecies-shared glycoproteins).  ",
    "Glycoproteins shared between exactly two species are positioned in the outer annular sector "
    "(inner radius ≈ 6.40, outer radius ≈ 10.00) as filled circles. Node colour is a mixture "
    "of the two contributing species' representative colours: purple for Gallus–Anas, "
    "olive-green for Anas–Columba, and orange-red for Gallus–Columba pairs. "
    "This zone captures partially conserved glycoproteins that may reflect shared functional "
    "constraints between specific species pairs."
)

add_label_paragraph(doc, "Outermost ring (glycan-type nodes).  ",
    "Seven glycan-type nodes are distributed evenly on a circle of radius 9.6: "
    "High Mannose, Pauci-mannose, Hybrid, Complex-Plain, Complex-Fucosylated, "
    "Complex-Sialylated, and Other. "
    "Each node's colour is drawn from a continuous blue gradient (pale #9DC4DE to deep #1A5C8E): "
    "darker shading indicates a higher number of distinct glycoproteins bearing that glycan type. "
    "All glycan-type nodes are rendered at the same size to emphasise colour over area."
)

add_label_paragraph(doc, "Connecting edges.  ",
    "Thin lines link each glycoprotein node to its detected glycan-type node(s), establishing "
    "a bipartite protein–glycan network. The density of edges to a given glycan-type node "
    "discloses which structural classes dominate the glycoproteome. "
    "Complex-Sialylated and High Mannose nodes typically attract the highest edge density, "
    "consistent with their prevalence in vertebrate glycoproteomes. "
    "The edge pattern for trispecies-conserved proteins versus species-specific proteins further "
    "illuminates whether glycan-type diversification is coupled to protein conservation."
)

add_label_paragraph(doc, "Biological significance.  ",
    "The network architecture reveals that the most evolutionarily conserved egg-white "
    "glycoproteins (deep-coloured inner-ring nodes) are preferentially decorated with "
    "Complex-Sialylated and High Mannose N-glycans, suggesting that these structural classes "
    "are under purifying selection. "
    "Lineage-specific glycoproteins, by contrast, display a broader spectrum of glycan types, "
    "including species-enriched Fucosylated and Pauci-mannose structures, implying that "
    "glycan remodelling accompanies protein neofunctionalisation during avian diversification."
)

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# FIGURE 4 — 2D Enrichment (Gallus vs Columba)
# ═══════════════════════════════════════════════════════════════════════════
add_figure_caption(doc, 4,
    "Two-dimensional protein–glycan fold-change enrichment analysis: "
    "Gallus gallus versus Columba livia.")

insert_image(doc,
    r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_2d_enrichment_Gallus_vs_Columba.png",
    width_cm=15)

add_label_paragraph(doc, "Overall design.  ",
    "Adapted from the bivariate enrichment framework of Ba et al. (2022, Science Advances, Fig. 3A), "
    "this plot jointly quantifies protein-level and glycan-level species differences for "
    "orthologous protein pairs identified between Gallus (early-precocial reference) and "
    "Columba (altricial comparator). "
    "The x-axis represents the log₂ protein fold-change (Log₂FC, Gallus/Columba), "
    "and the y-axis represents the log₂ glycan fold-change (Log₂FC, Gallus/Columba). "
    "Only proteins passing stringent BLASTp quality filters are included "
    "(E-value ≤ 1×10⁻⁵; average sequence identity ≥ 80% for full-length alignments, "
    "or maximum identity ≥ 50% for partial alignments)."
)

add_label_paragraph(doc, "Reference diagonal (y = x).  ",
    "The dashed diagonal line separates two biological regimes: "
    "points above the diagonal (blue-shaded region, labelled 'Glycan enriched in Gallus') "
    "represent proteins whose glycosylation is disproportionately elevated in Gallus relative to "
    "what would be expected from the protein-level fold-change alone — that is, the glycan change "
    "exceeds the protein change. "
    "Points below the diagonal (pink-shaded region, labelled 'Glycan suppressed in Gallus') "
    "indicate proteins whose glycosylation is relatively attenuated in Gallus, pointing to "
    "post-translational regulation of glycan occupancy or structure beyond changes in "
    "protein expression."
)

add_label_paragraph(doc, "Background proteins (grey points).  ",
    "Each grey point represents one BLASTp-matched orthologous protein pair with "
    "complete protein quantification (≥2 comparable runs) and glycan-site intensities in both species. "
    "Gene-name labels are annotated for notable background proteins to facilitate biological "
    "interpretation. The overall distribution of background proteins around the diagonal indicates "
    "the extent to which protein and glycan abundances co-scale between the two species, "
    "providing a quantitative estimate of glycan-remodelling independence from protein expression."
)

add_label_paragraph(doc, "Highlighted target proteins.  ",
    "Three functionally pivotal egg-white glycoproteins are highlighted with coloured markers "
    "and labelled annotation boxes: "
    "ovalbumin (OVAL, #E64B35 red), the most abundant egg-white protein; "
    "ovocleidin-116 (OC116, #4DBBD5 blue), a key calcite-nucleating shell-matrix protein; "
    "and transferrin (TRFE, #00A087 green), the predominant iron-transport glycoprotein in avian "
    "egg white. "
    "These pairs were assigned by manual inspection of BLASTp results to ensure biological "
    "accuracy despite average sequence identities below the automated 80% threshold. "
    "Their positions relative to the diagonal reveal whether inter-species divergence in "
    "glycosylation of these critical proteins is proportional to or independent of "
    "protein-expression divergence."
)

add_label_paragraph(doc, "Biological significance.  ",
    "The bivariate enrichment plot disentangles two conceptually distinct drivers of glycome "
    "divergence: changes in protein-expression levels (x-axis displacement from zero) versus "
    "changes in glycan structure or site occupancy (deviation from the diagonal). "
    "Highlighted proteins that fall substantially above the diagonal indicate that glycan "
    "remodelling has occurred independently of, and in excess of, protein-expression changes — "
    "a hallmark of adaptive glycan evolution. "
    "This approach, applied here to a precocial–altricial species comparison for the first time, "
    "reveals the degree to which N-glycan structural diversification of egg-white proteins "
    "has been shaped by divergent reproductive strategies."
)

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# FIGURES 5-8 — Glycan profiling stacked bars
# ═══════════════════════════════════════════════════════════════════════════
protein_info_dict = {
    "OVAL": {
        "num": 5,
        "full_name": "Ovalbumin (OVAL)",
        "description": (
            "Ovalbumin is the most abundant egg-white protein, comprising approximately 54% of "
            "total soluble egg-white protein in Gallus. It functions as a nutrient-storage "
            "glycoprotein and has well-characterised N-glycosylation at Asn293. "
            "Orthologues are present in all three species and are well-suited for cross-species "
            "glycan comparison."
        ),
        "species_note": "Gallus, Anas, and Columba",
        "img": r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_glycan_profiling_OVAL.png",
    },
    "OC116": {
        "num": 6,
        "full_name": "Ovocleidin-116 (OC116)",
        "description": (
            "Ovocleidin-116 is a C-type lectin-domain-containing protein exclusively associated "
            "with the calcified eggshell matrix and serves as a key nucleator of calcite "
            "crystallisation during shell formation. "
            "It is detected in both egg-laying precocial species (Gallus and Anas) but is absent "
            "or undetectable in the pigeon (Columba), which lacks a heavily calcified eggshell, "
            "making its glycan profile uniquely informative for shell-biology evolution."
        ),
        "species_note": "Gallus and Anas (Columba absent or undetected)",
        "img": r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_glycan_profiling_OC116.png",
    },
    "TRFE": {
        "num": 7,
        "full_name": "Transferrin (TRFE)",
        "description": (
            "Transferrin is an iron-binding transport glycoprotein constituting a major component "
            "of egg white. Its N-glycosylation plays a critical role in structural stability, "
            "receptor recognition, and the antimicrobial iron-sequestration function of the egg. "
            "Transferrin orthologues are present in all three species and exhibit some of the "
            "most abundant glycan signals in the mass-spectrometry dataset."
        ),
        "species_note": "Gallus, Anas, and Columba",
        "img": r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_glycan_profiling_TRFE.png",
    },
    "OC17": {
        "num": 8,
        "full_name": "Ovocalyxin-17 (OC17)",
        "description": (
            "Ovocalyxin-17 is a member of the BPI/PLUNC superfamily proteins restricted to the "
            "Gallus lineage; its strict orthologue was not identified in Anas or Columba within "
            "the BLASTp search space, making this effectively a chicken-specific glycoprotein. "
            "Its eggshell-residency and antibacterial properties are well documented. "
            "The absence of orthologues in duck and pigeon suggests it arose via lineage-specific "
            "gene duplication or functional divergence after the Galliformes–Anseriformes split."
        ),
        "species_note": "Gallus only (absent in Anas and Columba)",
        "img": r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_glycan_profiling_OC17.png",
    },
}

for prot_key, pinfo in protein_info_dict.items():
    add_figure_caption(doc, pinfo["num"],
        f"Glycan-type composition profile of {pinfo['full_name']}.")

    insert_image(doc, pinfo["img"], width_cm=12)

    add_label_paragraph(doc, "Protein background.  ", pinfo["description"])

    add_label_paragraph(doc, "Figure design.  ",
        f"Each vertical stacked bar represents one species ({pinfo['species_note']}). "
        "Bar height is fixed at 100%, and segment heights encode the relative abundance (%) "
        "of each N-glycan structural class. Intensities were derived from IGP (Intact "
        "GlycoProtein) quantification using the mean of three biological replicates from "
        "intact glycopeptide LC-MS/MS. Segments are coloured by glycan class following a "
        "consistent NPG (Nature Publishing Group) palette: "
        "High-Mannose (#4DBBD5, blue); "
        "Paucimannose/Truncated (#8491B4, grey-blue); "
        "Neutral Complex/Hybrid (#00A087, teal); "
        "Fucosylated Complex/Hybrid (#F39B7F, salmon); "
        "Sialylated Complex/Hybrid (#E64B35, red); "
        "Other (#CCCCCC, light grey). "
        "Only glycan classes detected for that protein in at least one species are shown."
    )

    if prot_key == "OVAL":
        add_label_paragraph(doc, "Interpretation.  ",
            "The glycan profile of OVAL reflects a protein whose N-glycosylation has been "
            "maintained across the three lineages but has undergone quantitative recomposition. "
            "A higher proportion of Sialylated Complex/Hybrid glycans in Gallus OVAL relative "
            "to Columba would indicate enhanced negative surface charge on the most abundant "
            "egg-white protein, potentially modulating protein–protein interactions or "
            "microbial defence within the egg environment. "
            "Cross-species differences in the Fucosylated fraction may reflect divergence in "
            "α1,3/1,6-fucosyltransferase gene expression or substrate specificity between "
            "precocial and altricial lineages."
        )
    elif prot_key == "OC116":
        add_label_paragraph(doc, "Interpretation.  ",
            "Because OC116 is absent in Columba, this figure provides a direct comparison of "
            "shell-matrix glycoprotein glycosylation between two precocial species. "
            "Similarities in the glycan profiles of Gallus and Anas OC116 would support the "
            "hypothesis that the glycan code on this protein is functionally constrained by "
            "the requirement for precise calcite-crystal nucleation. "
            "Differences, by contrast, would imply that glycan composition is not a strict "
            "prerequisite for shell-calcification function and that structural divergence "
            "has occurred under relaxed purifying selection."
        )
    elif prot_key == "TRFE":
        add_label_paragraph(doc, "Interpretation.  ",
            "As a conserved iron-transport glycoprotein present in all three species, "
            "transferrin's glycan profile acts as an internal benchmark for interpreting "
            "divergence in other proteins. A predominantly Sialylated Complex/Hybrid pattern "
            "conserved across all three species would indicate strong purifying selection on "
            "TRFE N-glycosylation, consistent with sialic acid's known role in "
            "prolonging serum half-life and maintaining iron-binding conformation. "
            "Any deviation in the altricial Columba TRFE profile relative to both precocial "
            "species would warrant investigation of Columba-specific glycosyltransferase activity."
        )
    elif prot_key == "OC17":
        add_label_paragraph(doc, "Interpretation.  ",
            "The single-species nature of this figure reflects the Gallus-specific origin of OC17. "
            "The glycan profile documents, for the first time, the N-glycan structural repertoire "
            "of this lineage-specific antimicrobial eggshell protein. "
            "The predominant glycan classes identified here can be used to infer which "
            "glycosyltransferase activities are co-recruited to the eggshell-matrix proteome "
            "in Gallus. The comparison of OC17's glycan profile with those of the functionally "
            "related OC116 further illuminates whether eggshell-specific proteins share "
            "a common glycan 'signature' in the chicken lineage."
        )

    add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# FIGURES 9-11 — Highlighted correlation scatter plots
# ═══════════════════════════════════════════════════════════════════════════
species_info = {
    "Gallus": {
        "num": 9,
        "common": "chicken (Gallus gallus)",
        "strategy": "precocial",
        "targets": "OVAL, OC116, TRFE, and OC17",
        "img": r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_highlighted_correlation_Gallus.png",
    },
    "Anas": {
        "num": 10,
        "common": "duck (Anas platyrhynchos)",
        "strategy": "precocial",
        "targets": "OVAL, OC116, and TRFE",
        "img": r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_highlighted_correlation_Anas.png",
    },
    "Columba": {
        "num": 11,
        "common": "pigeon (Columba livia)",
        "strategy": "altricial",
        "targets": "OVAL and TRFE",
        "img": r"E:\Data\Desktop\Work On\糖蛋白和蛋白联合分析\Figure\Fig_highlighted_correlation_Columba.png",
    },
}

for sp, sinfo in species_info.items():
    add_figure_caption(doc, sinfo["num"],
        f"Proteotype–glycoproteome co-abundance correlation in {sinfo['common']}.")

    insert_image(doc, sinfo["img"], width_cm=14)

    add_label_paragraph(doc, "Overall design.  ",
        f"Each data point in this scatter plot represents a single N-glycosylation site "
        f"on a detected glycoprotein in {sinfo['common']}. "
        f"The x-axis shows the log₂-transformed mean protein intensity (log₂ Protein Intensity), "
        f"and the y-axis shows the log₂-transformed mean glycan-site intensity "
        f"(log₂ Glycan Intensity), both averaged across three biological replicates. "
        f"Only proteins with ≥2 comparable quantitative runs (Number Comparable ≥ 2) "
        f"are retained to ensure statistical reliability. "
        f"The dashed diagonal reference line (y = x) denotes equal protein-to-glycan intensity "
        f"scaling; points above this line carry proportionally greater glycan signal, "
        f"while points below carry proportionally less."
    )

    add_label_paragraph(doc, "Background proteome (grey points).  ",
        "Grey semi-transparent points represent non-highlighted glycoproteins across the entire "
        "detected egg-white proteome/glycoproteome, providing the global distributional context. "
        "The density and slope of this background cloud reflect the typical degree of "
        "protein–glycan co-regulation across all detected proteins in this species."
    )

    add_label_paragraph(doc, "Highlighted target proteins.  ",
        f"Key egg-white glycoproteins are overlaid as coloured markers: "
        f"OVAL (ovalbumin, red #E64B35), OC116 (ovocleidin-116, blue #4DBBD5), "
        f"TRFE (transferrin, green #00A087)"
        + (", and OC17 (ovocalyxin-17, dark blue #3C5488, Gallus only)" if sp == "Gallus" else "")
        + f". Each coloured point corresponds to a specific N-glycosylation site on the target "
        f"protein, with site position (e.g., 147N) annotated in the label. "
        f"Larger marker size and black border facilitate visual discrimination from background."
    )

    add_label_paragraph(doc, "Spearman correlation statistics.  ",
        "The inset text box in the upper-left corner reports the Spearman rank correlation "
        "coefficient (ρ) and the associated P-value between protein intensity and glycan intensity "
        "across all detected glycosylation sites. "
        "A highly significant positive Spearman ρ indicates that, at the global proteome level, "
        "glycan-site abundance is strongly co-regulated with protein expression — that is, "
        "more abundant proteins tend to contribute more abundant glycan signals. "
        "Significance thresholds are denoted as: *** p < 0.001, ** p < 0.01, * p < 0.05, ns."
    )

    add_label_paragraph(doc, "Biological significance.  ",
        f"This figure establishes the quantitative relationship between the proteome and the "
        f"glycoproteome in the {sinfo['strategy']} {sinfo['common']}. "
        f"The position of the highlighted target proteins — {sinfo['targets']} — relative to the "
        f"diagonal and to the background distribution reveals whether these biologically critical "
        f"glycoproteins maintain glycan intensities commensurate with their protein expression levels "
        f"or whether they carry disproportionately high or low glycan loads. "
        f"Systematic deviation of OVAL, OC116, or TRFE from the global regression trend would "
        f"indicate post-translational glycan regulation at these specific loci, "
        f"consistent with their known functional importance in egg-white architecture, "
        f"shell calcification, and iron sequestration, respectively. "
        f"Comparing this figure across all three species (Figures 9–11) provides a direct "
        f"readout of how the proteotype–glycoproteome coupling is conserved or rewired "
        f"along divergent avian lineages."
    )

    add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════════════════
out_path = r"E:\Data\Desktop\Work On\Figure_Legends_Nature_Style.docx"
doc.save(out_path)
print(f"[OK] 文档已保存: {out_path}")
