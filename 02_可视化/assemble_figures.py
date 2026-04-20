"""
图片组装工具 (两步合一)

步骤1 — compose:  从 Sci_Adv_Figure/PNG/Fig{N}/ 中的子面板 PNG 合成大图
                   → 输出到 Sci_Adv_Figure/Composed/Fig{N}_composed.png
步骤2 — pack:     将合成大图 + 图注打包为单独的 docx 文档
                   → 输出 Science_Advances_Main_Figures.docx

用法:
    python assemble_figures.py           # 执行两步
    python assemble_figures.py compose   # 仅合成面板
    python assemble_figures.py pack      # 仅打包 docx
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont

# ── 公共路径 ──────────────────────────────────────────────────────
ROOT = Path(r"d:\system_folder\Desktop\Work On")
SOURCE_DIR = ROOT / "02_可视化" / "Sci_Adv_Figure" / "PNG"
COMPOSED_DIR = ROOT / "02_可视化" / "Sci_Adv_Figure" / "Composed"
LEGEND_FILE = ROOT / "03_文章撰写" / "Main_Figure_Legends_SA.txt"

# ═══════════════════════════════════════════════════════════════════
#  步骤 1: compose — PIL 面板合成
# ═══════════════════════════════════════════════════════════════════

CANVAS_MAX_WIDTH = 7200
OUTER_MARGIN = 180
PANEL_GAP = 70
ROW_GAP = 160
LABEL_BAND = 150
DIVIDER_WIDTH = 5
BACKGROUND = "white"
DIVIDER_COLOR = (150, 150, 150)
LABEL_COLOR = (0, 0, 0)
DEFAULT_DPI = (300, 300)

TARGET_ROW_HEIGHTS = {
    1: 2200,
    2: 1850,
    3: 1450,
    4: 1180,
    5: 1000,
    6: 900,
}

FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\timesbd.ttf"),
    Path(r"C:\Windows\Fonts\times.ttf"),
    Path(r"C:\Windows\Fonts\timesi.ttf"),
]


@dataclass(frozen=True)
class PanelImage:
    path: Path
    major: int
    minor: int


@dataclass(frozen=True)
class RowLayout:
    panels: list[PanelImage]
    scaled_sizes: list[tuple[int, int]]
    row_width: int
    row_height: int


def natural_panel_key(path: Path) -> tuple[int, int, str]:
    match = re.match(r"^(\d+)(?:-(\d+))?", path.stem)
    if not match:
        return (999, 999, path.name.lower())
    major = int(match.group(1))
    minor = int(match.group(2) or 0)
    return (major, minor, path.name.lower())


def iter_figure_dirs(base_dir: Path) -> Iterable[Path]:
    for folder in sorted(base_dir.iterdir(), key=lambda p: natural_panel_key(p)):
        if folder.is_dir() and folder.name.startswith("Fig"):
            yield folder


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for candidate in FONT_CANDIDATES:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def collect_panels(folder: Path) -> list[PanelImage]:
    panels: list[PanelImage] = []
    for image_path in sorted(folder.glob("*.png"), key=natural_panel_key):
        major, minor, _ = natural_panel_key(image_path)
        panels.append(PanelImage(path=image_path, major=major, minor=minor))
    if not panels:
        raise ValueError(f"No PNG panels found in {folder}")
    return panels


MAX_PANELS_PER_ROW = 4


def group_panels_by_major(panels: list[PanelImage]) -> list[list[PanelImage]]:
    groups: list[list[PanelImage]] = []
    current_major: int | None = None
    current_group: list[PanelImage] = []

    for panel in panels:
        if current_major is None or panel.major != current_major or len(current_group) >= MAX_PANELS_PER_ROW:
            if current_group:
                groups.append(current_group)
            current_major = panel.major
            current_group = [panel]
        else:
            current_group.append(panel)

    if current_group:
        groups.append(current_group)
    return groups


def compute_row_layout(panels: list[PanelImage]) -> RowLayout:
    sizes: list[tuple[int, int]] = []

    total_aspect = 0.0
    aspects = []
    for panel in panels:
        with Image.open(panel.path) as image:
            aspect = image.width / image.height
            aspects.append(aspect)
            total_aspect += aspect

    max_inner_width = CANVAS_MAX_WIDTH - 2 * OUTER_MARGIN
    gaps = PANEL_GAP * max(0, len(panels) - 1)

    if total_aspect > 0:
        target_height = (max_inner_width - gaps) / total_aspect
    else:
        target_height = 1000

    row_width = 0
    for aspect in aspects:
        w = max(1, round(target_height * aspect))
        h = max(1, round(target_height))
        sizes.append((w, h))
        row_width += w

    row_width += gaps
    row_height = max(1, round(target_height))

    return RowLayout(
        panels=panels,
        scaled_sizes=sizes,
        row_width=row_width,
        row_height=row_height,
    )


def compose_figure_folder(folder: Path) -> Path:
    panels = collect_panels(folder)
    row_groups = group_panels_by_major(panels)
    row_layouts = [compute_row_layout(group) for group in row_groups]

    canvas_width = min(
        CANVAS_MAX_WIDTH,
        max(layout.row_width for layout in row_layouts) + 2 * OUTER_MARGIN,
    )
    canvas_height = (
        2 * OUTER_MARGIN
        + sum(layout.row_height + LABEL_BAND for layout in row_layouts)
        + ROW_GAP * max(0, len(row_layouts) - 1)
    )

    canvas = Image.new("RGB", (canvas_width, canvas_height), BACKGROUND)
    draw = ImageDraw.Draw(canvas)
    label_font = load_font(108)

    panel_index = 0
    y_offset = OUTER_MARGIN

    for row_number, layout in enumerate(row_layouts):
        row_left = (canvas_width - layout.row_width) // 2
        image_top = y_offset + LABEL_BAND
        x_offset = row_left

        for panel, (panel_width, panel_height) in zip(layout.panels, layout.scaled_sizes):
            with Image.open(panel.path) as original:
                image = original.convert("RGB").resize((panel_width, panel_height), Image.Resampling.LANCZOS)
            canvas.paste(image, (x_offset, image_top))

            label = chr(ord("A") + panel_index)
            draw.text((x_offset, y_offset), label, fill=LABEL_COLOR, font=label_font)

            x_offset += panel_width
            if panel is not layout.panels[-1]:
                x_offset += PANEL_GAP

            panel_index += 1

        y_offset = image_top + layout.row_height + ROW_GAP

    output_path = COMPOSED_DIR / f"{folder.name}_composed.png"
    canvas.save(output_path, dpi=DEFAULT_DPI)
    return output_path


def step_compose() -> None:
    COMPOSED_DIR.mkdir(parents=True, exist_ok=True)
    figure_dirs = list(iter_figure_dirs(SOURCE_DIR))
    if not figure_dirs:
        raise FileNotFoundError(f"No Fig* directories found in {SOURCE_DIR}")

    outputs = []
    for folder in figure_dirs:
        output_path = compose_figure_folder(folder)
        outputs.append(output_path)
        print(f"[OK] {folder.name} -> {output_path.name}")

    print("\nComposed figures:")
    for output_path in outputs:
        print(output_path)


# ═══════════════════════════════════════════════════════════════════
#  步骤 2: pack — 将合成图打包为 docx
# ═══════════════════════════════════════════════════════════════════

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt as DPt

PACK_OUTPUT = ROOT / "Science_Advances_Main_Figures.docx"
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11.0)
MARGIN_LEFT = Inches(0.7)
MARGIN_RIGHT = Inches(0.7)
MARGIN_TOP = Inches(0.6)
MARGIN_BOTTOM = Inches(0.6)
CAPTION_SPACE = Inches(1.6)


def set_default_font(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = DPt(12)


def apply_section_layout(section, landscape: bool) -> None:
    if landscape:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = PAGE_HEIGHT
        section.page_height = PAGE_WIDTH
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM


def parse_legend_blocks(text: str) -> list[tuple[str, str]]:
    blocks = [block.strip() for block in text.strip().split("\n\n") if block.strip()]
    parsed = []
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        title = lines[0]
        body = " ".join(lines[1:])
        parsed.append((title, body))
    return parsed


def available_picture_size(section) -> tuple[int, int]:
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_height = section.page_height - section.top_margin - section.bottom_margin - CAPTION_SPACE
    return usable_width, usable_height


def figure_display_width(section, image_width: int, image_height: int):
    max_width, max_height = available_picture_size(section)
    width_from_height = int(max_height * image_width / image_height)
    return min(max_width, width_from_height)


def add_caption(document: Document, title: str, body: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = DPt(6)
    paragraph.paragraph_format.space_after = DPt(0)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False

    title_run = paragraph.add_run(title + " ")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = DPt(10)

    body_run = paragraph.add_run(body)
    body_run.font.name = "Times New Roman"
    body_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    body_run.font.size = DPt(10)


def add_figure_page(document: Document, image_path: Path, legend: tuple[str, str], add_section_break: bool) -> None:
    if add_section_break:
        section = document.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = document.sections[0]

    landscape = image_path.stem.lower() == "fig3_composed"
    apply_section_layout(section, landscape=landscape)

    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = DPt(0)
    picture_paragraph.paragraph_format.space_after = DPt(6)
    picture_paragraph.paragraph_format.keep_together = True
    picture_paragraph.paragraph_format.keep_with_next = True

    with Image.open(image_path) as img:
        width_px, height_px = img.size

    width_emu = figure_display_width(section, width_px, height_px)
    picture_paragraph.add_run().add_picture(str(image_path), width=width_emu)
    add_caption(document, *legend)


def step_pack() -> None:
    figure_paths = sorted(COMPOSED_DIR.glob("Fig*_composed.*"))
    legends = parse_legend_blocks(LEGEND_FILE.read_text(encoding="utf-8"))

    if len(figure_paths) != len(legends):
        raise ValueError(f"Expected the same number of figures and legend blocks, got {len(figure_paths)} and {len(legends)}")

    document = Document()
    set_default_font(document)
    apply_section_layout(document.sections[0], landscape=False)

    for index, (image_path, legend) in enumerate(zip(figure_paths, legends)):
        add_figure_page(document, image_path, legend, add_section_break=index > 0)

    document.save(PACK_OUTPUT)
    print(f"Saved: {PACK_OUTPUT}")


# ═══════════════════════════════════════════════════════════════════
#  主入口
# ═══════════════════════════════════════════════════════════════════

def main() -> None:
    args = [a.lower() for a in sys.argv[1:]]

    if not args or "compose" in args:
        print("=== Step 1: Compose panels → composed figures ===")
        step_compose()

    if not args or "pack" in args:
        print("\n=== Step 2: Pack composed figures → docx ===")
        step_pack()


if __name__ == "__main__":
    main()
