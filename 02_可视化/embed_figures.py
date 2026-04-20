"""
将合成图片嵌入手稿 (支持英文/中文版)
用法:
    python embed_figures.py en   # 英文版
    python embed_figures.py cn   # 中文版
    python embed_figures.py      # 默认处理两个版本
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"d:\system_folder\Desktop\Work On")
FIGURE_DIR = ROOT / "02_可视化" / "Sci_Adv_Figure" / "Composed"

CAPTION_SPACE_PT = 110
INSERT_AFTER = {
    18: "Fig1_composed.png",
    23: "Fig2_composed.jpg",
    26: "Fig3_composed.jpg",
    28: "Fig4_composed.png",
    30: "Fig5_composed.jpg",
    37: "Fig6_composed.jpg",
}

LANG_CONFIG = {
    "en": {
        "manuscript": ROOT / "03_文章撰写" / "manuscript20260412.docx",
        "legend": ROOT / "03_文章撰写" / "Main_Figure_Legends_SA.txt",
        "output": ROOT / "03_文章撰写" / "manuscript20260412_with_figures.docx",
        "fallback": ROOT / "03_文章撰写" / "manuscript20260412_with_figures_updated.docx",
        "east_asia_font": "Times New Roman",
    },
    "cn": {
        "manuscript": ROOT / "03_文章撰写" / "manuscript20260412_cn.docx",
        "legend": ROOT / "03_文章撰写" / "Main_Figure_Legends_SA_cn.txt",
        "output": ROOT / "03_文章撰写" / "manuscript20260412_cn_with_figures.docx",
        "fallback": ROOT / "03_文章撰写" / "manuscript20260412_cn_with_figures_updated.docx",
        "east_asia_font": "SimSun",
    },
}


def parse_legend_blocks(text: str) -> dict[str, tuple[str, str]]:
    blocks = [block.strip() for block in text.strip().split("\n\n") if block.strip()]
    parsed = {}
    for index, block in enumerate(blocks, start=1):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        parsed[f"Fig{index}_composed"] = (lines[0], " ".join(lines[1:]))
    return parsed


def new_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def available_space(section) -> tuple[int, int]:
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_height = section.page_height - section.top_margin - section.bottom_margin - Pt(CAPTION_SPACE_PT)
    return usable_width, usable_height


def scaled_width(section, image_width: int, image_height: int):
    max_width, max_height = available_space(section)
    width_from_height = int(max_height * image_width / image_height)
    return min(max_width, width_from_height)


def add_figure_block(anchor_paragraph, image_path: Path, title: str, legend_body: str, east_asia_font: str) -> None:
    from PIL import Image

    section = anchor_paragraph.part.document.sections[0]
    with Image.open(image_path) as image:
        image_width, image_height = image.size

    picture_paragraph = new_paragraph_after(anchor_paragraph)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(6)
    picture_paragraph.paragraph_format.space_after = Pt(6)
    picture_paragraph.paragraph_format.keep_together = True
    picture_paragraph.paragraph_format.keep_with_next = True
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=scaled_width(section, image_width, image_height))

    caption_paragraph = new_paragraph_after(picture_paragraph)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(6)
    caption_paragraph.paragraph_format.keep_together = True

    title_run = caption_paragraph.add_run(title + " ")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    title_run.font.size = Pt(10)

    body_run = caption_paragraph.add_run(legend_body)
    body_run.font.name = "Times New Roman"
    body_run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    body_run.font.size = Pt(10)


def embed(lang: str) -> None:
    cfg = LANG_CONFIG[lang]
    document = Document(cfg["manuscript"])
    legends = parse_legend_blocks(cfg["legend"].read_text(encoding="utf-8"))
    original_paragraphs = list(document.paragraphs)

    for paragraph_index in sorted(INSERT_AFTER.keys(), reverse=True):
        anchor = original_paragraphs[paragraph_index]
        figure_name = INSERT_AFTER[paragraph_index]
        title, legend_body = legends[Path(figure_name).stem]
        add_figure_block(anchor, FIGURE_DIR / figure_name, title, legend_body, cfg["east_asia_font"])

    try:
        document.save(cfg["output"])
        print(cfg["output"])
    except PermissionError:
        document.save(cfg["fallback"])
        print(cfg["fallback"])


def main() -> None:
    args = sys.argv[1:]
    if not args:
        langs = ["en", "cn"]
    else:
        langs = [a.lower() for a in args if a.lower() in LANG_CONFIG]
    for lang in langs:
        print(f"\n--- Embedding figures ({lang}) ---")
        embed(lang)


if __name__ == "__main__":
    main()
