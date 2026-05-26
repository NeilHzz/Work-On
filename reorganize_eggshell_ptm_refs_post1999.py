import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import Workbook


SOURCE_PATH = Path(r"e:\Data\Desktop\Work On\eggshell_matrix_ptm_direct_related_refs.docx")
TARGET_PATH = Path(r"e:\Data\Desktop\Work On\eggshell_matrix_ptm_post1999_by_modification.docx")
XLSX_PATH = Path(r"e:\Data\Desktop\Work On\eggshell_matrix_ptm_post1999_by_modification.xlsx")
ALL_XLSX_PATH = Path(r"e:\Data\Desktop\Work On\eggshell_matrix_ptm_all_literature.xlsx")


CATEGORY_METADATA = {
    "n_glycosylation": {
        "heading": "一、N-糖基化与糖蛋白鉴定",
        "note": "注：本节收录直接识别 N-糖基化位点、N-糖蛋白组成或以糖蛋白形式进行分子鉴定的研究。",
    },
    "o_glycosylation": {
        "heading": "二、O-糖基化",
        "note": "注：本节单列直接证明 O-糖基化存在的研究。",
    },
    "phosphorylation": {
        "heading": "三、磷酸化",
        "note": "注：本节聚焦蛋壳或鸡蛋蛋白中的磷酸化与磷蛋白组研究。",
    },
    "multi_ptm": {
        "heading": "四、多修饰联合组学",
        "note": "注：本节收录同时整合多种翻译后修饰信息的研究。",
    },
    "support": {
        "heading": "六、修饰讨论所需的背景与功能支持文献",
        "note": "注：以下文献不以单一修饰类型测定为主，但可为蛋壳基质蛋白修饰的功能解释、矿化背景和结构讨论提供支撑。",
    },
}


OTHER_MODIFICATION_HEADING = "五、除糖基化和磷酸化外的其他修饰或修饰加工事件"
OTHER_MODIFICATION_NOTE = "注：目前在这批蛋壳基质蛋白与鸡蛋蛋白文献中，除糖基化和磷酸化外，直接证据最明确的其他事件主要包括二硫键形成、蛋白聚糖化背景，以及去 N-糖基化这类修饰加工过程。"
OTHER_MODIFICATION_PARAGRAPHS = [
    "1. 二硫键形成：2002 年的 ovocleidin-116 研究除鉴定 glycosylated Asn 外，还明确解析了 disulfide bonds，说明部分蛋壳基质蛋白的成熟结构不仅依赖糖基化，也依赖二硫键稳定。该条目仍保留在本版文献正文中。",
    "2. 蛋白聚糖化背景：ovocleidin-116 最初被定义为 eggshell matrix proteoglycan 的核心蛋白，提示其与更大尺度的糖胺聚糖链修饰背景相关。不过这一直接定义来自 1999 年文献，在本版“仅保留 2000 年及以后文献”的规则下未纳入正文条目。",
    "3. 去 N-糖基化：ovalbumin 的位点特异性 de-N-glycosylation 说明鸡输卵管中存在新生糖蛋白质量控制过程，可视为修饰加工事件而非单纯加成型修饰。同样由于该文献发表于 1997 年，本版未将其保留在正文条目中。",
    "4. 证据边界：截至当前这批已整理文献，尚未看到足够直接且可稳定归类到蛋壳基质蛋白或鸡蛋蛋白体系中的乙酰化、甲基化、泛素化或 SUMO 化证据，因此不建议在正文中把这些类型写成已证实事实。",
]


DOI_TO_CATEGORY = {
    "10.1016/S0014-5793(99)01586-0": "multi_ptm",
    "10.1074/jbc.274.46.32915": "support",
    "10.1016/S0945-053X(02)00031-8": "n_glycosylation",
    "10.1021/acs.jafc.3c00708": "n_glycosylation",
    "10.1016/j.fbio.2020.100590": "n_glycosylation",
    "10.1016/j.ijbiomac.2020.08.193": "n_glycosylation",
    "10.1111/jfbc.14006": "n_glycosylation",
    "10.1021/jf048369l": "n_glycosylation",
    "10.1016/j.fbio.2024.103938": "o_glycosylation",
    "10.1073/pnas.94.12.6244": "n_glycosylation",
    "10.1042/BJ2400871": "phosphorylation",
    "10.1021/bi00298a027": "phosphorylation",
    "10.1002/pmic.200600635": "phosphorylation",
    "10.1021/acs.jafc.9b04638": "phosphorylation",
    "10.1016/j.foodchem.2020.127167": "multi_ptm",
    "10.2141/jpsa.009122": "support",
    "10.1007/s00018-009-0046-y": "support",
    "10.2741/3985": "support",
    "10.1186/s12860-021-00350-0": "support",
    "10.1186/s12953-015-0078-1": "support",
    "10.1126/sciadv.aar3219": "support",
    "10.1074/jbc.M610294200": "support",
    "10.1074/jbc.M406033200": "support",
}


ORDERED_CATEGORIES = [
    "n_glycosylation",
    "o_glycosylation",
    "phosphorylation",
    "multi_ptm",
    "support",
]


CATEGORY_LABELS = {
    "n_glycosylation": "N-糖基化与糖蛋白鉴定",
    "o_glycosylation": "O-糖基化",
    "phosphorylation": "磷酸化",
    "multi_ptm": "多修饰联合组学",
    "support": "背景/支持文献",
}


DIRECTNESS_LABELS = {
    "support": "背景支持文献",
}


ADDITIONAL_REFERENCES = [
    {
        "authors": "Geng F, Wang J, Liu D, Jin Y, Ma M",
        "year": 2017,
        "title": "Identification of N-Glycosites in Chicken Egg White Proteins Using an Omics Strategy",
        "doi": "10.1021/acs.jafc.7b01706",
        "category": "n_glycosylation",
        "journal": "Journal of Agricultural and Food Chemistry",
        "material": "蛋清",
    },
    {
        "authors": "Harvey DJ, Wing DR, Kuster B, Wilson IBH",
        "year": 2000,
        "title": "Composition of N-linked carbohydrates from ovalbumin and co-purified glycoproteins",
        "doi": "10.1016/S1044-0305(00)00122-7",
        "category": "n_glycosylation",
        "journal": "Journal of the American Society for Mass Spectrometry",
        "material": "蛋清（ovalbumin 及共纯化糖蛋白）",
    },
    {
        "authors": "Yamashita K, Tachibana Y, Nakayama T, Kitamura M, Ito Y, Kobata A",
        "year": 1982,
        "title": "Structural study of the carbohydrate moiety of hen ovomucoid. Occurrence of a series of pentaantennary complex-type asparagine-linked sugar chains.",
        "doi": "10.1016/S0021-9258(18)33585-3",
        "category": "n_glycosylation",
        "journal": "Journal of Biological Chemistry",
        "material": "蛋清（ovomucoid）",
    },
    {
        "authors": "Offengenden M, Fentabil MA, Wu J",
        "year": 2011,
        "title": "N-glycosylation of ovomucin from hen egg white",
        "doi": "10.1007/s10719-011-9328-3",
        "category": "n_glycosylation",
        "journal": "Glycoconjugate Journal",
        "material": "蛋清（ovomucin）",
    },
    {
        "authors": "Geng F, Xie Y, Wang J, Majumder K, Qiu N, Ma M",
        "year": 2018,
        "title": "N-Glycoproteomic Analysis of Chicken Egg Yolk",
        "doi": "10.1021/acs.jafc.8b04492",
        "category": "n_glycosylation",
        "journal": "Journal of Agricultural and Food Chemistry",
        "material": "卵黄",
    },
    {
        "authors": "Zhu F, Qiu N, Sun H, Meng Y, Zhou Y",
        "year": 2019,
        "title": "Integrated Proteomic and N-Glycoproteomic Analyses of Chicken Egg during Embryonic Development",
        "doi": "10.1021/acs.jafc.9b05133",
        "category": "multi_ptm",
        "journal": "Journal of Agricultural and Food Chemistry",
        "material": "鸡蛋整体（胚胎发育）",
    },
    {
        "authors": "Zhou Y, Qiu N, Mine Y, Keast R, Meng Y",
        "year": 2021,
        "title": "Comparative N-Glycoproteomic Analysis Provides Novel Insights into the Deterioration Mechanisms in Chicken Egg Vitelline Membrane during High-Temperature Storage",
        "doi": "10.1021/acs.jafc.0c07557",
        "category": "n_glycosylation",
        "journal": "Journal of Agricultural and Food Chemistry",
        "material": "卵黄膜",
    },
    {
        "authors": "Hirose J, Doi Y, Kitabatake N, Narita H",
        "year": 2006,
        "title": "Ovalbumin-Related Gene Y Protein Bears Carbohydrate Chains of the Ovomucoid Type",
        "doi": "10.1271/bbb.70.144",
        "category": "n_glycosylation",
        "journal": "Bioscience, Biotechnology, and Biochemistry",
        "material": "蛋清（ovalbumin-related gene Y）",
    },
    {
        "authors": "Qi Q, Shi D, Su W, Mu Y",
        "year": 2024,
        "title": "N-glycoproteomic profiling reveals structural and functional alterations in yellow primary preserved egg white under saline-alkali treatment",
        "doi": "10.1016/j.fochx.2024.101244",
        "category": "n_glycosylation",
        "journal": "Food Chemistry: X",
        "material": "蛋清（皮蛋加工体系）",
    },
    {
        "authors": "Dai D, Wang X, Wu K, Lan F, Jin J, Zhang W, Wen C, Li J, Yang N, Sun C",
        "year": 2025,
        "title": "Proteomic and N-glycosylation analysis of fertile egg white during storage and incubation in chickens",
        "doi": "10.1016/j.psj.2024.104526",
        "category": "n_glycosylation",
        "journal": "Poultry Science",
        "material": "受精蛋清",
    },
    {
        "authors": "Xiao D, Hu G, Ding Q, He H, Wang J, Geng F",
        "year": 2024,
        "title": "Research Note: Comprehensive proteomic, phosphoproteomic, and N-glycoproteomic analysis of chicken egg yolk plasma",
        "doi": "10.1016/j.psj.2024.104253",
        "category": "multi_ptm",
        "journal": "Poultry Science",
        "material": "卵黄浆",
    },
]


MATERIAL_BY_DOI = {
    "10.1016/S0014-5793(99)01586-0": "蛋壳基质蛋白（ovocleidin）",
    "10.1074/jbc.274.46.32915": "蛋壳基质蛋白（ovocleidin-116）",
    "10.1016/S0945-053X(02)00031-8": "蛋壳基质蛋白（ovocleidin-116）",
    "10.1002/pmic.200600635": "蛋壳钙化层",
    "10.1016/j.foodchem.2020.127167": "蛋壳基质",
    "10.1021/acs.jafc.3c00708": "蛋壳角质层/矿化层",
    "10.2141/jpsa.009122": "蛋壳形成/蛋壳基质",
    "10.1007/s00018-009-0046-y": "蛋壳基质蛋白",
    "10.2741/3985": "蛋壳整体",
    "10.1186/s12860-021-00350-0": "蛋壳整体",
    "10.1186/s12953-015-0078-1": "鹌鹑/鸡/火鸡蛋壳基质",
    "10.1126/sciadv.aar3219": "蛋壳钙化层/纳米结构",
    "10.1074/jbc.M610294200": "蛋壳/壳膜内侧（ovocalyxin-36）",
    "10.1074/jbc.M406033200": "蛋壳钙化层（ovocleidin-17）",
    "10.1016/j.fbio.2020.100590": "蛋清",
    "10.1016/j.fbio.2024.103938": "蛋清",
    "10.1016/j.ijbiomac.2020.08.193": "卵黄膜",
    "10.1021/acs.jafc.9b04638": "鸡蛋整体（孵化过程）",
    "10.1111/jfbc.14006": "受精蛋整体",
    "10.1042/BJ2400871": "卵黄高磷蛋白（phosvitin）",
    "10.1021/bi00298a027": "代表性鸡蛋蛋白（riboflavin-binding protein）",
    "10.1073/pnas.94.12.6244": "蛋清前体/输卵管 ovalbumin",
    "10.1021/jf048369l": "蛋清（ovalbumin gene Y）",
    "10.1021/acs.jafc.7b01706": "蛋清",
    "10.1016/S1044-0305(00)00122-7": "蛋清（ovalbumin 及共纯化糖蛋白）",
    "10.1016/S0021-9258(18)33585-3": "蛋清（ovomucoid）",
    "10.1007/s10719-011-9328-3": "蛋清（ovomucin）",
    "10.1021/acs.jafc.8b04492": "卵黄",
    "10.1021/acs.jafc.9b05133": "鸡蛋整体（胚胎发育）",
    "10.1021/acs.jafc.0c07557": "卵黄膜",
    "10.1271/bbb.70.144": "蛋清（ovalbumin-related gene Y）",
    "10.1016/j.fochx.2024.101244": "蛋清（皮蛋加工体系）",
    "10.1016/j.psj.2024.104526": "受精蛋清",
    "10.1016/j.psj.2024.104253": "卵黄浆",
}


DOI_PATTERN = re.compile(r"DOI:\s*(\S+)")
YEAR_PATTERN = re.compile(r"\.\s(\d{4})\.")
REFERENCE_PATTERN = re.compile(
    r"^(?P<authors>.+?)\.\s(?P<year>\d{4})\.\s(?P<title>.+?)\.\s(?P<journal>.+?)\.\sDOI:\s(?P<doi>\S+)$"
)


@dataclass
class Entry:
    reference: str
    summary: str
    quote: str
    doi: str
    year: int
    category: str
    authors: str
    title: str
    journal: str
    material: str


def normalize_doi(doi: str) -> str:
    return doi.strip().lower()


def directness_label(category: str) -> str:
    return DIRECTNESS_LABELS.get(category, "直接证据")


def parse_entries(document: Document) -> list[Entry]:
    entries: list[Entry] = []
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = (paragraph.text or "").strip()
        if "DOI:" not in text:
            continue

        doi_match = DOI_PATTERN.search(text)
        year_match = YEAR_PATTERN.search(text)
        if not doi_match or not year_match:
            raise ValueError(f"Unable to parse DOI or year from paragraph: {text}")

        doi = doi_match.group(1)
        reference_match = REFERENCE_PATTERN.match(text)
        if not reference_match:
            raise ValueError(f"Unable to parse reference fields from paragraph: {text}")

        summary = (paragraphs[index + 1].text or "").strip() if index + 1 < len(paragraphs) else ""
        quote = (paragraphs[index + 2].text or "").strip() if index + 2 < len(paragraphs) else ""

        if not summary.startswith("摘要要点：") or not quote.startswith("可引用表述（改写）："):
            raise ValueError(f"Entry notes are incomplete around DOI {doi}")

        year = int(year_match.group(1))
        entries.append(
            Entry(
                reference=text,
                summary=summary,
                quote=quote,
                doi=doi,
                year=year,
                category=DOI_TO_CATEGORY.get(doi, ""),
                authors=reference_match.group("authors"),
                title=reference_match.group("title"),
                journal=reference_match.group("journal"),
                material=MATERIAL_BY_DOI.get(doi, ""),
            )
        )
    return entries


def filter_and_group(entries: list[Entry]) -> dict[str, list[Entry]]:
    grouped = {category: [] for category in ORDERED_CATEGORIES}
    filtered = [entry for entry in entries if entry.year > 1999]

    missing_category = [entry.doi for entry in filtered if not entry.category]
    if missing_category:
        raise ValueError(f"Missing category mapping for DOIs: {missing_category}")

    for entry in filtered:
        grouped[entry.category].append(entry)

    for category_entries in grouped.values():
        category_entries.sort(key=lambda item: (item.year, item.reference))

    return grouped


def build_document(grouped_entries: dict[str, list[Entry]]) -> Document:
    document = Document()
    document.add_heading("蛋壳基质蛋白修饰相关文献清单", level=0)
    document.add_paragraph("筛选范围：仅保留 2000 年及以后文献，并按修饰类型重排")
    document.add_paragraph(
        "注：本版本以现有整理结果为基础，仅保留 2000 年及以后文献；正文优先按 N-糖基化、O-糖基化、磷酸化和多修饰联合组学编排，并单列“除糖基化和磷酸化外的其他修饰或修饰加工事件”说明，最后保留与修饰讨论直接相关的背景支持文献。文内保留 DOI，并延续原文档中的“摘要要点”和“可引用表述（改写）”。"
    )

    for category in ORDERED_CATEGORIES:
        if category == "support":
            document.add_heading(OTHER_MODIFICATION_HEADING, level=1)
            document.add_paragraph(OTHER_MODIFICATION_NOTE)
            for paragraph in OTHER_MODIFICATION_PARAGRAPHS:
                document.add_paragraph(paragraph)

        entries = grouped_entries[category]
        if not entries:
            continue

        metadata = CATEGORY_METADATA[category]
        document.add_heading(metadata["heading"], level=1)
        document.add_paragraph(metadata["note"])
        for entry in entries:
            document.add_paragraph(entry.reference)
            document.add_paragraph(entry.summary)
            document.add_paragraph(entry.quote)

    return document


def export_xlsx(grouped_entries: dict[str, list[Entry]]) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "PTM References"
    worksheet.append(["文章题目", "DOI", "修饰类型", "作者", "发表年限", "期刊", "样本来源/材料类型", "是否直接研究PTM"])

    for category in ORDERED_CATEGORIES:
        for entry in grouped_entries[category]:
            worksheet.append(
                [
                    entry.title,
                    entry.doi,
                    CATEGORY_LABELS[entry.category],
                    entry.authors,
                    entry.year,
                    entry.journal,
                    entry.material,
                    directness_label(entry.category),
                ]
            )

    worksheet.column_dimensions["A"].width = 80
    worksheet.column_dimensions["B"].width = 28
    worksheet.column_dimensions["C"].width = 20
    worksheet.column_dimensions["D"].width = 42
    worksheet.column_dimensions["E"].width = 12
    worksheet.column_dimensions["F"].width = 38
    worksheet.column_dimensions["G"].width = 28
    worksheet.column_dimensions["H"].width = 18
    workbook.save(str(XLSX_PATH))


def build_additional_entries() -> list[Entry]:
    entries: list[Entry] = []
    for item in ADDITIONAL_REFERENCES:
        entries.append(
            Entry(
                reference=f"{item['authors']}. {item['year']}. {item['title']}. DOI: {item['doi']}",
                summary="",
                quote="",
                doi=item["doi"],
                year=item["year"],
                category=item["category"],
                authors=item["authors"],
                title=item["title"],
                journal=item["journal"],
                material=item["material"],
            )
        )
    return entries


def merge_all_entries(parsed_entries: list[Entry]) -> list[Entry]:
    merged: dict[str, Entry] = {normalize_doi(entry.doi): entry for entry in parsed_entries}
    for entry in build_additional_entries():
        merged.setdefault(normalize_doi(entry.doi), entry)
    all_entries = list(merged.values())
    all_entries.sort(key=lambda item: (CATEGORY_LABELS[item.category], item.year, item.title))
    return all_entries


def export_all_literature_xlsx(entries: list[Entry]) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "All PTM Literature"
    worksheet.append(["文章题目", "DOI", "修饰类型", "作者", "发表年限", "期刊", "样本来源/材料类型", "是否直接研究PTM"])

    for entry in entries:
        worksheet.append(
            [
                entry.title,
                entry.doi,
                CATEGORY_LABELS[entry.category],
                entry.authors,
                entry.year,
                entry.journal,
                entry.material,
                directness_label(entry.category),
            ]
        )

    worksheet.column_dimensions["A"].width = 80
    worksheet.column_dimensions["B"].width = 28
    worksheet.column_dimensions["C"].width = 20
    worksheet.column_dimensions["D"].width = 46
    worksheet.column_dimensions["E"].width = 12
    worksheet.column_dimensions["F"].width = 38
    worksheet.column_dimensions["G"].width = 28
    worksheet.column_dimensions["H"].width = 18
    workbook.save(str(ALL_XLSX_PATH))


def main() -> None:
    source_document = Document(str(SOURCE_PATH))
    entries = parse_entries(source_document)
    grouped_entries = filter_and_group(entries)
    all_entries = merge_all_entries(entries)
    output_document = build_document(grouped_entries)
    output_document.save(str(TARGET_PATH))
    export_xlsx(grouped_entries)
    export_all_literature_xlsx(all_entries)

    removed = sorted((entry.year, entry.doi) for entry in entries if entry.year <= 1999)
    kept_count = sum(len(items) for items in grouped_entries.values())

    print(f"Created: {TARGET_PATH}")
    print(f"Created: {XLSX_PATH}")
    print(f"Created: {ALL_XLSX_PATH}")
    print(f"Kept entries: {kept_count}")
    print(f"All-literature entries: {len(all_entries)}")
    print("Removed entries (<=1999):")
    for year, doi in removed:
        print(f"- {year} | {doi}")
    print("Section counts:")
    for category in ORDERED_CATEGORIES:
        print(f"- {CATEGORY_METADATA[category]['heading']}: {len(grouped_entries[category])}")


if __name__ == "__main__":
    main()