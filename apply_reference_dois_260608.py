from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


BASE = Path.cwd()
WORK_DIR = next(BASE.glob("03_*"))
SRC = WORK_DIR / "manuscript260608.docx"
MATCHES = WORK_DIR / "manuscript260608_doi_matches.json"
OUT = WORK_DIR / "manuscript260608_citation_fixed.docx"
DOI_REPORT = WORK_DIR / "manuscript260608_doi_update_report.md"


MANUAL_DOIS = {
    3: "10.1016/j.cbpa.2008.11.020",
    11: "10.1002/jmor.21249",
}

TITLE_REPLACEMENTS = {
    3: (
        "Calcium transport in strongly calcifying laying birds: the role of calbindin and plasma membrane calcium ATPase.",
        "Calcium transport in strongly calcifying laying birds: mechanisms and regulation.",
    ),
    11: (
        "Do all geckos hatch in the same way? Histological and 3D studies of egg tooth morphogenesis in the geckos Eublepharis macularius and Lepidodactylus lugubris.",
        "Do all geckos hatch in the same way? Histological and 3D studies of egg tooth morphogenesis in the geckos Eublepharis macularius Blyth 1854 and Lepidodactylus lugubris Duméril & Bibron 1836.",
    ),
}

URL_ONLY = {
    7: "https://digitalcommons.usf.edu/wilson_bulletin/vol73/iss3/5",
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def append_doi(text: str, doi: str) -> str:
    if re.search(r"\bdoi:\s*10\.", text, re.I) or "https://doi.org/" in text.lower():
        return text
    text = text.rstrip()
    if text.endswith("."):
        text = text[:-1]
    return f"{text}. doi: {doi}."


def append_url(text: str, url: str) -> str:
    if url in text:
        return text
    text = text.rstrip()
    if text.endswith("."):
        text = text[:-1]
    return f"{text}. Available at: {url}."


def main() -> None:
    matches = json.loads(MATCHES.read_text(encoding="utf-8"))
    doi_by_ref = {
        item["number"]: item["match"]["doi"]
        for item in matches
        if item.get("accepted") and item.get("match") and item["match"].get("doi")
    }
    doi_by_ref.update(MANUAL_DOIS)

    doc = Document(str(SRC))
    updated = []
    url_added = []
    title_fixed = []
    in_refs = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "References":
            in_refs = True
            continue
        if not in_refs or not text:
            continue
        match = re.match(r"^(\d+)\.\s+(.+)$", text)
        if not match:
            continue
        number = int(match.group(1))
        new_text = text
        if number in TITLE_REPLACEMENTS:
            old, new = TITLE_REPLACEMENTS[number]
            if old in new_text:
                new_text = new_text.replace(old, new)
                title_fixed.append(number)
        if number in doi_by_ref:
            new_text = append_doi(new_text, doi_by_ref[number])
            updated.append(number)
        elif number in URL_ONLY:
            new_text = append_url(new_text, URL_ONLY[number])
            url_added.append(number)
        if new_text != text:
            replace_paragraph_text(paragraph, new_text)

    doc.save(str(OUT))
    DOI_REPORT.write_text(
        "\n".join(
            [
                "# DOI update report",
                "",
                f"Output: `{OUT.name}`",
                f"DOI added: {len(set(updated))} references",
                f"URL added for no-DOI legacy source: {url_added}",
                f"Title corrected: {title_fixed}",
                "",
                "No DOI was found for Ref. 7; a stable source page was added instead.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    print(OUT)
    print(DOI_REPORT)
    print("doi_added", len(set(updated)))
    print("url_added", url_added)
    print("title_fixed", title_fixed)


if __name__ == "__main__":
    main()
